import tkinter as tk
from tkinter import filedialog, messagebox, ttk, simpledialog
import os
import datetime
import subprocess
import logging
import logging.handlers  # Döngüsel log için gerekli
import json
from openpyxl import Workbook
from docx import Document
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
import threading
import sys
import webbrowser
import tempfile
import collections
import requests
import shutil
import random
import concurrent.futures  # For parallel processing
import multiprocessing  # For CPU-bound tasks (GIL bypass)
import time  # For performance measurement
import gc  # For garbage collection control
import platform  # For OS detection
from search_translations import search_translations, context_menu_translations, toggle_panel_translations

# Görüntü işleme kütüphaneleri
from PIL import Image, ImageTk, ImageDraw, ImageFont

# Log dosyası ayarları
def setup_logging():
    """
    Döngüsel log sistemi oluşturur. 
    Bu, log dosyasının belirli bir boyuta ulaştığında arşivlenmesini ve yeni bir log dosyası başlatılmasını sağlar.
    Böylece disk alanının dolması önlenir.
    """
    # Belgelerim klasörünü belirle (cross-platform desteği)
    documents_dir = os.path.join(os.path.expanduser('~'), 'Documents')

    # ListeKolay klasörü oluştur (yoksa)
    app_data_dir = os.path.join(documents_dir, 'ListeKolay')
    if not os.path.exists(app_data_dir):
        try:
            os.makedirs(app_data_dir)
        except Exception as e:
            # Oluşturulamazsa geçici dizini kullan
            import tempfile
            app_data_dir = tempfile.gettempdir()

    log_file = os.path.join(app_data_dir, "ListeKolay.log")

    # Döngüsel log yapılandırması
    log_formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s', datefmt='%Y-%m-%d %H:%M:%S')

    # Dosya tutucusu oluştur, maksimum boyut 5MB, 3 eski dosya arşivle
    file_handler = logging.handlers.RotatingFileHandler(
        log_file, 
        maxBytes=5*1024*1024,  # 5MB
        backupCount=3,  # 3 eski log dosyası sakla
        encoding='utf-8'
    )
    file_handler.setFormatter(log_formatter)

    # Kök log yapılandırıcısı
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)

    # Önceki tutucuları kaldır (eğer varsa)
    for handler in root_logger.handlers[:]:
        root_logger.removeHandler(handler)

    # Yeni tutucuyu ekle
    root_logger.addHandler(file_handler)

    # Log sisteminin başlatıldığını kaydet
    logging.info(f"Program başladı - Döngüsel log sistemi aktif (maks. 5MB, 3 arşiv) - Log dosyası: {log_file}")

# Log sistemini başlat
setup_logging()

# Basit sürükle-bırak desteği için sabit
DND_FILES = "DND_FILES"

# Tema renk sabitleri
# Açık Tema Renkleri
LIGHT_MODE_COLORS = {
    "bg": "#e9ecef",             # Açık gri arkaplan
    "text": "#000000",           # Tüm metin ve etiketler için siyah
    "secondary_text": "#000000", # İkincil metinler de siyah
    "accent": "#007bff",         # Mavi vurgu

    # Buton Renkleri (Açık mod)
    "folder_button": "#007bff",  # Klasör seç butonu: Mavi
    "exit_button": "#6c757d",    # Kapat butonu: Gri
    "cancel_button": "#dc3545",  # İptal butonu: Kırmızı  
    "start_button": "#28a745",   # Başlat butonu: Yeşil
    "filter_button": "#17a2b8",  # Filtrele butonu: Turkuaz

    # View mode butonları
    "active_view_button": "#17a2b8",   # Aktif görünüm butonu: Turkuaz
    "inactive_view_button": "#6c757d", # Pasif görünüm butonu: Koyu gri
    
    # Giriş alanları ve diyaloglar için renkler
    "entry_bg": "#ffffff",       # Giriş alanı arkaplan
    "btn_bg": "#007bff",         # Buton arkaplan
    "btn_fg": "#ffffff",         # Buton yazı rengi
    "btn_active_bg": "#0069d9",  # Buton aktif arkaplan
    "btn_active_fg": "#ffffff",  # Buton aktif yazı rengi

    # Ortak renkler
    "button_text": "#000000",     # Açık temada buton metinleri siyah
    "highlight": "#f8f9fa",       # Çok açık gri vurgu
    "border": "#ced4da",          # Açık gri kenarlık
    "error": "#dc3545",           # Kırmızı hata
    "success": "#28a745",         # Yeşil başarı
    "warning": "#ffc107"          # Sarı uyarı
}

# Koyu Tema Renkleri
DARK_MODE_COLORS = {
    "bg": "#212529",             # Koyu arkaplan
    "text": "#ffffff",           # Tüm metin ve etiketler için beyaz
    "secondary_text": "#ffffff", # İkincil metinler de beyaz
    "accent": "#0d6efd",         # Parlak mavi vurgu

    # Buton Renkleri (Koyu mod)
    "folder_button": "#007bff",  # Klasör seç butonu: Mavi
    "exit_button": "#6c757d",    # Kapat butonu: Gri
    "cancel_button": "#dc3545",  # İptal butonu: Kırmızı
    "start_button": "#28a745",   # Başlat butonu: Yeşil
    "filter_button": "#17a2b8",  # Filtrele butonu: Turkuaz

    # View mode butonları
    "active_view_button": "#17a2b8",   # Aktif görünüm butonu: Turkuaz
    "inactive_view_button": "#6c757d", # Pasif görünüm butonu: Koyu gri
    
    # Giriş alanları ve diyaloglar için renkler
    "entry_bg": "#343a40",       # Giriş alanı arkaplan
    "btn_bg": "#0d6efd",         # Buton arkaplan
    "btn_fg": "#ffffff",         # Buton yazı rengi
    "btn_active_bg": "#0b5ed7",  # Buton aktif arkaplan
    "btn_active_fg": "#ffffff",  # Buton aktif yazı rengi

    # Ortak renkler
    "button_text": "#ffffff",     # Koyu temada buton metinleri beyaz
    "highlight": "#2b3035",       # Hafif açık koyu gri vurgu
    "border": "#495057",          # Orta koyu gri kenarlık
    "error": "#dc3545",           # Kırmızı hata
    "success": "#28a745",         # Yeşil başarı
    "warning": "#ffc107"          # Sarı uyarı
}

# PIL konfigürasyonu
import warnings
from PIL import Image, ImageTk, ImageDraw
# Devre dışı bırak DecompressionBombWarning (EPS ve büyük resimler için)
warnings.simplefilter('ignore', Image.DecompressionBombWarning)
# PIL maksimum boyut limitini artır
Image.MAX_IMAGE_PIXELS = None

# Ön izlenebilir dosya uzantıları (küçük harflerle)
PREVIEWABLE_EXTENSIONS = [
    # Resim formatları
    '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.ico', '.svg', '.webp', '.tif', '.tiff',
    # Tasarım ve dokümantasyon
    '.pdf', '.eps', '.ai', '.psd',
    # Video formatları
    '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg', '.3gp',
    # RAW kamera formatları
    '.raw', '.cr2', '.nef', '.dng', '.arw',
    # HEIC/HEIF formatları
    '.heic', '.heif'
]

# PIL sürüm uyumluluğu için yardımcı fonksiyon
def get_pil_resize_method():
    """Farklı PIL sürümleri için tutarlı yeniden boyutlandırma yöntemi döndürür"""
    try:
        from PIL import Image, ImageFilter
        
        # Modern Pillow (9.1.0+) için Resampling sabitleri
        if hasattr(Image, 'Resampling') and hasattr(Image.Resampling, 'LANCZOS'):
            return Image.Resampling.LANCZOS
        
        # Pillow 4.0 - 8.x için
        if hasattr(Image, 'LANCZOS'):
            return Image.LANCZOS
            
        # PIL 1.1.3 - 3.x için
        if hasattr(Image, 'ANTIALIAS'):
            return Image.ANTIALIAS
        
        # Son çare - numeric value as that works in all versions
        return 3  # BICUBIC sabit değeri
    except ImportError:
        # If PIL is not available, return a fallback
        return 3

import fitz  # PyMuPDF
import io
import time
import gc
import pdf2image

# LOG AYARLARI
# Bu alan geriye dönük uyumluluk için tutuldu, ancak asıl log yapılandırması 
# setup_logging() fonksiyonu ile yapılıyor (uygulamanın en başında)
# Bu sayede log dosyası boyutu kontrol altında tutulacak

# Program başladı log mesajı setup_logging() içinde yazılıyor

# Import language dictionaries
try:
    from new_languages import de_dict, fr_dict, ru_dict, es_dict, it_dict, fa_dict, ur_dict, hi_dict, zh_dict, ja_dict
    from search_translations import search_translations
    
    # Create context_menu_translations if not found in search_translations
    if 'context_menu_translations' not in globals():
        context_menu_translations = {
            "preview_file": {
                "tr": "Dosyayı Önizle",
                "en": "Preview File",
                "de": "Datei-Vorschau",
                "fr": "Aperçu du fichier",
                "ru": "Предварительный просмотр файла",
                "es": "Vista previa del archivo",
                "it": "Anteprima file",
                "fa": "پیش نمایش فایل",
                "ur": "فائل کا پیش منظر",
                "hi": "फ़ाइल पूर्वावलोकन",
                "zh": "预览文件",
                "ja": "ファイルプレビュー",
                "ar": "معاينة الملف"
            },
            "delete_files": {
                "tr": "Dosyayı Sil",
                "en": "Delete File",
                "de": "Datei löschen",
                "fr": "Supprimer le fichier",
                "ru": "Удалить файл",
                "es": "Eliminar archivo",
                "it": "Elimina file",
                "fa": "حذف فایل",
                "ur": "فائل کو ڈیلیٹ کریں",
                "hi": "फ़ाइल हटाएं",
                "zh": "删除文件",
                "ja": "ファイルを削除",
                "ar": "حذف الملف"
            },
            "copy_files": {
                "tr": "Dosyayı Kopyala",
                "en": "Copy File",
                "de": "Datei kopieren",
                "fr": "Copier le fichier",
                "ru": "Копировать файл",
                "es": "Copiar archivo",
                "it": "Copia file",
                "zh": "复制文件",
                "ja": "ファイルをコピー",
                "ar": "نسخ الملف"
            },
            "move_files": {
                "tr": "Dosyayı Taşı",
                "en": "Move File",
                "de": "Datei verschieben",
                "fr": "Déplacer le fichier",
                "ru": "Переместить файл",
                "es": "Mover archivo",
                "it": "Sposta file",
                "zh": "移动文件",
                "ja": "ファイルを移动",
                "ar": "نقل الملف"
            },
            "rename_file": {
                "tr": "Yeniden Adlandır",
                "en": "Rename",
                "de": "Umbenennen",
                "fr": "Renommer",
                "ru": "Переименовать",
                "es": "Renombrar",
                "it": "Rinomina",
                "zh": "重命名",
                "ja": "名前を変更",
                "ar": "إعادة تسمية"
            },
            "select_all_files": {
                "tr": "Tümünü Seç",
                "en": "Select All",
                "de": "Alle auswählen",
                "fr": "Tout sélectionner",
                "ru": "Выбрать все",
                "es": "Seleccionar todo",
                "it": "Seleziona tutto",
                "zh": "全选",
                "ja": "すべて選択",
                "ar": "تحديد الكل"
            }
        }
except ImportError:
    # Try with direct import without path
    import new_languages
    from new_languages import de_dict, fr_dict, ru_dict, es_dict, it_dict, fa_dict, ur_dict, hi_dict, zh_dict, ja_dict
    
    # Try to import search_translations
    try:
        import search_translations
        from search_translations import search_translations
    except ImportError:
        # Fallback search translations
        search_translations = {
            "tr": "Dosyaları ara...",
            "en": "Search files..."
        }
    
    # Fallback context menu translations
    context_menu_translations = {
        "preview_file": {
            "tr": "Dosyayı Önizle",
            "en": "Preview File",
            "de": "Datei-Vorschau",
            "fr": "Aperçu du fichier",
            "ru": "Предварительный просмотр файла",
            "es": "Vista previa del archivo",
            "it": "Anteprima file",
            "zh": "预览文件",
            "ja": "ファイルプレビュー",
            "ar": "معاينة الملف"
        },
        "delete_files": {
            "tr": "Dosyayı Sil",
            "en": "Delete File",
            "de": "Datei löschen",
            "fr": "Supprimer le fichier",
            "ru": "Удалить файл",
            "es": "Eliminar archivo",
            "it": "Elimina file",
            "zh": "删除文件",
            "ja": "ファイルを削除",
            "ar": "حذف الملف"
        },
        "copy_files": {
            "tr": "Dosyayı Kopyala",
            "en": "Copy File",
            "de": "Datei kopieren",
            "fr": "Copier le fichier",
            "ru": "Копировать файл",
            "es": "Copiar archivo",
            "it": "Copia file",
            "zh": "复制文件",
            "ja": "ファイルをコピー",
            "ar": "نسخ الملف"
        },
        "move_files": {
            "tr": "Dosyayı Taşı",
            "en": "Move File",
            "de": "Datei verschieben",
            "fr": "Déplacer le fichier",
            "ru": "Переместить файл",
            "es": "Mover archivo",
            "it": "Sposta file",
            "zh": "移动文件",
            "ja": "ファイルを移动",
            "ar": "نقل الملف"
        },
        "rename_file": {
            "tr": "Yeniden Adlandır",
            "en": "Rename",
            "de": "Umbenennen",
            "fr": "Renommer",
            "ru": "Переименовать",
            "es": "Renombrar",
            "it": "Rinomina",
            "zh": "重命名",
            "ja": "名前を変更",
            "ar": "إعادة تسمية"
        },
        "select_all_files": {
            "tr": "Tümünü Seç",
            "en": "Select All",
            "de": "Alle auswählen",
            "fr": "Tout sélectionner",
            "ru": "Выбрать все",
            "es": "Seleccionar todo",
            "it": "Seleziona tutto",
            "zh": "全选",
            "ja": "すべて選択",
            "ar": "تحديد الكل"
        }
    }

translations = {
    "tr": {
        "open_file": "Dosyayı Aç",
        "open_file_location": "Dosya Konumunu Aç",
        "copy_filename": "Dosya Adını Kopyala",
        "copy_filepath": "Dosya Yolunu Kopyala",
        "no_files_to_select": "Seçilecek dosya bulunamadı",
        "selection_error": "Seçim hatası",
        "files_selected": "dosya seçildi",
        "select_folder": "📁 Klasör Seç",
        "no_folder_selected": "Henüz bir klasör seçilmedi",
        "start": "▶️ Başlat",
        "apply_filter": "🔍 Filtreyi Uygula",
        "cancel": "⏹️ İptal",
        "cancelling": "⏹️ İptal Ediliyor...",
        "exit": "✖️ Kapat",
        "select_all": "Tümünü Seç",
        "clear_all": "Temizle",
        "select_all_files": "Tümünü Seç",
        "delete_files": "Sil",
        "copy_files": "Kopyala",
        "move_files": "Taşı",
        "rename_file": "Yeniden Adlandır",
        "cut_files": "Kes",
        "file_operation_progress": "İşlem İlerlemesi: %{percent}% (%{current}/%{total})",
        "confirm_delete_title": "Silme Onayı",
        "confirm_delete_message": "Seçili dosyayı/dosyaları silmek istediğinizden emin misiniz?",
        "confirm_cut_title": "Kesme Onayı",
        "confirm_cut_message": "Seçili dosyayı/dosyaları kesmek istediğinizden emin misiniz?",
        "select_destination_folder": "Hedef Klasörü Seçin",
        "all_files_tip": "Tüm dosyalar uzantılarına bakılmaksızın listelenecek.",
        "filter_tip": "Filtreleme seçenekleri etkin, sadece seçilen uzantılara sahip dosyalar gösterilecek.",
        "calculating_statistics": "İstatistikler hesaplanıyor...",
        "loading_file_list": "Dosya listesi yükleniyor...",
        "tooltip_select": "Dosya listesi oluşturmak için bir klasör seçin",
        "tooltip_start": "Dosya listesini oluştur ve dışa aktar",
        "tooltip_apply": "Seçilen uzantılara göre dosyaları filtrele",
        "tooltip_cancel": "Devam eden işlemi iptal et",
        "tooltip_exit": "Uygulamayı kapat",
        "language": "Dil / Language",
        "view_mode_list": "Listele",
        "view_mode_preview": "Ön İzleme",
        "tooltip_list_view": "Dosyaları liste görünümünde göster",
        "tooltip_preview_view": "Dosyaları önizleme görünümünde göster",
        "files_and_previews": "Dosyalar ve Önizlemeler",
        "search_files": "Dosya ara...",
        "preview_mode_active": "Önizleme modu etkin",
        "loading_preview": "Önizlemeler yükleniyor...",
        "no_preview_available": "Önizlenebilir dosya bulunamadı",
        "preview_file": "Dosyayı Önizle",
        "open_file": "Dosyayı Aç",
        "open_file_location": "Dosya Konumunu Aç",
        "error_open_file": "Dosya açılırken bir hata oluştu",
        "error_open_location": "Dosya konumu açılırken bir hata oluştu",
        "error_open_url": "URL açılırken bir hata oluştu",
        "extension_not_found": "Uzantısı bulunamadı",
        "settings": "Ayarlar",
        "info": "Bilgi",
        "files_filtering": "Dosyalar filtreleniyor, lütfen bekleyin...",
        "files_gathering": "Dosyalar toplanıyor, lütfen bekleyin...",
        "files_loading": "Dosyalar Yükleniyor...",
        "filter_applying": "Filtre uygulanıyor...",
        "filter_cancelled": "Filtreleme işlemi iptal edildi",
        "excel_created": "Excel dosyası başarıyla oluşturuldu...",
        "error_occurred": "Bir hata oluştu: {0}",
        "error": "Hata",
        "ready": "Hazır",
        "operation_cancelled": "İşlem İptal Edildi",
        "file_processed": "İşlenen dosya: {0}/{1} • {2}",
        "prev_page": "Önceki",
        "next_page": "Sonraki",
        "page": "Sayfa",
        "processing": "İşleniyor: %{0:.1f}",
        "folder_loading": "Klasör yükleniyor...",
        "loading_subfolders": "Alt klasörler yükleniyor...",
        "folder_loaded_status": "📁 Klasör içeriği yüklendi. Liste oluşturmak için \"Başlat\" butonuna tıklayın.",
        "create_list_time": "🕒 Liste Oluşturulma Zamanı: {0}",
        "select_folder_first": "Lütfen önce bir klasör seçin!",
        "text_file_error": "Metin dosyası oluşturulamadı: {0}",
        "start_processing": "▶️ Liste oluşturuluyor...",
        "confirm_exit_title": "Çıkış",
        "confirm_exit_message": "Programdan çıkmak istediğinize emin misiniz?",
        "no_files_found": "Seçilen klasörde dosya bulunamadı.",
        "files_loaded_message": "{0} dosya yüklendi.",
        "files_filtered_message": "{0} dosya filtrelendi.",
        "filter_saved_message": "Filtreleme ayarları kaydedildi. Bir klasör seçtiğinizde uygulanacak.",
        "app_title": "ListeKolay - Dosya Listesi Oluşturucu",
        "app_subtitle": "Klasörlerinizdeki dosyaları hızlıca listeyin.",
        "full_window_title": "ListeKolay - Dosya Listesi Oluşturucu",
        "statistics_header": "İstatistikler",
        "total_files_label": "Toplam Dosya:",
        "folder_count_label": "Klasör Sayısı:",
        "total_size_label": "Toplam Boyut:",
        "tips_header": "İpuçları",
        "settings_header": "Ayarlar",
        "operation_status": "İşlem Durumu",
        "subfolders_label": "Alt Klasörler:",
        "include_label": "Dahil Et",
        "list_format_label": "Liste Formatı:",
        "save_location_label": "Kaydetme Yeri:",
        "desktop_label": "Masaüstü",
        "sort_criteria_label": "Sıralama Ölçütü:",
        "text_format_info": "Metin dosyası formatı",
        "excel_format_info": "Excel çalışma kitabı formatı",
        "word_format_info": "Word belgesi formatı",
        "html_format_info": "Web sayfası formatı",
        "filter_label": "Filtrele",
        "kategori_header": "Kategoriler",
        "extensions_header": "Dosya Uzantıları",
        "tip_1": "Alt klasör kutusunu işaretleyerek tüm alt dizinlerdeki dosyaları tarayabilirsiniz.",
        "tip_3": "Dosyaları isim, boyut veya uzantıya göre sıralayabilirsiniz.",
        "tip_4": "Oluşturulan listeler varsayılan olarak program klasörüne kaydedilir.",
        "tip_5": "Masaüstü seçeneğiyle dosyaları doğrudan masaüstüne kaydedebilirsiniz.",
        "tip_6": "İşlemi durdurmak için İptal düğmesini kullanın.",
        "tip_preview_formats": "PDF, JPG, PNG, GIF, PSD, AI, EPS dosyaları ön izleme özelliğine sahiptir.",
        "preview_file": "Dosyayı Önizle",
        "preview_window_title": "Dosya Önizleme",
        "preview_not_supported": "Bu dosya türü için önizleme desteklenmiyor.",
        "preview_error": "Dosya önizleme sırasında bir hata oluştu.",
        "file_list_section": "Dosyalar ve Önizlemeler",
        "list_view": "Listele",
        "preview_view": "Ön İzleme",
        "view_mode": "Görünüm Modu:",
        "preview_mode_active": "Ön izleme modu aktif",
        "no_preview_available": "Önizleme Yok",
        "loading_preview": "Önizleme yükleniyor...",
        "all_files": "Tüm Dosyalar",
        "image_files": "Görsel",
        "audio_files": "Ses",
        "video_files": "Video ve Ses",
        "text_files": "Metin",
        "compressed_files": "Sıkıştırılmış",
        "spreadsheet_files": "Hesap Tablosu",
        "presentation_files": "Sunum",
        "design_files": "Tasarım",
        "sort_name_asc": "Dosya Adı - Artan",
        "sort_name_desc": "Dosya Adı - Azalan",
        "sort_ext_asc": "Dosya Uzantısı - Artan",
        "sort_ext_desc": "Dosya Uzantısı - Azalan",
        "sort_size_asc": "Dosya Boyutu - Artan",
        "sort_size_desc": "Dosya Boyutu - Azalan",
        "sort_dir_asc": "Dosya Dizini - Artan",
        "selected_folder": "Seçilen Klasör:",
        "file_list": "Dosya Listesi",
        "row_number": "Sıra No",
        "file_name": "Dosya Adı",
        "file_type": "Dosya Türü",
        "file_path": "Dosya Yolu",
        "file_size": "Dosya Boyutu",
        "creation_date": "Oluşturulma Tarihi",
        "modification_date": "Değiştirilme Tarihi",
        "file_extension": "Dosya Uzantısı",
        "creation_time": "Oluşturulma Zamanı:",
        "sorted_by": "Sıralama Kriteri:",
        "excel_success": "Excel dosyası başarıyla oluşturuldu",
        "word_success": "Word belgesi başarıyla oluşturuldu",
        "html_success": "HTML dosyası başarıyla oluşturuldu",
        "text_success": "Metin dosyası başarıyla oluşturuldu",
        "open_file_title": "Dosyayı Aç",
        "open_file_message": "Dosyayı açmak ister misiniz?",
        "yes": "Evet",
        "no": "Hayır",
        "tooltip_subfolders": "Alt klasörlerdeki dosyaları da listeler",
        "tooltip_format": "Listenizin hangi formatta kaydedileceğini seçin",
        "tooltip_save_location": "Listenizin nereye kaydedileceğini seçin",
        "tooltip_sort_criteria": "Dosyaların hangi kritere göre sıralanacağını seçin",
        "tooltip_select_all": "Tüm dosya uzantılarını seçer",
        "tooltip_clear_all": "Tüm dosya uzantı seçimlerini temizler",
        "tooltip_file_category": "Dosya kategorilerini görüntülemek için tıklayın",
        "tooltip_file_extension": "Listelenecek dosya uzantılarını seçin",
        "tooltip_filter_apply": "Seçilen filtreleri uygulamak için tıklayın",
        "tooltip_all_files": "Tüm dosya uzantılarını seçer/temizler",
        "tooltip_category_expand": "Bu kategoriyi genişletmek/daraltmak için tıklayın",
        "tooltip_select_category": "Bu kategorideki tüm uzantıları seçer/temizler",
        "select_all_category": "Bu kategorideki tümünü seç",
        "media_files": "Medya Dosyaları (Ses ve Video)",
        "extension_search": "Ara-Bul",
        "no_files_to_select": "Seçilecek dosya bulunamadı",
        "selection_error": "Seçim hatası",
        "files_selected": "dosya seçildi",
        "delete_confirmation_title": "Silmeyi Onayla",
        "delete_confirmation_message": "dosya kalıcı olarak silinecek. Devam edilsin mi?",
        "files_deleted": "dosya silindi",
        "delete_error_title": "Silme Hatası",
        "select_target_folder": "Hedef Klasör Seçin",
        "files_copied": "dosya kopyalandı",
        "copy_error_title": "Kopyalama Hatası",
        "files_moved": "dosya taşındı",
        "move_error_title": "Taşıma Hatası",
        "rename_error_title": "Yeniden Adlandırma Hatası",
        "select_single_file": "Lütfen yeniden adlandırmak için tek bir dosya seçin",
        "rename_title": "Dosyayı Yeniden Adlandır",
        "enter_new_name": "Yeni dosya adını girin:",
        "file_renamed": "Dosya yeniden adlandırıldı",
        "copyright_footer": "© {year} Muallimun.Net - ListeKolay",
        "document_files": "Doküman",
        "code_files": "Kodlama",
        "program_files": "Program",
        "filtering_in_progress": "Filtreleme işlemi devam ediyor",
        "filter_complete": "Filtreleme tamamlandı",
        "filter_error": "Filtreleme hatası",
        "filter_error_details": "Filtreleme sırasında hata oluştu",
        "check_updates": "Güncellemeleri Kontrol Et",
        "update_available": "Güncelleme Mevcut",
        "update_available_message": "ListeKolay'ın yeni sürümü mevcut: {0}\nMevcut sürümünüz: {1}\n\nGüncellemeyi indirmek ister misiniz?",
        "no_update_available": "Güncelleme Yok",
        "no_update_available_message": "ListeKolay'ın en son sürümünü kullanıyorsunuz.",
        "update_check_error": "Güncelleme Kontrolü Hatası",
        "update_check_error_message": "Güncellemeler kontrol edilirken bir hata oluştu. Lütfen internet bağlantınızı kontrol edin ve tekrar deneyin.",
        "downloading_update": "Güncelleme İndiriliyor...",
        "downloading_update_message": "{0} sürümü indiriliyor...",
        "download_complete": "İndirme Tamamlandı",
        "download_complete_message": "Güncelleme başarıyla indirildi. Program, güncellemeyi uygulamak için yeniden başlatılacak.",
        "download_error": "İndirme Hatası",
        "download_error_message": "Güncelleme indirilirken hata oluştu: {0}",
        "download_button": "İndir",
        "cancel_button": "İptal",
        "light_mode": "Açık Mod",
        "dark_mode": "Koyu Mod",
        "theme_settings": "Tema Ayarları"


    },
    "en": {
        "select_folder": "📁 Select Folder",
        "extension_not_found": "No extension found",
        "copy_filename": "Copy File Name",
        "copy_filepath": "Copy File Path",
        "no_folder_selected": "No folder selected yet",
        "start": "▶️ Start",
        "apply_filter": "🔍 Apply Filter",
        "cancel": "⏹️ Cancel",
        "cancelling": "⏹️ Cancelling...",
        "exit": "✖️ Exit",
        "select_all": "Select All",
        "clear_all": "Clear All",
        "no_files_to_select": "No files to select",
        "selection_error": "Selection error",
        "files_selected": "files selected",
        "delete_confirmation_title": "Confirm Delete",
        "delete_confirmation_message": "files will be permanently deleted. Continue?",
        "files_deleted": "files deleted",
        "delete_error_title": "Delete Error",
        "select_target_folder": "Select Target Folder",
        "files_copied": "files copied",
        "copy_error_title": "Copy Error",
        "files_moved": "files moved",
        "move_error_title": "Move Error",
        "rename_error_title": "Rename Error",
        "select_single_file": "Please select only one file to rename",
        "rename_title": "Rename File",
        "enter_new_name": "Enter new file name:",
        "file_renamed": "File renamed",
        "all_files_tip": "All files will be listed regardless of extension.",
        "filter_tip": "Filter options enabled, only files with selected extensions will be shown.",
        "calculating_statistics": "Calculating statistics...",
        "loading_file_list": "Loading file list...",
        "tooltip_select": "Select a folder to generate the file list",
        "tooltip_start": "Create and export the file list",
        "tooltip_apply": "Filter files by selected extensions",
        "tooltip_cancel": "Cancel ongoing operation",
        "tooltip_exit": "Exit the application",
        "language": "Language / Dil",
        "view_mode_list": "List View",
        "view_mode_preview": "Preview",
        "tooltip_list_view": "Show files in list view",
        "tooltip_preview_view": "Show files with preview thumbnails",
        "files_and_previews": "Files and Previews",
        "search_files": "Search files...",
        "preview_mode_active": "Preview mode active",
        "loading_preview": "Loading previews...",
        "no_preview_available": "No previewable files found",
        "preview_file": "Preview File",
        "open_file": "Open File",
        "open_file_location": "Open File Location",
        "error_open_file": "Error opening file",
        "error_open_location": "Error opening file location",
        "error_open_url": "Error opening URL",
        "settings": "Settings",
        "info": "Info",
        "files_filtering": "Filtering files, please wait...",
        "files_gathering": "Gathering files, please wait...",
        "files_loading": "Loading files...",
        "filter_applying": "Applying filter...",
        "filter_cancelled": "Filtering operation cancelled",
        "excel_created": "Excel file created successfully...",
        "error_occurred": "An error occurred: {0}",
        "error": "Error",
        "ready": "Ready",
        "operation_cancelled": "Operation Cancelled",
        "file_processed": "Processing file: {0}/{1} • {2}",
        "prev_page": "Previous",
        "next_page": "Next",
        "page": "Page",
        "processing": "Processing: %{0:.1f}",
        "folder_loading": "Loading folder...",
        "loading_subfolders": "Loading subfolders...",
        "folder_loaded_status": "📁 Folder loaded. Click \"Start\" to generate the list.",
        "create_list_time": "🕒 List Creation Time: {0}",
        "select_folder_first": "Please select a folder first!",
        "text_file_error": "Text file could not be created: {0}",
        "start_processing": "▶️ Creating list...",
        "confirm_exit_title": "Exit",
        "confirm_exit_message": "Are you sure you want to exit the program?",
        "no_files_found": "No files found in the selected folder.",
        "files_loaded_message": "{0} files loaded.",
        "files_filtered_message": "{0} files filtered.",
        "filter_saved_message": "Filter settings saved. Will be applied when a folder is selected.",
        "app_title": "EasyLister - File List Generator",
        "app_subtitle": "Quickly list the files in your folders.",
        "full_window_title": "EasyLister - File List Generator",
        "statistics_header": "Statistics",
        "total_files_label": "Total Files:",
        "folder_count_label": "Number of Folders:",
        "total_size_label": "Total Size:",
        "tips_header": "Tips",
        "settings_header": "Settings",
        "operation_status": "Operation Status",
        "subfolders_label": "Include Subfolders:",
        "include_label": "Include",
        "list_format_label": "List Format:",
        "save_location_label": "Save Location:",
        "desktop_label": "Desktop",
        "sort_criteria_label": "Sorting Criteria:",
        "text_format_info": "Text file format",
        "excel_format_info": "Excel workbook format",
        "word_format_info": "Word document format",
        "html_format_info": "Web page format",
        "filter_label": "Filter",
        "kategori_header": "Categories",
        "extensions_header": "File Extensions",
        "tip_1": "Check the subfolder box to scan all files in subdirectories.",
        "tip_3": "Files can be sorted by name, size or extension.",
        "tip_4": "Lists are saved to the program folder by default.",
        "tip_5": "Use the desktop option to save files directly to your desktop.",
        "tip_6": "Use the Cancel button to stop any operation immediately.",
        "tip_preview_formats": "PDF, JPG, PNG, GIF, PSD, AI, EPS files support preview functionality.",
        "preview_file": "Preview File",
        "preview_window_title": "File Preview",
        "preview_not_supported": "Preview is not supported for this file type.",
        "preview_error": "An error occurred while previewing the file.",
        "list_view": "List",
        "preview_view": "Preview",
        "view_mode": "View Mode:",
        "preview_mode_active": "Preview mode active",
        "no_preview_available": "No Preview",
        "loading_preview": "Loading preview...",
        "all_files": "All Files",
        "image_files": "Images",
        "audio_files": "Audio",
        "video_files": "Video & Audio",
        "text_files": "Text",
        "program_files": "Program & Archive",
        "compressed_files": "Compressed",
        "document_files": "Documents",
        "spreadsheet_files": "Spreadsheets",
        "presentation_files": "Presentations",
        "sort_name_asc": "File Name - Ascending",
        "sort_name_desc": "File Name - Descending",
        "sort_ext_asc": "File Extension - Ascending",
        "sort_ext_desc": "File Extension - Descending",
        "sort_size_asc": "File Size - Ascending",
        "sort_size_desc": "File Size - Descending",
        "sort_dir_asc": "File Directory - Ascending",
        "selected_folder": "Selected Folder:",
        "file_list": "File List",
        "row_number": "Row No",
        "file_name": "File Name",
        "file_type": "File Type",
        "file_path": "File Path",
        "file_size": "File Size",
        "creation_date": "Creation Date",
        "modification_date": "Modification Date",
        "file_extension": "File Extension",
        "creation_time": "Creation Time:",
        "sorted_by": "Sorted by:",
        "excel_success": "Excel file successfully created",
        "word_success": "Word document successfully created",
        "html_success": "HTML file successfully created",
        "text_success": "Text file successfully created",
        "open_file_title": "Open File",
        "open_file_message": "Would you like to open the file?",
        "yes": "Yes",
        "no": "No",
        "file_list_section": "Files and Previews",
        "tooltip_subfolders": "Also lists files in subfolders",
        "tooltip_format": "Choose the format in which your list will be saved",
        "tooltip_save_location": "Choose where to save your list",
        "tooltip_sort_criteria": "Choose how files will be sorted",
        "tooltip_select_all": "Select all file extensions",
        "tooltip_clear_all": "Clear all file extension selections",
        "tooltip_file_category": "Click to view file categories",
        "tooltip_file_extension": "Select file extensions to be listed",
        "tooltip_filter_apply": "Click to apply selected filters",
        "tooltip_all_files": "Select/deselect all file extensions",
        "tooltip_category_expand": "Click to expand/collapse this category",
        "tooltip_select_category": "Select/deselect all extensions in this category",
        "select_all_category": "Select all in this category",
        "media_files": "Media",
        "code_files": "Code & Web",
        "design_files": "Design",
        "game_files": "Games",
        "extension_search": "Search:",
        "copyright_footer": "© {year} Muallimun.Net - ListeKolay",
        "filtering_in_progress": "Filtering in progress",
        "filter_complete": "Filtering complete",
        "filter_error": "Filtering error",
        "filter_error_details": "Error occurred during filtering",
        "check_updates": "Check for Updates",
        "update_available": "Update Available",
        "update_available_message": "A new version of ListeKolay is available: {0}\nYour current version: {1}\n\nWould you like to download the update?",
        "no_update_available": "No Update Available",
        "no_update_available_message": "You are using the latest version of ListeKolay.",
        "update_check_error": "Update Check Error",
        "update_check_error_message": "An error occurred while checking for updates. Please check your internet connection and try again.",
        "downloading_update": "Downloading Update...",
        "downloading_update_message": "Downloading version {0}...",
        "download_complete": "Download Complete",
        "download_complete_message": "Update has been downloaded successfully. The program will restart to apply the update.",
        "download_error": "Download Error",
        "download_error_message": "Failed to download the update: {0}",
        "download_button": "Download",
        "cancel_button": "Cancel",
        "light_mode": "Light Mode",
        "dark_mode": "Dark Mode",
        "theme_settings": "Theme Settings"

    },
    "ar": {
        "open_file": "فتح الملف",
        "open_file_location": "فتح موقع الملف",
        "copy_filename": "نسخ اسم الملف",
        "copy_filepath": "نسخ مسار الملف",
        "select_folder": "📁 اختر مجلد",
        "no_folder_selected": "لم يتم اختيار مجلد بعد",
        "start": "▶️ ابدأ",
        "apply_filter": "🔍 تطبيق التصفية",
        "cancel": "⏹️ إلغاء",
        "cancelling": "⏹️ جاري الإلغاء...",
        "exit": "✖️ خروج",
        "select_all": "تحديد الكل",
        "clear_all": "مسح الكل",
        "all_files_tip": "سيتم سرد جميع الملفات بغض النظر عن الامتداد.",
        "filter_tip": "خيارات التصفية مفعلة، سيتم عرض الملفات ذات الامتدادات المحددة فقط.",
        "calculating_statistics": "جاري حساب الإحصائيات...",
        "loading_file_list": "جاري تحميل قائمة الملفات...",
        "tooltip_select": "حدد مجلد لإنشاء قائمة الملفات",
        "tooltip_start": "إنشاء وتصدير قائمة الملفات",
        "tooltip_apply": "تصفية الملفات حسب الامتدادات المحددة",
        "tooltip_cancel": "إلغاء العملية الجارية",
        "tooltip_exit": "الخروج من التطبيق",
        "language": "اللغة / Language",
        "search_files": "البحث عن الملفات...",
        "settings": "إعدادات",
        "info": "معلومات",
        "files_filtering": "جاري تصفية الملفات، يرجى الانتظار...",
        "files_gathering": "جاري تجميع الملفات، يرجى الانتظار...",
        "files_loading": "جاري تحميل الملفات...",
        "filter_applying": "جاري تطبيق التصفية...",
        "filter_cancelled": "تم إلغاء عملية التصفية",
        "excel_created": "تم إنشاء ملف إكسل بنجاح...",
        "error_occurred": "حدث خطأ: {0}",
        "error": "خطأ",
        "ready": "جاهز",
        "operation_cancelled": "تم إلغاء العملية",
        "file_processed": "معالجة الملف: {0}/{1} • {2}",
        "prev_page": "السابق",
        "next_page": "التالي",
        "page": "صفحة",
        "processing": "المعالجة: %{0:.1f}",
        "folder_loading": "جاري تحميل المجلد...",
        "loading_subfolders": "جاري تحميل المجلدات الفرعية...",
        "folder_loaded_status": "📁 تم تحميل المجلد. انقر على \"ابدأ\" لإنشاء القائمة.",
        "create_list_time": "🕒 وقت إنشاء القائمة: {0}",
        "select_folder_first": "الرجاء تحديد مجلد أولاً!",
        "text_file_error": "تعذر إنشاء ملف نصي: {0}",
        "start_processing": "▶️ جاري إنشاء القائمة...",
        "confirm_exit_title": "خروج",
        "confirm_exit_message": "هل أنت متأكد أنك تريد الخروج من البرنامج؟",
        "no_files_found": "لم يتم العثور على ملفات في المجلد المحدد.",
        "files_loaded_message": "تم تحميل {0} ملف.",
        "files_filtered_message": "تم تصفية {0} ملف.",
        "filter_saved_message": "تم حفظ إعدادات التصفية. سيتم تطبيقها عند تحديد مجلد.",
        "app_title": "قوائم سهلة - منشئ قوائم الملفات",
        "app_subtitle": "قم بسرد الملفات في مجلداتك بسرعة.",
        "full_window_title": "قوائم سهلة - منشئ قوائم الملفات",
        "statistics_header": "إحصائيات",
        "total_files_label": "إجمالي الملفات:",
        "folder_count_label": "عدد المجلدات:",
        "total_size_label": "الحجم الإجمالي:",
        "tips_header": "نصائح",
        "settings_header": "إعدادات",
        "operation_status": "حالة العملية",
        "subfolders_label": "تضمين المجلدات الفرعية:",
        "include_label": "تضمين",
        "list_format_label": "تنسيق القائمة:",
        "save_location_label": "موقع الحفظ:",
        "desktop_label": "سطح المكتب",
        "sort_criteria_label": "معيار الترتيب:",
        "text_format_info": "تنسيق ملف نصي",
        "excel_format_info": "تنسيق مصنف إكسل",
        "word_format_info": "تنسيق مستند وورد",
        "html_format_info": "تنسيق صفحة ويب",
        "filter_label": "تصفية",
        "kategori_header": "الفئات",
        "extensions_header": "امتدادات الملفات",
        "tip_1": "حدد مربع المجلدات الفرعية لمسح جميع الملفات في الدلائل الفرعية.",
        "tip_3": "يمكن فرز الملفات حسب الاسم أو الحجم أو الامتداد.",
        "tip_4": "يتم حفظ القوائم في مجلد البرنامج بشكل افتراضي.",
        "tip_5": "استخدم خيار سطح المكتب لحفظ الملفات مباشرة على سطح المكتب.",
        "tip_6": "استخدم زر الإلغاء لإيقاف أي عملية على الفور.",
        "tip_preview_formats": "ملفات PDF و JPG و PNG و GIF و PSD و AI و EPS تدعم وظيفة المعاينة.",
        "all_files": "جميع الملفات",
        "image_files": "صور",
        "audio_files": "صوت",
        "video_files": "فيديو و صوت",
        "text_files": "نصوص",
        "code_files": "برمجة و ويب",
        "data_files": "بيانات",
        "document_files": "مستندات",
        "spreadsheet_files": "جداول بيانات",
        "presentation_files": "عروض تقديمية",
        "program_files": "برامج و أرشيف",
        "compressed_files": "ملفات مضغوطة",
        "sort_name_asc": "اسم الملف - تصاعدي",
        "sort_name_desc": "اسم الملف - تنازلي",
        "sort_ext_asc": "امتداد الملف - تصاعدي",
        "sort_ext_desc": "امتداد الملف - تنازلي",
        "sort_size_asc": "حجم الملف - تصاعدي",
        "sort_size_desc": "حجم الملف - تنازلي",
        "sort_dir_asc": "دليل الملف - تصاعدي",
        "selected_folder": "المجلد المحدد:",
        "file_list": "قائمة الملفات",
        "row_number": "رقم الصف",
        "file_name": "اسم الملف",
        "file_type": "نوع الملف",
        "file_path": "مسار الملف",
        "file_size": "حجم الملف",
        "creation_date": "تاريخ الإنشاء",
        "modification_date": "تاريخ التعديل",
        "file_extension": "امتداد الملف",
        "creation_time": "وقت الإنشاء:",
        "sorted_by": "تم الفرز حسب:",
        "excel_success": "تم إنشاء ملف إكسل بنجاح",
        "word_success": "تم إنشاء مستند وورد بنجاح",
        "html_success": "تم إنشاء ملف HTML بنجاح",
        "text_success": "تم إنشاء ملف نصي بنجاح",
        "open_file_title": "فتح الملف",
        "open_file_message": "هل ترغب في فتح الملف؟",
        "yes": "نعم",
        "no": "لا",
        "tooltip_subfolders": "يسرد أيضًا الملفات في المجلدات الفرعية",
        "tooltip_format": "اختر التنسيق الذي سيتم حفظ قائمتك به",
        "tooltip_save_location": "اختر مكان حفظ قائمتك",
        "tooltip_sort_criteria": "اختر كيفية فرز الملفات",
        "tooltip_select_all": "تحديد جميع امتدادات الملفات",
        "tooltip_clear_all": "مسح جميع تحديدات امتداد الملفات",
        "tooltip_file_category": "انقر لعرض فئات الملفات",
        "tooltip_file_extension": "حدد امتدادات الملفات المراد سردها",
        "tooltip_filter_apply": "انقر لتطبيق الفلاتر المحددة",
        "tooltip_all_files": "تحديد/إلغاء تحديد جميع امتدادات الملفات",
        "tooltip_category_expand": "انقر لتوسيع/طي هذه الفئة",
        "tooltip_select_category": "تحديد/إلغاء تحديد جميع الامتدادات في هذه الفئة",
        "select_all_category": "تحديد الكل في هذه الفئة",
        "media_files": "وسائط",
        "code_files": "برمجة و ويب",
        "data_files": "بيانات",
        "design_files": "تصميم",
        "game_files": "ألعاب",
        "extension_search": "بحث:",
        "copyright_footer": "© {year} معلمون.نت - قوائم لسهلة"
    },
    "de": de_dict,
    "fr": fr_dict,
    "ru": ru_dict,
    "es": es_dict,
    "it": it_dict,
    "fa": fa_dict,
    "ur": ur_dict,
    "hi": hi_dict,
    "zh": zh_dict,
    "ja": ja_dict
}


class FileManagerApp:
    def __init__(self, root):
        """
        ListeKolay uygulamasının ana sınıfı.
        Bu sınıf dosya listeleme, önizleme ve dışa aktarma işlemlerini yönetir.

        Args:
            root: Ana Tkinter penceresi
        """
        # Koruma bayrakları - birden fazla dil/tema değişikliği olayını önlemek için
        self.config_loading_in_progress = False
        self.theme_change_in_progress = False 
        self.theme_update_in_progress = False

        # Geçici dosya izleme
        self.temp_files = []

        # İptal işaretçileri
        self.cancel_flag = False  # Eski uyumluluk için
        self.cancel_event = threading.Event()  # Thread-safe iptal mekanizması
        self.root = root
        self.current_language = "tr"  # Default language is Turkish

        # Uygulama sürüm bilgisi
        self.current_version = "5.3.0"
        self.github_version_url = "https://github.com/muallimun/listekolay/raw/main/listekolay_version.txt"
        self.github_download_url = "https://github.com/muallimun/listekolay/releases/latest"

        # İptal mekanizması için gelişmiş thread-safe yapılar
        self.cancel_event = threading.Event()  # Thread-safe iptal event nesnesi
        self.cancel_flag = False  # Geriye dönük uyumluluk için flag

        # Geçici dosyaların izlenmesi için liste
        self.temp_files = []

        # İşlem durumu için animasyonlu simgeler
        self.spinner_chars = ["⟳", "⟲", "↻", "↺"]
        self.progress_icons = ["⏳", "🔄", "⚙️", "📊"]

        # Tema ayarları (açık/koyu mod)
        self.is_dark_mode = tk.BooleanVar(value=False)  # Varsayılan olarak açık mod

        # Add custom translations for pagination
        self.pagination_translations = {
            "tr": {"page": "Sayfa", "prev_page": "Önceki", "next_page": "Sonraki"},
            "en": {"page": "Page", "prev_page": "Previous", "next_page": "Next"},
            "ar": {"page": "صفحة", "prev_page": "السابق", "next_page": "التالي"},
            "de": {"page": "Seite", "prev_page": "Zurück", "next_page": "Weiter"},
            "fr": {"page": "Page", "prev_page": "Précédent", "next_page": "Suivant"},
            "ru": {"page": "Страница", "prev_page": "Предыдущая", "next_page": "Следующая"},
            "es": {"page": "Página", "prev_page": "Anterior", "next_page": "Siguiente"},
            "it": {"page": "Pagina", "prev_page": "Precedente", "next_page": "Successiva"},
            "zh": {"page": "页面", "prev_page": "上一页", "next_page": "下一页"},
            "ja": {"page": "ページ", "prev_page": "前へ", "next_page": "次へ"}
        }

        # Initialize translations from the global translations dict
        self.languages = translations

        self.root.title(self.get_text("full_window_title"))
        self.root.geometry("1024x768")
        # Use normal state instead of zoomed (which doesn't work on some platforms)
        # self.root.state("zoomed")      
        self.root.minsize(800, 600)       # Minimum window size
        self.root.configure(bg="#e9ecef")
        self.root.resizable(True, True)   # Window can be resized
        self.files = []
        self.filtered_files = []          # Store filtered files separately
        self.include_subfolders = tk.BooleanVar(value=False)
        self.selected_folder_path = ""
        self.save_to_desktop = tk.BooleanVar(value=False)
        self.cancel_flag = False

        # For file view mode (list or preview)
        self.view_mode_var = tk.StringVar(value="list")  # Default to list view

        # OPTIMIZATION: Enhanced preview caching system with LRU (Least Recently Used) behavior
        self.preview_cache = {}  # Cache for previews to improve performance
        self.preview_cache_keys = []  # Keep track of cache access order for LRU implementation

        # Önizleme gezinme değişkenleri
        self.current_preview_files = []  # Tüm önizleme dosyalarını saklar
        self.current_preview_index = -1  # Şu anda gösterilen dosyanın indeksi

        # OPTIMIZATION: Performance settings for large folders
        # OPTIMIZATION: Increased batch sizes for better performance with large folders
        self.file_loading_batch_size = 1000    # How many files to process at once during loading (increased from 500)
        # OPTIMIZATION: Increase batch sizes for faster processing of large file lists
        self.file_filtering_batch_size = 5000  # How many files to filter at once (increased from 2000)
        self.file_display_batch_size = 2000    # How many files to add to UI at once (increased from 1000)
        # OPTIMIZATION: Increase preview batch size for better parallel processing
        self.preview_batch_size = 200          # How many previews to generate at once (increased from 100)
        self.max_preview_cache_size = 750     # Maximum number of thumbnails to cache (increased from 500)
        self.preview_items_per_page = 150     # Number of preview items per page (increased from 100)
        self.preview_page = 1                 # Current preview page

        # İlerleme çubuğu değişkenleri
        self.progress_var = tk.DoubleVar(value=0)
        self.cancel_progress = False

        self.export_formats = {
            "text": tk.BooleanVar(value=True),
            "excel": tk.BooleanVar(value=False),
            "word": tk.BooleanVar(value=False),
            "html": tk.BooleanVar(value=False)
        }

        # Category variables for selecting/deselecting all extensions
        self.category_vars = {}
        self.all_files_var = tk.BooleanVar(value=False)
        self.sort_options = [
            "sort_name_asc",
            "sort_name_desc",
            "sort_ext_asc",
            "sort_ext_desc",
            "sort_size_asc",
            "sort_size_desc",
            "sort_dir_asc"
        ]
        self.selected_sort = tk.StringVar(value=self.sort_options[0])

# Final: Genişletilmiş ve mantıksal olarak gruplanmış 6 dosya kategorisi
        self.file_categories = {
            # 1. Belge Dosyaları (dokümanlar, tablolar, sunumlar, metinler)
            "document_files": [
                # Belgeler
                ".doc", ".docx", ".rtf", ".odt", ".pdf", ".txt", ".epub", ".mobi", ".tex", 
                ".pages", ".md", ".csv", ".log", ".udf",
                # Tablolar
                ".xls", ".xlsx", ".xlsm", ".ods", ".numbers",
                # Sunumlar
                ".ppt", ".pptx", ".odp", ".key", ".pps", ".ppsx"
            ],

            # 2. Görsel Dosyalar (resim formatları, raster-vektörel)
            "image_files": [
                ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".svg", ".ico", 
                ".raw", ".heif", ".cr2",  ".psd", ".ai", ".eps",
            ],

            # 3. Video ve Ses Dosyaları (medya formatları)
            "video_files": [
                # Video
                ".mp4", ".avi", ".mov", ".wmv", ".flv", ".mkv", ".webm", ".m4v", ".mpg", ".mpeg", 
                ".3gp", ".ts", ".vob", ".asf", ".ogv", ".m2v",
                # Ses
                ".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".alac", ".aiff", ".opus", ".m4a"
            ],

            # 4. Tasarım ve Eğitim İçeriği Dosyaları (grafik, CAD, font, etkileşimli içerik)
            "design_files": [
                # Grafik tasarım ve vektör
                ".psd", ".ai", ".eps", ".xd", ".indd", ".cdr", ".fig", ".afdesign", ".afphoto",
                # Yazı tipleri
                ".ttf", ".otf", ".woff", ".woff2", ".eot", ".fon",
                # 3D / CAD
                ".dwg", ".dxf", ".skp", ".3ds", ".max", ".c4d", ".blend", ".fbx", ".obj", ".stl", ".step", ".stp",
                # Animasyon / Hareketli medya
                ".ae", ".swf",
                # Eğitim / Etkileşimli içerik
                ".h5p", ".scorm", ".xar", ".cptx", ".story", ".ismp", ".quiz", ".interact", ".ao"
            ],

            # 5. Kod ve Web Dosyaları (programlama ve betik dosyaları)
            "code_files": [
                # Programlama dilleri
                ".py", ".java", ".c", ".cpp", ".cs", ".php", ".rb", ".go", ".swift", ".ts", ".js",
                # Web dilleri
                ".html", ".css", ".vue", ".jsx", ".ini", ".dat",
                # Veritabanı, yapılandırma
                ".sql", ".json", ".xml", ".yaml", ".yml", ".config", ".mdb", ".mde", ".accdb", ".accdt", ".accde",
                # Scriptler
                ".sh", ".bat", ".cer",
                # Sunucu-tarayıcı
                ".asp", ".aspx", ".jsp",
                # Diğer
                ".h5p"  # Eğer eğitim aracı olarak değilse web içerik olarak burada da olabilir (ama yukarıda da var)
            ],

            # 6. Program ve Arşiv Dosyaları (uygulamalar, kurulum ve sıkıştırılmış dosyalar)
            "program_files": [
                # Çalıştırılabilir ve kurulum
                ".exe", ".dll", ".msi", ".app", ".jar", ".dmg", ".apk", ".deb", ".rpm", ".apk",
                # Arşivleme
                ".zip", ".rar", ".7z", ".tar", ".gz", ".iso"
            ]
        }

        # Keep track of which extensions are selected
        self.selected_extensions = {}
        for category, extensions in self.file_categories.items():
            for ext in extensions:
                self.selected_extensions[ext] = tk.BooleanVar(value=True)

        # Create GUI
        self.create_gui()
        self.update_ui_state()

        # Ayarları yükle
        self.load_config()

        # Bind close event
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        # Show startup message
        self.update_status(self.get_text("ready"))

    def show_error(self, error_title, error_message, exception=None):
        """Kullanıcıya hata göster ve loglama yap"""
        # Hata detaylarını logla
        if exception:
            logging.error(f"{error_title}: {str(exception)}")
        else:
            logging.error(error_title)

        # Kullanıcıya hata mesajı göster
        messagebox.showerror(
            self.get_text("error"), 
            error_message
        )

        # Durum çubuğunda da göster
        self.update_status(error_message)

    def get_text(self, key):
        """Localization helper function"""
        # Special case for search_files - use our centralized search translations
        if key == "search_files" and self.current_language in search_translations:
            return search_translations[self.current_language]

        # Special case for pagination - use our custom pagination translations
        if key in ["page", "prev_page", "next_page"] and hasattr(self, 'pagination_translations'):
            if self.current_language in self.pagination_translations and key in self.pagination_translations[self.current_language]:
                return self.pagination_translations[self.current_language][key]
            elif "en" in self.pagination_translations and key in self.pagination_translations["en"]:
                return self.pagination_translations["en"][key]
                
        # Special case for context menu items - use our centralized context menu translations
        if key in ["preview_file", "delete_files", "copy_files", "move_files", "rename_file", "select_all_files", 
                  "delete_file", "copy_file", "move_file", "cut_file", "open_file", "open_file_location",
                  "copy_filename", "copy_filepath", "updating_preview", "do_you_want_to_delete", "copied_to_clipboard", "rename_error",
                  "large_file_warning", "large_file_slow", "loading_large_file", "view_changed_to_list", 
                  "preview_not_available"]:
            # Varsayılan değerler (fallback) tanımla - herhangi bir hata durumunda bunlar kullanılacak
            defaults = {
                "preview_file": "Preview File",
                "delete_files": "Delete Files",
                "copy_files": "Copy Files",
                "move_files": "Move Files",
                "delete_file": "Delete File",
                "copy_file": "Copy File",
                "move_file": "Move File",
                "cut_file": "Cut File",
                "open_file": "Open File",
                "open_file_location": "Open File Location",
                "copy_filename": "Copy Filename",
                "copy_filepath": "Copy File Path",
                "rename_file": "Rename File",
                "select_all_files": "Select All Files",
                "updating_preview": "Updating preview...",
                "do_you_want_to_delete": "Are you sure you want to delete this file",
                "copied_to_clipboard": "Copied to clipboard",
                "rename_error": "Rename error",
                "large_file_warning": "Large file warning",
                "large_file_slow": "This file is very large and may take time to load",
                "loading_large_file": "Loading large file...",
                "view_changed_to_list": "View changed to list mode",
                "preview_not_available": "No Preview Available"
            }
            
            try:
                # search_translations modülündeki global sözlükten çeviri al
                from search_translations import search_translations as st_dict
                if key in st_dict and self.current_language in st_dict[key]:
                    return st_dict[key][self.current_language]
                
                # Çeviri varsa kullan
                if key in context_menu_translations:
                    # Mevcut dilde varsa onu kullan
                    if self.current_language in context_menu_translations[key]:
                        return context_menu_translations[key][self.current_language]
                    # İngilizce varsa onu kullan
                    elif "en" in context_menu_translations[key]:
                        return context_menu_translations[key]["en"]
                    # Türkçe varsa onu kullan
                    elif "tr" in context_menu_translations[key]:
                        return context_menu_translations[key]["tr"]
            except (NameError, AttributeError, TypeError, KeyError):
                # Herhangi bir hata durumunda varsayılan değeri döndür
                pass
                
            # Hiçbir çeviri bulunamazsa varsayılan değer döndür
            return defaults.get(key, key)

        if key in self.languages[self.current_language]:
            return self.languages[self.current_language][key]
        # Fallback to English if key not found in current language
        elif key in self.languages["en"]:
            return self.languages["en"][key]
        # Return the key itself if not found in any language
        return key

    def create_gui(self):
        # Create base frame with padding
        self.main_frame = tk.Frame(
            self.root, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            padx=10, 
            pady=10
        )
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # Title and subtitle in a more compact layout
        title_frame = tk.Frame(
            self.main_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        title_frame.pack(fill=tk.X, pady=(0, 5))  # Reduced padding

        # Header container to align items horizontally
        header_container = tk.Frame(
            title_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        header_container.pack(side=tk.LEFT, fill=tk.X)

        # Program title - left aligned
        title_label = tk.Label(
            header_container, 
            text=self.get_text("app_title"), 
            font=("Segoe UI", 16, "bold"),  # Slightly smaller font
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor="w"  # Left aligned text
        )
        title_label.pack(side=tk.LEFT, pady=(0, 2))

        # Subtitle - next to title with separator
        self.subtitle_label = tk.Label(
            header_container, 
            text=" - " + self.get_text("app_subtitle"),  # Add separator
            font=("Segoe UI", 10, "italic"),  # Italic for style
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        self.subtitle_label.pack(side=tk.LEFT, padx=(5, 0), pady=(4, 0))  # Align vertically

        # Theme settings container - placed on the right side of the title frame
        theme_container = tk.Frame(
            title_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        theme_container.pack(side=tk.RIGHT)

        # Theme mode label
        theme_label = tk.Label(
            theme_container, 
            text=self.get_text("theme_settings"),
            font=("Segoe UI", 9),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        theme_label.pack(side=tk.LEFT, padx=(0, 5))

        # Light mode (sun icon) radio button
        light_radio = tk.Radiobutton(
            theme_container,
            text="☀️",  # Sun emoji
            variable=self.is_dark_mode,
            value=False,
            command=self.toggle_theme_mode,
            font=("Segoe UI", 9),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            selectcolor="#f8f9fa",
            indicatoron=False,
            width=2,
            bd=1,
            relief=tk.RAISED if not self.is_dark_mode.get() else tk.FLAT
        )
        self.create_tooltip(light_radio, self.get_text("light_mode"))
        light_radio.pack(side=tk.LEFT)

        # Dark mode (moon icon) radio button
        dark_radio = tk.Radiobutton(
            theme_container,
            text="🌙",  # Moon emoji
            variable=self.is_dark_mode,
            value=True,
            command=self.toggle_theme_mode,
            font=("Segoe UI", 9),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            selectcolor="#212529",
            indicatoron=False,
            width=2,
            bd=1,
            relief=tk.RAISED if self.is_dark_mode.get() else tk.FLAT
        )
        self.create_tooltip(dark_radio, self.get_text("dark_mode"))
        dark_radio.pack(side=tk.LEFT)



        # Top controls frame (folder selection, action buttons)
        top_frame = tk.Frame(
            self.main_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        top_frame.pack(fill=tk.X, pady=10)

        # Folder selection with integrated subfolder option
        folder_frame = tk.Frame(
            top_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        folder_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        folder_label = tk.Label(
            folder_frame, 
            text=self.get_text("selected_folder"), 
            font=("Segoe UI", 9), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        folder_label.pack(side=tk.LEFT, padx=(0, 5))

        self.folder_path_var = tk.StringVar(value=self.get_text("no_folder_selected"))
        self.folder_path_entry = tk.Entry(
            folder_frame, 
            textvariable=self.folder_path_var, 
            width=35,  # Reduced width
            state="readonly", 
            font=("Segoe UI", 9), 
            readonlybackground=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.folder_path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))

        self.select_folder_btn = tk.Button(
            folder_frame, 
            text=self.get_text("select_folder"), 
            command=self.select_folder, 
            font=("Segoe UI", 9), 
            bg="#007bff",  # Mavi (klasör seçme butonu için uygun)
            fg="white",
            activebackground="#0069d9",
            activeforeground="white", 
            bd=0,
            padx=10
        )
        self.create_tooltip(self.select_folder_btn, self.get_text("tooltip_select"))
        self.select_folder_btn.pack(side=tk.LEFT, padx=(0, 10))

        # Move subfolder option to here beside folder selection
        self.subfolder_cb = tk.Checkbutton(
            folder_frame, 
            text=self.get_text("subfolders_label") + " " + self.get_text("include_label"), 
            variable=self.include_subfolders,
            font=("Segoe UI", 9), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
            selectcolor="#ffffff",  # Beyaz tik kutuları
            command=self.on_subfolder_changed
        )
        self.create_tooltip(self.subfolder_cb, self.get_text("tooltip_subfolders"))
        self.subfolder_cb.pack(side=tk.LEFT)

        # Buttons frame
        button_frame = tk.Frame(
            top_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        button_frame.pack(side=tk.RIGHT, padx=(10, 0))

        # Buttons frame is now cleaner after moving search to files and previews section

        # Language selection
        language_frame = tk.Frame(
            button_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        language_frame.pack(side=tk.LEFT, padx=(0, 10))

        language_label = tk.Label(
            language_frame, 
            text=self.get_text("language") + ":", 
            font=("Segoe UI", 9), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        language_label.pack(side=tk.LEFT, padx=(0, 5))

        self.language_var = tk.StringVar(value=self.current_language)
        self.language_dropdown = ttk.Combobox(
            language_frame, 
            textvariable=self.language_var, 
            values=["tr", "en", "ar", "de", "fr", "ru", "es", "it", "fa", "ur", "hi", "zh", "ja"], 
            state="readonly", 
            width=5
        )
        self.language_dropdown.bind("<<ComboboxSelected>>", self.change_language)
        self.language_dropdown.pack(side=tk.LEFT, padx=(0, 10))



        # Action buttons
        self.start_btn = tk.Button(
            button_frame, 
            text=self.get_text("start"), 
            command=self.start_processing, 
            font=("Segoe UI", 9), 
            bg="#28a745",  # Yeşil (işlemi başlat butonu)
            fg="white", 
            activebackground="#218838",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.start_btn, self.get_text("tooltip_start"))
        self.start_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.cancel_btn = tk.Button(
            button_frame, 
            text=self.get_text("cancel"), 
            command=self.cancel_operation, 
            font=("Segoe UI", 9), 
            bg="#dc3545",  # Kırmızı (iptal butonu)
            fg="white",
            activebackground="#c82333",
            activeforeground="white",
            bd=0,
            padx=10,
            state=tk.DISABLED
        )
        self.create_tooltip(self.cancel_btn, self.get_text("tooltip_cancel"))
        self.cancel_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.exit_btn = tk.Button(
            button_frame, 
            text=self.get_text("exit"), 
            command=self.on_close, 
            font=("Segoe UI", 9), 
            bg="#6c757d",  # Gri (çıkış butonu için uygun)
            fg="white",
            activebackground="#5a6268",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.exit_btn, self.get_text("tooltip_exit"))
        self.exit_btn.pack(side=tk.LEFT)

        # Main content frame (2 columns)
        content_frame = tk.Frame(self.main_frame, bg="#e9ecef")
        content_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        # Sol panel toggle butonu için container
        toggle_container = tk.Frame(content_frame, bg="#e9ecef")
        toggle_container.pack(side=tk.LEFT, fill=tk.Y)

        # Sol panel toggle butonu
        self.left_panel_visible = tk.BooleanVar(value=True)
        self.toggle_left_panel_btn = tk.Button(
            toggle_container,
            text="◀",  # Sol ok işareti (paneli gizle)
            command=self.toggle_left_panel,
            font=("Segoe UI", 12, "bold"),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            activebackground=LIGHT_MODE_COLORS["btn_active_bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_active_bg"],
            activeforeground=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            bd=1,
            width=2,
            pady=5,
            relief=tk.RAISED
        )
        self.toggle_left_panel_btn.pack(fill=tk.Y, padx=(0, 5))
        self.create_tooltip(self.toggle_left_panel_btn, toggle_panel_translations.get(self.current_language, "Sol paneli aç/kapat"))

        # Left column (Settings, Filter, Tips) - Now with increased width
        self.left_column = tk.Frame(content_frame, bg="#e9ecef", width=450)  # Increased width
        self.left_column.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10))
        self.left_column.pack_propagate(False)  # Prevent shrinking

        # Settings panel
        self.settings_frame = tk.LabelFrame(
            self.left_column, 
            text=self.get_text("settings_header"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=10,
            pady=10
        )
        self.settings_frame.pack(fill=tk.X, pady=(0, 10))

        # Subfolder setting removed from here - moved to folder selection area

        # Export format setting - improved layout with label inline with checkboxes
        format_container = tk.Frame(self.settings_frame, bg="#e9ecef")
        format_container.pack(fill=tk.X, pady=(0, 5))

        format_label = tk.Label(
            format_container, 
            text=self.get_text("list_format_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        format_label.pack(side=tk.LEFT, padx=(0, 5))

        # Formats frame is now part of the same container
        formats_frame = tk.Frame(format_container, bg="#e9ecef")
        formats_frame.pack(side=tk.LEFT, fill=tk.X)

        # Text format
        text_cb = tk.Checkbutton(
            formats_frame, 
            text="TXT", 
            variable=self.export_formats["text"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(text_cb, self.get_text("text_format_info"))
        text_cb.pack(side=tk.LEFT, padx=(0, 5))

        # Excel format
        excel_cb = tk.Checkbutton(
            formats_frame, 
            text="Excel", 
            variable=self.export_formats["excel"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(excel_cb, self.get_text("excel_format_info"))
        excel_cb.pack(side=tk.LEFT, padx=(0, 5))

        # Word format
        word_cb = tk.Checkbutton(
            formats_frame, 
            text="Word", 
            variable=self.export_formats["word"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(word_cb, self.get_text("word_format_info"))
        word_cb.pack(side=tk.LEFT, padx=(0, 5))

        # HTML format
        html_cb = tk.Checkbutton(
            formats_frame, 
            text="HTML", 
            variable=self.export_formats["html"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(html_cb, self.get_text("html_format_info"))
        html_cb.pack(side=tk.LEFT)

        # Combined save location and sort criteria in one row
        combined_settings_frame = tk.Frame(self.settings_frame, bg="#e9ecef")
        combined_settings_frame.pack(fill=tk.X, pady=(0, 5))

        # Save location setting - left side
        save_frame = tk.Frame(combined_settings_frame, bg="#e9ecef")
        save_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        save_label = tk.Label(
            save_frame, 
            text=self.get_text("save_location_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        save_label.pack(side=tk.LEFT)

        self.desktop_cb = tk.Checkbutton(
            save_frame, 
            text=self.get_text("desktop_label"), 
            variable=self.save_to_desktop,
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(self.desktop_cb, self.get_text("tooltip_save_location"))
        self.desktop_cb.pack(side=tk.LEFT, padx=(5, 0))

        # Small spacer between settings
        spacer = tk.Frame(combined_settings_frame, bg="#e9ecef", width=20)
        spacer.pack(side=tk.LEFT)

        # Sort criteria setting - right side
        sort_frame = tk.Frame(combined_settings_frame, bg="#e9ecef")
        sort_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        sort_label = tk.Label(
            sort_frame, 
            text=self.get_text("sort_criteria_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        sort_label.pack(side=tk.LEFT, padx=(0, 5))

        self.sort_dropdown = ttk.Combobox(
            sort_frame, 
            textvariable=self.selected_sort, 
            state="readonly", 
            font=("Segoe UI", 9), 
            width=15  # Reduced width
        )
        self.create_tooltip(self.sort_dropdown, self.get_text("tooltip_sort_criteria"))
        self.populate_sort_dropdown()
        self.sort_dropdown.pack(side=tk.LEFT)

        # Add Filter button to settings section
        filter_btn_frame = tk.Frame(self.settings_frame, bg="#e9ecef")
        filter_btn_frame.pack(fill=tk.X, pady=(5, 0))

        # First row: Filter button and search
        self.show_filter_btn = tk.Button(
            filter_btn_frame, 
            text=self.get_text("filter_label"), 
            command=self.toggle_filter_section, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=10
        )
        self.create_tooltip(self.show_filter_btn, self.get_text("tooltip_file_category"))
        self.show_filter_btn.pack(side=tk.LEFT, padx=(0, 10))

        # Extension search variable still needed for filter panel functionality
        self.extension_search_var = tk.StringVar()

        # Filter panel - initially hidden
        self.filter_frame = tk.LabelFrame(
            self.left_column, 
            text=self.get_text("filter_label"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=10,
            pady=10
        )
        # Don't pack it initially (hidden by default)

        # Filter controls
        filter_controls_frame = tk.Frame(self.filter_frame, bg="#e9ecef")
        filter_controls_frame.pack(fill=tk.X, pady=(0, 10))

        self.select_all_btn = tk.Button(
            filter_controls_frame, 
            text=self.get_text("select_all"), 
            command=self.select_all_extensions, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.create_tooltip(self.select_all_btn, self.get_text("tooltip_select_all"))
        self.select_all_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.clear_all_btn = tk.Button(
            filter_controls_frame, 
            text=self.get_text("clear_all"), 
            command=self.clear_all_extensions, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.create_tooltip(self.clear_all_btn, self.get_text("tooltip_clear_all"))
        self.clear_all_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.apply_filter_btn = tk.Button(
            filter_controls_frame, 
            text=self.get_text("apply_filter"), 
            command=self.apply_filter, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.create_tooltip(self.apply_filter_btn, self.get_text("tooltip_filter_apply"))
        self.apply_filter_btn.pack(side=tk.LEFT)

        # Add search box for extensions
        search_frame = tk.Frame(self.filter_frame, bg="#e9ecef")
        search_frame.pack(fill=tk.X, pady=(5, 0))

        self.extension_search_label = tk.Label(
            search_frame,
            text="🔍 " + self.get_text("extension_search"),
            font=("Segoe UI", 9),
            bg="#e9ecef",
            fg="#495057"
        )
        self.extension_search_label.pack(side=tk.LEFT, padx=(0, 5))

        self.extension_search_var = tk.StringVar()
        # Remove auto-filter on typing

        self.extension_search_entry = tk.Entry(
            search_frame,
            textvariable=self.extension_search_var,
            font=("Segoe UI", 9),
            width=15  # Shorter width
        )
        self.extension_search_entry.pack(side=tk.LEFT, padx=(0, 5))

        # Add search button
        self.search_button = tk.Button(
            search_frame,
            text="🔍",
            command=self.filter_extensions,
            font=("Segoe UI", 9, "bold"),
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.search_button.pack(side=tk.LEFT)

        # Removed category header to save space

        # Filter categories
        self.filter_categories_frame = tk.Frame(self.filter_frame, bg="#e9ecef")
        self.filter_categories_frame.pack(fill=tk.BOTH, expand=True)

        # Create a canvas with scrollbar for categories (increased height for full screen usage)
        self.category_canvas = tk.Canvas(self.filter_categories_frame, bg="#e9ecef", highlightthickness=0, height=300)
        category_scrollbar = ttk.Scrollbar(self.filter_categories_frame, orient=tk.VERTICAL, command=self.category_canvas.yview)
        self.category_canvas.configure(yscrollcommand=category_scrollbar.set)

        category_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.category_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Frame inside the canvas for categories
        self.category_inner_frame = tk.Frame(self.category_canvas, bg="#e9ecef")
        self.category_canvas.create_window((0, 0), window=self.category_inner_frame, anchor=tk.NW)

        # Populate categories
        self.populate_categories()

        # Configure canvas
        self.category_inner_frame.bind("<Configure>", lambda e: self.category_canvas.configure(scrollregion=self.category_canvas.bbox("all")))
        self.category_canvas.bind("<Configure>", self.on_category_canvas_configure)

        # Right column (Statistics, File List)
        right_column = tk.Frame(content_frame, bg="#e9ecef")
        right_column.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Statistics Panel
        self.stats_frame = tk.LabelFrame(
            right_column, 
            text=self.get_text("statistics_header"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=10,
            pady=10
        )
        self.stats_frame.pack(fill=tk.X, pady=(0, 10))

        # Statistics content - Now in a single row
        stats_content_frame = tk.Frame(self.stats_frame, bg="#e9ecef")
        stats_content_frame.pack(fill=tk.X)

        # Total files
        total_files_label = tk.Label(
            stats_content_frame, 
            text=self.get_text("total_files_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W
        )
        total_files_label.pack(side=tk.LEFT, padx=(0, 5))

        self.total_files_var = tk.StringVar(value="0")
        total_files_value = tk.Label(
            stats_content_frame, 
            textvariable=self.total_files_var, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        total_files_value.pack(side=tk.LEFT, padx=(0, 15))

        # Folder count
        folder_count_label = tk.Label(
            stats_content_frame, 
            text=self.get_text("folder_count_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W
        )
        folder_count_label.pack(side=tk.LEFT, padx=(0, 5))

        self.folder_count_var = tk.StringVar(value="0")
        folder_count_value = tk.Label(
            stats_content_frame, 
            textvariable=self.folder_count_var, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        folder_count_value.pack(side=tk.LEFT, padx=(0, 15))

        # Total size
        total_size_label = tk.Label(
            stats_content_frame, 
            text=self.get_text("total_size_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W
        )
        total_size_label.pack(side=tk.LEFT, padx=(0, 5))

        self.total_size_var = tk.StringVar(value="0 MB")
        total_size_value = tk.Label(
            stats_content_frame, 
            textvariable=self.total_size_var, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        total_size_value.pack(side=tk.LEFT)

        # File list - with top bar for title and view mode toggle
        file_list_header_frame = tk.Frame(right_column, bg="#e9ecef", pady=5)
        file_list_header_frame.pack(fill=tk.X)

        # Add title and view mode toggles in the same row
        file_list_header_label = tk.Label(
            file_list_header_frame, 
            text=self.get_text("file_list_section"),
            font=("Segoe UI", 10, "bold"), 
            bg="#e9ecef", 
            fg="#343a40"
        )
        file_list_header_label.pack(side=tk.LEFT, padx=(5, 10))

        # Add a search frame on the right side before view mode buttons
        search_frame = tk.Frame(file_list_header_frame, bg="#e9ecef")
        search_frame.pack(side=tk.RIGHT, padx=(10, 10))

        # Search label
        self.file_search_label = tk.Label(
            search_frame, 
            text=self.get_text("extension_search") + ":", 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        self.file_search_label.pack(side=tk.LEFT, padx=(0, 5))

        # Arama girişi için container - Entry ve temizleme butonu içerir
        search_entry_container = tk.Frame(search_frame, bg="#e9ecef")
        search_entry_container.pack(side=tk.LEFT)

        # Add the search entry field
        self.file_search_var = tk.StringVar()
        self.file_search_var.trace("w", self.filter_file_list)
        self.file_search_entry = tk.Entry(
            search_entry_container,
            textvariable=self.file_search_var,
            font=("Segoe UI", 9),
            width=15,  # Biraz daha geniş
            fg="#000000"  # Metin rengini her zaman siyah olarak ayarla
        )
        self.file_search_entry.pack(side=tk.LEFT)

        # Temizleme butonu
        self.clear_search_button = tk.Button(
            search_entry_container,
            text="✕",
            font=("Segoe UI", 7),
            bg="#e9ecef",
            fg="#666666",
            relief=tk.FLAT,
            command=self.clear_search_field,
            padx=0,
            pady=0,
            width=2
        )
        self.clear_search_button.pack(side=tk.LEFT, padx=(0, 2))
        self.clear_search_button.config(state=tk.DISABLED)  # Başlangıçta devre dışı

        # Create a placeholder for the search entry
        self.file_search_entry.insert(0, self.get_text("search_files"))
        # Change color to gray for placeholder
        self.file_search_entry.config(fg='gray')

        # Bind focus events to handle placeholder text
        self.file_search_entry.bind("<FocusIn>", self.on_search_focus_in)
        self.file_search_entry.bind("<FocusOut>", self.on_search_focus_out)

        # Arama ipucu ekle
        search_tooltip = "Filtre dosya adları ve uzantılara göre yapılır"
        self.create_tooltip(self.file_search_entry, search_tooltip)

        # View mode buttons - compact design
        controls_container = tk.Frame(file_list_header_frame, bg="#e9ecef")
        controls_container.pack(side=tk.RIGHT, padx=(10, 10))

        # View mode buttons - compact design (reverted from radio buttons to normal buttons)
        view_modes_container = tk.Frame(controls_container, bg="#e9ecef")
        view_modes_container.pack(side=tk.LEFT)

        # Button frame
        button_frame = tk.Frame(view_modes_container, bg="#e9ecef")
        button_frame.pack(side=tk.LEFT)

        # List view button
        self.list_view_btn = tk.Button(
            button_frame,
            text=self.get_text("list_view"),
            command=lambda: self.set_view_mode("list"),
            font=("Segoe UI", 9),
            bg="#17a2b8" if self.view_mode_var.get() == "list" else "#e9ecef",  # Turkuaz (aktif) veya gri (pasif)
            fg="white" if self.view_mode_var.get() == "list" else "#495057",
            activebackground="#138496",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.list_view_btn, self.get_text("tooltip_list_view"))
        self.list_view_btn.pack(side=tk.LEFT, padx=(0, 10))

        # Preview view button
        self.preview_view_btn = tk.Button(
            button_frame,
            text=self.get_text("preview_view"),
            command=lambda: self.set_view_mode("preview"),
            font=("Segoe UI", 9),
            bg="#17a2b8" if self.view_mode_var.get() == "preview" else "#e9ecef",  # Turkuaz (aktif) veya gri (pasif)
            fg="white" if self.view_mode_var.get() == "preview" else "#495057",
            activebackground="#138496",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.preview_view_btn, self.get_text("tooltip_preview_view"))
        self.preview_view_btn.pack(side=tk.LEFT)

        # We removed the duplicate header and search is now in the first header

        # File list frame
        file_list_frame = tk.LabelFrame(
            right_column, 
            text="",  # No text here anymore
            font=("Segoe UI", 10, "bold"), 
            bg="#e9ecef", 
            fg="#343a40",
            padx=10,
            pady=10
        )
        file_list_frame.pack(fill=tk.BOTH, expand=True)

        # Cache for preview images

        # Cache for preview images
        self.preview_cache = {}

        # Create Treeview for file list - it will be placed in the container
        self.create_file_list_treeview(file_list_frame)

        # Status Bar
        status_frame = tk.Frame(self.main_frame, bg="#e9ecef", height=25)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))

        # İlerleme çubuğu için frame - önce tanımlıyoruz, görünmez durumda
        self.progress_frame = tk.Frame(status_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])

        # İlerleme çubuğu - progress_var ile ilişkilendirilmiş
        self.progress_bar = ttk.Progressbar(
            self.progress_frame, 
            orient=tk.HORIZONTAL, 
            length=100, 
            mode='determinate',
            variable=self.progress_var
        )
        # Pack işlemini dosya işlemleri sırasında yapacağız

        # Info icon button for website link - sol altta
        info_btn = tk.Button(
            status_frame,
            text="ℹ️",  # Info emoji
            command=lambda: self.open_website("https://www.muallimun.com/listekolay/"),
            font=("Segoe UI", 14, "bold"),  # Daha büyük font
            bg="#e9ecef",
            fg="#007bff",  # Mavi renk
            activebackground="#e9ecef",
            activeforeground="#0056b3",  # Koyu mavi hover
            bd=0,
            width=3,
            relief=tk.FLAT
        )
        self.create_tooltip(info_btn, "ListeKolay web sitesi")
        info_btn.pack(side=tk.LEFT, padx=(10, 5))

        self.status_var = tk.StringVar(value=self.get_text("ready"))
        status_label = tk.Label(
            status_frame, 
            textvariable=self.status_var, 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W,
            padx=10,
            pady=5
        )
        status_label.pack(side=tk.LEFT, fill=tk.X, expand=True)



        # Tips panel - now at the bottom of left column with enhanced styling
        self.tips_frame = tk.LabelFrame(
            self.left_column, 
            text=self.get_text("tips_header"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],  # Theme-aware background
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=12,
            pady=12,
            height=225,  # Reduced height for better UI proportions
            relief=tk.GROOVE,  # Nicer relief style
            bd=2  # Slightly thicker border
        )
        self.tips_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))
        self.tips_frame.pack_propagate(False)  # Prevent the frame from shrinking to fit content

        # Inner frame for tips with padding
        tips_inner_frame = tk.Frame(
            self.tips_frame,
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            padx=2
        )
        tips_inner_frame.pack(fill=tk.BOTH, expand=True)

        # Tips - reduced to 6 tips with shorter format
        tip_labels = [
            self.get_text("tip_1"),
            self.get_text("tip_3"),
            self.get_text("tip_4"),
            self.get_text("tip_5"),
            self.get_text("tip_6"),
            self.get_text("tip_preview_formats") # Özel ipucu: Desteklenen ön izleme formatları (dil desteğiyle)
        ]

        # Calculate maximum width for wrapping - use maximum available width
        frame_width = self.left_column.winfo_reqwidth() or 300  # Use actual width or default to 300
        wrap_width = frame_width - 20  # Minimum padding for better appearance

        for i, tip in enumerate(tip_labels):
            # Create a separator line between tips (except before the first tip)
            if i > 0:
                separator = tk.Frame(
                    tips_inner_frame,
                    height=1,
                    bg="#a0afc0" if not self.is_dark_mode.get() else "#4a5568"  # Theme-aware separator color
                )
                separator.pack(fill=tk.X, pady=3)

            tip_label = tk.Label(
                tips_inner_frame, 
                text=f"• {tip}", 
                font=("Segoe UI", 9), 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],  # Match theme background
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],  # Tema rengini kullan
                justify=tk.LEFT,
                anchor=tk.W,
                wraplength=wrap_width  # Use calculated width to make text fill the width
            )
            tip_label.pack(fill=tk.X, anchor=tk.W, pady=(2, 0))

        # Add footer with copyright and website link
        footer_frame = tk.Frame(self.main_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(5, 0), before=status_frame)

        # Get current year for copyright notice
        current_year = datetime.datetime.now().year

        # Left side - Program name and copyright
        program_text = f"© {current_year} Muallimun.Net - ListeKolay"

        program_label = tk.Label(
            footer_frame, 
            text=program_text,
            font=("Segoe UI", 8), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        program_label.pack(side=tk.LEFT, padx=10)

        # Version number next to program name
        version_label = tk.Label(
            footer_frame,
            text=f"v{self.current_version}",
            font=("Segoe UI", 8),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        version_label.pack(side=tk.LEFT, padx=(2, 0))

        # Güncelleme ikonu (Label olarak) - Kesinlikle renkli olacak
        update_icon = tk.Label(
            footer_frame,
            text="⟳",
            font=("Segoe UI", 13, "bold"),  # Daha da büyük font
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg="#FF4500" if not self.is_dark_mode.get() else "#FF9800",  # Turuncu-kırmızı renk
            cursor="hand2",  # El işareti
            padx=2
        )
        update_icon.bind("<Button-1>", lambda e: self.check_for_updates(False))  # Tıklama olayını bağla
        update_icon.pack(side=tk.LEFT, padx=(5, 0))
        self.create_tooltip(update_icon, self.get_text("check_updates"))

        # Right side - Website link
        website_link = tk.Label(
            footer_frame,
            text="www.muallimun.net",
            font=("Segoe UI", 8, "underline"),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg=LIGHT_MODE_COLORS["accent"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["accent"],
            cursor="hand2"  # Change cursor to hand when hovering
        )
        website_link.pack(side=tk.RIGHT, padx=10)

        # Make the link clickable
        website_link.bind("<Button-1>", lambda e: self.open_website("http://www.muallimun.net"))

        # İlerleme çubuğu için frame
        self.progress_frame = tk.Frame(self.main_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])

        # İlerleme çubuğu - varsayılan olarak gizli
        self.progress_bar = ttk.Progressbar(
            self.progress_frame, 
            orient=tk.HORIZONTAL, 
            length=100, 
            mode='determinate',
            variable=self.progress_var
        )
        # Not: pack_forget ile gizleniyor, gerektiğinde gösterilecek
        self.progress_bar["value"] = 0

    def create_file_list_treeview(self, parent):
        # Create frame for file list
        self.file_list_frame = tk.Frame(
            parent, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.file_list_frame.pack(fill=tk.BOTH, expand=True)

        # Create container for the file view (will hold both Tree and Preview panels)
        self.file_view_container = tk.Frame(
            self.file_list_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.file_view_container.pack(fill=tk.BOTH, expand=True)

        # Create frame for Treeview with scrollbars
        tree_frame = tk.Frame(
            self.file_view_container, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        tree_frame.pack(fill=tk.BOTH, expand=True)

        # Create scrollbars
        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        h_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)

        # Create Treeview
        self.file_tree = ttk.Treeview(
            tree_frame,
            columns=("name", "extension", "path", "size", "created", "modified"),
            show="headings",
            selectmode="extended",
            yscrollcommand=v_scrollbar.set,
            xscrollcommand=h_scrollbar.set
        )

        # Configure scrollbars
        v_scrollbar.config(command=self.file_tree.yview)
        h_scrollbar.config(command=self.file_tree.xview)

        # Pack scrollbars
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        # Pack Treeview
        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Configure columns with click-to-sort functionality
        self.file_tree.heading("name", text=self.get_text("file_name"), 
                              command=lambda: self.treeview_sort_column("name", False))
        self.file_tree.heading("extension", text=self.get_text("file_extension"),
                              command=lambda: self.treeview_sort_column("extension", False))
        self.file_tree.heading("path", text=self.get_text("file_path"),
                              command=lambda: self.treeview_sort_column("path", False))
        self.file_tree.heading("size", text=self.get_text("file_size"),
                              command=lambda: self.treeview_sort_column("size", False))
        self.file_tree.heading("created", text=self.get_text("creation_date"),
                              command=lambda: self.treeview_sort_column("created", False))
        self.file_tree.heading("modified", text=self.get_text("modification_date"),
                              command=lambda: self.treeview_sort_column("modified", False))

        self.file_tree.column("name", width=150)
        self.file_tree.column("extension", width=80)
        self.file_tree.column("path", width=200)
        self.file_tree.column("size", width=80)
        self.file_tree.column("created", width=120)
        self.file_tree.column("modified", width=120)

        # Bind right-click event for context menu
        self.file_tree.bind("<Button-3>", self.show_context_menu)

        # Bind double-click event to open file
        # Bind double-click event to open file
        self.file_tree.bind("<Double-1>", lambda event: self.open_selected_file())

        # Create right-click context menu (enhanced with more options)
        self.context_menu = tk.Menu(
            self.root, 
            tearoff=0, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )

        # Open file option
        self.context_menu.add_command(
            label=self.get_text("open_file"),
            command=self.open_selected_file
        )

        # Open file location option
        self.context_menu.add_command(
            label=self.get_text("open_file_location"),
            command=self.open_file_location
        )

        # Add separator
        self.context_menu.add_separator()

        # Preview file option
        self.context_menu.add_command(
            label=self.get_text("preview_file"),
            command=self.preview_selected_file
        )

        # Add separator
        self.context_menu.add_separator()

        # Copy filename option
        self.context_menu.add_command(
            label=self.get_text("copy_filename"),
            command=self.copy_filename_to_clipboard
        )

        # Copy file path option
        self.context_menu.add_command(
            label=self.get_text("copy_filepath"),
            command=self.copy_filepath_to_clipboard
        )

        # Add separator for file operations
        self.context_menu.add_separator()

        # Select all files option
        self.context_menu.add_command(
            label=self.get_text("select_all_files"),
            command=self.select_all_files
        )

        # Delete files option
        self.context_menu.add_command(
            label=self.get_text("delete_file"),
            command=self.delete_selected_files
        )

        # Copy files option
        self.context_menu.add_command(
            label=self.get_text("copy_file"),
            command=self.copy_selected_files
        )

        # Move files option
        self.context_menu.add_command(
            label=self.get_text("move_file"),
            command=self.move_selected_files
        )

        # Cut files option
        self.context_menu.add_command(
            label=self.get_text("cut_file"),
            command=self.cut_selected_files
        )

        # Rename file option
        self.context_menu.add_command(
            label=self.get_text("rename_file"),
            command=self.rename_selected_file
        )

        # Set up context menu for preview mode thumbnails
        self.preview_context_menu = tk.Menu(
            self.root, 
            tearoff=0, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )

        # Open file option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("open_file"),
            command=self.open_preview_file
        )

        # Open file location option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("open_file_location"),
            command=self.open_preview_file_location
        )

        # Add separator
        self.preview_context_menu.add_separator()

        # Copy filename option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("copy_filename"),
            command=self.copy_preview_filename_to_clipboard
        )

        # Copy file path option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("copy_filepath"),
            command=self.copy_preview_filepath_to_clipboard
        )

    def show_context_menu(self, event):
        """Show context menu on right-click in the file treeview"""
        # Identify the item under the cursor
        item = self.file_tree.identify_row(event.y)

        if item:
            # Check if the item under cursor is already selected
            already_selected = item in self.file_tree.selection()

            # If the item is not in the current selection, clear selection and select only this item
            if not already_selected:
                # If holding Ctrl or Shift key, add to selection instead of replacing it
                if not (event.state & 0x0004) and not (event.state & 0x0001):  # Ctrl or Shift not pressed
                    self.file_tree.selection_set(item)

            # Show the context menu
            try:
                self.context_menu.tk_popup(event.x_root, event.y_root)
            finally:
                # Make sure to release the grab
                self.context_menu.grab_release()

    def preview_selected_file(self):
        """Open a preview window for the selected file"""
        # Get the selected item
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection

        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")

        if not values:
            return  # No values found

        # Extract file name and path
        file_name = values[0]
        file_ext = values[1]
        file_dir_path = values[2]

        # Construct full file path
        # For Windows paths that already include filename, use as-is
        if os.path.basename(file_dir_path) == file_name:
            file_path = file_dir_path
        else:
            # Otherwise join directory and filename
            file_path = os.path.join(file_dir_path, file_name)

        # Check if the file exists
        if not os.path.isfile(file_path):
            messagebox.showerror(
                self.get_text("error"),
                f"{file_path} not found."
            )
            return

        # Open the preview window
        self.create_file_preview_window(file_path)

    def open_selected_file(self):
        """Open the selected file with the default application"""
        # Get the selected item
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection

        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")

        if not values:
            return  # No values found

        # Extract file path and name
        file_name = values[0]
        file_ext = values[1]
        file_dir_path = values[2]

        # Construct full file path
        # For Windows paths that already include filename, use as-is
        if os.path.basename(file_dir_path) == file_name:
            file_path = file_dir_path
        else:
            # Otherwise join directory and filename
            file_path = os.path.join(file_dir_path, file_name)

        # Check if the file exists
        if not os.path.isfile(file_path):
            messagebox.showerror(
                self.get_text("error"),
                f"{file_path} not found."
            )
            return

        # Open the file
        self.open_file(file_path)

    def toggle_filter_section(self):
        """Toggle the visibility of the filter section"""
        if hasattr(self, "filter_frame"):
            # Check if the filter panel is currently visible
            is_visible = False
            for widget in self.filter_frame.master.winfo_children():
                if widget == self.filter_frame and widget.winfo_ismapped():
                    is_visible = True
                    break

            if is_visible:
                # Hide the filter panel
                self.filter_frame.pack_forget()
                self.show_filter_btn.config(bg="#17a2b8")  # Reset button color

                # Show the tips frame if it was hidden
                if hasattr(self, 'tips_frame'):
                    self.tips_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))
            else:
                # ALWAYS hide the tips panel when filter is shown - per user request
                if hasattr(self, 'tips_frame'):
                    self.tips_frame.pack_forget()

                # Show the filter panel - expand to use the full available height
                self.filter_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0), after=self.settings_frame)
                self.show_filter_btn.config(bg="#138496")  # Change button color to indicate active state

                # Clear all extension selections when the filter section is opened
                self.clear_all_extensions()

    def filter_extensions(self, *args):
        """Filter the displayed extensions based on search text"""
        search_text = self.extension_search_var.get().lower()

        # If search text is empty, just refresh without search
        if not search_text:
            self.populate_categories()

            # Hide any previous "not found" message
            if hasattr(self, 'not_found_label') and self.not_found_label:
                self.not_found_label.pack_forget()
            return

        # Clear all extension selections when a new search begins
        for ext in self.selected_extensions:
            self.selected_extensions[ext].set(False)
        # Repopulate categories with the search filter
        self.populate_categories(search_filter=search_text)

        # Track if we found any matching extensions
        found_match = False
        matching_category = None

        # Check each category for matching extensions
        for category in self.file_categories:
            # Get extensions for this category
            category_extensions = self.file_categories[category]

            # Check if any extension in this category matches the search
            matching_extensions = [ext for ext in category_extensions if search_text in ext.lower()]

            if matching_extensions:
                found_match = True
                matching_category = category

                # Set this category as active and expand it
                self.toggle_category_display(category)
                break

        # Display "not found" message if no matches
        if not found_match:
            # Create or update "not found" message
            if not hasattr(self, 'not_found_label') or not self.not_found_label:
                self.not_found_label = tk.Label(
                    self.extensions_container,
                    text=f"'{search_text}' not found in any file extension",
                    font=("Segoe UI", 9, "italic"),
                    bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                    fg="#dc3545"  # Keep red text for error message in both themes
                )
            else:
                self.not_found_label.config(text=f"'{search_text}' not found in any file extension")

            # Display the message
            self.not_found_label.pack(pady=10)

            # Hide any previously shown extension frames
            self.active_category = None
            for cat in self.file_categories.keys():
                if hasattr(self, f"{cat}_extensions_frame"):
                    ext_frame = getattr(self, f"{cat}_extensions_frame")
                    ext_frame.pack_forget()

                    # Reset button appearance
                    if hasattr(self, f"{cat}_btn"):
                        btn = getattr(self, f"{cat}_btn")
                        btn.config(bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], relief=tk.RAISED)
        else:
            # Hide any previous "not found" message
            if hasattr(self, 'not_found_label') and self.not_found_label:
                self.not_found_label.pack_forget()

        # Restore the selection state for category checkboxes
        # Create a dictionary to track which categories are visible
        visible_categories = {cat: False for cat in self.file_categories.keys()}

        # If we have matching extensions, mark their categories as visible
        for category in self.file_categories.keys():
            # Check if this category has extensions that match the search
            if search_text:  # Using search_text instead of undefined search_filter
                cat_exts = self.file_categories[category]
                matching_exts = [ext for ext in cat_exts if search_text in ext.lower()]
                visible_categories[category] = bool(matching_exts)
            else:
                visible_categories[category] = True

        # Update category checkbox state based on visibility
        for category, is_visible in visible_categories.items():
            if category in self.category_vars:
                # Only set to false if not visible, preserve user selection if visible
                if not is_visible:
                    self.category_vars[category].set(False)

    def populate_categories(self, search_filter=None):
        # Clear existing widgets
        for widget in self.category_inner_frame.winfo_children():
            widget.destroy()

        # Create a frame for the categories - using 3x2 grid layout
        categories_container = tk.Frame(
            self.category_inner_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        categories_container.pack(fill=tk.BOTH, expand=True)

        # Process each category from the file_categories dictionary
        all_categories = list(self.file_categories.keys())

        # Create category variables if not already created
        if not hasattr(self, 'category_vars'):
            self.category_vars = {}

        # Create extensions container (will hold all extension panels)
        self.extensions_container = tk.Frame(
            self.category_inner_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.extensions_container.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

        # Track currently active category
        if not hasattr(self, 'active_category'):
            self.active_category = None

        # Create checkboxes in a 3x2 grid
        for i, category in enumerate(all_categories):
            # Determine grid position (3 columns, 2 rows)
            row = i // 3
            col = i % 3

            # Create a variable for the category checkbox if not exists
            if category not in self.category_vars:
                self.category_vars[category] = tk.BooleanVar(value=False)

            # Create a frame for this category in the grid
            category_frame = tk.Frame(
                categories_container, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
                padx=5, 
                pady=5
            )
            category_frame.grid(row=row, column=col, sticky="nsew")

            # Create the category button with checkbox appearance
            # Doğrudan kategori adlarını kullan - eşleştirmeye gerek yok
            category_mapped = category

            category_label_text = self.get_text(category_mapped)
            category_btn = tk.Button(
                category_frame,
                text=category_label_text,
                font=("Segoe UI", 9, "bold"),
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                bd=1,
                relief=tk.RAISED,
                padx=8,
                pady=3,
                anchor=tk.W,
                command=lambda cat=category: self.toggle_category_display(cat)
            )
            self.create_tooltip(category_btn, self.get_text("tooltip_category_expand"))
            category_btn.pack(side=tk.TOP, anchor=tk.W, fill=tk.X)

            # Store reference to category button
            setattr(self, f"{category}_btn", category_btn)

            # Create extensions frame for this category (initially hidden)
            extensions_frame = tk.Frame(
                self.extensions_container, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
            )
            setattr(self, f"{category}_extensions_frame", extensions_frame)

            # Add a select all checkbox at the top of extensions
            select_all_frame = tk.Frame(
                extensions_frame, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
            )
            select_all_frame.pack(fill=tk.X, pady=(0, 5))

            select_all_cb = tk.Checkbutton(
                select_all_frame,
                text=self.get_text("select_all_category"),
                variable=self.category_vars[category],
                font=("Segoe UI", 9, "italic"),
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                selectcolor="#ffffff",  # Beyaz tik kutuları
                command=lambda cat=category: self.toggle_all_category_extensions(cat)
            )
            self.create_tooltip(select_all_cb, self.get_text("tooltip_select_category"))
            select_all_cb.pack(side=tk.LEFT, padx=(5, 0))

            # Create the extensions within this frame
            self._create_extensions_list(category, extensions_frame, search_filter=search_filter)

    def _create_extensions_list(self, category, parent_frame, search_filter=None):
        """Create the list of file extensions for a category"""
        extensions = self.file_categories[category]

        # Filter extensions if search is active
        highlighted_extensions = []
        if search_filter:
            # Find extensions that match the search filter
            highlighted_extensions = [ext for ext in extensions if search_filter in ext.lower()]

            # If we're not clearing search (empty string), use the filtered list
            if search_filter:
                extensions = highlighted_extensions

        # Skip if no extensions match search (but not if clearing search)
        if search_filter and not extensions:
            return

        # Create grid layout for extensions
        ext_grid_frame = tk.Frame(parent_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
        ext_grid_frame.pack(fill=tk.X, expand=True)

        # Determine optimal number of columns based on extension count
        num_extensions = len(extensions)
        # Daha fazla sütun kullan (3-5) ve her sütunda daha az öğe olsun
        num_columns = min(5, max(2, num_extensions // 4))  # Use 2-5 columns
        extensions_per_column = max(3, (num_extensions + num_columns - 1) // num_columns)  # En az 3 öğe

        # Create columns
        column_frames = []
        for i in range(num_columns):
            column_frame = tk.Frame(ext_grid_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
            column_frame.pack(side=tk.LEFT, fill=tk.Y, expand=True)
            column_frames.append(column_frame)

        # Dictionary to store references to highlighted checkboxes
        highlighted_checkboxes = {}

        # Add checkboxes for each extension in columns
        for i, ext in enumerate(extensions):
            column_idx = i // extensions_per_column
            if column_idx >= num_columns:
                column_idx = num_columns - 1

            ext_frame = tk.Frame(column_frames[column_idx], bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
            ext_frame.pack(fill=tk.X)

            # Create checkbox if the extension variable exists, otherwise create it
            if ext not in self.selected_extensions:
                self.selected_extensions[ext] = tk.BooleanVar(value=False)

            # Determine if this extension should be highlighted
            should_highlight = search_filter and ext in highlighted_extensions

            ext_cb = tk.Checkbutton(
                ext_frame, 
                text=ext, 
                variable=self.selected_extensions[ext],
                font=("Segoe UI", 9, "bold" if should_highlight else "normal"), 
                bg="#e8f4f8" if should_highlight else (LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]),  # Highlight or theme-appropriate background color
                fg="#0056b3" if should_highlight else (LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]),  # Blue for highlights, otherwise theme-appropriate text color
                selectcolor="#ffffff",  # Beyaz tik kutuları
                anchor=tk.W
            )
            self.create_tooltip(ext_cb, self.get_text("tooltip_file_extension"))
            ext_cb.pack(fill=tk.X)

            # Store reference to checkbox if highlighted
            if should_highlight:
                highlighted_checkboxes[ext] = ext_cb

                # Make this the last checkbox in the list so it's shown
                self.last_highlighted_checkbox = ext_cb

        # Check if any of the extensions in this category are selected
        any_selected = any(self.selected_extensions.get(ext, tk.BooleanVar(value=False)).get() 
                         for ext in self.file_categories.get(category, []))
        if any_selected:
            # Update the category checkbox if any extensions are selected
            self.category_vars[category].set(True)

    def toggle_category_display(self, category):
        """Show one category's extensions and hide others"""
        # Check if there's an active category and it's the same as the clicked one
        toggle_off = self.active_category == category

        # Hide all extension frames first
        for cat in self.file_categories.keys():
            if hasattr(self, f"{cat}_extensions_frame"):
                ext_frame = getattr(self, f"{cat}_extensions_frame")
                ext_frame.pack_forget()

                # Reset button appearance
                if hasattr(self, f"{cat}_btn"):
                    btn = getattr(self, f"{cat}_btn")
                    btn.config(bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], relief=tk.RAISED)

        # If we're not toggling off, show the clicked category
        if not toggle_off:
            # Set as active category
            self.active_category = category

            # Show extensions frame
            ext_frame = getattr(self, f"{category}_extensions_frame")
            ext_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

            # Update button appearance to indicate active state
            btn = getattr(self, f"{category}_btn")
            # Use a light blue for light mode and a darker blue for dark mode as active state
            active_bg = "#d1ecf1" if not self.is_dark_mode.get() else "#0d4b66"
            btn.config(bg=active_bg, relief=tk.SUNKEN)
        else:
            # We're toggling off the current category
            self.active_category = None

    def update_tips(self):
        """Update tips with the current language"""
        if hasattr(self, 'tips_frame'):
            # Clear existing tips
            for widget in self.tips_frame.winfo_children():
                widget.destroy()

            # Add new tips with updated language
            tip_labels = [
                self.get_text("tip_1"),
                self.get_text("tip_3"),
                self.get_text("tip_4"),
                self.get_text("tip_5"),
                self.get_text("tip_6"),
                self.get_text("tip_preview_formats")
            ]

            for i, tip in enumerate(tip_labels):
                tip_label = tk.Label(
                    self.tips_frame, 
                    text=f"• {tip}", 
                    font=("Segoe UI", 9), 
                    bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
                    fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                    justify=tk.LEFT if self.current_language != "ar" else tk.RIGHT,
                    wraplength=280
                )
                tip_label.pack(anchor=tk.W, pady=(0 if i > 0 else 0, 2))

    def toggle_all_category_extensions(self, category):
        """Toggle all extensions in a category"""
        # Get the selection state
        select_all = self.category_vars[category].get()

        # Find the extensions for this category
        category_extensions = []

        # Yalnızca normal kategorileri kullan, çünkü zaten mantıksal gruplamayı file_categories içinde yaptık
        if category in self.file_categories:
            category_extensions = self.file_categories[category]

        # Apply the selection to all extensions in this category
        for ext in category_extensions:
            if ext in self.selected_extensions:
                self.selected_extensions[ext].set(select_all)

        # Apply filter
        self.apply_filter()

    def toggle_category(self, category, show=None):
        """Toggle category accordion visibility"""
        # Get the button, select frame and extensions frame
        btn = getattr(self, f"{category}_btn")
        select_frame = getattr(self, f"{category}_select_frame")
        frame = getattr(self, f"{category}_frame")
        select_btn = getattr(self, f"{category}_select_btn")

        # Check if the frame is packed (visible)
        is_visible = frame.winfo_ismapped()

        # If show parameter is provided, use it instead of toggling
        if show is not None:
            should_show = show
        else:
            should_show = not is_visible

        # If we're showing this category and not restoring state after filtering
        if should_show and show is None:
            # First, collapse all other categories - selections will be preserved
            for cat in self.category_vars.keys():
                if cat != category:  # Don't affect the clicked category yet
                    try:
                        cat_btn = getattr(self, f"{cat}_btn")
                        cat_select_frame = getattr(self, f"{cat}_select_frame")
                        cat_frame = getattr(self, f"{cat}_frame")

                        # Only hide if currently visible
                        if cat_frame.winfo_ismapped():
                            cat_select_frame.pack_forget()
                            cat_frame.pack_forget()
                            cat_btn.config(text="▶ " + self.get_text(cat))
                    except (AttributeError, Exception) as e:
                        logging.warning(f"Failed to process category {cat}: {str(e)}")

        # Now toggle the clicked category
        if not should_show:
            select_frame.pack_forget()
            frame.pack_forget()
            btn.config(text="▶ " + self.get_text(category))
        else:
            select_frame.pack(fill=tk.X, padx=15)
            select_btn.pack(side=tk.LEFT, fill=tk.X)
            frame.pack(fill=tk.X, pady=(0, 5))
            btn.config(text="▼ " + self.get_text(category))

        # Update the canvas scrollregion
        self.category_canvas.configure(scrollregion=self.category_canvas.bbox("all"))

    def on_category_canvas_configure(self, event):
        # Update the canvas inner frame width to match the canvas
        self.category_canvas.itemconfig("all", width=event.width)

    def populate_sort_dropdown(self):
        # Clear current items
        self.sort_dropdown['values'] = []

        # Add localized options
        localized_options = [self.get_text(option) for option in self.sort_options]
        self.sort_dropdown['values'] = localized_options

        # Select the first option by default
        self.sort_dropdown.current(0)

    def on_search_focus_in(self, event):
        """Handle focus in event for search entry"""
        if self.file_search_entry.get() == self.get_text("search_files"):
            self.file_search_entry.delete(0, tk.END)
            # Set text color to always black for better visibility
            self.file_search_entry.config(fg="#000000")
            # Temizleme butonunu devre dışı bırak (placeholder vardı)
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.DISABLED)
        else:
            # İçerik varsa temizleme butonunu etkinleştir
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.NORMAL)

        # Highlight search field with a light blue background when focused
        self.file_search_entry.config(bg="#e6f2ff")

    def on_search_focus_out(self, event):
        """Handle focus out event for search entry"""
        if not self.file_search_entry.get():
            self.file_search_entry.insert(0, self.get_text("search_files"))
            self.file_search_entry.config(fg='gray')
            # Temizleme butonunu devre dışı bırak
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.DISABLED)
        else:
            # İçerik varsa temizleme butonunu aktif tut
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.NORMAL)

        # Reset background color when focus is lost
        self.file_search_entry.config(bg="white")

    def clear_search_field(self):
        """Arama alanını temizler ve tüm dosyaları gösterir"""
        # Placeholder metni olmadan alanı temizle
        self.file_search_entry.delete(0, tk.END)
        # Odağı arama kutusuna getir
        self.file_search_entry.focus_set()
        # Arama alanını odaklandığında olduğu gibi işaretle
        self.file_search_entry.config(bg="#e6f2ff", fg="#000000")
        # Temizleme butonunu devre dışı bırak
        self.clear_search_button.config(state=tk.DISABLED)
        # Dosya listesini güncelle
        self.filter_file_list()

    def filter_file_list(self, *args):
        """Filter the file list based on search text"""
        # Debounce zaman kontrolü - fazla sık aramaları önle
        current_time = time.time()
        debounce_wait = 0.3  # 300ms debounce

        if hasattr(self, 'last_search_time') and (current_time - self.last_search_time < debounce_wait):
            # Çok sık aramaları engelle - aramaları biriktirip tek seferde yap
            if not hasattr(self, 'search_pending') or not self.search_pending:
                self.search_pending = True
                self.root.after(int(debounce_wait * 1000), self._execute_pending_search)
            return

        # Zaman damgasını güncelle
        self.last_search_time = current_time
        self.search_pending = False

        search_text = self.file_search_var.get().lower()

        # Skip filtering if the text is the placeholder
        if search_text == self.get_text("search_files").lower():
            return

        # Clear the file list
        self.clear_file_list()

        # Temizleme butonunun durumunu güncelle (yeni)
        if hasattr(self, 'clear_search_button'):
            if search_text and search_text != self.get_text("search_files").lower():
                self.clear_search_button.config(state=tk.NORMAL)
            else:
                self.clear_search_button.config(state=tk.DISABLED)

        # If search text is empty or placeholder, show all files
        if not search_text:
            if hasattr(self, 'all_files'):
                # Büyük dosya listeleri için toplu işleme yap
                if len(self.all_files) > 1000:
                    # Özellikle büyük dosya listeleri için kullanıcı arayüzünü dondurmamak için
                    # dosyaları küçük gruplar halinde göster (her grupta 100 dosya)
                    self._update_file_list_chunk(self.all_files[:200])  # İlk 200 dosyayı hemen göster

                    # Arayüzde "Dosyalar yükleniyor..." gibi bir durum göster
                    self.update_status(f"{self.get_text('files_loading')} ({len(self.all_files)} {self.get_text('files')})")

                    # Kalan dosyaları arka planda yükle
                    self.root.after(50, lambda: self._load_remaining_files(self.all_files[200:]))
                else:
                    # Küçük dosya listeleri için tüm dosyaları doğrudan ekle
                    for file_info in self.all_files:
                        self._add_file_to_list(file_info)
            return

        # Filter files based on the search text
        if hasattr(self, 'all_files'):
            # Büyük listeler için arama göstergesi
            is_large_search = len(self.all_files) > 1000
            if is_large_search:
                self.update_status(f"{self.get_text('searching')}...")
                self.progress_bar.start(10)  # Animasyonlu ilerleme çubuğu başlat
                # Aramayı arka planda yap ve UI'ı bloklamaktan kaçın
                self.root.after(10, lambda: self._perform_search(search_text))
            else:
                # Küçük listeler için doğrudan ara
                self._perform_search(search_text)

    def _execute_pending_search(self):
        """Bekleyen arama isteğini yürüt (debounce mekanizmasının parçası)"""
        self.search_pending = False
        self.filter_file_list()

    def _perform_search(self, search_text):
        """Asıl arama işlemini gerçekleştir (filtreleme işlevi için yardımcı metod)"""
        found_files = []
        try:
            # Arama işlemini gerçekleştir
            for file_info in self.all_files:
                file_name = file_info["name"].lower()
                file_ext = file_info["extension"].lower()

                # If the search text is in file name or extension, add it to the list
                if search_text in file_name or search_text in file_ext:
                    self._add_file_to_list(file_info)
                    found_files.append(file_info)

            # If in preview mode, refresh the preview panel with filtered files
            is_preview_mode = hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview"
            if is_preview_mode:
                # Büyük listeler için önizleme yumuşak geçişi
                if len(found_files) > 100:
                    # Önce ilk grup dosyaları göster
                    first_batch = found_files[:50]
                    self._build_preview_panel(first_batch)

                    # Sonra kalan dosyaları biraz gecikmeyle ekle (UI yanıt vermeye devam etsin)
                    self.root.after(100, lambda: self._append_preview_files(found_files[50:]))
                else:
                    # Az sayıda dosya için doğrudan göster
                    self._build_preview_panel(found_files)

        finally:
            # Her durumda ilerleme çubuğunu durdur
            self.progress_bar.stop()
            # İstatistikleri güncelle
            found_count = len(found_files)
            self.update_status(f"{found_count} {self.get_text('files_found')}")

    def _append_preview_files(self, files):
        """Önizleme paneline daha fazla dosya ekle (aşamalı yükleme için)"""
        if not files or not hasattr(self, 'view_mode_var') or self.view_mode_var.get() != "preview":
            return

        # Önizleme ekranını güncelle, mevcut içeriği koruyarak
        self._build_preview_panel(files, append=True)

    def _load_remaining_files(self, files, chunk_size=200):
        """Kalan dosyaları parçalar halinde yükle - UI'yi bloklamadan büyük listeleri göster"""
        if not files:
            self.update_status(f"{len(self.all_files)} {self.get_text('files_loaded_message')}")
            return

        # Bir sonraki parçayı işle
        chunk = files[:chunk_size]
        remaining = files[chunk_size:]

        # Chunks for this iteration
        self._update_file_list_chunk(chunk)

        # Eğer daha fazla dosya varsa, bir sonraki parçaya geç
        if remaining:
            # İlerleme bilgisini güncelle
            processed = len(self.all_files) - len(remaining)
            percent = (processed / len(self.all_files)) * 100
            self.update_status(f"{self.get_text('files_loading')} ({percent:.1f}%)")

            # Bir sonraki parça için zamanlama yap (UI'nin donmasını önler)
            self.root.after(50, lambda: self._load_remaining_files(remaining, chunk_size))

    def toggle_theme_mode(self):
        """Koyu/açık mod arası geçiş yap ve temayı uygula"""
        # Tema değiştirme işlemi zaten devam ediyorsa, gereksiz yinelenen çağrıları önle
        if hasattr(self, 'theme_change_in_progress') and self.theme_change_in_progress:
            logging.info("Tema değişikliği zaten devam ediyor, tekrarlayan çağrı engellendi")
            return

        # Değişiklik işaretçisini ayarla
        self.theme_change_in_progress = True

        try:
            is_dark = self.is_dark_mode.get()

            # Tema renk sabitlerini seç
            theme = DARK_MODE_COLORS if is_dark else LIGHT_MODE_COLORS

            logging.info(f"Tema değiştiriliyor: {'Koyu Mod' if is_dark else 'Açık Mod'}")

            # Ana arka plan rengi
            self.root.configure(bg=theme["bg"])
            self.main_frame.configure(bg=theme["bg"])

            # Gizli sol filtreleme bölümünün temaya uygun renklerini ayarla
            if hasattr(self, 'filter_frame') and self.filter_frame:
                self.filter_frame.configure(bg=theme["bg"])

            # Uygulama genelinde tüm widget'ları güncelle
            self._update_widget_colors(self.main_frame, theme)

            # Entry widget'ların metin renklerini özellikle güncelle
            if hasattr(self, 'file_search_entry'):
                if self.file_search_entry.get() == self.get_text("search_files"):
                    # Placeholder metin gri kalmalı
                    self.file_search_entry.config(fg='gray')
                else:
                    # Normal metin tema rengine uymalı
                    self.file_search_entry.config(fg=theme["text"])

            # Toggle butonunu özellikle güncelle
            if hasattr(self, 'toggle_left_panel_btn'):
                self.toggle_left_panel_btn.configure(
                    bg=theme["bg"],
                    fg=theme["text"],
                    activebackground=theme["btn_active_bg"],
                    activeforeground=theme["text"]
                )

            # TÜM BUTONLARI ve METİNLERİ ZORLA GÜNCELLE - Tema geçiş sorununu çözer
            # Bu bölüm hem açık hem koyu mod için çalışır
            all_buttons = []
            all_labels = []
            
            # Tüm buton ve label widget'ları topla
            def collect_widgets(widget):
                for child in widget.winfo_children():
                    if isinstance(child, tk.Button):
                        all_buttons.append(child)
                    elif isinstance(child, tk.Label):
                        all_labels.append(child)
                    # Alt widget'ları da tara
                    collect_widgets(child)
            
            collect_widgets(self.main_frame)
            
            # Tüm butonları zorla güncelle
            for btn in all_buttons:
                try:
                    btn.configure(fg=theme["button_text"], 
                                activeforeground=theme["button_text"],
                                bg=theme.get("button", theme["bg"]))
                except:
                    pass
            
            # Tüm label'ları zorla güncelle  
            for label in all_labels:
                try:
                    current_fg = label.cget("fg")
                    # Gri placeholder metinler hariç, diğer tüm metinleri güncelle
                    if current_fg != "gray" and current_fg != "grey":
                        label.configure(fg=theme["text"])
                except:
                    pass
            
            # Önemli butonları özellikle güncelle
            for btn_name in ['start_btn', 'select_folder_btn', 'cancel_btn', 'apply_filter_btn', 
                           'select_all_btn', 'clear_all_btn', 'exit_btn', 'show_filter_btn']:
                if hasattr(self, btn_name):
                    btn = getattr(self, btn_name)
                    try:
                        # Buton türüne özel arka plan rengi belirle
                        if btn_name == 'start_btn':
                            btn.configure(bg=theme["start_button"], fg=theme["button_text"],
                                        activeforeground=theme["button_text"])
                        elif btn_name == 'select_folder_btn': 
                            btn.configure(bg=theme["folder_button"], fg=theme["button_text"],
                                        activeforeground=theme["button_text"])
                        elif btn_name == 'cancel_btn':
                            btn.configure(bg=theme["cancel_button"], fg=theme["button_text"],
                                        activeforeground=theme["button_text"])
                        elif btn_name == 'apply_filter_btn':
                            btn.configure(bg=theme["filter_button"], fg=theme["button_text"],
                                        activeforeground=theme["button_text"])
                        else:
                            # Genel buton rengi
                            btn.configure(bg=theme["button"], fg=theme["button_text"],
                                        activeforeground=theme["button_text"])
                    except:
                        pass

            # Config dosyasına kaydet
            self.save_config()

            logging.info("Tema başarıyla değiştirildi")

        except Exception as e:
            logging.error(f"Tema değiştirilirken hata oluştu: {str(e)}")

        finally:
            # Her durumda işaretçiyi sıfırla
            self.theme_change_in_progress = False
            logging.info("Tema değiştirme işlemi tamamlandı")

    def _update_widget_colors(self, parent, theme):
        """Belirtilen parent widget'ın altındaki tüm widget'ların renklerini güncelle"""
        # Parent widget'ın kendisini güncelle
        if isinstance(parent, (tk.Frame, tk.LabelFrame, tk.Label, tk.Button)):
            # OPTIMIZASYON: Önce mevcut rengi kontrol et, gerekli olmayan renk güncellemelerini önle
            if parent.cget("bg") != theme["bg"]:
                parent.configure(bg=theme["bg"])

            # LabelFrame başlıkları için özel işlem
            if isinstance(parent, tk.LabelFrame):
                # LabelFrame başlık metninin rengini güncelle
                parent.configure(fg=theme["text"])

            # Özel renk ayarları
            elif isinstance(parent, tk.Label):
                # Tüm etiketleri ve tüm metin içeren widget'ları güncelle
                if "fg" in parent.configure():
                    # Açık temada tüm metinler siyah, koyu temada tüm metinler beyaz
                    # Label metnini tema rengine ayarla (gri, koyu gri, açık gri, vs dikkate almadan)
                    # Tooltip rengi veya buton özel renkleri olmadığı sürece tüm metinleri güncelle
                    if parent.cget("background") != "#ffffcc":  # Tooltip rengini kontrol et
                        if parent.cget("fg") != theme["text"]:
                            parent.configure(fg=theme["text"])

            # Butonlar için özel ayarlar
            elif isinstance(parent, tk.Button):
                # Metin rengini güncelle - Filtreleme bölümündeki butonlar için özel kontrol
                button_text = parent.cget("text")

                # SORUN ÇÖZÜMÜ: Tema renkleri ve buton renkleri arasındaki tutarsızlıkları önlemek için
                # Tema geçişlerinde butonların görünüm sorunu düzeltildi

                # AYDIRLIK MOD: Beyaz metinli saydam butonlar sorununu çöz
                is_dark_mode = self.is_dark_mode.get()

                # KRITIK DÜZELTME: Tüm butonların metin renklerini ZORUNLU olarak güncelle
                # Açık temada beyaz metin sorununu tamamen çözer
                
                if button_text == self.get_text("select_all") or button_text == self.get_text("clear_all") or button_text == self.get_text("apply_filter") or button_text == "🔍" or button_text == self.get_text("filter_label"):
                    # Bu butonlar için siyah/beyaz metin rengi (temaya bağlı) - ZORUNLU GÜNCELLEME
                    parent.configure(fg=theme["text"], activeforeground=theme["text"])
                else:
                    # Diğer butonlar için standart buton metin rengi - ZORUNLU GÜNCELLEME
                    parent.configure(fg=theme["button_text"], activeforeground=theme["button_text"])
                
                # EKSTRA GÜVENCE: Açık temada beyaz metin kalmasını önle
                if not is_dark_mode and parent.cget("fg") in ["white", "#ffffff", "#FFFFFF"]:
                    parent.configure(fg=theme["button_text"], activeforeground=theme["button_text"])

                # Buton türlerine göre renk atamaları
                if "✖️ Kapat" in button_text or "❌" in button_text:
                    parent.configure(bg=theme["exit_button"])
                elif "📁 Klasör" in button_text:
                    parent.configure(bg=theme["folder_button"])
                elif "⏹️ İptal" in button_text:
                    parent.configure(bg=theme["cancel_button"])
                elif "▶️ Başlat" in button_text:
                    parent.configure(bg=theme["start_button"])
                elif "🔍 Filtre" in button_text or "Filtrele" in button_text:
                    parent.configure(bg=theme["filter_button"])
                elif "Listele" in button_text:
                    # Görünüm modu butonları için özel işlem
                    if self.view_mode_var.get() == "list":
                        parent.configure(bg=theme["active_view_button"])
                    else:
                        parent.configure(bg=theme["inactive_view_button"])
                elif "Ön İzleme" in button_text or "Preview" in button_text:
                    # Görünüm modu butonları için özel işlem
                    if self.view_mode_var.get() == "preview":
                        parent.configure(bg=theme["active_view_button"])
                    else:
                        parent.configure(bg=theme["inactive_view_button"])

                # Düz butonlar için (simge butonlar)
                if "relief" in parent.configure() and parent.cget("relief") == tk.FLAT:
                    if "fg" in parent.configure() and parent.cget("fg") == "#007bff":
                        parent.configure(fg=theme["accent"])

        # Tüm alt widget'ları yinelemeli olarak güncelle
        for child in parent.winfo_children():
            self._update_widget_colors(child, theme)

    def _delayed_theme_update(self):
        """
        Geciktirilmiş tema güncellemesi - yükleme/dil değişikliği sonrası kullanılır.
        Planlanan güncelleme bittiğinde işaretleri temizler.
        """
        try:
            # Tema güncellemesini uygula
            self._force_theme_update()
        finally:
            # Planlama işaretçisini temizle
            if hasattr(self, 'theme_update_scheduled'):
                self.theme_update_scheduled = False

    def _force_theme_update(self):
        """
        Tema değişikliklerini yeniden uygula. Bu metod, dil değişikliğinden sonra 
        butonların doğru renklenmesi için kullanılır.
        """
        # Tema güncelleme işlemi zaten devam ediyorsa, gereksiz yinelenen çağrıları önle
        if hasattr(self, 'theme_update_in_progress') and self.theme_update_in_progress:
            logging.info("_force_theme_update zaten çalışıyor, tekrarlayan çağrı engellendi")
            return

        # Güncelleme işaretçisini ayarla
        self.theme_update_in_progress = True

        try:
            logging.info("Tema değişiklikleri yeniden uygulanıyor (dil değişikliğinden sonra)")
            # Mevcut tema renkleri
            theme = DARK_MODE_COLORS if self.is_dark_mode.get() else LIGHT_MODE_COLORS

            # Önemli butonların renklerini doğrudan güncelle
            for btn_name in ['start_btn', 'select_folder_btn', 'cancel_btn', 'apply_filter_btn']:
                if hasattr(self, btn_name):
                    btn = getattr(self, btn_name)
                    if btn_name == 'start_btn':
                        btn.configure(bg=theme["start_button"], fg=theme["button_text"])
                    elif btn_name == 'select_folder_btn': 
                        btn.configure(bg=theme["folder_button"], fg=theme["button_text"])
                    elif btn_name == 'cancel_btn':
                        btn.configure(bg=theme["cancel_button"], fg=theme["button_text"])
                    elif btn_name == 'apply_filter_btn':
                        btn.configure(bg=theme["filter_button"], fg=theme["button_text"])

            # Önemli butonları hemen güncelle (listede olmayanlar)
            # Özel kontrolleri genel taramaya bırakmak yerine doğrudan işleyelim
            if hasattr(self, 'view_frame') and self.view_frame:
                for child in self.view_frame.winfo_children():
                    if isinstance(child, tk.Button):
                        button_text = child.cget("text")
                        # Görünüm modu butonları
                        if "Listele" in button_text or "List" in button_text:
                            if self.view_mode_var.get() == "list":
                                child.configure(bg=theme["active_view_button"])
                            else:
                                child.configure(bg=theme["inactive_view_button"])
                            child.configure(fg=theme["button_text"])
                        elif "Ön İzleme" in button_text or "Preview" in button_text:
                            if self.view_mode_var.get() == "preview":
                                child.configure(bg=theme["active_view_button"])
                            else:
                                child.configure(bg=theme["inactive_view_button"])
                            child.configure(fg=theme["button_text"])

            logging.info("Tema değişiklikleri başarıyla yeniden uygulandı")

        except Exception as e:
            logging.error(f"_force_theme_update sırasında hata oluştu: {str(e)}")

        finally:
            # Her durumda işaretçiyi sıfırla
            self.theme_update_in_progress = False

    def update_ui_state(self):
        # Update UI state based on the current application state
        if not self.selected_folder_path:
            # Disable buttons that need a folder
            self.start_btn.config(state=tk.DISABLED)
            self.subfolder_cb.config(state=tk.DISABLED)
            self.apply_filter_btn.config(state=tk.DISABLED)
        else:
            # Enable buttons when a folder is selected
            self.start_btn.config(state=tk.NORMAL)
            self.subfolder_cb.config(state=tk.NORMAL)
            self.apply_filter_btn.config(state=tk.NORMAL)

    def create_tooltip(self, widget, text):
        def enter(event):
            # Store tooltip as attribute of widget to avoid global reference issues
            widget.tooltip = tk.Toplevel(widget)
            widget.tooltip.overrideredirect(True)
            widget.tooltip.geometry(f"+{event.x_root + 15}+{event.y_root + 10}")

            tooltip_label = tk.Label(
                widget.tooltip,
                text=text,
                justify=tk.LEFT,
                background="#ffffcc",
                relief="solid",
                borderwidth=1,
                font=("Segoe UI", 8),
                wraplength=250
            )
            tooltip_label.pack(padx=2, pady=2)

            # For Arabic language, switch text direction
            if self.current_language == "ar":
                tooltip_label.config(justify=tk.RIGHT)

        def leave(event):
            if hasattr(widget, 'tooltip') and widget.tooltip.winfo_exists():
                widget.tooltip.destroy()
                delattr(widget, 'tooltip')

        def motion(event):
            if hasattr(widget, 'tooltip') and widget.tooltip.winfo_exists():
                widget.tooltip.geometry(f"+{event.x_root + 15}+{event.y_root + 10}")

        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)
        widget.bind("<Motion>", motion)

    def change_language(self, event=None):
        # Get the selected language
        new_lang = self.language_var.get()

        # Update the current language
        if new_lang in self.languages:
            self.current_language = new_lang

            # Update the UI with the new language
            self.update_ui_language()

            # Force update of all categories and extensions
            self.populate_categories()

            # Update main elements and titles
            self.update_main_titles()

            # Log the language change
            logging.info(f"Language changed to: {new_lang}")

    def update_main_titles(self):
        # Update window title
        self.root.title(self.get_text("full_window_title"))

        # Doğrudan alt başlık referansını güncelle
        if hasattr(self, 'subtitle_label'):
            self.subtitle_label.config(text=" - " + self.get_text("app_subtitle"))

        # Update title and subtitle
        for widget in self.main_frame.winfo_children():
            if isinstance(widget, tk.Frame):
                for child in widget.winfo_children():
                    # Güvenli bir şekilde widget'ın "text" özelliğine erişmeye çalış
                    try:
                        # Sadece Label widget'ları için
                        if isinstance(child, tk.Label):
                            # Label'ın mevcut metni
                            text = child.cget("text")

                            # Başlık güncelleme
                            if "ListeKolay" in text or "EasyLister" in text:
                                child.config(text=self.get_text("app_title"))
                            # Diğer metinleri güncelleme
                            elif text and not text.startswith(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9")) and not "©" in text:
                                # İlgili çeviri anahtarını bul
                                found_key = None

                                # Tüm dillerde anahtar ara
                                for key in self.languages[self.current_language].keys():
                                    for lang_code in self.languages.keys():
                                        if self.languages[lang_code].get(key, "") == text:
                                            found_key = key
                                            break

                                    if found_key:
                                        break

                                # Çeviriyi güncelle
                                if found_key:
                                    child.config(text=self.get_text(found_key))
                    except tk.TclError:
                        # Widget'ta "text" özelliği yoksa sessizce devam et
                        continue

    def treeview_sort_column(self, column, reverse):
        """Sort the treeview content when a column header is clicked"""
        # OPTIMIZATION: Enhanced sorting algorithm for large file lists
        # This improved implementation uses better algorithms and data structures
        # to significantly speed up sorting operations

        # OPTIMIZATION: Progress bar and cursor change for large lists 
        item_count = len(self.file_tree.get_children(''))
        show_progress = item_count > 2000  # Show progress for moderately large lists
        use_optimized_sort = item_count > 5000  # Use highly optimized algorithm for very large lists

        if show_progress:
            self.root.config(cursor="watch")
            self.update_status(self.get_text("sorting_files"))
            self.progress_bar.start(10)
            # Allow UI to update
            self.root.update_idletasks()

        try:
            # OPTIMIZATION: Special handling for columns that need special sorting
            # Date columns need datetime parsing
            if column in ["created", "modified"]:
                # Try to handle different date formats
                try:
                    # Use a separate function to parse dates more efficiently
                    def parse_date(date_str):
                        try:
                            # Try common formats
                            for fmt in ["%Y-%m-%d %H:%M:%S", "%d.%m.%Y %H:%M:%S", "%m/%d/%Y %H:%M:%S"]:
                                try:
                                    return datetime.datetime.strptime(date_str, fmt)
                                except ValueError:
                                    continue
                            # If none of the formats match, use a simple string comparison
                            return date_str
                        except Exception:
                            return date_str

                    # Use list comprehension for better performance
                    l = [(parse_date(self.file_tree.set(k, column)), k) for k in self.file_tree.get_children('')]
                except Exception:
                    # Fallback to string sorting if date parsing fails
                    l = [(self.file_tree.set(k, column), k) for k in self.file_tree.get_children('')]

            # Size column needs numeric conversion
            elif column == "size":
                # Cache the size conversion function for reuse
                def extract_size_bytes(size_text):
                    if "KB" in size_text:
                        return float(size_text.replace(" KB", "")) * 1024
                    elif "MB" in size_text:
                        return float(size_text.replace(" MB", "")) * 1024 * 1024
                    elif "GB" in size_text:
                        return float(size_text.replace(" GB", "")) * 1024 * 1024 * 1024
                    elif "B" in size_text:
                        return float(size_text.replace(" B", ""))
                    return 0

                # OPTIMIZATION: Pre-allocate the list to avoid resizing
                l = [(extract_size_bytes(self.file_tree.set(k, column)), k) for k in self.file_tree.get_children('')]

            # Default string sorting
            else:
                l = [(self.file_tree.set(k, column), k) for k in self.file_tree.get_children('')]

            # OPTIMIZATION: Use natural sort for filenames and paths
            if column in ["name", "path"]:
                # Natural sort for filenames (1, 2, 10 instead of 1, 10, 2)
                import re
                def natural_sort_key(s):
                    return [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', str(s[0]))]

                l.sort(key=natural_sort_key, reverse=reverse)
            else:
                # Regular sort for other columns
                l.sort(reverse=reverse)

            # OPTIMIZATION: Batch move items for better performance
            # First detach all items
            if item_count > 1000:
                # For very large lists, detach and reattach all at once
                items = [k for _, k in l]
                self.file_tree.detach(*self.file_tree.get_children(''))

                # Reattach in the new order
                for idx, item in enumerate(items):
                    self.file_tree.move(item, '', idx)
            else:
                # For smaller lists, just move items
                for index, (val, k) in enumerate(l):
                    self.file_tree.move(k, '', index)

            # Reverse sort next time
            self.file_tree.heading(column, command=lambda: self.treeview_sort_column(column, not reverse))

            # Update column headers to show sort indication
            for col in ["name", "extension", "path", "size", "created", "modified"]:
                if col == column:
                    if reverse:
                        self.file_tree.heading(col, text=f"▼ {self.get_text(f'file_{col}' if col != 'extension' else 'file_extension')}")
                    else:
                        self.file_tree.heading(col, text=f"▲ {self.get_text(f'file_{col}' if col != 'extension' else 'file_extension')}")
                else:
                    self.file_tree.heading(col, text=self.get_text(f"file_{col}" if col != 'extension' else "file_extension"))

        finally:
            # Always restore cursor and progress bar
            if show_progress:
                self.root.config(cursor="")
                self.progress_bar.stop()
                self.update_status(self.get_text("ready"))

    def set_view_mode(self, mode):
        """Set the view mode (list or preview)"""
        # Only change if we're switching modes
        if self.view_mode_var.get() == mode:
            return

        self.view_mode_var.set(mode)

        # Update button appearance for the view mode buttons
        if hasattr(self, 'list_view_btn') and hasattr(self, 'preview_view_btn'):
            # Update list view button
            self.list_view_btn.config(
                bg="#17a2b8" if mode == "list" else "#e9ecef",
                fg="white" if mode == "list" else "#495057"
            )

            # Update preview view button
            self.preview_view_btn.config(
                bg="#17a2b8" if mode == "preview" else "#e9ecef",
                fg="white" if mode == "preview" else "#495057"
            )

        # Switch view based on mode
        if mode == "list":
            # Show the normal file list view
            self._switch_to_list_view()
        else:  # preview mode
            # Show the preview view
            self._switch_to_preview_view()

    def _switch_to_list_view(self):
        """Switch to normal list view"""
        # Hide preview frame if it exists
        if hasattr(self, 'preview_frame') and self.preview_frame.winfo_exists():
            self.preview_frame.pack_forget()

        # Show the file tree
        self.file_tree.master.pack(fill=tk.BOTH, expand=True)

        # Update status
        self.update_status(self.get_text("ready"))

    def _switch_to_preview_view(self):
        """Switch to preview view mode"""
        # Hide the tree view
        self.file_tree.master.pack_forget()

        # Create preview container if needed
        if not hasattr(self, 'preview_frame') or not self.preview_frame.winfo_exists():
            # Main preview area - contains both thumbnail view and pagination
            self.preview_frame = tk.Frame(self.file_view_container, bg="#e9ecef")
            self.preview_frame.pack(fill=tk.BOTH, expand=True)

            # Add the main content area first (without pagination)
            # This ensures thumbnails take up most of the space
            self.thumbnails_area = tk.Frame(self.preview_frame, bg="#e9ecef")
            self.thumbnails_area.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

            # Add pagination frame at bottom - make it more visible with a subtle border and background
            self.pagination_frame = tk.Frame(
                self.preview_frame, 
                bg="#f0f0f0", 
                height=50,  # Higher height for better visibility
                bd=1,       # Light border
                relief=tk.GROOVE  # Subtle raised effect
            )
            self.pagination_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 5))
            self.pagination_frame.pack_propagate(False)  # Maintain fixed height

            # Previous page button - nicer styling
            self.prev_page_btn = tk.Button(
                self.pagination_frame,
                text="◄ " + self.get_text("prev_page"),
                command=self._go_to_prev_page,
                bg="#e0e0e0",
                fg="#495057",
                font=("Segoe UI", 9, "bold"),  # Bold font
                relief=tk.RAISED,
                padx=15,  # More padding
                pady=3,
                bd=1,
                state=tk.DISABLED
            )
            self.prev_page_btn.pack(side=tk.LEFT, padx=15, pady=8)  # More padding

            # Page info label - centered with better styling
            self.page_info_label = tk.Label(
                self.pagination_frame,
                text=f"{self.get_text('page')} 1/1",
                bg="#f0f0f0",
                fg="#495057",
                font=("Segoe UI", 10)  # Slightly larger font
            )
            self.page_info_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

            # Next page button - matching style with prev button
            self.next_page_btn = tk.Button(
                self.pagination_frame,
                text=self.get_text("next_page") + " ►",
                command=self._go_to_next_page,
                bg="#e0e0e0",
                fg="#495057",
                font=("Segoe UI", 9, "bold"),  # Bold font
                relief=tk.RAISED,
                padx=15,  # More padding
                pady=3,
                bd=1,
                state=tk.DISABLED
            )
            self.next_page_btn.pack(side=tk.RIGHT, padx=15, pady=8)  # More padding

            # Create a canvas with scrollbar for thumbnails
            self.preview_canvas = tk.Canvas(self.thumbnails_area, bg="#e9ecef")
            preview_scrollbar = ttk.Scrollbar(self.thumbnails_area, orient=tk.VERTICAL, command=self.preview_canvas.yview)
            self.preview_canvas.config(yscrollcommand=preview_scrollbar.set)

            # Pack scrollbar and canvas
            preview_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            self.preview_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            # Create frame inside canvas for thumbnails
            self.thumbnail_container = tk.Frame(self.preview_canvas, bg="#e9ecef")
            self.preview_canvas.create_window((0, 0), window=self.thumbnail_container, anchor=tk.NW)

            # Configure canvas to resize with frame
            self.thumbnail_container.bind("<Configure>", lambda e: self.preview_canvas.configure(
                scrollregion=self.preview_canvas.bbox("all")
            ))
            
            # Create preview context menu for thumbnail items
            self.preview_context_menu = tk.Menu(
                self.root, 
                tearoff=0, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
            )
            
            # Open file option for preview
            self.preview_context_menu.add_command(
                label=self.get_text("open_file"),
                command=self.open_preview_file
            )

            # Open file location option for preview
            self.preview_context_menu.add_command(
                label=self.get_text("open_file_location"),
                command=self.open_preview_file_location
            )

            # Add separator
            self.preview_context_menu.add_separator()

            # Preview file option 
            self.preview_context_menu.add_command(
                label=self.get_text("preview_file"),
                command=self.preview_selected_preview_file
            )

            # Add separator
            self.preview_context_menu.add_separator()

            # Copy filename option for preview
            self.preview_context_menu.add_command(
                label=self.get_text("copy_filename"),
                command=self.copy_preview_filename_to_clipboard
            )

            # Copy file path option for preview
            self.preview_context_menu.add_command(
                label=self.get_text("copy_filepath"),
                command=self.copy_preview_filepath_to_clipboard
            )
            
            # Add separator for file operations
            self.preview_context_menu.add_separator()
            
            # Delete files option
            self.preview_context_menu.add_command(
                label=self.get_text("delete_files"),
                command=self.delete_preview_file
            )
            
            # Copy files option
            self.preview_context_menu.add_command(
                label=self.get_text("copy_files"),
                command=self.copy_preview_file
            )
            
            # Move files option
            self.preview_context_menu.add_command(
                label=self.get_text("move_files"),
                command=self.move_preview_file
            )
            
            # Rename file option
            self.preview_context_menu.add_command(
                label=self.get_text("rename_file"),
                command=self.rename_preview_file
            )
            
            # Add separator - no more menu items
            # Note: "Tümünü Seç" (Select All) option removed as it doesn't work well in preview mode

            # Bind mousewheel event for scrolling
            self.preview_canvas.bind_all("<MouseWheel>", lambda e: self.preview_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        else:
            # Show existing preview frame
            self.preview_frame.pack(fill=tk.BOTH, expand=True)

        # Update status
        self.update_status(self.get_text("preview_mode_active"))

        # Update with current files
        self._update_preview_panel()

    def _update_preview_panel(self):
        """Update the preview panel with current files"""
        # Only update if in preview mode
        if self.view_mode_var.get() != "preview":
            return

        # Check if we have files to display
        if not hasattr(self, 'filtered_files') or not self.filtered_files:
            # No files to display
            if hasattr(self, 'thumbnail_container'):
                try:
                    # Güvenli şekilde mevcut widget'ları temizle
                    try:
                        for widget in self.thumbnail_container.winfo_children():
                            try:
                                widget.destroy()
                            except Exception as widget_error:
                                logging.error(f"Error destroying widget: {str(widget_error)}")
                    except Exception as children_error:
                        logging.error(f"Error getting thumbnail children: {str(children_error)}")
                    
                    # Show message - güvenli şekilde yap
                    try:
                        msg_label = tk.Label(
                            self.thumbnail_container,
                            text=self.get_text("no_preview_available"),
                            font=("Segoe UI", 12),
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
                        )
                        msg_label.pack(pady=50)
                    except Exception as label_error:
                        logging.error(f"Error creating no-preview message: {str(label_error)}")
                except Exception as container_error:
                    logging.error(f"Error with thumbnail container: {str(container_error)}")
                
                # Reset preview files list
                self.current_preview_files = []
            return

        # Filter only previewable files for performance optimization
        previewable_files = []

        # Count file types for statistics (only previewable types)
        self.file_type_stats = {}

        for file_info in self.filtered_files:
            if file_info.get("is_folder", False):
                continue

            # Get file path and extension
            file_path = ""
            if "path" in file_info and "name" in file_info:
                file_path = os.path.join(file_info["path"], file_info["name"])
            else:
                continue

            file_ext = os.path.splitext(file_path)[1].lower()

            # Only include previewable file extensions
            if file_ext in PREVIEWABLE_EXTENSIONS:
                # Add this file to previewable files
                previewable_files.append(file_info)

                # Normalize extension for statistics (remove dot and make jpg/jpeg consistent)
                ext = file_ext.replace(".", "")
                if ext == "jpeg":
                    ext = "jpg"

                # Add to statistics
                if ext in self.file_type_stats:
                    self.file_type_stats[ext] += 1
                else:
                    self.file_type_stats[ext] = 1

        # Use only previewable files for the rest of the function
        self.previewable_files = previewable_files

        # Clear existing thumbnails - güvenli şekilde yap
        if hasattr(self, 'thumbnail_container'):
            try:
                for widget in self.thumbnail_container.winfo_children():
                    widget.destroy()
            except Exception as e:
                logging.error(f"Error clearing thumbnail_container: {str(e)}")
                # Widget hatası varsa, liste görünümüne geç
                self.set_view_mode("list")
                return

        # Start building thumbnails
        self.update_status(self.get_text("loading_preview"))

        # This code section is no longer needed since we create the pagination frame in _switch_to_preview_view
        # We'll just update the pagination text
        if hasattr(self, 'pagination_frame'):
            # Update pagination text based on current language
            self.prev_page_btn.config(text=f"◀ {self.get_text('prev_page')}")
            self.next_page_btn.config(text=f"{self.get_text('next_page')} ▶")

        # Get the current filtered files and rebuild the preview panel
        if hasattr(self, 'filtered_files') and self.filtered_files:
            # Store all files for preview navigation, but only those that can be previewed
            self.current_preview_files = []
            for file_info in self.filtered_files:
                # Only add files that can be previewed (not folders and only previewable file types)
                if not file_info.get("is_folder", False):
                    # Get file extension
                    file_ext = file_info.get("extension", "").lower().replace(".", "")
                    if not file_ext and "path" in file_info and "name" in file_info:
                        file_path = os.path.join(file_info["path"], file_info["name"])
                        file_ext = os.path.splitext(file_path)[1].lower().replace(".", "")

                    # Only include files that we can actually preview
                    if file_ext in ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "pdf", "psd", "ai", "eps"]:
                        self.current_preview_files.append(file_info)

            # Initialize current preview index if needed
            if not hasattr(self, 'current_preview_index'):
                self.current_preview_index = -1

            # Store all files
            self.all_preview_files = self.filtered_files

            # Calculate total pages
            total_items = len(self.filtered_files)
            self.total_preview_pages = max(1, (total_items + self.preview_items_per_page - 1) // self.preview_items_per_page)

            # Adjust current page if needed
            if self.preview_page > self.total_preview_pages:
                self.preview_page = 1

            # Update page info
            self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")

            # Set button states based on current page
            if self.preview_page <= 1:
                self.prev_page_btn.config(state=tk.DISABLED)
            else:
                self.prev_page_btn.config(state=tk.NORMAL)

            if self.preview_page >= self.total_preview_pages:
                self.next_page_btn.config(state=tk.DISABLED)
            else:
                self.next_page_btn.config(state=tk.NORMAL)

            # Get current page items from previewable files if available, otherwise from filtered files
            if hasattr(self, 'previewable_files') and self.previewable_files:
                preview_source = self.previewable_files
                total_preview_items = len(self.previewable_files)
                self.total_preview_pages = max(1, (total_preview_items + self.preview_items_per_page - 1) // self.preview_items_per_page)

                # Adjust current page if needed
                if self.preview_page > self.total_preview_pages:
                    self.preview_page = 1

                # Update page info
                self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")

                start_idx = (self.preview_page - 1) * self.preview_items_per_page
                end_idx = min(start_idx + self.preview_items_per_page, total_preview_items)
                current_page_files = preview_source[start_idx:end_idx]

                # Log for debugging
                logging.info(f"Showing preview page {self.preview_page}/{self.total_preview_pages}, items {start_idx+1}-{end_idx} of {total_preview_items} (optimized previewable files)")
            else:
                # Fallback to filtered files
                start_idx = (self.preview_page - 1) * self.preview_items_per_page
                end_idx = min(start_idx + self.preview_items_per_page, total_items)
                current_page_files = self.filtered_files[start_idx:end_idx]

                # Log for debugging
                logging.info(f"Showing preview page {self.preview_page}/{self.total_preview_pages}, items {start_idx+1}-{end_idx} of {total_items}")

            # Build preview with current page files
            self._build_preview_panel(current_page_files)

    def _build_preview_panel(self, files, append=False):
        """Build a preview panel showing file thumbnails

        Args:
            files: List of file info dictionaries to display
            append: If True, append files to existing preview panel instead of clearing it
        """
        # Only proceed if we're in preview mode
        if self.view_mode_var.get() != "preview":
            return

        # Initialize pagination attributes if not already set
        if not hasattr(self, 'preview_page'):
            self.preview_page = 1
            self.preview_items_per_page = 100  # Show 100 items per page (increased from 50)

        # Initialize or reset cancel flag
        self.cancel_flag = False
        self.enable_cancel_button()

        # Filtreleme durumunu izlemek için özel bir değişken ekliyoruz
        self.filtering_complete = False

        # "Filtreleme işlemleri devam ediyor" yazısını dönen simge ile göster
        self.update_status(self.get_text("filtering_in_progress") + " ⟳")

        # Update status
        self.update_status(self.get_text("loading_preview"))

        # OPTIMIZATION: Use batch processing for thumbnails to improve performance with large folders
        self.thumb_batch_size = self.preview_batch_size  # Use the batch size defined in initialization

        # Clear existing thumbnails (or keep them if appending)
        if hasattr(self, 'thumbnail_container'):
            if not append:
                for widget in self.thumbnail_container.winfo_children():
                    widget.destroy()

        # Create image references holder
        if not hasattr(self, 'preview_images'):
            self.preview_images = []
        elif not append:
            # Only clear if not in append mode
            self.preview_images.clear()

        # Determine preview sizes
        preview_width = 150  # Width for preview images
        preview_height = 150  # Height for preview images
        padding = 10  # Padding between thumbnails

        # Increase height for buttons
        preview_frame_height = preview_height + 60  # Add extra height for filename and buttons

        # Calculate max columns based on container width
        container_width = self.file_view_container.winfo_width()
        if container_width > 0:
            # Calculate how many thumbnails fit in the available width
            # Each thumbnail takes width + padding on each side
            max_columns = max(1, container_width // (preview_width + padding * 2))
        else:
            # Default if width not yet available
            max_columns = 4  # Default number of columns

        # Use the files directly as they should already be filtered for previewable types
        # in _update_preview_panel
        preview_files = files

        # Set up progress tracking
        total_files = len(preview_files)
        self.progress_bar["value"] = 0
        self.progress_bar["maximum"] = 100

        # If no previewable files, show message
        if not preview_files:
            msg_label = tk.Label(
                self.thumbnail_container,
                text=self.get_text("no_preview_available"),
                font=("Segoe UI", 12),
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
            )
            msg_label.pack(pady=50)
            self.update_status(self.get_text("preview_mode_active"))
            return

        # Process thumbnails in a separate thread
        def process_thumbnails():
            try:
                # Dönen simge göster
                self.spinner_chars = ["⟳", "⟲", "↻", "↺"]
                self.spinner_index = 0

                # Simge güncelleme fonksiyonu
                def update_spinner():
                    if self.cancel_flag or hasattr(self, 'filtering_complete') and self.filtering_complete:
                        return

                    self.spinner_index = (self.spinner_index + 1) % len(self.spinner_chars)
                    spinner_char = self.spinner_chars[self.spinner_index]
                    self.root.after(0, lambda: self.update_status(f"{self.get_text('filtering_in_progress')} {spinner_char}"))

                    # 200ms sonra tekrar güncelle
                    self.root.after(200, update_spinner)

                # Dönen simgeyi başlat
                update_spinner()

                # Create a grid for thumbnails
                row = 0
                col = 0

                # Initialize cancel flag for this operation specifically
                self.cancel_flag = False
                self.root.after(0, self.enable_cancel_button)

                # Define a check for cancellation in the main thread
                def check_cancel():
                    return hasattr(self, 'cancel_flag') and self.cancel_flag

                # OPTIMIZATION: Process files in batches for better performance
                batches = [preview_files[i:i+self.thumb_batch_size] for i in range(0, len(preview_files), self.thumb_batch_size)]
                batch_count = 0

                # Process each batch with cancellation checks
                for batch_idx, batch in enumerate(batches):
                    # Frequently check if operation is cancelled
                    if check_cancel():
                        self.root.after(0, lambda: self.update_status(self.get_text("operation_cancelled")))
                        self.root.after(0, lambda: self.progress_bar.config(value=0))
                        # Provide visual feedback that cancellation is in progress
                        self.root.after(0, lambda: self.cancel_btn.config(text=self.get_text("cancelling")))
                        self.root.after(0, self.disable_cancel_button)
                        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))
                        return

                    # Process each file in the current batch
                    batch_files = []

                    # Update progress for this batch
                    processed_count = batch_idx * self.thumb_batch_size
                    batch_progress = processed_count / total_files * 100
                    self.root.after(0, lambda p=batch_progress: self.progress_bar.config(value=p))
                    self.root.after(0, lambda c=processed_count, t=total_files: 
                                  self.update_status(f"{self.get_text('generating_previews')} ({c}/{t})"))

                    # Process all files in this batch
                    for file_info in batch:
                        file_name = file_info.get("name", "")
                        file_path = os.path.join(file_info.get("path", ""), file_name)
                        extension = file_info.get("extension", "").lower()

                        # Store thumbnail info for grid placement
                        batch_files.append((file_name, file_path, extension, file_info))

                    # Create thumbnail frame with more details
                    def create_frame(row, col):
                        # Increased height to fit more file details
                        frame = tk.Frame(
                            self.thumbnail_container,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            width=preview_width + padding,
                            height=preview_frame_height + 40,  # Increased height for additional details
                            highlightbackground=LIGHT_MODE_COLORS["border"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["border"],
                            highlightthickness=1
                        )
                        frame.grid(row=row, column=col, padx=padding//2, pady=padding//2, sticky="nsew")
                        frame.grid_propagate(False)  # Keep frame size fixed
                        return frame

                    # PERFORMANCE OPTIMIZATION: Process thumbnails for all files in the batch using parallel processing
                    thumbnails = []
                    errors = []

                    # Track current position for grid layout
                    idx = processed_count
                    row = idx // max_columns
                    col = idx % max_columns

                    # Parallel file preview function that doesn't rely on self
                    def generate_preview_for_file(file_data):
                        try:
                            file_name, file_path, extension, file_info, idx_position = file_data

                            # Calculate grid position
                            r = idx_position // max_columns
                            c = idx_position % max_columns

                            # OPTIMIZATION: Use lower resolution previews for the general preview page
                            preview_width_reduced = int(preview_width * 0.7)  # 70% of original size for overview
                            preview_height_reduced = int(preview_height * 0.7)  # 70% of original size for overview

                            # Import required modules first
                            try:
                                import os
                                import io
                                import tempfile
                                import logging
                                from PIL import Image, ImageTk
                                
                                # Check if file exists
                                if not os.path.exists(file_path):
                                    logging.error(f"File not found: {file_path}")
                                    return "error", (r, c, file_path)
                            except ImportError as imp_err:
                                logging.error(f"Import error in generate_preview: {str(imp_err)}")
                                return "error", (r, c, file_path)
                            
                            # Try to generate preview
                            try:
                                # Log file preview generation with standard logging
                                print(f"Generating preview for file: {file_path} (type: {extension}, position: {r},{c})")
                                
                                # Call the file preview function
                                try:
                                    preview_img = self._create_file_preview(file_path, preview_width_reduced, preview_height_reduced)
                                    
                                    if preview_img:
                                        print(f"Preview successfully created for: {file_path}")
                                        return ("success", (r, c, preview_img, file_path, file_name, file_info.get("size", 0)))
                                    else:
                                        print(f"_create_file_preview returned None for {file_path}")
                                        return ("error", (r, c, file_path))
                                except Exception as preview_err:
                                    print(f"Error in _create_file_preview for {file_path}: {str(preview_err)}")
                                    return ("error", (r, c, file_path))
                            except Exception as e:
                                logging.error(f"Error creating thumbnail for {file_path}: {str(e)}")
                                return ("error", (r, c, file_path))
                        except Exception as main_e:
                            logging.error(f"Main preview generator error: {str(main_e)}")
                            return ("error", (-1, -1, "Unknown file"))

                    # Determine the number of files to process
                    num_files = len(batch_files)

                    # Prepare file data for processing
                    file_data_list = []
                    for i, (file_name, file_path, extension, file_info) in enumerate(batch_files):
                        idx_position = processed_count + i
                        file_data_list.append((file_name, file_path, extension, file_info, idx_position))

                    # Use ThreadPoolExecutor for IO-bound preview generation
                    # This is better than ProcessPoolExecutor for this task since most operations are IO-bound
                    # and the GIL is frequently released during file operations
                    with concurrent.futures.ThreadPoolExecutor(max_workers=min(8, num_files)) as executor:
                        future_to_file = {executor.submit(generate_preview_for_file, file_data): file_data 
                                         for file_data in file_data_list}

                        # Collect results as they complete
                        for future in concurrent.futures.as_completed(future_to_file):
                            if self.cancel_flag:
                                executor.shutdown(wait=False)
                                return

                            try:
                                result = future.result()
                                if len(result) >= 2:  # Ensure we have at least two items in the result tuple
                                    result_type, result_data = result[0], result[1]
                                    if result_type == "success" and result_data:
                                        file_info = result_data[3] if len(result_data) > 3 else "Unknown file"
                                        logging.info(f"Successfully added thumbnail to UI queue: {file_info}")
                                        thumbnails.append(result_data)
                                    else:
                                        error_info = result_data[2] if len(result_data) > 2 else "Unknown error"
                                        logging.warning(f"Failed to create preview, adding to errors list: {error_info}")
                                        errors.append(result_data)
                                else:
                                    logging.error(f"Invalid result format received: {result}")
                                    # Create a generic error entry
                                    errors.append((-1, -1, "Invalid result format"))
                            except Exception as e:
                                logging.error(f"Error collecting preview result: {str(e)}")

                    # Add all thumbnails to UI in main thread with improved details
                    def add_thumbnail(r, c, img, path, name, size):
                        frame = create_frame(r, c)

                        # File image/icon at the top
                        if img:
                            # Keep reference
                            self.preview_images.append(img)

                            # Create label for image
                            img_label = tk.Label(frame, image=img, 
                                                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
                            img_label.place(relx=0.5, rely=0.35, anchor=tk.CENTER)  # Moved up to make room for details

                            # Add click event - get file index from current preview files
                            file_index = next((i for i, f in enumerate(self.current_preview_files) if f.get("path") == os.path.dirname(path) and f.get("name") == os.path.basename(path)), -1)
                            img_label.bind("<Button-1>", lambda e, p=path, idx=file_index: self.create_file_preview_window(p, idx))
                            # Add right-click context menu
                            img_label.bind("<Button-3>", lambda e, p=path: self.show_preview_context_menu(e, p))
                        else:
                            # No preview, show extension icon with improved styling
                            ext = os.path.splitext(path)[1].lower().replace(".", "")
                            ext_label = tk.Label(
                                frame,
                                text=ext.upper(),
                                font=("Segoe UI", 16, "bold"),
                                bg="#4285F4",  # Google blue for better visibility
                                fg="white",
                                width=4,
                                height=2,
                                relief=tk.RAISED  # 3D effect
                            )
                            ext_label.place(relx=0.5, rely=0.35, anchor=tk.CENTER)  # Moved up

                            # Add click event - get file index from current preview files
                            file_index = next((i for i, f in enumerate(self.current_preview_files) if f.get("path") == os.path.dirname(path) and f.get("name") == os.path.basename(path)), -1)
                            ext_label.bind("<Button-1>", lambda e, p=path, idx=file_index: self.create_file_preview_window(p, idx))
                            # Add right-click context menu
                            ext_label.bind("<Button-3>", lambda e, p=path: self.show_preview_context_menu(e, p))

                        # File details container to organize information
                        details_frame = tk.Frame(
                            frame, 
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
                        )
                        details_frame.place(relx=0.5, rely=0.75, anchor=tk.CENTER, width=preview_width-10, height=70)

                        # Add filename with better styling
                        display_name = name
                        if len(name) > 20:
                            display_name = name[:17] + "..."

                        name_label = tk.Label(
                            details_frame,
                            text=display_name,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                            font=("Segoe UI", 9, "bold"),  # Bold for emphasis
                            wraplength=preview_width-10
                        )
                        name_label.pack(pady=(0, 2))

                        # Add click event to filename - get file index from current preview files
                        file_index = next((i for i, f in enumerate(self.current_preview_files) if f.get("path") == os.path.dirname(path) and f.get("name") == os.path.basename(path)), -1)
                        name_label.bind("<Button-1>", lambda e, p=path, idx=file_index: self.create_file_preview_window(p, idx))

                        # Add file size information
                        size_str = self.format_file_size(size)
                        size_label = tk.Label(
                            details_frame,
                            text=size_str,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8)
                        )
                        size_label.pack(pady=(0, 2))

                        # Add file extension info
                        ext = os.path.splitext(path)[1].lower()
                        ext_info_label = tk.Label(
                            details_frame,
                            text=ext,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8)
                        )
                        ext_info_label.pack(pady=(0, 2))

                    # Function to show error thumbnail with improved styling
                    def show_error(r, c, path):
                        frame = create_frame(r, c)

                        # Improved error icon
                        error_icon_frame = tk.Frame(frame, bg="#dc3545", width=60, height=60, relief=tk.RAISED, bd=2)
                        error_icon_frame.place(relx=0.5, rely=0.35, anchor=tk.CENTER)

                        err_label = tk.Label(
                            error_icon_frame,
                            text="!",
                            font=("Segoe UI", 20, "bold"),
                            bg="#dc3545",
                            fg="white",
                            width=3,
                            height=2
                        )
                        err_label.pack(fill=tk.BOTH, expand=True)

                        # File details container similar to normal thumbnails
                        details_frame = tk.Frame(
                            frame, 
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
                        )
                        details_frame.place(relx=0.5, rely=0.75, anchor=tk.CENTER, width=preview_width-10, height=70)

                        # Add filename with error indication
                        name = os.path.basename(path)
                        if len(name) > 20:
                            name = name[:17] + "..."

                        name_label = tk.Label(
                            details_frame,
                            text=name,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg="#dc3545",  # Red text for error (always red regardless of theme)
                            font=("Segoe UI", 9, "bold"),
                            wraplength=preview_width-10
                        )
                        name_label.pack(pady=(0, 2))

                        # Add error message
                        error_label = tk.Label(
                            details_frame,
                            text="Önizleme kullanılamıyor",  # Preview not available
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8),
                            wraplength=preview_width-10
                        )
                        error_label.pack(pady=(0, 2))

                        # Add file extension
                        ext = os.path.splitext(path)[1].lower()
                        ext_label = tk.Label(
                            details_frame,
                            text=ext,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8)
                        )
                        ext_label.pack(pady=(0, 2))

                    # Update UI with all thumbnails in this batch
                    for r, c, img, path, name, size in thumbnails:
                        self.root.after(0, lambda r=r, c=c, img=img, path=path, name=name, size=size: 
                                       add_thumbnail(r, c, img, path, name, size))

                    # Update UI with all errors in this batch
                    for r, c, path in errors:
                        self.root.after(0, lambda r=r, c=c, path=path: show_error(r, c, path))

                    # İlk batch yüklendiyse (ilk sayfa görüntülenecek durumdaysa)
                    # istatistikleri göster ve filtreleme işleminin tamamlandığını belirt
                    if batch_idx == 0:
                        self.filtering_complete = True
                        stats_message = self._calculate_file_type_statistics()
                        if stats_message:
                            self.root.after(0, lambda msg=stats_message: self.status_var.set(msg))
                            logging.info(f"Displaying initial file stats: {stats_message}")

                    # Update progress after processing this batch
                    batch_count += 1

                    # Update grid position for next batch (if any)
                    processed_count += len(batch_files)
                    row = processed_count // max_columns
                    col = processed_count % max_columns

                    # Brief pause to keep UI responsive
                    self.root.update_idletasks()
                    time.sleep(0.05)  # Small delay to prevent UI freezing

                # Configure the scroll region
                self.root.after(0, lambda: self.thumbnail_container.update_idletasks())
                self.root.after(0, lambda: self.preview_canvas.config(scrollregion=self.preview_canvas.bbox("all")))

                # Filtreleme ve önizleme yüklenmesi tamamlandı, bayrak ayarla
                self.filtering_complete = True

                # İlerleme çubuğunu sıfırla
                self.root.after(0, lambda: self.progress_bar.config(value=0))

                # İstatistikleri hesapla ve göster
                stats_message = self._calculate_file_type_statistics()
                if stats_message:
                    self.root.after(0, lambda msg=stats_message: self.status_var.set(msg))
                    logging.info(f"Displaying file stats when thumbnails loaded: {stats_message}")
                else:
                    # Eğer istatistikler hesaplanamazsa, varsayılan durumu göster
                    self.root.after(0, lambda: self.update_status(self.get_text("preview_mode_active")))

            except Exception as e:
                logging.error(f"Error building preview panel: {str(e)}")
                self.root.after(0, lambda: self.update_status(f"Error building preview: {str(e)}"))
                self.filtering_complete = True  # Hata olsa bile tamamlandı olarak işaretle

        # Start processing in background thread
        threading.Thread(target=process_thumbnails, daemon=True).start()

    def _create_eps_preview(self, file_path, max_width, max_height):
        """Specialized function to create a preview for EPS files.
        Uses multiple methods with performance optimizations for large files."""

        # Fallback function to create placeholder
        def create_eps_placeholder():
            color = "#8BC34A"  # Green for EPS
            img = Image.new('RGB', (max_width, max_height), color)
            draw = ImageDraw.Draw(img)
            # Add a border and text
            draw.rectangle([(0, 0), (max_width-1, max_height-1)], outline="white", width=2)
            draw.text((max_width//2, max_height//2), "EPS", fill="white", anchor="mm")
            return ImageTk.PhotoImage(img)

        # Dosya boyutu kontrolü
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)  # MB cinsinden
            is_large_eps = file_size_mb > 10  # 10MB'dan büyük EPS dosyaları büyük olarak kabul edilir
        except Exception:
            is_large_eps = False
            
        # Büyük dosyalar için uyarı göster
        if is_large_eps and hasattr(self, 'update_status'):
            self.update_status(self.get_text("loading_large_file"))
            
        # Büyük dosyalar için DPI ve çözünürlük ayarlarını düzenle
        dpi = 72 if is_large_eps else 150
        density = '72' if is_large_eps else '150'
        scale_factor = 0.25 if is_large_eps else 0.5

        # We'll try several methods in sequence, from most reliable to least reliable
        preview_image = None

        try:
            # Create a temporary directory for conversion files
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_pdf_path = os.path.join(temp_dir, "temp_eps_preview.pdf")

                # METHOD 1: Directly use PIL to open EPS - only try for small files
                if not is_large_eps:
                    try:
                        # Set a timeout to prevent hanging on large files
                        img = Image.open(file_path)
                        # Use a smaller target size to prevent decompression bombs
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image
                    except Exception as e:
                        logging.info(f"Direct EPS loading failed: {str(e)}")
                        # PIL doğrudan EPS yüklemesi başarısız oldu, hafızayı temizle
                        if 'img' in locals():
                            del img
                            
                # METHOD 2: Use cairosvg for EPS/SVG files - works for many EPS files
                try:
                    # SVG ve EPS dosyaları benzer formatlardır, cairosvg bazen işe yarar
                    from cairosvg import svg2png
                    import io
                    
                    # EPS dosyasını oku ve SVG olarak işlemeyi dene
                    with open(file_path, 'rb') as eps_file:
                        eps_data = eps_file.read()
                    
                    # Büyük dosyalar için daha küçük boyut hedefle
                    target_width = int(max_width * 1.5)
                    target_height = int(max_height * 1.5)
                    
                    # SVG olarak dönüştürmeyi dene
                    png_data = svg2png(bytestring=eps_data, output_width=target_width, output_height=target_height)
                    
                    # PNG verilerini bir PIL görüntüsüne dönüştür
                    img = Image.open(io.BytesIO(png_data))
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                    return preview_image
                except Exception as e:
                    logging.info(f"cairosvg EPS conversion failed: {str(e)}")
                    # cairosvg hata verirse hafızayı temizle
                    if 'img' in locals():
                        del img
                    if 'eps_data' in locals():
                        del eps_data
                    if 'png_data' in locals():
                        del png_data

                # METHOD 3: Use pdf2image with performance optimizations
                try:
                    # Try to convert EPS directly to image
                    from pdf2image import convert_from_path

                    # Define poppler path to ensure we can find the tools - dinamik olarak ara
                    import shutil
                    # Poppler araçlarını ara
                    pdftoppm_path = shutil.which('pdftoppm')
                    poppler_path = os.path.dirname(pdftoppm_path) if pdftoppm_path else ""

                    # OPTIMIZASYON: Büyük dosyalar için daha düşük DPI ve daha uzun timeout
                    images = convert_from_path(
                        file_path, 
                        first_page=1, 
                        last_page=1,
                        dpi=dpi,  # Büyük dosyalar için düşük DPI
                        size=(max_width, max_height),
                        use_cropbox=True,
                        fmt='ppm',  # PPM formatı daha güvenilir
                        poppler_path=poppler_path,
                        use_pdftocairo=True,  # pdftocairo, pdftoppm'den daha iyi çalışır
                        timeout=30 if is_large_eps else 10  # Büyük dosyalar için daha uzun timeout
                    )

                    if images and len(images) > 0:
                        img = images[0]
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        # Hafızayı temizle
                        del images[0]
                        del images
                        return preview_image
                except Exception as e:
                    logging.info(f"pdf2image EPS conversion failed: {str(e)}")
                    # Hafızayı temizle
                    if 'images' in locals():
                        del images
                    if 'img' in locals():
                        del img

                # METHOD 4: Use PyMuPDF (fitz) with reduced scale factor for large files
                try:
                    pdf_doc = fitz.open(file_path)
                    if pdf_doc.page_count > 0:
                        page = pdf_doc[0]
                        # Büyük dosyalar için daha düşük zoom faktörü kullan
                        pix = page.get_pixmap(matrix=fitz.Matrix(scale_factor, scale_factor))
                        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        # Belleği serbest bırak
                        pix = None
                        
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        pdf_doc.close()
                        return preview_image
                    pdf_doc.close()
                except Exception as e:
                    logging.info(f"PyMuPDF EPS loading failed: {str(e)}")
                    # Hafızayı temizle
                    if 'pdf_doc' in locals() and hasattr(pdf_doc, 'close'):
                        pdf_doc.close()
                    if 'pix' in locals():
                        pix = None
                    if 'img' in locals():
                        del img

                # METHOD 5: Try using ImageMagick with optimized parameters for file size
                try:
                    # Use ImageMagick convert command - sistemden bul
                    import shutil
                    convert_path = shutil.which('convert') or 'convert'

                    # Create temporary output image path
                    temp_image_path = os.path.join(temp_dir, "temp_eps_preview.png")

                    # OPTIMIZASYON: Büyük dosyalar için daha düşük yoğunluk ve basitleştirilmiş dönüşüm
                    # Convert EPS to PNG using ImageMagick with density parameter optimized for file size
                    subprocess.run(
                        [convert_path, '-density', density, '-background', 'white', '-flatten', 
                         file_path, temp_image_path],
                        stdout=subprocess.PIPE, stderr=subprocess.PIPE, 
                        timeout=30 if is_large_eps else 10  # Büyük dosyalar için daha uzun timeout
                    )

                    # Check if image was created successfully
                    if os.path.exists(temp_image_path) and os.path.getsize(temp_image_path) > 0:
                        img = Image.open(temp_image_path)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image

                    # If that failed, try a simpler conversion method with lower quality for large files
                    quality_param = ['-quality', '50'] if is_large_eps else []
                    subprocess.run(
                        [convert_path] + quality_param + [file_path, temp_image_path],
                        stdout=subprocess.PIPE, stderr=subprocess.PIPE, 
                        timeout=30 if is_large_eps else 10
                    )

                    # Check if image was created with the simpler method
                    if os.path.exists(temp_image_path) and os.path.getsize(temp_image_path) > 0:
                        img = Image.open(temp_image_path)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image

                except Exception as e:
                    logging.info(f"ImageMagick EPS conversion failed: {str(e)}")
                    # Hafızayı temizle
                    if 'img' in locals():
                        del img

        except Exception as e:
            logging.error(f"All EPS preview methods failed: {str(e)}")
            
            # Büyük EPS dosyalarının başarısız olması durumunda uyarı mesajını güncelle
            if is_large_eps and hasattr(self, 'update_status'):
                self.update_status(self.get_text("preview_not_available"))

        # If all methods fail or exceptions occur, create a placeholder
        return create_eps_placeholder()

    def _create_file_preview(self, file_path, max_width=150, max_height=150):
        """Create a thumbnail preview for a file based on its type"""
        # İçe aktarma - kesinlikle gerekli
        import os
        import tempfile
        import io
        import logging
        import subprocess
        from PIL import Image, ImageTk, ImageDraw, ImageFont
        
        # Normalize file path to avoid Windows/Unix path issues
        file_path = os.path.normpath(file_path)

        # PERFORMANCE OPTIMIZATION: Implement a thread-safe LRU caching mechanism for thumbnails
        # This prevents regenerating the same thumbnails multiple times
        # Create a cache key based on the file path and requested dimensions
        cache_key = f"{file_path}_{max_width}_{max_height}"

        # Initialize cache structures if they don't exist
        if not hasattr(self, '_preview_cache_lock'):
            self._preview_cache_lock = threading.RLock()

        # Thread-safe cache operations
        with self._preview_cache_lock:
            # Initialize cache if it doesn't exist (with LRU behavior to limit memory usage)
            if not hasattr(self, '_preview_cache'):
                # Use OrderedDict for efficient LRU cache behavior
                self._preview_cache = collections.OrderedDict()
                self._preview_cache_max_size = 200  # Limit cache size to avoid memory issues

            # Check if we've already generated this thumbnail
            if cache_key in self._preview_cache:
                # Move item to the end to mark as recently used
                value = self._preview_cache.pop(cache_key)
                self._preview_cache[cache_key] = value
                return value

            # PERFORMANCE OPTIMIZATION: Limit cache size with LRU eviction policy
            # When the cache gets full, remove the oldest (least recently used) items first
            if len(self._preview_cache) >= self._preview_cache_max_size:
                # Remove oldest item (first item in OrderedDict)
                self._preview_cache.popitem(last=False)

        # Check the file extension
        file_ext = os.path.splitext(file_path)[1].lower()

        preview_image = None

        try:
            # SVG files require special handling
            if file_ext == '.svg':
                try:
                    # SVG işleme için cairosvg modülünü kullan
                    import io
                    from cairosvg import svg2png
                    
                    # SVG boyutunu belirle (dosyayı açarak)
                    try:
                        import xml.etree.ElementTree as ET
                        tree = ET.parse(file_path)
                        root = tree.getroot()
                        
                        # SVG boyutlarını al (varsayılan değerler 300x300)
                        width = int(float(root.get('width', '300').replace('px', '').strip()))
                        height = int(float(root.get('height', '300').replace('px', '').strip()))
                        
                        # Boyut oranını koru
                        scale = min(max_width / width, max_height / height)
                        new_width = int(width * scale)
                        new_height = int(height * scale)
                    except Exception as svg_size_error:
                        logging.error(f"SVG boyutu belirlenemedi: {svg_size_error}")
                        new_width, new_height = max_width, max_height
                    
                    # SVG'yi PNG'ye dönüştür
                    png_data = svg2png(url=file_path, output_width=new_width, output_height=new_height)
                    
                    # PNG verilerini bir PIL görüntüsüne dönüştür
                    img = Image.open(io.BytesIO(png_data))
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as svg_error:
                    logging.error(f"SVG önizleme oluşturulamadı: {str(svg_error)}")
                    # Hata durumunda ikonla göster
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", "SVG")
            
            # WebP files may need special handling for animation
            elif file_ext == '.webp':
                try:
                    # Önce dosyanın animasyonlu olup olmadığını kontrol et
                    img = Image.open(file_path)
                    
                    # WebP'nin animasyonlu olup olmadığını kontrol et
                    try:
                        is_animated = hasattr(img, "is_animated") and img.is_animated
                    except Exception:
                        is_animated = False
                    
                    if is_animated:
                        # Animasyonlu WebP için ilk kareyi al
                        img.seek(0)  # İlk kareye git
                    
                    # Yeniden boyutlandır
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as webp_error:
                    logging.error(f"WebP önizleme oluşturulamadı: {str(webp_error)}")
                    # Hata durumunda ikonla göster
                    preview_image = self._create_styled_icon(max_width, max_height, "#009688", "WEBP")
                
            # TIFF files need careful handling due to potential multi-page nature
            elif file_ext in ['.tiff', '.tif']:
                try:
                    # TIFF dosyasını aç
                    img = Image.open(file_path)
                    
                    # TIFF'in çok sayfalı olup olmadığını kontrol et
                    try:
                        is_multipage = hasattr(img, "n_frames") and img.n_frames > 1
                    except Exception:
                        is_multipage = False
                    
                    if is_multipage:
                        # Çok sayfalı TIFF için ilk sayfayı al
                        img.seek(0)  # İlk sayfaya git
                    
                    # Yeniden boyutlandır
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as tiff_error:
                    logging.error(f"TIFF önizleme oluşturulamadı: {str(tiff_error)}")
                    # Hata durumunda ikonla göster
                    preview_image = self._create_styled_icon(max_width, max_height, "#795548", "TIFF")
            
            # Video files - create thumbnail preview
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg', '.3gp']:
                preview_image = None
                try:
                    # Video önizlemesi için ffmpeg kullan
                    import tempfile
                    import subprocess
                    import os
                    import shutil
                    
                    # Geçici dosya oluştur
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                        temp_output = temp_file.name
                    
                    # ffmpeg yolu - sistemden dinamik olarak bul
                    ffmpeg_path = shutil.which('ffmpeg') or 'ffmpeg'
                    
                    # Dosya boyutunu kontrol et
                    file_size = 0
                    is_large_video = False
                    try:
                        file_size = os.path.getsize(file_path)
                        is_large_video = file_size > 1024 * 1024 * 1024  # 1GB'dan büyük mü?
                        print(f"Video dosya boyutu: {file_path} - {file_size / (1024*1024):.2f} MB")
                    except Exception as e:
                        print(f"Dosya boyutu alınamadı: {str(e)}")
                    
                    thumbnail_created = False
                    
                    # Büyük video dosyaları için (1GB+)
                    if is_large_video:
                        print(f"Büyük video dosyası algılandı, süper optimize edilmiş yöntem kullanılıyor: {file_path}")
                        thumbnail_created = False
                        
                        # İlk deneme: Doğrudan küçük resim oluştur, çok düşük kalite
                        try:
                            # Windows'ta terminal penceresi gizlemek için STARTUPINFO kullan
                            startupinfo = None
                            if os.name == 'nt':  # Windows
                                startupinfo = subprocess.STARTUPINFO()
                                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                                startupinfo.wShowWindow = subprocess.SW_HIDE
                            
                            subprocess.run([
                                ffmpeg_path, '-y', '-ss', '00:00:00.1', '-i', file_path,
                                '-vframes', '1', '-q:v', '10', 
                                '-vf', f'scale={max_width}:{max_height}',
                                '-analyzeduration', '10000',  # 10 saniye analiz
                                '-probesize', '1000000',  # 1MB analiz
                                temp_output
                            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=3, startupinfo=startupinfo)
                            
                            if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                                thumbnail_created = True
                                print(f"Büyük video için ilk yöntem başarılı: {file_path}")
                        except Exception as e:
                            print(f"Süper hızlı yöntem başarısız: {e}, alternatif deneniyor")
                        
                        # İkinci deneme: Sadece videonun ilk birkaç MB'ını işle
                        if not thumbnail_created:
                            try:
                                # Windows'ta terminal penceresi gizlemek için STARTUPINFO kullan
                                startupinfo = None
                                if os.name == 'nt':  # Windows
                                    startupinfo = subprocess.STARTUPINFO()
                                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                                    startupinfo.wShowWindow = subprocess.SW_HIDE
                                
                                # Doğrudan ilk kareyi çıkar, analiz süresini ve boyutunu sınırla
                                subprocess.run([
                                    ffmpeg_path, '-y', '-i', file_path,
                                    '-vframes', '1', '-q:v', '20',  # Çok düşük kalite = çok hızlı
                                    '-vf', f'scale={max_width//2}:{max_height//2}',  # Daha küçük ölçek
                                    '-analyzeduration', '1000',  # 1 saniye analiz
                                    '-probesize', '500000',  # 500KB analiz
                                    temp_output
                                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=2, startupinfo=startupinfo)
                                
                                if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                                    thumbnail_created = True
                                    print(f"Büyük video için ikinci yöntem başarılı: {file_path}")
                            except Exception as e2:
                                print(f"İkinci yöntem de başarısız: {e2}, son çare deneniyor")

                        # Üçüncü deneme: Dosyaya en hızlı erişim
                        if not thumbnail_created:
                            try:
                                # Windows'ta terminal penceresi gizlemek için STARTUPINFO kullan
                                startupinfo = None
                                if os.name == 'nt':  # Windows
                                    startupinfo = subprocess.STARTUPINFO()
                                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                                    startupinfo.wShowWindow = subprocess.SW_HIDE
                                
                                # En agresif yöntem - ffmpeg'in en hızlı ayarlarını kullan
                                subprocess.run([
                                    ffmpeg_path, '-y',
                                    '-analyzeduration', '100',  # Minimum analiz süresi
                                    '-probesize', '1000',  # Çok küçük probe
                                    '-i', file_path,
                                    '-frames:v', '1',  # Sadece 1 kare
                                    '-q:v', '31',  # En düşük kalite
                                    '-vf', f'scale=48:48',  # Çok küçük resim
                                    temp_output
                                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=1, startupinfo=startupinfo)
                                
                                if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                                    thumbnail_created = True
                                    print(f"Büyük video için acil yöntem başarılı: {file_path}")
                            except Exception as e3:
                                print(f"Tüm video önizleme yöntemleri başarısız: {e3}")
                    else:
                        # Normal boyutlu dosyalar için
                        try:
                            # Windows'ta terminal penceresi gizlemek için STARTUPINFO kullan
                            startupinfo = None
                            if os.name == 'nt':  # Windows
                                startupinfo = subprocess.STARTUPINFO()
                                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                                startupinfo.wShowWindow = subprocess.SW_HIDE
                            
                            # İlk kareyi çıkar (1. saniye)
                            subprocess.run([
                                ffmpeg_path, '-y', '-ss', '00:00:01', '-i', file_path,
                                '-vframes', '1', '-q:v', '2', temp_output
                            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10, startupinfo=startupinfo)
                            
                            # Başarısız olursa başlangıçtan al
                            if not os.path.exists(temp_output) or os.path.getsize(temp_output) == 0:
                                subprocess.run([
                                    ffmpeg_path, '-y', '-ss', '00:00:00', '-i', file_path,
                                    '-vframes', '1', '-q:v', '2', temp_output
                                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10, startupinfo=startupinfo)
                            
                            thumbnail_created = (os.path.exists(temp_output) and os.path.getsize(temp_output) > 0)
                        except:
                            print(f"Normal video dosyası için ilk yöntem başarısız, alternatif deneniyor: {file_path}")
                            try:
                                # Windows'ta terminal penceresi gizlemek için STARTUPINFO kullan
                                startupinfo = None
                                if os.name == 'nt':  # Windows
                                    startupinfo = subprocess.STARTUPINFO()
                                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                                    startupinfo.wShowWindow = subprocess.SW_HIDE
                                
                                # Daha basit bir yaklaşım
                                subprocess.run([
                                    ffmpeg_path, '-y', '-i', file_path,
                                    '-vframes', '1', '-q:v', '5', temp_output
                                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=5, startupinfo=startupinfo)
                                thumbnail_created = (os.path.exists(temp_output) and os.path.getsize(temp_output) > 0)
                            except:
                                print(f"Video dosyası için tüm yöntemler başarısız: {file_path}")
                    
                    # Thumbnail oluştur (her iki yol için de)
                    if thumbnail_created:
                        try:
                            img = Image.open(temp_output)
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            
                            # Play simgesi ekle
                            draw = ImageDraw.Draw(img)
                            center_x, center_y = img.width // 2, img.height // 2
                            triangle_size = min(img.width, img.height) // 4
                            
                            # Daire arka plan
                            draw.ellipse([
                                center_x - triangle_size, center_y - triangle_size,
                                center_x + triangle_size, center_y + triangle_size
                            ], fill=(0, 0, 0, 128))
                            
                            # Üçgen play ikonu
                            draw.polygon([
                                (center_x - triangle_size//2, center_y - triangle_size//2),
                                (center_x - triangle_size//2, center_y + triangle_size//2),
                                (center_x + triangle_size//2, center_y)
                            ], fill=(255, 255, 255, 220))
                            
                            preview_image = ImageTk.PhotoImage(img)
                            
                            # Geçici dosyayı temizle
                            try:
                                os.unlink(temp_output)
                            except:
                                pass
                        except Exception as img_error:
                            print(f"Thumbnail oluşturma hatası: {str(img_error)}")
                
                except Exception as e:
                    print(f"Video önizleme oluşturulamadı: {str(e)}")
                
                # Eğer önizleme oluşturulamadıysa, bir video ikonu göster
                if not preview_image:
                    preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", "VIDEO")
            
            # RAW camera files - add basic support
            elif file_ext in ['.raw', '.cr2', '.nef', '.dng', '.arw']:
                try:
                    # RAW dosyaları için Pillow/PIL'in sadece bazı formatları desteklediğini not et
                    img = Image.open(file_path)
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as raw_error:
                    logging.error(f"RAW önizleme oluşturulamadı: {str(raw_error)}")
                    # Hata durumunda ikonla göster
                    preview_image = self._create_styled_icon(max_width, max_height, "#673AB7", "RAW")
                
            # HEIC/HEIF format - special handling
            elif file_ext in ['.heic', '.heif']:
                try:
                    # HEIC/HEIF dosyaları için özel destek
                    # Pillow yeni sürümlerde destekleyebilir, ancak çoğu durumda heif-convert gerekir
                    
                    # İlk olarak direk PIL ile deniyoruz
                    try:
                        img = Image.open(file_path)
                        
                        # Check if it's a large image (>10MP) for optimization
                        try:
                            mp = (img.width * img.height) / 1000000  # Megapixels
                            is_large = mp > 10
                        except:
                            is_large = False
                            
                        # For large images, use more aggressive downsampling first
                        if is_large:
                            # Calculate intermediate size to improve performance
                            scale_factor = 0.5
                            intermediate_w = int(img.width * scale_factor)
                            intermediate_h = int(img.height * scale_factor)
                            img = img.resize((intermediate_w, intermediate_h), get_pil_resize_method())
                        
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image
                    except Exception as pil_heic_error:
                        logging.info(f"PIL ile HEIC açılamadı, alternatif yöntemler deneniyor: {str(pil_heic_error)}")
                    
                    # PIL başarısız olduysa, FFmpeg ile dönüştürme deneyelim
                    import tempfile
                    import subprocess
                    import os
                    
                    # Geçici dosya oluştur
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                        temp_output = temp_file.name
                    
                    # ffmpeg yolu - sistemden dinamik olarak bul
                    import shutil
                    ffmpeg_path = shutil.which('ffmpeg') or 'ffmpeg'
                    
                    # HEIC'i JPG'ye dönüştür
                    subprocess.run([
                        ffmpeg_path, '-y', '-i', file_path, temp_output
                    ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                    
                    # Dönüştürülmüş dosyayı yükle
                    if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                        img = Image.open(temp_output)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        
                        # Geçici dosyayı temizle
                        try:
                            os.unlink(temp_output)
                        except:
                            pass
                        
                        return preview_image
                    
                    # Yine başarısız olursa, son çare olarak yer tutucu oluştur
                    logging.error(f"HEIC önizleme oluşturulamadı: {file_path}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", "HEIC")
                    
                except Exception as heic_error:
                    logging.error(f"HEIC önizleme oluşturulamadı: {str(heic_error)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", "HEIC")
            
            # ICO format - special handling
            elif file_ext == '.ico':
                try:
                    # ICO dosyaları birden fazla boyutta ikon içerebilir
                    # En büyük olanı seçmek için özel işleme yapılabilir
                    img = Image.open(file_path)
                    
                    # ICO dosyasının tüm boyutlarını al
                    if hasattr(img, 'ico_sizes'):
                        try:
                            # En büyük ikon boyutunu seç
                            sizes = list(img.ico_sizes())
                            if sizes:
                                largest_size = max(sizes, key=lambda size: size[0] * size[1])
                                img = img.ico_getimage(largest_size)
                        except Exception as ico_size_err:
                            logging.info(f"ICO boyutu seçilemedi: {str(ico_size_err)}")
                    
                    # Şeffaf arkaplan üzerine ikon yerleştir
                    if img.mode == 'RGBA':
                        # Şeffaf kısmı görünür yapmak için kontrastlı bir arkaplan kullan
                        background = Image.new('RGB', img.size, (240, 240, 240))
                        background.paste(img, (0, 0), img)
                        img = background
                    
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as ico_error:
                    logging.error(f"ICO önizleme oluşturulamadı: {str(ico_error)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#009688", "ICO")
            
            # HEIC/HEIF files (Apple format) - special handling
            elif file_ext in ['.heic', '.heif']:
                try:
                    # Önce pillow-heif ile açmayı dene (daha hızlı)
                    try:
                        img = Image.open(file_path)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                    except Exception as pillow_heif_error:
                        logging.info(f"HEIC açılamadı (pillow-heif): {str(pillow_heif_error)}")
                        
                        # Eğer pillow-heif yoksa, ffmpeg ile dönüştürmeyi dene
                        import tempfile
                        import subprocess
                        
                        # Geçici dosya oluştur
                        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                            temp_output = temp_file.name
                        
                        # ffmpeg yolu - sistemden bul
                        import shutil
                        ffmpeg_path = shutil.which('ffmpeg') or 'ffmpeg'
                        
                        # HEIC'i JPG'ye dönüştür
                        subprocess.run([
                            ffmpeg_path, '-y', '-i', file_path, temp_output
                        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                        
                        # Dönüştürülmüş dosyayı yükle
                        if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                            img = Image.open(temp_output)
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            preview_image = ImageTk.PhotoImage(img)
                            
                            # Geçici dosyayı temizle
                            try:
                                os.unlink(temp_output)
                            except:
                                pass
                        else:
                            # Dönüştürme başarısız olursa, ikonla göster
                            preview_image = self._create_styled_icon(max_width, max_height, "#4CAF50", "HEIC")
                except Exception as e:
                    logging.error(f"HEIC önizleme oluşturulamadı: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#4CAF50", "HEIC")
            
            # Video files (MP4, AVI, MOV, etc.) - create thumbnail preview
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg', '.3gp']:
                try:
                    # Video önizlemesi için ffmpeg kullan
                    import tempfile
                    import subprocess
                    import os
                    
                    # Geçici dosya oluştur
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                        temp_output = temp_file.name
                    
                    # ffmpeg yolu - sistemden bul
                    import shutil
                    ffmpeg_path = shutil.which('ffmpeg') or 'ffmpeg'
                    
                    # Video'nun ilk karesini çıkar (00:00:01 zamanından)
                    try:
                        subprocess.run([
                            ffmpeg_path, '-y', '-ss', '00:00:01', '-i', file_path,
                            '-vframes', '1', '-q:v', '2', temp_output
                        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                        
                        # İlk kare alınamazsa, 00:00:00 zamanını dene
                        if not os.path.exists(temp_output) or os.path.getsize(temp_output) == 0:
                            subprocess.run([
                                ffmpeg_path, '-y', '-ss', '00:00:00', '-i', file_path,
                                '-vframes', '1', '-q:v', '2', temp_output
                            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                        
                        # Thumbnail oluştur
                        if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                            img = Image.open(temp_output)
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            
                            # Play simgesi ekle thumbnail'in üzerine
                            draw = ImageDraw.Draw(img)
                            
                            # Oynatma üçgeni çiz
                            center_x, center_y = img.width // 2, img.height // 2
                            triangle_size = min(img.width, img.height) // 4
                            
                            # Yarı saydam arka plan dairesi
                            draw.ellipse([
                                center_x - triangle_size, center_y - triangle_size,
                                center_x + triangle_size, center_y + triangle_size
                            ], fill=(0, 0, 0, 128))
                            
                            # Oynatma üçgeni (sağa bakan)
                            triangle_points = [
                                (center_x - triangle_size//2, center_y - triangle_size//2),
                                (center_x - triangle_size//2, center_y + triangle_size//2),
                                (center_x + triangle_size//2, center_y)
                            ]
                            draw.polygon(triangle_points, fill=(255, 255, 255, 220))
                            
                            preview_image = ImageTk.PhotoImage(img)
                            
                            # Geçici dosyayı temizle
                            try:
                                os.unlink(temp_output)
                            except:
                                pass
                        else:
                            # Thumbnail oluşturulamazsa, video ikonu göster
                            preview_image = self._create_styled_icon(max_width, max_height, "#E53935", "VIDEO")
                    except Exception as ffmpeg_error:
                        logging.error(f"FFmpeg ile video thumbnail oluşturulamadı: {str(ffmpeg_error)}")
                        preview_image = self._create_styled_icon(max_width, max_height, "#E53935", "VIDEO")
                except Exception as e:
                    logging.error(f"Video önizleme oluşturulamadı: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#E53935", "VIDEO")
            
            # Other regular image formats
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                try:
                    # Open and resize the image
                    img = Image.open(file_path)
                    
                    # Check if it's a large image (>10MP) for optimization
                    try:
                        mp = (img.width * img.height) / 1000000  # Megapixels
                        is_large = mp > 10
                    except:
                        is_large = False
                        
                    # For large images, use more aggressive downsampling first
                    if is_large:
                        # Calculate intermediate size to improve performance
                        scale_factor = 0.5
                        intermediate_w = int(img.width * scale_factor)
                        intermediate_h = int(img.height * scale_factor)
                        img = img.resize((intermediate_w, intermediate_h), get_pil_resize_method())
                    
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as e:
                    logging.error(f"Error creating preview for {file_path}: {str(e)}")
                    # Create a placeholder for failed image with file extension
                    img = Image.new("RGB", (max_width, max_height), color="#f0f0f0")
                    draw = ImageDraw.Draw(img)
                    draw.rectangle([10, 10, max_width-10, max_height-10], outline="#dc3545", width=2)
                    # Show only extension without dot and make it uppercase (with fallback to IMG if too long)
                    ext_text = file_ext.upper()[1:] if len(file_ext[1:]) <= 5 else "IMG"
                    draw.text((max_width//2, max_height//2), ext_text, fill="#dc3545", anchor="mm")
                    preview_image = ImageTk.PhotoImage(img)

            # PDF files with progressive loading for large files
            elif file_ext == '.pdf':
                # Get PDF file size for optimization decisions
                try:
                    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)  # MB cinsinden
                    is_large_pdf = file_size_mb > 20  # 20MB'dan büyük PDF'ler büyük olarak kabul edilir
                except Exception:
                    is_large_pdf = False
                
                # Get the first page of PDF
                try:
                    # OPTIMIZASYON: Büyük PDF'ler için düşük çözünürlüklü önizleme kullan
                    # Bu, hafıza tüketimini ve işleme süresini azaltır
                    scale_factor = 0.25 if is_large_pdf else 0.5
                    
                    # Büyük PDF'ler için uyarı göster
                    if is_large_pdf and hasattr(self, 'update_status'):
                        self.update_status(self.get_text("loading_large_file"))
                    
                    pdf_doc = fitz.open(file_path)
                    if pdf_doc.page_count > 0:
                        # OPTIMIZASYON: Büyük PDF için 1. sayfanın düşük çözünürlükte önizlemesini oluştur
                        page = pdf_doc[0]
                        
                        # Okuma işlemi için thread safe timeout uygula
                        # Bu, çok büyük veya karmaşık PDF'lerin UI'yi dondurmasını önler
                        pix = page.get_pixmap(matrix=fitz.Matrix(scale_factor, scale_factor))
                        
                        # Hafıza optimizasyonu - gereksiz verileri temizle
                        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        pix = None  # Hafızayı temizle
                        
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                    
                    # Açık dosya tanıtıcılarını temizle
                    pdf_doc.close()
                    
                except Exception as e:
                    logging.error(f"Error with PyMuPDF for {file_path}: {str(e)}")
                    # Fall back to pdf2image if fitz fails
                    try:
                        # Windows sistemlerde poppler'ın manuel olarak kurulu olması gerekir
                        # Windows olmayan sistemlerde, varsayılan yolu kullanır
                        # EXE paketlemesi için poppler bağımlılığını eklemeyi unutmayın
                        if os.name == 'nt':  # Windows sistemi
                            poppler_path = ""  # Boş string, default kullanılır
                        else:
                            poppler_path = ""  # Boş string, default kullanılır

                        # OPTIMIZASYON: Büyük PDF'ler için okuma performansını ayarla
                        dpi = 72 if is_large_pdf else 150  # Düşük DPI daha hızlı işlenir
                        
                        # Use pdftocairo which often produces better quality
                        images = pdf2image.convert_from_path(
                            file_path, 
                            first_page=1, 
                            last_page=1, 
                            size=(max_width, max_height),
                            dpi=dpi,
                            # Poppler path parametresi sadece gerekli olduğunda kullan
                            # PDF işleme çoğu durumda poppler olmadan da çalışır
                            use_pdftocairo=True
                        )
                        if images:
                            preview_image = ImageTk.PhotoImage(images[0])
                    except Exception as e2:
                        logging.error(f"Error with pdf2image for {file_path}: {str(e2)}")
                        # Create a placeholder
                        preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", "PDF")

            # EPS files - use our specialized function
            elif file_ext == '.eps':
                preview_image = self._create_eps_preview(file_path, max_width, max_height)

            # Design files (PSD, AI) - with optimizations for large files
            elif file_ext in ['.psd', '.ai']:
                # Set default placeholder color based on file type
                color = "#1976D2" if file_ext == '.psd' else "#FF5722"  # Blue for PSD, Orange for AI
                file_type = file_ext[1:].upper()
                
                # Get file size for optimization decisions
                try:
                    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)  # MB cinsinden
                    is_large_design_file = file_size_mb > 15  # 15MB'dan büyük design dosyaları büyük olarak kabul edilir
                except Exception:
                    is_large_design_file = False
                
                # Büyük dosyalar için uyarı göster
                if is_large_design_file and hasattr(self, 'update_status'):
                    self.update_status(self.get_text("loading_large_file"))

                try:
                    # PSD dosyalarını PIL ile aç (büyük dosyalar için optimize edilmiş)
                    if file_ext == '.psd':
                        # OPTIMIZASYON: Büyük PSD dosyalarında bellek yönetimi
                        if is_large_design_file:
                            # Geçici dosya temizliğini kolaylaştırmak için with bloğu kullan 
                            img = Image.open(file_path)
                            
                            # Bellek kullanımını azaltmak için büyük PSD dosyalarını daha agresif küçült
                            if img.width > 1000 or img.height > 1000:
                                # İlk önce agresif bir şekilde boyutunu küçült, sonra thumbnail oluştur
                                scale_factor = 0.25 if is_large_design_file else 0.5
                                new_width = int(img.width * scale_factor)
                                new_height = int(img.height * scale_factor)
                                img = img.resize((new_width, new_height), get_pil_resize_method())
                            
                            # Son olarak önizleme boyutlarına getir
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            preview_image = ImageTk.PhotoImage(img)
                        else:
                            # Normal boyutlu PSD için standart işlem
                            img = Image.open(file_path)
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            preview_image = ImageTk.PhotoImage(img)
                        
                        # Hemen hafızayı temizle
                        if 'img' in locals():
                            del img
                            
                        return preview_image
                    
                    # AI dosyalarını PyMuPDF ile aç (onlar genellikle PDF uyumludur)
                    elif file_ext == '.ai':
                        # AI dosyaları için PDF render optimizasyonları
                        scale_factor = 0.25 if is_large_design_file else 0.5
                        
                        pdf_doc = fitz.open(file_path)
                        if pdf_doc.page_count > 0:
                            page = pdf_doc[0]
                            pix = page.get_pixmap(matrix=fitz.Matrix(scale_factor, scale_factor))
                            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                            
                            # Hafızayı serbest bırak
                            pix = None
                            
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            preview_image = ImageTk.PhotoImage(img)
                            
                            # Hafızayı temizle
                            if 'img' in locals():
                                del img
                                
                            pdf_doc.close()
                            return preview_image
                        pdf_doc.close()
                except Exception as e:
                    logging.error(f"Error with design file {file_path}: {str(e)}")

                # Herhangi bir hata durumunda yer tutucu simge oluştur
                if not preview_image:
                    preview_image = self._create_styled_icon(max_width, max_height, color, file_type)

            # Office documents - Word (expanded with more formats)
            elif file_ext in ['.doc', '.docx', '.dot', '.dotx', '.dotm', '.rtf', '.odt', '.wpd']:
                try:
                    # Try to generate a thumbnail from Word using docx, if available
                    if file_ext == '.docx' and 'docx' in sys.modules:
                        try:
                            # This is just for potential future expansion - we'll use placeholders for now
                            pass
                        except Exception as word_e:
                            logging.info(f"Failed to generate Word preview: {str(word_e)}")

                    # Create a nice styled Word icon
                    if file_ext in ['.rtf', '.odt', '.wpd']:
                        # Label with the file extension for non-Word formats
                        word_label = file_ext[1:].upper()
                    else:
                        # Just use DOC for all Word formats
                        word_label = "DOC"

                    preview_image = self._create_styled_icon(max_width, max_height, "#2B579A", word_label)  # Word blue
                except Exception as e:
                    logging.error(f"Error with document file {file_path}: {str(e)}")
                    # Fallback icon
                    word_label = file_ext[1:].upper() if len(file_ext) <= 5 else "DOC"
                    preview_image = self._create_styled_icon(max_width, max_height, "#2B579A", word_label)

            # Office documents - Excel (expanded with more formats)
            elif file_ext in ['.xls', '.xlsx', '.xlsm', '.xlsb', '.xlt', '.xltx', '.xltm', '.csv', '.ods', '.tsv', '.numbers']:
                try:
                    # For CSV/TSV files, we could potentially show a preview of the data
                    if file_ext in ['.csv', '.tsv']:
                        # Just a placeholder for now
                        pass

                    # Create a nice styled Excel icon
                    if file_ext in ['.csv', '.tsv', '.ods', '.numbers']:
                        # Label with file extension for non-Excel formats
                        excel_label = file_ext[1:].upper()
                    else:
                        # Just use XLS for all Excel formats
                        excel_label = "XLS"

                    preview_image = self._create_styled_icon(max_width, max_height, "#217346", excel_label)  # Excel green
                except Exception as e:
                    logging.error(f"Error with spreadsheet file {file_path}: {str(e)}")
                    # Fallback icon
                    excel_label = file_ext[1:].upper() if len(file_ext) <= 5 else "XLS"
                    preview_image = self._create_styled_icon(max_width, max_height, "#217346", excel_label)

            # Office documents - PowerPoint (expanded with more formats)
            elif file_ext in ['.ppt', '.pptx', '.pptm', '.pps', '.ppsx', '.ppsm', '.pot', '.potx', '.potm', '.odp', '.key']:
                try:
                    # Create a nice styled PowerPoint icon
                    if file_ext in ['.odp', '.key']:
                        # Label with file extension for non-PowerPoint formats
                        ppt_label = file_ext[1:].upper()
                    else:
                        # Just use PPT for all PowerPoint formats
                        ppt_label = "PPT"

                    preview_image = self._create_styled_icon(max_width, max_height, "#D24726", ppt_label)  # PowerPoint orange
                except Exception as e:
                    logging.error(f"Error with presentation file {file_path}: {str(e)}")
                    # Fallback icon
                    ppt_label = file_ext[1:].upper() if len(file_ext) <= 5 else "PPT"
                    preview_image = self._create_styled_icon(max_width, max_height, "#D24726", ppt_label)

            # Text and code files (expanded with more formats)
            elif file_ext in ['.txt', '.md', '.json', '.xml', '.html', '.htm', '.css', '.js', '.jsx', '.ts', '.tsx', 
                           '.py', '.java', '.c', '.cpp', '.cs', '.h', '.hpp', '.rb', '.php', '.swift', '.go', '.rs',
                           '.pl', '.lua', '.r', '.sh', '.bat', '.ps1', '.yaml', '.yml', '.toml', '.ini', '.cfg',
                           '.conf', '.log', '.sql', '.asm', '.tex', '.srt', '.vtt', '.csv', '.m', '.f90', '.kt']:
                try:
                    # Group files into categories with similar colors
                    if file_ext in ['.html', '.htm', '.xml', '.jsx', '.tsx']:
                        # Markup files - light blue
                        color = "#03A9F4"
                    elif file_ext in ['.js', '.ts', '.py', '.rb', '.php', '.swift', '.java', '.cs', '.go', '.rs', '.kt']:
                        # Popular programming languages - indigo
                        color = "#3F51B5"
                    elif file_ext in ['.c', '.cpp', '.h', '.hpp', '.asm', '.m', '.f90']:
                        # Systems programming - deep blue
                        color = "#1A237E"
                    elif file_ext in ['.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf']:
                        # Config files - teal
                        color = "#009688"
                    elif file_ext in ['.log', '.txt', '.md', '.csv', '.srt', '.vtt']:
                        # Plain text files - gray
                        color = "#607D8B"
                    else:
                        # Other text files - blue-grey
                        color = "#607D8B"

                    # We could potentially show a preview of text content, but for now just an icon
                    txt_label = file_ext[1:].upper() if len(file_ext) <= 5 else "TXT"
                    preview_image = self._create_styled_icon(max_width, max_height, color, txt_label)
                except Exception as e:
                    logging.error(f"Error with text file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#607D8B", "TXT")

            # Archive files (expanded with more formats)
            elif file_ext in ['.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz', '.tgz', '.lzma', '.iso', '.cab', 
                           '.dmg', '.img', '.jar', '.war', '.ear', '.bzip2', '.tbz2', '.gzip', '.z', '.lz', '.lz4']:
                try:
                    # Group by archive type
                    if file_ext in ['.iso', '.dmg', '.img']:
                        # Disk images - amber-red
                        color = "#FF8F00"
                        archive_label = file_ext[1:].upper() if len(file_ext) <= 5 else "DISK"
                    elif file_ext in ['.jar', '.war', '.ear']:
                        # Java archives - red
                        color = "#D32F2F"
                        archive_label = file_ext[1:].upper() if len(file_ext) <= 5 else "JAR"
                    else:
                        # Regular archives - amber
                        color = "#FFC107"
                        archive_label = file_ext[1:].upper() if len(file_ext) <= 5 else "ZIP"

                    # Create a styled icon for archives
                    preview_image = self._create_styled_icon(max_width, max_height, color, archive_label)
                except Exception as e:
                    logging.error(f"Error with archive file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#FFC107", "ZIP")

            # Executable and installable files (expanded with more formats)
            elif file_ext in ['.exe', '.msi', '.bat', '.cmd', '.ps1', '.sh', '.bash', '.app', '.run', '.bin', '.deb', '.rpm', 
                           '.dmg', '.pkg', '.appimage', '.apk', '.ipa', '.xap', '.msix', '.dll', '.so', '.dylib', '.com', '.vbs']:
                try:
                    # Group by platform/type for better color coding
                    if file_ext in ['.sh', '.bash', '.run', '.bin']:
                        # Unix executables - deep red
                        color = "#B71C1C"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "UNIX"
                    elif file_ext in ['.app', '.dmg', '.pkg']:
                        # macOS executables - dark red
                        color = "#C62828"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "MAC"
                    elif file_ext in ['.deb', '.rpm', '.appimage']:
                        # Linux packages - brick red
                        color = "#D32F2F"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "LINUX"
                    elif file_ext in ['.apk', '.ipa', '.xap', '.msix']:
                        # Mobile apps - light red
                        color = "#E53935"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "MOBILE"
                    elif file_ext in ['.dll', '.so', '.dylib']:
                        # Library files - orange-red
                        color = "#E64A19"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "LIB"
                    elif file_ext in ['.bat', '.cmd', '.ps1', '.vbs', '.com']:
                        # Scripts and command files - orange
                        color = "#EF6C00"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "SCRIPT"
                    else:
                        # Windows executables - standard red
                        color = "#F44336"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "EXE"

                    # Create a styled icon for executables
                    preview_image = self._create_styled_icon(max_width, max_height, color, exe_label)
                except Exception as e:
                    logging.error(f"Error with executable file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#F44336", "EXE")

            # Audio files - expanded with more formats
            elif file_ext in ['.mp3', '.wav', '.ogg', '.flac', '.aac', '.wma', '.m4a', '.aiff', '.ape', '.midi', '.mid', '.amr', '.opus', '.alac', '.aif']:
                try:
                    # Create a styled icon for audio files
                    audio_label = file_ext[1:].upper() if len(file_ext) <= 5 else "AUD"
                    preview_image = self._create_styled_icon(max_width, max_height, "#9C27B0", audio_label)  # Purple
                except Exception as e:
                    logging.error(f"Error with audio file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#9C27B0", "AUD")

            # Video files - expanded with more formats
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg', '.3gp', '.vob', '.ts', '.mts', '.m2ts', '.divx', '.asf', '.ogv']:
                try:
                    # Create a styled icon for video files
                    video_label = file_ext[1:].upper() if len(file_ext) <= 5 else "VID"
                    preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", video_label)  # Deep orange
                except Exception as e:
                    logging.error(f"Error with video file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", "VID")

            # Font files
            elif file_ext in ['.ttf', '.otf', '.woff', '.woff2', '.eot']:
                try:
                    # Create a styled icon for font files
                    font_label = file_ext[1:].upper() if len(file_ext) <= 5 else "FONT"
                    preview_image = self._create_styled_icon(max_width, max_height, "#009688", font_label)  # Teal
                except Exception as e:
                    logging.error(f"Error with font file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#009688", "FONT")

            # Database files
            elif file_ext in ['.db', '.sqlite', '.sqlite3', '.mdb', '.accdb', '.sql', '.dbf']:
                try:
                    # Create a styled icon for database files
                    db_label = file_ext[1:].upper() if len(file_ext) <= 5 else "DB"
                    preview_image = self._create_styled_icon(max_width, max_height, "#1565C0", db_label)  # Dark blue
                except Exception as e:
                    logging.error(f"Error with database file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#1565C0", "DB")

            # E-book files
            elif file_ext in ['.epub', '.mobi', '.azw', '.azw3', '.fb2', '.cbz', '.cbr']:
                try:
                    # Create a styled icon for e-book files
                    book_label = file_ext[1:].upper() if len(file_ext) <= 5 else "EBOOK"
                    preview_image = self._create_styled_icon(max_width, max_height, "#4CAF50", book_label)  # Green
                except Exception as e:
                    logging.error(f"Error with e-book file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#4CAF50", "EBOOK")

            # 3D and CAD files
            elif file_ext in ['.obj', '.stl', '.fbx', '.blend', '.3ds', '.dae', '.dwg', '.dxf', '.max', '.c4d']:
                try:
                    # Create a styled icon for 3D files
                    model_label = file_ext[1:].upper() if len(file_ext) <= 5 else "3D"
                    preview_image = self._create_styled_icon(max_width, max_height, "#673AB7", model_label)  # Deep purple
                except Exception as e:
                    logging.error(f"Error with 3D/CAD file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#673AB7", "3D")

            # GIS and Map files
            elif file_ext in ['.shp', '.kml', '.kmz', '.gpx', '.geojson', '.osm', '.mbtiles', '.dem', '.tiff', '.asc']:
                try:
                    # Create a styled icon for GIS files
                    gis_label = file_ext[1:].upper() if len(file_ext) <= 5 else "GIS"
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", gis_label)  # Indigo
                except Exception as e:
                    logging.error(f"Error with GIS file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", "GIS")

            # Certificate and key files
            elif file_ext in ['.pem', '.crt', '.cer', '.key', '.p12', '.pfx']:
                try:
                    # Create a styled icon for certificate files
                    cert_label = file_ext[1:].upper() if len(file_ext) <= 5 else "CERT"
                    preview_image = self._create_styled_icon(max_width, max_height, "#795548", cert_label)  # Brown
                except Exception as e:
                    logging.error(f"Error with certificate file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#795548", "CERT")

            # If no specific handler, create a generic icon
            elif not preview_image:
                # Get just the extension without the dot
                ext = file_ext[1:].upper() if len(file_ext) > 1 else "?"
                # Truncate if too long
                if len(ext) > 5:
                    ext = ext[:4] + "…"
                preview_image = self._create_styled_icon(max_width, max_height, "#9E9E9E", ext)  # Grey

        except Exception as e:
            print(f"Error creating preview for {file_path}: {str(e)}")
            # Return a placeholder for errors
            img = Image.new('RGB', (max_width, max_height), "#F44336")
            draw = ImageDraw.Draw(img)
            draw.rectangle([(0, 0), (max_width-1, max_height-1)], outline="white", width=1)
            draw.text((max_width//2, max_height//2 - 10), "!", fill="white", anchor="mm")
            preview_image = ImageTk.PhotoImage(img)

        # If we got a preview, cache it for future use
        if preview_image:
            try:
                # Store in cache (thread-safe operation)
                with self._preview_cache_lock:
                    self._preview_cache[cache_key] = preview_image

                # Let OrderedDict handle the LRU order
                # No need for manual tracking


                # Debug log
                if len(self._preview_cache) % 100 == 0:
                    logging.info(f"Preview cache size: {len(self._preview_cache)} items")
            except Exception as e:
                logging.error(f"Error updating preview cache: {str(e)}")

            return preview_image
        else:
            return None

    def _create_styled_icon(self, width, height, color, text):
        """
        Create a styled icon with the given color and text
        This is used for file types that don't have actual visual previews
        """
        try:
            # Create a new image with the specified color
            img = Image.new('RGB', (width, height), color)
            draw = ImageDraw.Draw(img)

            # Log for debugging
            logging.info(f"Creating styled icon with color {color} and text {text}")

            # Add an outer border
            draw.rectangle([(0, 0), (width-1, height-1)], outline="white", width=2)

            # Draw a file icon shape (a rectangle with a folded corner)
            padding = 10
            draw.rectangle(
                [(padding, padding), (width-padding, height-padding)], 
                fill="white", outline=color, width=1
            )

            # Fold the top-right corner
            corner_size = 15
            draw.polygon(
                [
                    (width-padding-corner_size, padding),
                    (width-padding, padding),
                    (width-padding, padding+corner_size)
                ],
                fill=color
            )

            # Draw lines to represent text in the "document"
            line_padding = 8
            line_height = 5
            num_lines = 3
            line_width = width - (padding*2 + line_padding*2)
            line_start_y = padding + 20

            for i in range(num_lines):
                y_pos = line_start_y + (i * (line_height + 5))
                draw.rectangle(
                    [(padding + line_padding, y_pos), 
                     (padding + line_padding + line_width, y_pos + line_height)], 
                    fill=color
                )

            # Add text centered at the bottom
            # In older versions of PIL, anchor="mm" might not be supported
            # So let's calculate the text position manually if needed
            try:
                draw.text(
                    (width//2, height - padding - 8), 
                    text, 
                    fill=color, 
                    anchor="mm"
                )
            except TypeError:
                # Older PIL versions don't support anchor
                # We'll need to center manually
                text_bbox = draw.textbbox((0, 0), text)
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                text_x = (width - text_width) // 2
                text_y = height - padding - 8 - text_height // 2
                draw.text((text_x, text_y), text, fill=color)

            photo_img = ImageTk.PhotoImage(img)
            return photo_img

        except Exception as e:
            logging.error(f"Error creating styled icon: {str(e)}")
            # Fallback to a simple colored rectangle with text
            try:
                img = Image.new('RGB', (width, height), color)
                draw = ImageDraw.Draw(img)
                draw.rectangle([(0, 0), (width-1, height-1)], outline="white", width=1)
                try:
                    draw.text((width//2, height//2), text, fill="white", anchor="mm")
                except TypeError:
                    # Center manually for older PIL versions
                    text_bbox = draw.textbbox((0, 0), text)
                    text_width = text_bbox[2] - text_bbox[0]
                    text_height = text_bbox[3] - text_bbox[1]
                    text_x = (width - text_width) // 2
                    text_y = (height - text_height) // 2
                    draw.text((text_x, text_y), text, fill="white")
                return ImageTk.PhotoImage(img)
            except Exception as e2:
                logging.error(f"Fallback icon creation also failed: {str(e2)}")
                # Last resort fallback
                img = Image.new('RGB', (width, height), "#F44336")  # Red
                return ImageTk.PhotoImage(img)

    def update_ui_language(self):
        # Update window title
        self.root.title(self.get_text("full_window_title"))

        # Update all widgets with text
        # Title section
        for widget in self.main_frame.winfo_children():
            if isinstance(widget, tk.LabelFrame):
                try:
                    current_text = widget.cget("text")
                    # Try to find a matching key in the language dictionary
                    found_key = None
                    for key in self.languages[self.current_language].keys():
                        # Check if this text matches any language's value for this key
                        for lang_code in self.languages.keys():
                            if self.languages[lang_code].get(key, "") == current_text:
                                found_key = key
                                break

                        if found_key:
                            break

                    # If found, update with current language version
                    if found_key:
                        widget.config(text=self.get_text(found_key))
                except Exception as e:
                    logging.warning(f"Failed to update label text: {str(e)}")

        # Update buttons
        self.select_folder_btn.config(text=self.get_text("select_folder"))
        self.start_btn.config(text=self.get_text("start"))
        self.cancel_btn.config(text=self.get_text("cancel"))
        self.exit_btn.config(text=self.get_text("exit"))
        self.select_all_btn.config(text=self.get_text("select_all"))
        self.clear_all_btn.config(text=self.get_text("clear_all"))
        self.apply_filter_btn.config(text=self.get_text("apply_filter"))
        self.show_filter_btn.config(text=self.get_text("filter_label"))

        # Update view mode buttons if they exist
        if hasattr(self, 'list_view_btn'):
            self.list_view_btn.config(text=self.get_text("list_view"))
        if hasattr(self, 'preview_view_btn'):
            self.preview_view_btn.config(text=self.get_text("preview_view"))

        # Update search labels if they exist
        if hasattr(self, 'extension_search_label'):
            self.extension_search_label.config(text="🔍 " + self.get_text("extension_search"))
        if hasattr(self, 'file_search_label'):
            self.file_search_label.config(text=self.get_text("extension_search") + ":")

        # Update search placeholder text if entry exists
        if hasattr(self, 'file_search_entry') and not self.file_search_var.get():
            self.file_search_entry.delete(0, 'end')
            self.file_search_entry.insert(0, self.get_text("search_files"))
            self.file_search_entry.config(fg='gray')

        # Update desktop checkbox
        if hasattr(self, 'desktop_cb'):
            self.desktop_cb.config(text=self.get_text("desktop_label"))

        # Update checkbox texts
        self.subfolder_cb.config(text=self.get_text("include_label"))

        # Update setting labels
        for frame_name in ["settings_frame", "filter_frame", "tips_frame", "stats_frame", "file_list_frame"]:
            if hasattr(self, frame_name):
                frame = getattr(self, frame_name)

                # Check if frame is a LabelFrame before trying to access text property
                try:
                    if isinstance(frame, tk.LabelFrame):
                        # Update frame title
                        current_text = frame.cget("text")
                        found_key = None

                        # Look for this text in all languages
                        for key in ["settings_header", "filter_label", "tips_header", "statistics_header", "file_list"]:
                            # Check if this text matches any language's value for this key
                            for lang_code in self.languages.keys():
                                if self.languages[lang_code].get(key, "") == current_text:
                                    found_key = key
                                    break

                            if found_key:
                                break

                        # If found, update with current language version
                        if found_key:
                            frame.config(text=self.get_text(found_key))
                except Exception as e:
                    logging.warning(f"Failed to update frame title: {str(e)}")

        # Update labels in settings section
        for widget in self.settings_frame.winfo_children():
            if isinstance(widget, tk.Frame):
                for child in widget.winfo_children():
                    if isinstance(child, tk.Label):
                        current_text = child.cget("text")
                        found_key = None

                        # Check against all languages
                        for key in ["subfolders_label", "list_format_label", "save_location_label", "sort_criteria_label"]:
                            # Check if this text matches any language's value for this key
                            for lang_code in self.languages.keys():
                                if self.languages[lang_code].get(key, "") == current_text:
                                    found_key = key
                                    break

                            if found_key:
                                break

                        # If found, update with current language version
                        if found_key:
                            child.config(text=self.get_text(found_key))
                    elif isinstance(child, tk.Checkbutton) and child == self.subfolder_cb:
                        child.config(text=self.get_text("include_label"))

        # Update all static labels
        self.update_all_static_labels(self.main_frame)

        # Update tooltip texts
        self.create_tooltip(self.select_folder_btn, self.get_text("tooltip_select"))
        self.create_tooltip(self.start_btn, self.get_text("tooltip_start"))
        self.create_tooltip(self.cancel_btn, self.get_text("tooltip_cancel"))
        self.create_tooltip(self.exit_btn, self.get_text("tooltip_exit"))
        self.create_tooltip(self.subfolder_cb, self.get_text("tooltip_subfolders"))
        self.create_tooltip(self.select_all_btn, self.get_text("tooltip_select_all"))
        self.create_tooltip(self.clear_all_btn, self.get_text("tooltip_clear_all"))
        self.create_tooltip(self.apply_filter_btn, self.get_text("tooltip_filter_apply"))

        # Update tip texts
        self.update_tips()

        # Update footer text with current language and year
        for widget in self.main_frame.winfo_children():
            if isinstance(widget, tk.Frame) and widget.winfo_children():
                for child in widget.winfo_children():
                    if isinstance(child, tk.Label) and "©" in child.cget("text"):
                        # Update copyright footer text
                        current_year = datetime.datetime.now().year
                        copyright_text = self.get_text("copyright_footer").format(year=current_year)
                        child.config(text=copyright_text)

        # Update categories
        self.populate_categories()

        # Update sorting options
        self.populate_sort_dropdown()

        # Update file tree headers with click-to-sort functionality
        self.file_tree.heading("name", text=self.get_text("file_name"), 
                              command=lambda: self.treeview_sort_column("name", False))
        self.file_tree.heading("extension", text=self.get_text("file_extension"),
                              command=lambda: self.treeview_sort_column("extension", False))
        self.file_tree.heading("path", text=self.get_text("file_path"),
                              command=lambda: self.treeview_sort_column("path", False))
        self.file_tree.heading("size", text=self.get_text("file_size"),
                              command=lambda: self.treeview_sort_column("size", False))
        self.file_tree.heading("created", text=self.get_text("creation_date"),
                              command=lambda: self.treeview_sort_column("created", False))
        self.file_tree.heading("modified", text=self.get_text("modification_date"),
                              command=lambda: self.treeview_sort_column("modified", False))

        # Update status text if it's the default ready message
        if self.status_var.get() == self.get_text("ready") or self.status_var.get() == "Ready":
            self.update_status(self.get_text("ready"))

        # Update the folder path display if it's not set
        if self.folder_path_var.get() == self.get_text("no_folder_selected") or self.folder_path_var.get() == "No folder selected":
            self.folder_path_var.set(self.get_text("no_folder_selected"))

        # Update toggle button tooltip for current language
        if hasattr(self, 'toggle_left_panel_btn'):
            self.create_tooltip(self.toggle_left_panel_btn, toggle_panel_translations.get(self.current_language, "Sol paneli aç/kapat"))

        # Force a redraw
        self.root.update_idletasks()

    def update_all_static_labels(self, parent_widget):
        """Recursively update all Label widgets with translated text"""
        for widget in parent_widget.winfo_children():
            try:
                if isinstance(widget, tk.Label) and hasattr(widget, 'cget'):
                    try:
                        current_text = widget.cget("text")
                        if current_text:  # Ensure text is not empty
                            # Skip dynamic content like file counts, etc.
                            if current_text.isdigit() or current_text.startswith(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9")):
                                continue

                            # Try to find a matching key in the language dictionary
                            found_key = None
                            for key in self.languages[self.current_language].keys():
                                # Check if this text matches any language's value for this key
                                for lang_code in self.languages.keys():
                                    if self.languages[lang_code].get(key, "") == current_text:
                                        found_key = key
                                        break

                                if found_key:
                                    break

                            # If found, update with current language version
                            if found_key:
                                widget.config(text=self.get_text(found_key))
                    except Exception as e:
                        logging.warning(f"Error updating label text: {str(e)}")

                # Recursively process child widgets
                if widget.winfo_children():
                    self.update_all_static_labels(widget)
            except Exception as e:
                logging.warning(f"Error processing widget: {str(e)}")

    def update_tips(self):
        """Update tips with the current language"""
        # Find tips frame and update all tip labels
        if hasattr(self, "tips_frame"):
            # Update tip labels
            tip_texts = [
                self.get_text("tip_1"),
                self.get_text("tip_3"),
                self.get_text("tip_4"),
                self.get_text("tip_5"),
                self.get_text("tip_6"),
                self.get_text("tip_preview_formats")
            ]

            # Find all tip labels recursively
            tip_labels = []

            def find_tip_labels(parent):
                for widget in parent.winfo_children():
                    try:
                        if isinstance(widget, tk.Label) and hasattr(widget, "cget") and (
                            widget.cget("text").startswith("• ") or 
                            any(widget.cget("text").endswith(t) for t in self.languages["en"].values())
                        ):
                            tip_labels.append(widget)
                    except Exception as e:
                        logging.warning(f"Error checking tip label: {str(e)}")

                    if widget.winfo_children():
                        find_tip_labels(widget)

            find_tip_labels(self.tips_frame)

            # Update the tip texts
            for i, label in enumerate(tip_labels[:len(tip_texts)]):
                label.config(text=f"• {tip_texts[i]}")

    def select_folder(self):
        folder_path = filedialog.askdirectory(title=self.get_text("select_folder"))
        if folder_path:
            # Always switch back to list mode automatically when selecting a new folder
            # This ensures a consistent experience and better performance
            self.view_mode_var.set("list")
            self._switch_to_list_view()

            # Store the selected folder and update display
            self.selected_folder_path = folder_path
            self.folder_path_var.set(folder_path)

            # Update UI state now that a folder is selected
            self.update_ui_state()

            # Load files from the selected folder
            self.load_files_thread()

    def load_files_thread(self):
        # Clear current files
        self.files = []
        self.filtered_files = []

        # Ensure any memory from previous operations is cleaned up
        self._cleanup_memory()

        # Update UI
        self.clear_file_list()
        self.update_status(self.get_text("folder_loading"))

        # Start loading files in a separate thread
        self.cancel_flag = False

        # Before starting a new thread, check if we're running in EXE mode
        # and/or dealing with a large directory to avoid spawning multiple instances
        is_frozen = getattr(sys, 'frozen', False)
        if is_frozen:
            # In EXE mode, do a quick size check first
            try:
                # Quick check if this is a large directory
                top_files_count = len([f for f in os.listdir(self.selected_folder_path) 
                                    if os.path.isfile(os.path.join(self.selected_folder_path, f))])

                # If top directory has many files and we're including subfolders,
                # use a more aggressive memory management approach
                if (top_files_count > 500 and self.include_subfolders.get()) or top_files_count > 1000:
                    logging.info(f"Large directory detected: {top_files_count} files. Using optimized loading.")

                    # Make sure to clean up memory before proceeding
                    self._cleanup_memory()

                    # Sleep briefly to allow system to stabilize
                    time.sleep(0.1)
            except Exception as e:
                logging.warning(f"Error checking directory size: {str(e)}")

        # Now start the loading thread
        loading_thread = threading.Thread(target=self.load_files)
        loading_thread.daemon = True
        loading_thread.start()

    def load_files(self):
        try:
            # Clear current files
            self.files = []
            folder_count = 0
            total_size = 0

            # Update the UI to show loading state
            self.root.after(0, lambda: self.update_status(self.get_text("files_loading")))
            self.root.after(0, lambda: self.progress_bar.start(5))
            self.root.after(0, lambda: self.enable_cancel_button())

            # OPTIMIZATION: Create a batch processing approach
            file_batch = []
            total_processed = 0
            total_estimated_files = 0
            last_ui_update_time = time.time()

            # OPTIMIZATION: Pre-count files for better progress indication (for large directories)
            # Only count if top folder has more than certain number of direct files
            try:
                # Quick check of top directory first
                top_files_count = len([f for f in os.listdir(self.selected_folder_path) 
                                    if os.path.isfile(os.path.join(self.selected_folder_path, f))])

                # If we have many files or include subfolders, do a rough estimation
                if top_files_count > 1000 or self.include_subfolders.get():
                    self.update_status(f"{self.get_text('files_gathering')}...")
                    total_estimated_files = self._estimate_file_count()
                    logging.info(f"Estimated file count: {total_estimated_files}")
            except Exception as e:
                logging.warning(f"Error estimating file count: {str(e)}")
                total_estimated_files = 0

            # Create a deque for efficient batch processing
            file_batch = collections.deque(maxlen=self.file_loading_batch_size)

            # Walk through the folder structure
            for root, dirs, files in os.walk(self.selected_folder_path):
                # Check if the operation was cancelled
                if self.cancel_flag:
                    return self.handle_cancellation()

                # If we're only processing the top level, don't go into subdirectories
                if not self.include_subfolders.get() and root != self.selected_folder_path:
                    continue

                folder_count += 1

                # PERFORMANCE OPTIMIZATION: Parallel processing to calculate file sizes for faster sorting
                # For large directories, we now use ProcessPoolExecutor to bypass GIL and utilize multiple CPU cores
                try:
                    files_with_sizes = []

                    # Distribute work across multiple cores 
                    def get_file_size(file):
                        if self.cancel_flag:
                            return None
                        try:
                            file_path = os.path.join(root, file)
                            size = os.path.getsize(file_path)
                            return (file, size)
                        except:
                            return (file, 0)

                    # PERFORMANCE BOOST: Use ThreadPoolExecutor for file operations
                    # Avoiding ProcessPoolExecutor for better compatibility with exe compilation
                    # This ensures consistent behavior between script mode and exe deployments

                    # Detect if we're running in compiled mode vs script mode
                    is_frozen = getattr(sys, 'frozen', False)

                    # Always use ThreadPoolExecutor in exe mode to avoid multiprocessing issues
                    # For script mode, we can still use ProcessPoolExecutor for large file counts
                    if not is_frozen and len(files) > 500 and multiprocessing.current_process().name == 'MainProcess':
                        # In script mode, for large file lists, use ProcessPoolExecutor
                        try:
                            cpu_count = multiprocessing.cpu_count()
                            process_count = max(4, min(cpu_count, 16))  # Use between 4 and 16 processes

                            logging.info(f"Using ProcessPoolExecutor with {process_count} workers for {len(files)} files")

                            with concurrent.futures.ProcessPoolExecutor(max_workers=process_count) as executor:
                                # Create full file paths to pass to executor
                                file_paths = [(file, os.path.join(root, file)) for file in files]

                                # Define a worker-friendly function that doesn't rely on self
                                def get_size_process(file_tuple):
                                    filename, filepath = file_tuple
                                    try:
                                        size = os.path.getsize(filepath)
                                        return (filename, size)
                                    except:
                                        return (filename, 0)

                                # Submit all files for processing
                                future_to_file = {executor.submit(get_size_process, file_tuple): file_tuple[0] 
                                                for file_tuple in file_paths}

                                # Collect results as they complete
                                for future in concurrent.futures.as_completed(future_to_file):
                                    if self.cancel_flag:
                                        executor.shutdown(wait=False)
                                        return self.handle_cancellation()

                                    result = future.result()
                                    if result is not None:
                                        files_with_sizes.append(result)
                        except Exception as pe:
                            logging.warning(f"ProcessPoolExecutor failed, falling back to ThreadPoolExecutor: {str(pe)}")
                            # Fall back to ThreadPoolExecutor if ProcessPoolExecutor fails
                            with concurrent.futures.ThreadPoolExecutor(max_workers=32) as executor:
                                future_to_file = {executor.submit(get_file_size, file): file for file in files}
                                for future in concurrent.futures.as_completed(future_to_file):
                                    if self.cancel_flag:
                                        executor.shutdown(wait=False)
                                        return self.handle_cancellation()
                                    result = future.result()
                                    if result is not None:
                                        files_with_sizes.append(result)
                    else:
                        # For smaller file lists or in exe mode, thread pool is more reliable
                        with concurrent.futures.ThreadPoolExecutor(max_workers=16) as executor:
                            # Submit all files for processing
                            future_to_file = {executor.submit(get_file_size, file): file for file in files}

                            # Collect results as they complete
                            for future in concurrent.futures.as_completed(future_to_file):
                                if self.cancel_flag:
                                    executor.shutdown(wait=False)
                                    return self.handle_cancellation()

                                result = future.result()
                                if result is not None:
                                    files_with_sizes.append(result)
                except Exception as file_e:
                    logging.warning(f"Error getting file size: {str(file_e)}")
                    # Skip this file from the sorting process

                # Sort by size, smallest first
                try:
                    if files_with_sizes:  # Verify we have something to sort
                        files_with_sizes.sort(key=lambda x: x[1])
                        files = [f[0] for f in files_with_sizes]
                except Exception as e:
                    logging.warning(f"Error sorting files by size: {str(e)}")
                    # Continue with unsorted files

                # Process files in the current directory
                for file in files:
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        return self.handle_cancellation()

                    file_path = os.path.join(root, file)
                    try:
                        # Get file info
                        file_stats = os.stat(file_path)
                        file_size = file_stats.st_size
                        total_size += file_size

                        # Extract file extension
                        _, file_extension = os.path.splitext(file)

                        # Format dates - OPTIMIZATION: lazy loading of date formatting
                        # Defer expensive datetime operations
                        file_info = {
                            "name": file,
                            "extension": file_extension.lower(),
                            "path": root,
                            "full_path": file_path,
                            "size": file_size,
                            "ctime": file_stats.st_ctime,  # Store raw timestamp
                            "mtime": file_stats.st_mtime,  # Store raw timestamp
                            "created": None,  # Will be formatted when needed
                            "modified": None  # Will be formatted when needed
                        }

                        # Format the dates now (we need them for the UI)
                        file_info["created"] = datetime.datetime.fromtimestamp(file_info["ctime"]).strftime('%Y-%m-%d %H:%M:%S')
                        file_info["modified"] = datetime.datetime.fromtimestamp(file_info["mtime"]).strftime('%Y-%m-%d %H:%M:%S')

                        # Add file to the batch
                        file_batch.append(file_info)
                        total_processed += 1

                        # When batch is full, update UI and clear batch
                        if len(file_batch) >= self.file_loading_batch_size:
                            # Update the file list in batches for better performance
                            batch_list = list(file_batch)
                            self.files.extend(batch_list)

                            # Update progress calculations
                            progress_percentage = 0
                            if total_estimated_files > 0:
                                progress_percentage = min(95, (total_processed / total_estimated_files) * 100)
                                self.root.after(0, lambda p=progress_percentage: self.progress_bar.config(value=p))

                            # Only update UI every 0.5 seconds to avoid UI freeze with rapid updates
                            current_time = time.time()
                            if current_time - last_ui_update_time >= 0.5:
                                last_ui_update_time = current_time
                                # Update progress periodically to keep UI responsive
                                progress_text = f"{self.get_text('files_loading')} ({total_processed})"
                                if total_estimated_files > 0:
                                    progress_text += f" - {progress_percentage:.1f}%"
                                self.root.after(0, lambda msg=progress_text: self.update_status(msg))
                                self.root.update_idletasks()  # Force UI update

                            # Clear the batch for the next round
                            file_batch.clear()

                    except Exception as e:
                        logging.error(f"Error processing file {file_path}: {str(e)}")

            # Add any remaining files in the final batch
            if file_batch:
                batch_list = list(file_batch)
                self.files.extend(batch_list)

            # Save a copy of all files for search functionality
            # We need a copy here since filtering modifies elements
            self.all_files = self.files.copy()

            # Sort files before applying filter to ensure consistent results
            self.sort_files()

            # Set flag to indicate this is the first load
            self.is_first_load = True

            # Apply the initial filter immediately
            self.apply_filter_internal()

            # Update statistics
            self.root.after(0, lambda: self.update_statistics(len(self.filtered_files), folder_count, total_size))

            # Update UI
            self.root.after(0, lambda: self.update_status(
                self.get_text("folder_loaded_status")
            ))
            self.root.after(0, lambda: self.progress_bar.stop())
            def update_progress_value():
                self.progress_bar["value"] = 100
            self.root.after(0, update_progress_value)
            self.root.after(0, lambda: self.disable_cancel_button())

            # Preview button reference no longer needed with radio buttons

            # Run garbage collection to free memory
            self._cleanup_memory()

            # Log the operation
            logging.info(f"Loaded {len(self.files)} files from {self.selected_folder_path}")

        except Exception as e:
            # Handle errors
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            self.root.after(0, lambda: self.progress_bar.stop())
            self.root.after(0, lambda: self.disable_cancel_button())
            logging.error(f"Error loading files: {str(e)}")
            messagebox.showerror(self.get_text("error"), error_message)

    def _estimate_file_count(self):
        """
        OPTIMIZATION: Quickly estimate the number of files to be processed.
        This provides better progress indication for very large folders.
        Uses optimized sampling and memory-efficient processing to avoid full directory traversal.
        """
        try:
            # Check if we're running in EXE mode
            is_frozen = getattr(sys, 'frozen', False)

            # Check if cancel is requested
            if hasattr(self, 'cancel_flag') and self.cancel_flag:
                return 100  # Default estimate on cancellation

            # Use the selected folder path rather than looking it up again
            selected_path = getattr(self, 'selected_folder_path', self.folder_path_var.get())

            # Quick check if path is valid
            if not os.path.isdir(selected_path):
                return 100  # Default estimate for invalid path

            # For non-recursive mode, count top-level files efficiently
            if not self.include_subfolders.get():
                # Use scandir() instead of listdir() for better performance with large directories
                file_count = 0
                try:
                    with os.scandir(selected_path) as entries:
                        for entry in entries:
                            # Check cancellation periodically for responsiveness
                            if file_count % 1000 == 0 and hasattr(self, 'cancel_flag') and self.cancel_flag:
                                return 100

                            if entry.is_file():
                                file_count += 1
                                # Cap the estimation time for very large directories
                                # In EXE mode, be more conservative with large counts
                                cap = 5000 if is_frozen else 10000
                                if file_count > cap:
                                    # Apply a buffer factor based on what we've seen so far
                                    return int(file_count * 1.2)
                except (PermissionError, FileNotFoundError, OSError):
                    return 100

                return file_count

            # For recursive mode, use advanced sampling for better efficiency
            total_files = 0
            total_dirs = 0
            sampled_dirs = 0

            # In EXE mode, be more conservative with time spent on estimation
            max_time = 0.3 if is_frozen else 0.5  # Max seconds to spend on estimation
            start_time = time.time()

            # Use an efficient sample-based approach for large directories
            # Start with the top level
            top_level_files = 0
            top_level_dirs = []

            # Scan top level efficiently
            try:
                with os.scandir(selected_path) as entries:
                    for entry in entries:
                        if hasattr(self, 'cancel_flag') and self.cancel_flag:
                            return 100  # Cancelled, return reasonable default

                        if entry.is_file():
                            top_level_files += 1
                        elif entry.is_dir():
                            top_level_dirs.append(entry.path)

                        # In EXE mode, check more frequently if we should stop sampling
                        if is_frozen and (top_level_files + len(top_level_dirs)) % 500 == 0:
                            if time.time() - start_time > max_time * 0.4:  # Use 40% of budget for top level
                                break
            except (PermissionError, FileNotFoundError, OSError):
                # Handle access errors gracefully
                return max(100, top_level_files)

            # Add top level files to our total
            total_files += top_level_files
            total_dirs += 1

            # If we have lots of subdirectories, sample a subset for efficiency
            # Use a lower threshold in EXE mode
            dir_threshold = 15 if is_frozen else 20
            if len(top_level_dirs) > dir_threshold:
                # Prioritize directories with common names that often have many files
                # Include both English and Turkish common folder names
                common_large_folders = ["documents", "downloads", "pictures", "videos", "music", 
                                       "photos", "images", "docs", "media", "belgeler", "resimler",
                                       "indirilenler", "müzik", "muzik", "videolar", "dosyalar", 
                                       "dökümanlar", "dokumanlar", "fotoğraflar", "fotograflar"]

                # Sort directories to prioritize sampling known large folder types
                priority_dirs = []
                other_dirs = []

                for dir_path in top_level_dirs:
                    dirname = os.path.basename(dir_path).lower()
                    if any(common in dirname for common in common_large_folders):
                        priority_dirs.append(dir_path)
                    else:
                        other_dirs.append(dir_path)

                # Take some random samples to avoid bias
                # Be more conservative in EXE mode
                if len(priority_dirs) > 3:
                    random.shuffle(priority_dirs)

                if len(other_dirs) > 5:
                    random.shuffle(other_dirs)

                # Sample fewer directories in EXE mode to reduce memory pressure
                max_priority = 3 if is_frozen else 5
                max_total = 7 if is_frozen else 10

                # Take priority directories first, then fill with others
                sample_dirs = priority_dirs[:max_priority]
                remaining_slots = max_total - len(sample_dirs)

                if remaining_slots > 0:
                    sample_dirs.extend(other_dirs[:remaining_slots])
            else:
                # If few directories, process all of them
                sample_dirs = top_level_dirs

            # Sample the selected directories to depth 1 only
            sample_dir_files = 0
            sample_dir_subdirs = 0

            # Process each directory in our sample
            for dir_idx, dir_path in enumerate(sample_dirs):
                # Check timing constraint more frequently
                if time.time() - start_time > max_time:
                    # Time limit reached, extrapolate from what we've seen
                    break

                # Check cancellation
                if hasattr(self, 'cancel_flag') and self.cancel_flag:
                    break

                # In EXE mode, periodically force memory cleanup
                if is_frozen and dir_idx > 0 and dir_idx % 4 == 0:
                    # Quick GC pass to ensure memory is available
                    gc.collect()

                # Count files and subdirectories in this directory
                dir_file_count = 0
                dir_subdir_count = 0

                try:
                    with os.scandir(dir_path) as entries:
                        for entry in entries:
                            if entry.is_file():
                                dir_file_count += 1
                            elif entry.is_dir():
                                dir_subdir_count += 1

                            # Check limits periodically
                            if (dir_file_count + dir_subdir_count) % 500 == 0:
                                if hasattr(self, 'cancel_flag') and self.cancel_flag:
                                    break
                                if time.time() - start_time > max_time:
                                    break
                except (PermissionError, FileNotFoundError, OSError):
                    # Skip inaccessible directories
                    continue

                # Update counts
                sample_dir_files += dir_file_count
                sample_dir_subdirs += dir_subdir_count
                total_dirs += 1
                total_files += dir_file_count

            # Use our sample to extrapolate the full estimate
            # Base case - what we've counted so far
            estimate = total_files

            # If we have subfolders, extrapolate based on our sample
            if top_level_dirs and sample_dirs:  # If we have directories and sampled some
                # Calculate average files per sampled directory
                avg_files_per_dir = sample_dir_files / max(1, len(sample_dirs))

                # Estimate for unsampled top-level directories
                unsampled_count = len(top_level_dirs) - len(sample_dirs)
                if unsampled_count > 0:
                    estimate += unsampled_count * avg_files_per_dir

                # Add estimate for subdirectories (second level)
                if sample_dir_subdirs > 0:
                    # Assume each second-level directory has similar file count as first level
                    # but with a dampening factor since deeper directories often have fewer files
                    # Be more conservative in EXE mode to prevent overestimation
                    dampening = 0.6 if is_frozen else 0.7
                    avg_files_per_subdir = avg_files_per_dir * dampening
                    estimate += sample_dir_subdirs * avg_files_per_subdir

                # Add a buffer for very deep directory structures
                if self.include_subfolders.get() and (sample_dir_subdirs > 20 or len(top_level_dirs) > 20):
                    # More conservative multiplier in EXE mode
                    if is_frozen:
                        depth_multiplier = 1.0 + min(0.5, (sample_dir_subdirs + len(top_level_dirs)) / 200)
                    else:
                        depth_multiplier = 1.0 + min(1.0, (sample_dir_subdirs + len(top_level_dirs)) / 100)
                    estimate *= depth_multiplier

            # Cap maximum estimated files to prevent excessive UI progress bar issues
            max_estimate = 50000 if is_frozen else 100000
            estimate = min(max_estimate, estimate)

            # Ensure we return a reasonable minimum
            return max(100, int(estimate))

        except Exception as e:
            logging.error(f"Error estimating file count: {str(e)}")
            return 100  # Default estimate on error

    def enable_cancel_button(self):
        self.cancel_btn.config(state=tk.NORMAL)

    def disable_cancel_button(self):
        """Disable the cancel button to prevent multiple clicks"""
        self.cancel_btn.config(state=tk.DISABLED)

    def handle_cancellation(self):
        """
        Common method to handle cancellation UI updates across the application
        Centralizes all cancellation logic for consistent behavior
        """
        # Schedule re-enabling of the UI after a short delay
        self.root.after(500, self.disable_cancel_button)

        # Update the status bar with cancellation message
        self.root.after(0, lambda: self.update_status(self.get_text("operation_cancelled")))

        # Reset UI state once operation is fully cancelled
        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))

        # Reset cancellation flags for future operations
        def reset_cancel_flags():
            self.cancel_flag = False
            self.cancel_event.clear()
            # Clean up any temporary files created during the operation
            self._cleanup_temp_files()

        # Schedule flag reset for after UI updates complete
        self.root.after(1200, reset_cancel_flags)

    def _cleanup_temp_files(self):
        """
        Clean up any temporary files created during operations.
        This prevents accumulation of temp files that could consume disk space.
        """
        temp_files_removed = 0

        try:
            # Process each temp file in our tracking list
            for temp_file in self.temp_files[:]:
                if os.path.exists(temp_file):
                    try:
                        # Remove the file
                        os.remove(temp_file)
                        temp_files_removed += 1
                        # Remove from our tracking list
                        self.temp_files.remove(temp_file)
                    except (PermissionError, OSError) as e:
                        logging.warning(f"Could not remove temp file {temp_file}: {str(e)}")
                else:
                    # File doesn't exist, just remove from tracking
                    self.temp_files.remove(temp_file)

            if temp_files_removed > 0:
                logging.info(f"Cleaned up {temp_files_removed} temporary files")

        except Exception as e:
            # Non-critical error, just log it
            logging.error(f"Error during temp file cleanup: {str(e)}")

        # OPTIMIZATION: Clean up memory after operations
        self.root.after(1500, self._cleanup_memory)

        return False  # Common return value for operation cancellation

    def _cleanup_memory(self):
        """
        OPTIMIZATION: Enhanced memory management for large operations
        This helps prevent memory leaks and improves performance
        for operations that process many files by implementing a more
        aggressive and proactive memory management strategy.
        """
        try:
            # Detect if we're running in compiled mode vs script mode
            is_frozen = getattr(sys, 'frozen', False)

            # Thread-safety for cache operations
            if hasattr(self, 'preview_cache_lock'):
                cache_lock = self.preview_cache_lock
            else:
                cache_lock = threading.RLock()
                self.preview_cache_lock = cache_lock

            with cache_lock:
                # OPTIMIZATION: More aggressive cache management 
                # Clear the preview cache more aggressively based on usage patterns
                if hasattr(self, 'preview_cache'):
                    cache_size = len(self.preview_cache)

                    # If cache is very large, be more aggressive in pruning
                    # In EXE mode, be even more aggressive with memory management
                    threshold = self.max_preview_cache_size * (0.6 if is_frozen else 0.8)
                    if cache_size > threshold:
                        # Determine how many items to keep based on current memory pressure
                        # Keep fewer items when cache is larger (adaptive strategy)
                        # In EXE mode, keep even fewer items
                        keep_ratio = 0.1 if is_frozen else 0.15
                        keep_count = min(50 if is_frozen else 100, 
                                        max(10 if is_frozen else 25, 
                                            int(self.max_preview_cache_size * keep_ratio)))

                        if hasattr(self, 'preview_cache_keys') and self.preview_cache_keys:
                            # Keep the most recently used items, scaled by current cache size
                            keys_to_keep = self.preview_cache_keys[-keep_count:] if len(self.preview_cache_keys) > keep_count else self.preview_cache_keys

                            # Create new cache with only those items (more efficient than deleting)
                            new_cache = {}
                            for key in keys_to_keep:
                                if key in self.preview_cache:
                                    new_cache[key] = self.preview_cache[key]

                            # Log memory cleanup for better diagnostics
                            items_removed = cache_size - len(new_cache)
                            if items_removed > 0:
                                logging.info(f"Memory optimization: Released {items_removed} cached previews from memory")

                            # Replace old cache with smaller one
                            self.preview_cache = new_cache
                            self.preview_cache_keys = keys_to_keep
                        else:
                            # If no LRU tracking, clear most of the cache 
                            self.preview_cache = {}

            # Clear any temporary references that might be holding large objects
            if hasattr(self, 'preview_images'):
                self.preview_images = []

            # In EXE mode with large file lists, be more aggressive in memory management
            if is_frozen and hasattr(self, 'files') and len(self.files) > 1000:
                # Create simplified file list data to reduce memory usage
                simplified_files = []
                for file_info in self.files:
                    # Create a simplified copy with only essential data
                    simple_info = {
                        "name": file_info["name"],
                        "extension": file_info["extension"],
                        "path": file_info["path"],
                        "full_path": file_info["full_path"],
                        "size": file_info["size"],
                        "created": file_info["created"],
                        "modified": file_info["modified"]
                    }
                    simplified_files.append(simple_info)

                # Replace with simplified version
                self.files = simplified_files

                # Log memory optimization
                logging.info(f"EXE mode memory optimization: Simplified {len(simplified_files)} file records")

                # If all_files is also large, simplify it too
                if hasattr(self, 'all_files') and len(self.all_files) > 1000:
                    # Create simplified all_files list
                    simplified_all_files = []
                    for file_info in self.all_files:
                        simple_info = {
                            "name": file_info["name"],
                            "extension": file_info["extension"],
                            "path": file_info["path"],
                            "full_path": file_info["full_path"],
                            "size": file_info["size"],
                            "created": file_info["created"],
                            "modified": file_info["modified"]
                        }
                        simplified_all_files.append(simple_info)

                    # Replace with simplified version
                    self.all_files = simplified_all_files
                    logging.info(f"EXE mode memory optimization: Simplified {len(simplified_all_files)} all_files records")

            # Clear references to large file lists when appropriate
            if hasattr(self, 'cancel_flag') and self.cancel_flag:
                if hasattr(self, 'files'):
                    self.files = []
                if hasattr(self, 'filtered_files'):
                    self.filtered_files = []
                if hasattr(self, 'all_files'):
                    self.all_files = []

            # Clear thumbnail references
            if hasattr(self, 'preview_thumbnails'):
                self.preview_thumbnails = []

            # Explicitly force multiple passes of garbage collection to free memory
            for i in range(3):  # Multiple passes are more effective for complex reference cycles
                gc.collect(i)  # Collect specific generation

            # Encourage Python to return memory to OS
            if is_frozen and platform.system() == 'Windows':
                try:
                    # Windows-specific memory optimization for EXE mode
                    import ctypes
                    # Use EmptyWorkingSet to aggressively return memory to the OS
                    try:
                        ctypes.windll.psapi.EmptyWorkingSet(ctypes.windll.kernel32.GetCurrentProcess())
                        logging.info("Windows memory working set emptied")
                    except Exception as ws_e:
                        logging.warning(f"Windows EmptyWorkingSet failed: {str(ws_e)}")

                    # Try additional Python-specific memory optimizations
                    try:
                        ctypes.pythonapi.PyGC_Collect()
                    except Exception:
                        pass

                    # Try malloc_trim if available (glibc)
                    try:
                        ctypes.cdll.LoadLibrary('libc.so.6')
                        ctypes.CDLL('libc.so.6').malloc_trim(0)
                    except Exception:
                        pass
                except ImportError:
                    pass
            elif hasattr(sys, 'pypy_version_info'):
                # PyPy specific
                if hasattr(gc, 'collect'):
                    gc.collect()

            logging.info("Enhanced memory cleanup completed")
        except Exception as e:
            logging.error(f"Error during memory cleanup: {str(e)}")

    def cancel_operation(self):
        """
        Cancel the current operation and update UI accordingly.
        Uses both cancel_flag (legacy) and thread-safe cancel_event for robust cancellation.
        """
        # Set the cancellation flags - both legacy and new thread-safe mechanism
        self.cancel_flag = True
        self.cancel_event.set()

        # Update the status and progress immediately
        self.update_status(self.get_text("operation_cancelled"))
        self.progress_bar.stop()
        self.progress_bar["value"] = 0

        # Temporarily disable further cancellations to prevent multiple clicks
        self.cancel_btn.config(state=tk.DISABLED)

        # Visual feedback during cancellation with animated icon
        spinner_idx = random.randint(0, len(self.spinner_chars)-1)
        cancel_text = f"{self.get_text('cancelling')} {self.spinner_chars[spinner_idx]}"
        self.cancel_btn.config(text=cancel_text)
        self.root.update_idletasks()  # Force UI update

        # Apply common cancellation handling
        self.handle_cancellation()

        # Log the action with more detailed information
        logging.info("Operation cancelled by user - setting cancel flags and events")

    def _calculate_file_type_statistics(self):
        """Dosya uzantı istatistiklerini hesaplar ve statusbar için formatlı metni döndürür"""
        if not hasattr(self, 'filtered_files') or not self.filtered_files:
            return None

        # Her tür istatistik için yeni bir sözlük oluştur
        temp_stats = {}

        # Önizlemesi desteklenen dosya türleri
        supported_preview_extensions = [
            "jpg", "jpeg", "png", "gif", "bmp", "tiff", "tif", "svg", "ico", 
            "pdf", "eps", "psd", "ai"
        ]

        # Tüm dosyaları tekrar sayalım
        for file_info in self.filtered_files:
            if file_info.get("is_folder", False):
                continue

            # Dosya uzantısını al
            file_ext = file_info.get("extension", "").lower().replace(".", "")
            if not file_ext and "path" in file_info and "name" in file_info:
                file_path = os.path.join(file_info["path"], file_info["name"])
                file_ext = os.path.splitext(file_path)[1].lower().replace(".", "")

            # jpeg ve jpg uzantılarını birleştir
            if file_ext == "jpeg":
                file_ext = "jpg"

            # Sadece önizleme desteği olan uzantıları say
            if file_ext in supported_preview_extensions:
                # İstatistikleri güncelle
                if file_ext in temp_stats:
                    temp_stats[file_ext] += 1
                else:
                    temp_stats[file_ext] = 1

        # Geçici istatistikleri ana sözlüğe atayalım ve formatlı metni hazırlayalım
        self.file_type_stats = temp_stats

        if not self.file_type_stats:
            return None

        # Format statistics: sort by count (descending)
        stats_sorted = sorted(self.file_type_stats.items(), key=lambda x: x[1], reverse=True)

        # Ön izleme modunda top 10, liste modunda ya hiç gösterme ya da tüm dosya türleri
        if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview":
            # Ön izleme modunda en çok bulunan 10 dosya türünü göster
            top_stats = stats_sorted[:10]

            stats_text = []
            for ext, count in top_stats:
                # Uppercase the extension for better visibility
                stats_text.append(f"{ext.upper()}: {count}")

            # Create a nice statistics message including file count
            file_count_text = f"{len(self.filtered_files)} {self.get_text('files')}"
            stats_message = f"{file_count_text} | " + " | ".join(stats_text)
        else:
            # Liste modunda sadece dosya sayısı bilgisini göster
            file_count_text = f"{len(self.filtered_files)} {self.get_text('files')}"
            stats_message = file_count_text

        logging.info(f"Calculated stats: {stats_message}")
        return stats_message

    def update_status(self, message):
        """Update the status message in the status bar"""
        # If filtering is complete and we're trying to show filtering message, show statistics instead
        if hasattr(self, 'filtering_complete') and self.filtering_complete:
            # Skip filtering messages if filtering is already complete
            if message.startswith(self.get_text("filtering_in_progress")) or message == self.get_text("filtering"):
                # Filtreleme tamamlandıysa, istatistikleri göster
                stats_message = self._calculate_file_type_statistics()
                if stats_message:
                    self.status_var.set(stats_message)
                    logging.info(f"Showing statistics instead of filtering message: {stats_message}")
                    return

        # Check if we have file type statistics to display when in preview mode
        if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview" and message == self.get_text("preview_mode_active"):
            # Loglama ekleyelim
            logging.info("Updating status with preview mode stats")

            # İstatistik hesaplayıp gösterelim
            stats_message = self._calculate_file_type_statistics()
            if stats_message:
                self.status_var.set(stats_message)
                return

        # Default behavior (no statistics or not in preview mode)
        self.status_var.set(message)
        self.root.update_idletasks()

    def update_statistics(self, file_count, folder_count, total_size):
        # Update file count with optional subfolder indicator
        self.total_files_var.set(str(file_count))

        # Update folder count with clearer wording if subfolders are included
        if self.include_subfolders.get() and folder_count > 1:
            subfolder_count = folder_count - 1  # Main folder + subfolders
            folder_label = f"{folder_count} ({subfolder_count} {self.get_text('subfolders_label')})"
            self.folder_count_var.set(folder_label)
        else:
            self.folder_count_var.set(str(folder_count))

        # Format total size with appropriate units for better readability
        if total_size < 1024:
            size_str = f"{total_size} B"
        elif total_size < 1024 * 1024:
            size_str = f"{total_size/1024:.2f} KB"
        elif total_size < 1024 * 1024 * 1024:
            size_str = f"{total_size/(1024*1024):.2f} MB"
        else:
            size_str = f"{total_size/(1024*1024*1024):.2f} GB"

        self.total_size_var.set(size_str)

        # Log statistics for debugging
        logging.info(f"Statistics updated: {file_count} files, {folder_count} folders, {size_str} total size")

    def clear_file_list(self):
        # Clear all items in the file tree
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)

    def update_file_list(self, files):
        # Clear current list
        self.clear_file_list()

        # Store the original files list for filtering
        self.all_files = files.copy()

        # No files to display
        if not files:
            return

        # OPTIMIZATION: Insert files in batches for better performance with large file lists
        total_files = len(files)
        # Use the batch size defined in initialization
        batch_size = self.file_display_batch_size

        # Process in batches
        for i in range(0, total_files, batch_size):
            # Get the current batch
            batch = files[i:min(i+batch_size, total_files)]

            # Create a temporary list for batch updates
            batch_values = []
            for file_info in batch:
                values = (
                    file_info["name"],
                    file_info["extension"],
                    file_info["path"],
                    self.format_file_size(file_info["size"]),
                    file_info["created"],
                    file_info["modified"]
                )
                batch_values.append(values)

            # Insert the entire batch at once
            for values in batch_values:
                self.file_tree.insert("", tk.END, values=values)

            # Update the UI after each batch to maintain responsiveness
            self.root.update_idletasks()

        # If in preview mode, refresh the preview panel
        if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview":
            self._build_preview_panel(files)

    def _add_file_to_list(self, file_info):
        """Helper method to add a file to the list"""
        values = (
            file_info["name"],
            file_info["extension"],
            file_info["path"],
            self.format_file_size(file_info["size"]),
            file_info["created"],
            file_info["modified"]
        )
        self.file_tree.insert("", tk.END, values=values)

    def format_file_size(self, size):
        # Format file size with appropriate units
        if size < 1024:
            return f"{size} B"
        elif size < 1024 * 1024:
            return f"{size/1024:.2f} KB"
        elif size < 1024 * 1024 * 1024:
            return f"{size/(1024*1024):.2f} MB"
        else:
            return f"{size/(1024*1024*1024):.2f} GB"

    def select_all_extensions(self):
        # Set all extension checkboxes to True
        for ext in self.selected_extensions:
            self.selected_extensions[ext].set(True)

    def select_all_files(self):
        """Treeview'daki tüm dosyaları seç"""
        try:
            # Mevcut tüm öğeleri al
            children = self.file_tree.get_children()
            if children:
                # Tüm öğeleri seç
                self.file_tree.selection_set(children)
                # Durum bilgisini güncelle
                self.update_status(f"{len(children)} {self.get_text('files_selected')}")
            else:
                # Liste boşsa bilgi mesajı göster
                self.update_status(self.get_text("no_files_to_select"))
        except Exception as e:
            logging.error(f"Select all files error: {str(e)}")
            self.show_error("Error", f"{self.get_text('selection_error')}: {str(e)}")

    def clear_all_extensions(self):
        """Clear all file extension selections and reset the search filter"""
        # Clear all extension checkboxes
        for ext in self.selected_extensions:
            self.selected_extensions[ext].set(False)

        # Clear category selection checkboxes
        for category in self.category_vars:
            if category in self.category_vars:
                self.category_vars[category].set(False)

        # Reset search filter if it exists
        if hasattr(self, 'extension_search_var'):
            self.extension_search_var.set("")
            # Reset categories display based on cleared search
            self.filter_extensions()

    def apply_filter(self):
        # Start filtering in a separate thread
        self.cancel_flag = False
        self.update_status(self.get_text("filter_applying"))
        self.progress_bar.start(5)
        self.enable_cancel_button()

        filtering_thread = threading.Thread(target=self.apply_filter_internal)
        filtering_thread.daemon = True
        filtering_thread.start()

    def apply_filter_internal(self):
        try:
            # Set filtering status flag at the beginning
            self.filtering_complete = False

            # Filtreleme işlemi başladığında durum çubuğunda bir dönen simge göster
            self.update_status(self.get_text("filtering_in_progress") + " ⟳")
            self.root.config(cursor="watch")  # İmleç değiştir

            # Güncelleme fonksiyonu - dönen simge efekti için
            self.spinner_chars = ["⟳", "⟲", "↻", "↺"]
            self.spinner_index = 0

            def update_spinner():
                if self.cancel_flag:
                    return

                # Dönen simge karakterini değiştir
                self.spinner_index = (self.spinner_index + 1) % len(self.spinner_chars)
                self.update_status(self.get_text("filtering_in_progress") + " " + 
                                 self.spinner_chars[self.spinner_index])

                # Her 200ms'de bir güncelle
                self.root.after(200, update_spinner)

            # Dönen simgeyi başlat
            update_spinner()

            # Check if any extensions are selected
            any_selected = any(self.selected_extensions[ext].get() for ext in self.selected_extensions)

            # OPTIMIZATION: If this is the first load (right after load_files), show all files
            if hasattr(self, 'is_first_load') and self.is_first_load:
                self.is_first_load = False  # Reset flag after first use
                any_selected = False  # Force showing all files at first load

            # Create a set of selected extensions for faster lookup
            # OPTIMIZATION: Use set for O(1) lookup time instead of linear search
            selected_extensions_set = {ext.lower() for ext in self.selected_extensions 
                                      if ext in self.selected_extensions and self.selected_extensions[ext].get()}

            # If no extensions are selected, don't filter anything
            if not any_selected:
                self.filtered_files = self.files.copy()
            else:
                # OPTIMIZATION: Process in batches to improve responsiveness with large file lists
                self.filtered_files = []
                total_files = len(self.files)
                processed_count = 0
                last_ui_update_time = time.time()

                # OPTIMIZATION: Use deque for more efficient append operations
                filtered_batch = collections.deque(maxlen=self.file_filtering_batch_size * 2)

                # Process each file with efficient batching
                for i in range(0, total_files, self.file_filtering_batch_size):
                    # Get the current batch
                    end_idx = min(i + self.file_filtering_batch_size, total_files)
                    batch = self.files[i:end_idx]

                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        # Use special filter cancelled text for this particular case
                        self.root.after(0, lambda: self.update_status(self.get_text("filter_cancelled")))
                        self.root.after(0, lambda: self.cancel_btn.config(text=self.get_text("cancelling")))
                        self.root.after(0, lambda: self.progress_bar.stop())
                        self.root.after(500, self.disable_cancel_button)
                        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))
                        return

                    # Process each file in the current batch
                    # OPTIMIZATION: Avoid function calls in inner loop
                    for file_info in batch:
                        ext = file_info["extension"].lower()

                        # Check if the extension is in the selected extensions set (or if we're showing all files)
                        if not any_selected:
                            # If no extension is selected, show all files including ones without extension
                            if not ext:
                                # Handle files with no extension - create a copy to avoid modifying original
                                file_copy = file_info.copy()
                                file_copy["extension"] = self.get_text("extension_not_found")
                                filtered_batch.append(file_copy)
                            else:
                                filtered_batch.append(file_info)
                        else:
                            # If extensions are selected for filtering, only show files with those extensions
                            if ext in selected_extensions_set:
                                filtered_batch.append(file_info)
                            # Files with no extension are excluded when filtering by extension

                    # Only update UI every 0.3 seconds to improve performance
                    processed_count += len(batch)
                    current_time = time.time()
                    if current_time - last_ui_update_time >= 0.3:
                        last_ui_update_time = current_time
                        progress = (processed_count / total_files) * 100
                        self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                        self.root.after(0, lambda c=processed_count, t=total_files: 
                                      self.update_status(f"{self.get_text('filter_applying')} ({c}/{t})"))
                        self.root.update_idletasks()

                # Convert deque to list for the filtered files
                self.filtered_files = list(filtered_batch)

            # Sort files according to selected criteria
            self.sort_files()

            # OPTIMIZATION: Incrementally update UI for very large file lists
            # This avoids UI freezes when updating the treeview with thousands of items
            if len(self.filtered_files) > 5000:
                # For extremely large datasets, update in chunks
                self.update_status(f"{self.get_text('files_filtered_message').format(len(self.filtered_files))} - {self.get_text('loading_file_list')}")

                # First clear the list
                self.root.after(0, self.clear_file_list)
                self.root.update_idletasks()

                # Then update in chunks of the display batch size
                chunk_size = self.file_display_batch_size
                for i in range(0, len(self.filtered_files), chunk_size):
                    end_idx = min(i + chunk_size, len(self.filtered_files))
                    chunk = self.filtered_files[i:end_idx]
                    # Pass a copy of the chunk to avoid reference issues
                    self.root.after(0, lambda files=chunk.copy(): self._update_file_list_chunk(files))
                    # Allow small pauses for UI to remain responsive
                    self.root.update_idletasks()
            else:
                # For smaller datasets, update all at once
                self.root.after(0, lambda: self.update_file_list(self.filtered_files))

            # Update statistics
            # OPTIMIZATION: Calculate statistics in a more efficient way
            total_size = 0
            folder_paths = set()

            # Reset file type statistics for the filtered files
            self.file_type_stats = {}

            for file in self.filtered_files:
                total_size += file["size"]
                folder_paths.add(file["path"])

                # Update file type statistics
                if not file.get("is_folder", False):
                    # Get file extension, use a placeholder for files without extension
                    file_ext = file.get("extension", "").lower()
                    if not file_ext:
                        file_ext = self.get_text("extension_not_found")

                    # Update the count for this extension
                    if file_ext in self.file_type_stats:
                        self.file_type_stats[file_ext] += 1
                    else:
                        self.file_type_stats[file_ext] = 1

            folder_count = len(folder_paths)
            self.root.after(0, lambda: self.update_statistics(len(self.filtered_files), folder_count, total_size))

            # Update status
            self.root.after(0, lambda: self.update_status(
                self.get_text("files_filtered_message").format(len(self.filtered_files))
            ))
            self.root.after(0, lambda: self.progress_bar.stop())
            def update_progress_value():
                self.progress_bar["value"] = 100
            self.root.after(0, update_progress_value)
            self.root.after(0, lambda: self.disable_cancel_button())

            # Run garbage collection to free memory after filtering
            self._cleanup_memory()

            # İşlem bitince imleci normale döndür ve statüyü güncelle
            self.root.config(cursor="")
            # Set filtering complete flag to true
            self.filtering_complete = True
            # Update status with completion message
            self.update_status(self.get_text("filter_complete"))

        except Exception as e:
            # Hata göster ve loglama yap
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            self.root.after(0, lambda: self.progress_bar.stop())
            self.root.after(0, lambda: self.disable_cancel_button())
            logging.error(f"Error applying filter: {str(e)}")

            # Hata mesajını göster
            self.show_error(
                self.get_text("filter_error"), 
                f"{self.get_text('filter_error_details')}: {str(e)}", 
                e
            )

            # İmleci normale döndür
            self.root.config(cursor="")
            # Ensure we mark filtering as complete, even in case of error
            self.filtering_complete = True

    def _update_file_list_chunk(self, files_chunk):
        """Helper method to update the file list with a chunk of files
        Used for incremental updates with very large file lists"""
        # OPTIMIZATION: Batch insert records using a list comprehension
        # This is faster than inserting one by one for large chunks

        # First, prepare all values
        values_list = []
        for file_info in files_chunk:
            values = (
                file_info["name"],
                file_info["extension"],
                file_info["path"],
                self.format_file_size(file_info["size"]),
                file_info["created"],
                file_info["modified"]
            )
            values_list.append(values)

        # If we have a small number of files, insert directly
        if len(values_list) <= 100:
            for values in values_list:
                self.file_tree.insert("", tk.END, values=values)
        else:
            # For larger chunks, disable UI updates temporarily for better performance
            self.file_tree.config(takefocus=0)  # Temporarily disable focus

            try:
                # Insert all records
                for values in values_list:
                    self.file_tree.insert("", tk.END, values=values)
            finally:
                # Re-enable focus
                self.file_tree.config(takefocus=1)
                # Update the view
                self.file_tree.update()

    def sort_files(self):
        # Get the sort criteria
        sort_option = self.sort_options[self.sort_dropdown.current()]

        # Sort the filtered files
        if sort_option == "sort_name_asc":
            self.filtered_files.sort(key=lambda x: x["name"].lower())
        elif sort_option == "sort_name_desc":
            self.filtered_files.sort(key=lambda x: x["name"].lower(), reverse=True)
        elif sort_option == "sort_ext_asc":
            self.filtered_files.sort(key=lambda x: x["extension"].lower())
        elif sort_option == "sort_ext_desc":
            self.filtered_files.sort(key=lambda x: x["extension"].lower(), reverse=True)
        elif sort_option == "sort_size_asc":
            self.filtered_files.sort(key=lambda x: x["size"])
        elif sort_option == "sort_size_desc":
            self.filtered_files.sort(key=lambda x: x["size"], reverse=True)
        elif sort_option == "sort_dir_asc":
            self.filtered_files.sort(key=lambda x: x["path"].lower())

    def start_processing(self):
        # Check if a folder is selected
        if not self.selected_folder_path:
            messagebox.showwarning(
                self.get_text("error"),
                self.get_text("select_folder_first")
            )
            return

        # Check if any export format is selected
        if not any(self.export_formats[fmt].get() for fmt in self.export_formats):
            messagebox.showwarning(
                self.get_text("error"),
                self.get_text("select_format_first")
            )
            return

        # Start processing in a separate thread
        self.cancel_flag = False
        self.update_status(self.get_text("start_processing"))
        self.progress_bar["value"] = 0
        self.enable_cancel_button()

        processing_thread = threading.Thread(target=self.process_files)
        processing_thread.daemon = True
        processing_thread.start()

    def process_files(self):
        try:
            # Check if we have files to process
            if not self.filtered_files:
                self.root.after(0, lambda: messagebox.showinfo(
                    self.get_text("info"),
                    self.get_text("no_files_found")
                ))
                self.root.after(0, lambda: self.update_status(self.get_text("ready")))
                self.root.after(0, lambda: self.disable_cancel_button())
                return

            # Create timestamp for file names
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            folder_name = os.path.basename(self.selected_folder_path)
            base_filename = f"ListeKolay_{folder_name}_{timestamp}"

            # Determine save location
            if self.save_to_desktop.get():
                desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                # Create ListeKolay folder on desktop if it doesn't exist
                listekolay_folder = os.path.join(desktop_path, "ListeKolay")
                if not os.path.exists(listekolay_folder):
                    try:
                        os.makedirs(listekolay_folder)
                        logging.info(f"Created ListeKolay folder at {listekolay_folder}")
                    except Exception as e:
                        logging.error(f"Failed to create ListeKolay folder: {str(e)}")
                save_path = listekolay_folder
            else:
                # Yeni varsayılan dizin: Documents/ListeKolay/List
                documents_path = os.path.join(os.path.expanduser("~"), "Documents")
                listekolay_folder = os.path.join(documents_path, "ListeKolay")
                list_folder = os.path.join(listekolay_folder, "List")
                
                # Gerekli klasörleri oluştur
                for folder in [listekolay_folder, list_folder]:
                    if not os.path.exists(folder):
                        try:
                            os.makedirs(folder)
                            logging.info(f"Created folder: {folder}")
                        except Exception as e:
                            logging.error(f"Failed to create folder {folder}: {str(e)}")
                            # Oluşturulamazsa, uygulama dizinine kaydet
                            list_folder = os.path.dirname(os.path.abspath(__file__))
                            break
                
                save_path = list_folder

            created_files = []

            # Process text format
            if self.export_formats["text"].get():
                text_file = os.path.join(save_path, base_filename + ".txt")
                if self.export_text_file(text_file):
                    created_files.append(("text", text_file))

            # Process Excel format
            if self.export_formats["excel"].get():
                excel_file = os.path.join(save_path, base_filename + ".xlsx")
                if self.export_excel_file(excel_file):
                    created_files.append(("excel", excel_file))

            # Process Word format
            if self.export_formats["word"].get():
                word_file = os.path.join(save_path, base_filename + ".docx")
                if self.export_word_file(word_file):
                    created_files.append(("word", word_file))

            # Process HTML format
            if self.export_formats["html"].get():
                html_file = os.path.join(save_path, base_filename + ".html")
                if self.export_html_file(html_file):
                    created_files.append(("html", html_file))

            # Update UI
            def update_progress_value():
                self.progress_bar["value"] = 100
            self.root.after(0, update_progress_value)
            self.root.after(0, lambda: self.update_status(self.get_text("create_list_time").format(
                datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )))
            self.root.after(0, lambda: self.disable_cancel_button())

            # Ask user if they want to open the first created file
            if created_files:
                file_type, file_path = created_files[0]
                self.root.after(0, lambda: self.ask_to_open_file(file_type, file_path))

        except Exception as e:
            # Handle errors
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            self.root.after(0, lambda: self.disable_cancel_button())
            logging.error(f"Error processing files: {str(e)}")
            self.root.after(0, lambda: messagebox.showerror(self.get_text("error"), error_message))

    def export_text_file(self, file_path):
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write(f"{self.get_text('file_list')} - {self.selected_folder_path}\n")
                f.write(f"{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

                # Add sort information
                sort_option = self.selected_sort.get()
                sort_name = self.get_text(sort_option)
                f.write(f"{self.get_text('sorted_by')}: {sort_name}\n")

                f.write("=" * 80 + "\n\n")

                # Make sure files are sorted according to the selected criteria
                self.sort_files()

                # Write file information
                for i, file_info in enumerate(self.filtered_files, 1):
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        return self.handle_cancellation()

                    # Update progress
                    progress = (i / len(self.filtered_files)) * 100
                    self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                    self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                                 self.update_status(self.get_text("file_processed").format(i, total, name)))

                    # Write file details
                    f.write(f"{i}. {file_info['name']}\n")
                    f.write(f"   {self.get_text('file_path')}: {file_info['path']}\n")
                    f.write(f"   {self.get_text('file_extension')}: {file_info['extension']}\n")
                    f.write(f"   {self.get_text('file_size')}: {self.format_file_size(file_info['size'])}\n")
                    f.write(f"   {self.get_text('creation_date')}: {file_info['created']}\n")
                    f.write(f"   {self.get_text('modification_date')}: {file_info['modified']}\n")
                    f.write("\n")

            # Log success
            logging.info(f"Created text file: {file_path}")
            return True

        except Exception as e:
            # Log error
            error_message = self.get_text("text_file_error").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating text file: {str(e)}")
            return False

    def export_excel_file(self, file_path):
        try:
            # Create a new workbook and select the active sheet
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = self.get_text("file_list")

            # Create a more attractive header
            # Title row
            title_cell = sheet.cell(row=1, column=1)
            title_cell.value = self.get_text("file_list")
            title_cell.font = Font(name='Arial', size=14, bold=True, color="0000FF")
            sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
            title_cell.alignment = Alignment(horizontal="center", vertical="center")

            # Folder path row
            folder_cell = sheet.cell(row=2, column=1)
            folder_cell.value = f"{self.get_text('selected_folder')}: {self.selected_folder_path}"
            folder_cell.font = Font(name='Arial', size=10, bold=True)
            sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=7)

            # Creation time row
            time_cell = sheet.cell(row=3, column=1)
            time_cell.value = f"{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            time_cell.font = Font(name='Arial', size=10)
            sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=7)

            # Sort information row
            sort_option = self.selected_sort.get()
            sort_name = self.get_text(sort_option)
            sort_cell = sheet.cell(row=4, column=1)
            sort_cell.value = f"{self.get_text('sorted_by')} {sort_name}"
            sort_cell.font = Font(name='Arial', size=10, italic=True)
            sheet.merge_cells(start_row=4, start_column=1, end_row=4, end_column=7)

            # Add a bit of space before the table
            # Row 5 is empty for spacing

            # Make sure files are sorted according to the selected criteria
            self.sort_files()

            # Add header row
            headers = [
                self.get_text("row_number"),
                self.get_text("file_name"),
                self.get_text("file_extension"),
                self.get_text("file_path"),
                self.get_text("file_size"),
                self.get_text("creation_date"),
                self.get_text("modification_date")
            ]

            # Style for the header row
            header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            header_border = Border(
                left=Side(style='thin'), 
                right=Side(style='thin'), 
                top=Side(style='thin'), 
                bottom=Side(style='thin')
            )

            # Apply header styles
            for col, header in enumerate(headers, 1):
                cell = sheet.cell(row=6, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = header_border

            # Style for data rows
            data_border = Border(
                left=Side(style='thin'), 
                right=Side(style='thin'), 
                top=Side(style='thin'), 
                bottom=Side(style='thin')
            )
            even_row_fill = PatternFill(start_color="E9EDF4", end_color="E9EDF4", fill_type="solid")
            odd_row_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

            # Add file information
            for i, file_info in enumerate(self.filtered_files, 1):
                # Check if the operation was cancelled
                if self.cancel_flag:
                    return self.handle_cancellation()

                # Update progress
                progress = (i / len(self.filtered_files)) * 100
                self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                             self.update_status(self.get_text("file_processed").format(i, total, name)))

                # Add file details
                row = i + 6  # Data starts at row 7 (6+1)

                # Apply alternating row colors
                row_fill = even_row_fill if i % 2 == 0 else odd_row_fill

                # Set cell values and styles
                for col in range(1, 8):
                    cell = sheet.cell(row=row, column=col)
                    cell.border = data_border
                    cell.fill = row_fill

                    # Set alignment based on column
                    if col == 1:  # Row number - center align
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

                # Set cell values
                sheet.cell(row=row, column=1).value = i
                sheet.cell(row=row, column=2).value = file_info["name"]
                sheet.cell(row=row, column=3).value = file_info["extension"]
                sheet.cell(row=row, column=4).value = file_info["path"]
                sheet.cell(row=row, column=5).value = self.format_file_size(file_info["size"])
                sheet.cell(row=row, column=6).value = file_info["created"]
                sheet.cell(row=row, column=7).value = file_info["modified"]

            # Auto-adjust column widths
            for col in range(1, len(headers) + 1):
                max_length = 0
                for row in range(1, len(self.filtered_files) + 7):  # +7 for the header rows
                    cell_value = sheet.cell(row=row, column=col).value
                    if cell_value:
                        max_length = max(max_length, len(str(cell_value)))
                # Set minimum and maximum widths
                adjusted_width = min(50, max(12, max_length + 2))
                sheet.column_dimensions[chr(64 + col)].width = adjusted_width

            # Apply auto-filter to the header row to make columns filterable
            sheet.auto_filter.ref = f"A6:G{len(self.filtered_files) + 6}"

            # Freeze the header row so it stays visible when scrolling
            sheet.freeze_panes = "A7"

            # Save the workbook
            workbook.save(file_path)

            # Log success
            logging.info(f"Created Excel file: {file_path}")
            return True

        except Exception as e:
            # Log error
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating Excel file: {str(e)}")
            return False

    def export_word_file(self, file_path):
        try:
            # Create a new document
            document = Document()

            # Add title
            document.add_heading(self.get_text("file_list"), level=1)

            # Add folder information
            document.add_paragraph(f"{self.get_text('selected_folder')}: {self.selected_folder_path}")
            document.add_paragraph(f"{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # Add sort information
            sort_option = self.selected_sort.get()
            sort_name = self.get_text(sort_option)
            document.add_paragraph(f"{self.get_text('sorted_by')}: {sort_name}")

            # Make sure files are sorted according to the selected criteria
            self.sort_files()

            # Add a horizontal line
            document.add_paragraph("_" * 50)

            # Add file information
            for i, file_info in enumerate(self.filtered_files, 1):
                # Check if the operation was cancelled
                if self.cancel_flag:
                    return self.handle_cancellation()

                # Update progress
                progress = (i / len(self.filtered_files)) * 100
                self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                             self.update_status(self.get_text("file_processed").format(i, total, name)))

                # Add file details
                document.add_heading(f"{i}. {file_info['name']}", level=2)
                table = document.add_table(rows=5, cols=2)
                table.style = "Table Grid"

                # Fill the table with file details
                cells = table.rows[0].cells
                cells[0].text = self.get_text("file_path")
                cells[1].text = file_info["path"]

                cells = table.rows[1].cells
                cells[0].text = self.get_text("file_extension")
                cells[1].text = file_info["extension"]

                cells = table.rows[2].cells
                cells[0].text = self.get_text("file_size")
                cells[1].text = self.format_file_size(file_info["size"])

                cells = table.rows[3].cells
                cells[0].text = self.get_text("creation_date")
                cells[1].text = file_info["created"]

                cells = table.rows[4].cells
                cells[0].text = self.get_text("modification_date")
                cells[1].text = file_info["modified"]

                # Add a space after each file
                document.add_paragraph()

            # Save the document
            document.save(file_path)

            # Log success
            logging.info(f"Created Word file: {file_path}")
            return True

        except Exception as e:
            # Log error
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating Word file: {str(e)}")
            return False

    def export_html_file(self, file_path):
        try:
            # Make sure files are sorted according to the selected criteria
            self.sort_files()

            with open(file_path, 'w', encoding='utf-8') as f:
                # Write HTML header
                f.write(f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{self.get_text('file_list')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1, h2 {{ color: #333; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ padding: 8px; text-align: left; border: 1px solid #ddd; }}
        th {{ background-color: #f2f2f2; }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        .info {{ color: #666; margin-bottom: 20px; }}
    </style>
</head>
<body>
    <h1>{self.get_text('file_list')}</h1>
    <div class="info">
        <p>{self.get_text('selected_folder')}: {self.selected_folder_path}</p>
        <p>{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p>{self.get_text('sorted_by')}: {self.get_text(self.selected_sort.get())}</p>
    </div>
    <table>
        <tr>
            <th>{self.get_text('row_number')}</th>
            <th>{self.get_text('file_name')}</th>
            <th>{self.get_text('file_extension')}</th>
            <th>{self.get_text('file_path')}</th>
            <th>{self.get_text('file_size')}</th>
            <th>{self.get_text('creation_date')}</th>
            <th>{self.get_text('modification_date')}</th>
        </tr>
""")

                # Add file information
                for i, file_info in enumerate(self.filtered_files, 1):
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        return self.handle_cancellation()

                    # Update progress
                    progress = (i / len(self.filtered_files)) * 100
                    self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                    self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                                 self.update_status(self.get_text("file_processed").format(i, total, name)))

                    # Write file details
                    f.write(f"""
        <tr>
            <td>{i}</td>
            <td>{file_info['name']}</td>
            <td>{file_info['extension']}</td>
            <td>{file_info['path']}</td>
            <td>{self.format_file_size(file_info['size'])}</td>
            <td>{file_info['created']}</td>
            <td>{file_info['modified']}</td>
        </tr>""")

                # Write HTML footer
                f.write("""
    </table>
</body>
</html>""")

            # Log success
            logging.info(f"Created HTML file: {file_path}")
            return True

        except Exception as e:
            # Log error
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating HTML file: {str(e)}")
            return False

    def ask_to_open_file(self, file_type, file_path):
        # Format success message based on file type
        success_message = self.get_text(f"{file_type}_success")
        self.update_status(success_message)

        # Ask user if they want to open the file
        response = messagebox.askyesno(
            self.get_text("open_file_title"),
            self.get_text("open_file_message"),
            icon=messagebox.QUESTION
        )

        if response:
            try:
                # Open the file with the default application
                if sys.platform == "win32":
                    os.startfile(file_path)
                elif sys.platform == "darwin":  # macOS
                    subprocess.call(["open", file_path])
                else:  # Linux and other Unix-like
                    subprocess.call(["xdg-open", file_path])

                logging.info(f"Opened file: {file_path}")
            except Exception as e:
                logging.error(f"Error opening file {file_path}: {str(e)}")
                messagebox.showerror(
                    self.get_text("error"),
                    self.get_text("error_occurred").format(str(e))
                )

    def _go_to_prev_page(self):
        """Navigate to the previous page of thumbnails"""
        if hasattr(self, 'preview_page') and self.preview_page > 1:
            # Update status
            self.update_status(self.get_text("loading_preview"))

            # Show loading indicator
            self.progress_bar.start(10)

            # Change page
            self.preview_page -= 1

            # Log for debugging
            logging.info(f"Moving to previous page: {self.preview_page}")

            # Update preview panel
            self._update_preview_panel()

            # Update page info
            if hasattr(self, 'page_info_label') and hasattr(self, 'total_preview_pages'):
                self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")

            # Disable prev button if on first page
            if self.preview_page == 1:
                self.prev_page_btn.config(state=tk.DISABLED)

            # Always enable next button if not on last page
            if hasattr(self, 'total_preview_pages') and self.preview_page < self.total_preview_pages:
                self.next_page_btn.config(state=tk.NORMAL)

    def _go_to_next_page(self):
        """Navigate to the next page of thumbnails"""
        if hasattr(self, 'preview_page') and hasattr(self, 'total_preview_pages') and self.preview_page < self.total_preview_pages:
            # Update status
            self.update_status(self.get_text("loading_preview"))

            # Show loading indicator
            self.progress_bar.start(10)

            # Change page
            self.preview_page += 1

            # Log for debugging
            logging.info(f"Moving to next page: {self.preview_page}/{self.total_preview_pages}")

            # Update preview panel
            self._update_preview_panel()

            # Update page info
            if hasattr(self, 'page_info_label'):
                self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")

            # Disable next button if on last page
            if self.preview_page == self.total_preview_pages:
                self.next_page_btn.config(state=tk.DISABLED)

            # Always enable prev button if not on first page
            if self.preview_page > 1:
                self.prev_page_btn.config(state=tk.NORMAL)

    def _update_preview_content(self, file_path):
        """Update the content of the preview window without destroying it"""
        if not hasattr(self, 'preview_window') or not self.preview_window:
            return False

        try:
            # Get file details
            file_name = os.path.basename(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            # Get modification and creation dates
            file_mod_date = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S') if os.path.exists(file_path) else ""

            # Try to get creation date (platform specific)
            try:
                if os.name == 'nt':  # Windows
                    file_creation_date = datetime.datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
                else:  # Unix/Linux/Mac
                    # On Unix, getctime returns status change time, not creation time, since Unix doesn't track creation time
                    # Using stat to get the best approximation
                    stat_info = os.stat(file_path)
                    file_creation_date = datetime.datetime.fromtimestamp(stat_info.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            except:
                file_creation_date = ""

            # Update window title (optional)
            self.preview_window.title(self.get_text("preview_window_title"))

            # Find and update all widgets in the info frame
            if hasattr(self, '_preview_info_widgets'):
                # Update file name
                if 'name_label' in self._preview_info_widgets:
                    self._preview_info_widgets['name_label'].config(text=file_name)

                # Update file type icon
                if 'icon_label' in self._preview_info_widgets:
                    # File type icon or text based on extension
                    if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                        icon_text = "IMG"
                        icon_bg = "#28a745"  # Green for images
                    elif file_extension in ['.pdf']:
                        icon_text = "PDF"
                        icon_bg = "#dc3545"  # Red for PDFs 
                    elif file_extension in ['.psd', '.ai', '.eps']:
                        icon_text = "PSD"
                        icon_bg = "#6610f2"  # Purple for design files
                    elif file_extension in ['.doc', '.docx', '.txt', '.rtf']:
                        icon_text = "DOC"
                        icon_bg = "#007bff"  # Blue for documents
                    elif file_extension in ['.mp3', '.wav', '.flac', '.aac', '.ogg']:
                        icon_text = "AUD"
                        icon_bg = "#fd7e14"  # Orange for audio
                    elif file_extension in ['.mp4', '.mov', '.mkv', '.avi']:
                        icon_text = "VID"
                        icon_bg = "#6f42c1"  # Purple for video
                    else:
                        # Get the file extension without dot and make it uppercase
                        icon_text = file_extension.upper().replace(".", "") if file_extension else "FILE"
                        icon_bg = "#6c757d"  # Gray for other types

                    self._preview_info_widgets['icon_label'].config(text=icon_text, bg=icon_bg)

                # Update size label
                if 'size_label' in self._preview_info_widgets:
                    self._preview_info_widgets['size_label'].config(text=f"{self.get_text('file_size')}: {self.format_file_size(file_size)}")

                # Update extension label
                if 'ext_label' in self._preview_info_widgets:
                    self._preview_info_widgets['ext_label'].config(text=f"{self.get_text('file_extension')}: {file_extension}")

                # Update created date label
                if 'created_label' in self._preview_info_widgets:
                    self._preview_info_widgets['created_label'].config(text=f"{self.get_text('creation_date')}: {file_creation_date}")

                # Update modified date label
                if 'modified_label' in self._preview_info_widgets:
                    self._preview_info_widgets['modified_label'].config(text=f"{self.get_text('modification_date')}: {file_mod_date}")

                # Update navigation label
                if 'nav_label' in self._preview_info_widgets and hasattr(self, 'current_preview_files') and len(self.current_preview_files) > 1 and self.current_preview_index >= 0:
                    nav_text = f"{self.current_preview_index + 1} / {len(self.current_preview_files)}"
                    self._preview_info_widgets['nav_label'].config(text=nav_text)

                # Update navigation buttons state
                if 'prev_button' in self._preview_info_widgets:
                    self._preview_info_widgets['prev_button'].config(state=tk.NORMAL if self.current_preview_index > 0 else tk.DISABLED)

                if 'next_button' in self._preview_info_widgets:
                    self._preview_info_widgets['next_button'].config(state=tk.NORMAL if self.current_preview_index < len(self.current_preview_files) - 1 else tk.DISABLED)

            # Clear preview frame and add new content
            if hasattr(self, '_preview_content_frame'):
                # Remove all widgets from preview frame
                for widget in self._preview_content_frame.winfo_children():
                    widget.destroy()

                # Check file type and generate preview
                if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                    self.preview_image(self._preview_content_frame, file_path)
                elif file_extension in ['.pdf']:
                    self.preview_pdf(self._preview_content_frame, file_path)
                elif file_extension in ['.psd', '.ai', '.eps']:
                    self.preview_design_file(self._preview_content_frame, file_path)
                else:
                    # Show unsupported message
                    label = tk.Label(
                        self._preview_content_frame, 
                        text=self.get_text("preview_not_supported"),
                        font=("Segoe UI", 12),
                        bg="#e9ecef"
                    )
                    label.pack(pady=20)

            return True
        except Exception as e:
            logging.error(f"Error updating preview content: {str(e)}")
            return False

    def _go_to_prev_file(self):
        """Navigate to the previous file in the preview window"""
        if not hasattr(self, 'current_preview_files') or not self.current_preview_files:
            return

        if self.current_preview_index > 0:
            # Get previous file
            self.current_preview_index -= 1
            prev_file = self.current_preview_files[self.current_preview_index]

            # Open preview for that file
            if "path" in prev_file and "name" in prev_file:
                # Get full file path
                file_path = os.path.join(prev_file["path"], prev_file["name"])

                # Always try to update existing preview window first
                result = self._update_preview_content(file_path)
                if not result:
                    # Only if updating fails, create a new window
                    logging.debug("Falling back to creating a new preview window")
                    self.create_file_preview_window(file_path, self.current_preview_index)

    def _go_to_next_file(self):
        """Navigate to the next file in the preview window"""
        if not hasattr(self, 'current_preview_files') or not self.current_preview_files:
            return

        if self.current_preview_index < len(self.current_preview_files) - 1:
            # Get next file
            self.current_preview_index += 1
            next_file = self.current_preview_files[self.current_preview_index]

            # Open preview for that file
            if "path" in next_file and "name" in next_file:
                # Get full file path
                file_path = os.path.join(next_file["path"], next_file["name"])

                # Always try to update existing preview window first
                result = self._update_preview_content(file_path)
                if not result:
                    # Only if updating fails, create a new window
                    logging.debug("Falling back to creating a new preview window")
                    self.create_file_preview_window(file_path, self.current_preview_index)

    def on_subfolder_changed(self):
        """Called when the Include Subfolders checkbox state changes"""
        # If a folder is selected, reload files when subfolder setting changes
        if self.selected_folder_path:
            # Show loading indicator
            self.progress_bar.start(10)

            # Update status bar with appropriate message
            if self.include_subfolders.get():
                status_msg = self.get_text("loading_subfolders") if "loading_subfolders" in self.languages[self.current_language] else "Loading subfolders..."
            else:
                status_msg = self.get_text("folder_loading")
            self.update_status(status_msg)

            # Start a new thread to reload the files with the new subfolder setting
            self.load_files_thread()

            # Log the change for debugging
            logging.info(f"Subfolder option changed to: {self.include_subfolders.get()}")

            # Change the tooltip to reflect current state
            subfolders_tooltip = self.get_text("tooltip_subfolders")
            if self.include_subfolders.get():
                subfolders_tooltip += " ✓"  # Add a checkmark to indicate it's enabled
            self.create_tooltip(self.subfolder_cb, subfolders_tooltip)

    def open_website(self, url):
        """Open a website URL in the default browser"""
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception as e:
            logging.error(f"Failed to open website: {e}")
            messagebox.showerror("Error", f"Failed to open website: {e}")

    def check_for_updates(self, silent=False):
        """
        GitHub'dan son sürüm bilgisini alıp mevcut sürümle karşılaştırır
        silent=True ise sadece güncelleme varsa bildirim yapar
        """
        try:
            # Sürüm kontrolü için HTTP isteği gönder
            response = requests.get(self.github_version_url, timeout=5)

            if response.status_code == 200:
                # Uzaktaki sürüm bilgisini al (boşlukları temizle)
                github_version = response.text.strip()

                # Sürüm karşılaştırması yap
                if github_version != self.current_version:
                    # Güncelleme mevcut
                    if messagebox.askyesno(
                        self.get_text("update_available"),
                        self.get_text("update_available_message").format(github_version, self.current_version)
                    ):
                        # Replit ortamında doğrudan indirme devre dışı, sayfayı aç
                        self.open_website(self.github_download_url)
                        # Normal sistemlerde indirme aşağıdaki şekilde olacak
                        # self.download_update(github_version)
                elif not silent:  
                    # Zaten son sürüm kullanılıyor ve sessiz mod değilse bildirim yap
                    messagebox.showinfo(
                        self.get_text("no_update_available"),
                        self.get_text("no_update_available_message")
                    )

                return github_version != self.current_version
            else:
                if not silent:
                    # Hata durumunda bildirim yap (sessiz mod değilse)
                    self.show_error(
                        self.get_text("update_check_error"),
                        self.get_text("update_check_error_message")
                    )
                return False

        except Exception as e:
            if not silent:
                # Bağlantı hatası bildirim yap (sessiz mod değilse)
                self.show_error(
                    self.get_text("update_check_error"),
                    f"{self.get_text('update_check_error_message')} ({str(e)})"
                )
            logging.error(f"{self.get_text('update_check_error')}: {str(e)}")
            return False

    def download_update(self, new_version):
        """
        Yeni sürümü GitHub'dan indir ve otomatik olarak güncelle
        """
        try:
            # İndirme URL'ini oluştur (GitHub releases sayfasından indirme)
            download_url = f"{self.github_download_url}/download/v{new_version}/ListeKolay_v{new_version}.zip"

            # İndirme ilerleme penceresini oluştur
            download_window = tk.Toplevel(self.root)
            download_window.title(self.get_text("downloading_update"))
            download_window.geometry("400x150")
            download_window.resizable(False, False)
            download_window.configure(bg="#e9ecef")
            download_window.transient(self.root)  # Ana pencereye bağlı
            download_window.grab_set()  # Diğer işlemleri engelle

            # Pencere merkezi konumlandırma
            download_window.update_idletasks()
            width = download_window.winfo_width()
            height = download_window.winfo_height()
            x = (download_window.winfo_screenwidth() // 2) - (width // 2)
            y = (download_window.winfo_screenheight() // 2) - (height // 2)
            download_window.geometry(f"+{x}+{y}")

            # İndirme durumu etiketi
            status_label = tk.Label(
                download_window, 
                text=self.get_text("downloading_update_message").format(new_version),
                font=("Segoe UI", 10),
                bg="#e9ecef",
                fg="#212529"
            )
            status_label.pack(pady=(20, 10))

            # İndirme ilerleme çubuğu
            progress_bar = ttk.Progressbar(
                download_window, 
                orient=tk.HORIZONTAL, 
                length=350, 
                mode='indeterminate'
            )
            progress_bar.pack(pady=10, padx=25)
            progress_bar.start(10)

            # İptal butonu
            cancel_button = tk.Button(
                download_window,
                text=self.get_text("cancel_button"),
                command=download_window.destroy,
                font=("Segoe UI", 9),
                bg="#dc3545",
                fg="white",
                activebackground="#c82333",
                activeforeground="white",
                bd=0,
                padx=10
            )
            cancel_button.pack(pady=10)

            # İndirme fonksiyonu (thread içinde çalışacak)
            def download_thread():
                try:
                    # Geçici dosya oluştur
                    temp_dir = tempfile.gettempdir()
                    output_file = os.path.join(temp_dir, f"ListeKolay_v{new_version}.zip")

                    # İndirme işlemini başlat
                    response = requests.get(download_url, stream=True)

                    if response.status_code == 200:
                        # Dosyayı kaydet
                        with open(output_file, "wb") as f:
                            for chunk in response.iter_content(chunk_size=1024):
                                if chunk:  # Boş paketleri filtrele
                                    f.write(chunk)

                        # İndirme penceresini kapat
                        download_window.after(0, download_window.destroy)

                        # İndirme tamamlandı mesajı
                        if messagebox.showinfo(
                            self.get_text("download_complete"),
                            self.get_text("download_complete_message")
                        ):
                            # Yeni sürümü başlat
                            self.launch_updated_version(output_file)
                    else:
                        # İndirme hatası
                        download_window.after(0, download_window.destroy)
                        messagebox.showerror(
                            self.get_text("download_error"),
                            self.get_text("download_error_message").format(f"HTTP {response.status_code}")
                        )
                except Exception as e:
                    # İndirme sırasında hata
                    logging.error(f"Download error: {e}")
                    download_window.after(0, download_window.destroy)
                    messagebox.showerror(
                        self.get_text("download_error"),
                        self.get_text("download_error_message").format(str(e))
                    )

            # İndirme thread'ini başlat
            download_thread_obj = threading.Thread(target=download_thread)
            download_thread_obj.daemon = True
            download_thread_obj.start()

        except Exception as e:
            # Genel hata durumu
            logging.error(f"Update download error: {e}")
            messagebox.showerror(
                self.get_text("download_error"),
                self.get_text("download_error_message").format(str(e))
            )

    def launch_updated_version(self, zip_file):
        """
        İndirilen zip dosyasını çıkart ve yeni sürümü başlat
        """
        try:
            # Programın mevcut konumunu al
            current_dir = os.path.dirname(os.path.abspath(sys.argv[0]))

            # Zip dosyasını çıkartma ve yeni sürümü başlatma işlemleri için yardımcı betik oluştur
            updater_script = os.path.join(tempfile.gettempdir(), "listekolay_updater.py")

            # Yolları uygun şekilde formatlayalım
            safe_zip_path = zip_file.replace('\\', '\\\\')
            safe_current_dir = current_dir.replace('\\', '\\\\')
            safe_python_path = os.path.join(current_dir, 'listekolay.py').replace('\\', '\\\\')

            # Dosyayı normal string oluşturarak yazalım
            updater_content = """
import os
import sys
import time
import zipfile
import shutil
import subprocess

def update_app():
    # Orijinal uygulamanın kapanması için bekle
    time.sleep(2)

    try:
        # Zip dosyasını çıkart
        with zipfile.ZipFile(r"{0}", "r") as zip_ref:
            zip_ref.extractall(r"{1}")

        # Yeni sürümü başlat
        subprocess.Popen(["python", r"{2}"])

        return True
    except Exception as e:
        print(f"Update error: {{e}}")
        return False

if __name__ == "__main__":
    update_app()
"""
            # Format ile değerleri ekle
            formatted_content = updater_content.format(safe_zip_path, safe_current_dir, safe_python_path)

            # Dosyaya yaz
            with open(updater_script, "w", encoding="utf-8") as f:
                f.write(formatted_content)

            # Yardımcı betiği başlat
            subprocess.Popen([sys.executable, updater_script])

            # Mevcut uygulamayı kapat
            self.on_close()

        except Exception as e:
            # Güncelleme hatası
            logging.error(f"Update launch error: {e}")
            messagebox.showerror(
                self.get_text("download_error"),
                self.get_text("download_error_message").format(str(e))
            )
    def get_app_data_dir(self):
        """
        Uygulama verilerinin kaydedileceği dizini belirler.
        Belgelerim klasörü altında 'ListeKolay' klasörünü kullanır.
        """
        # Belgelerim klasörünü belirle (cross-platform desteği)
        documents_dir = os.path.join(os.path.expanduser('~'), 'Documents')

        # ListeKolay klasörü oluştur (yoksa)
        app_data_dir = os.path.join(documents_dir, 'ListeKolay')
        if not os.path.exists(app_data_dir):
            try:
                os.makedirs(app_data_dir)
                logging.info(f"ListeKolay veri klasörü oluşturuldu: {app_data_dir}")
            except Exception as e:
                logging.error(f"ListeKolay veri klasörü oluşturulamadı: {str(e)}")
                # Oluşturulamazsa geçici dizini kullan
                import tempfile
                app_data_dir = tempfile.gettempdir()
                logging.info(f"Alternatif olarak geçici dizin kullanılacak: {app_data_dir}")

        logging.info(f"Yapılandırma dosyaları şuraya kaydedilecek: {app_data_dir}")
        return app_data_dir


    def save_config(self):
        """Kullanıcı ayarlarını config.json dosyasına kaydet"""
        try:
            config = {
                "language": self.current_language,
                # "last_folder" değeri artık kaydedilmiyor
                "include_subfolders": self.include_subfolders.get(),
                "export_formats": {
                    "text": self.export_formats["text"].get(),
                    "excel": self.export_formats["excel"].get(),
                    "word": self.export_formats["word"].get(),
                    "html": self.export_formats["html"].get()
                },
                "save_to_desktop": self.save_to_desktop.get(),
                "sort_criteria": self.selected_sort.get(),
                # "view_mode" değeri artık kaydedilmiyor
                "is_dark_mode": self.is_dark_mode.get()
            }

            # Config dosyası için uygun dizini belirle (exe veya script modu)
            app_data_dir = self.get_app_data_dir()
            config_path = os.path.join(app_data_dir, "config.json")

            # Config dosyasının bir yedeğini oluştur (kaydetmeden önce)
            try:
                backup_path = os.path.join(app_data_dir, "config.json.bak")
                if os.path.exists(config_path):
                    shutil.copy2(config_path, backup_path)
                    logging.info(f"Config dosyası yedeklendi: {backup_path}")
            except Exception as backup_error:
                logging.error(f"Config dosyası yedeklenirken hata oluştu: {str(backup_error)}")

            # Dizine yazma erişimi kontrolü
            if not os.access(app_data_dir, os.W_OK):
                logging.warning(f"Dizine yazma erişimi yok: {app_data_dir}")
                # Son çare olarak geçici dosyalar dizinine yaz
                import tempfile
                temp_dir = tempfile.gettempdir()
                config_path = os.path.join(temp_dir, "config.json")
                logging.info(f"Alternatif olarak geçici dizine yazılıyor: {config_path}")

            # Asıl kayıt işlemi
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=4, ensure_ascii=False)

            # Oluşturulan config dosyasının konumunu logla
            logging.info(f"Config dosyası şuraya kaydedildi: {config_path}")

            logging.info("Ayarlar başarıyla kaydedildi")
        except Exception as e:
            logging.error(f"Ayarları kaydederken hata oluştu: {str(e)}")

            # 1. Yedekten geri yüklemeyi dene
            try:
                app_data_dir = self.get_app_data_dir()
                backup_path = os.path.join(app_data_dir, "config.json.bak")
                config_path = os.path.join(app_data_dir, "config.json")

                if os.path.exists(backup_path):
                    if os.path.exists(config_path):
                        os.remove(config_path)
                    shutil.copy2(backup_path, config_path)
                    logging.info("Config dosyası yedekten geri yüklendi")
                    return
            except Exception as restore_error:
                logging.error(f"Config yedekten geri yüklenirken hata oluştu: {str(restore_error)}")

            # 2. Yöntem: Dosya yazma hatası olursa, tekrar deneme yaparak veri kaybını önleyelim
            try:
                # Önce temp dosyaya yaz, sonra adını değiştir (daha güvenli yaklaşım)
                temp_path = os.path.join(app_data_dir, "config_temp.json")
                with open(temp_path, 'w', encoding='utf-8') as f:
                    json.dump(config, f, indent=4, ensure_ascii=False)

                # Başarıyla yazıldıysa, asıl dosyanın yerine koy
                if os.path.exists(temp_path):
                    if os.path.exists(config_path):
                        os.remove(config_path)
                    os.rename(temp_path, config_path)
                    logging.info("İkinci denemede ayarlar başarıyla kaydedildi")
            except Exception as e2:
                logging.error(f"Ayarları tekrar kaydederken ikinci hata oluştu: {str(e2)}")

    def load_config(self):
        """config.json dosyasından kullanıcı ayarlarını yükle"""
        # Yükleme sırasında gereksiz döngüsel çağrıları önlemek için flag ekle
        if hasattr(self, 'config_loading_in_progress') and self.config_loading_in_progress:
            logging.info("Config yükleme işlemi zaten devam ediyor, tekrarlayan çağrı engellendi")
            return

        # Config yükleme işlemini başlat
        self.config_loading_in_progress = True

        try:
            # Önce uygulama veri dizinini belirle 
            app_data_dir = self.get_app_data_dir()
            config_path = os.path.join(app_data_dir, "config.json")

            logging.info(f"Config dosyası aranıyor: {config_path}")

            # Config dosyası yoksa oluştur
            if not os.path.exists(config_path):
                logging.info("Yapılandırma dosyası bulunamadı, varsayılan ayarlarla oluşturuluyor")
                self.save_config()

                # Yeniden kontrol et
                if not os.path.exists(config_path):
                    logging.error("Config dosyası oluşturulamadı, varsayılan ayarlar yüklenecek")
                    return  # Alternatif konuma gitme, çünkü artık sadece app_data_dir kullanılmalı


            # Config dosyasını oku
            logging.info(f"Config dosyası okunuyor: {config_path}")
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                logging.info("Config dosyası başarıyla okundu")

            # ÖNEMLİ: Tema modunu önce ayarla, böylece diğer ayarlar tema üzerine uygulanır
            # Tema modu (açık/koyu)
            has_theme_setting = False
            if "is_dark_mode" in config:
                has_theme_setting = True
                self.is_dark_mode.set(config["is_dark_mode"])
                # Tema modunu hemen uygula
                self.toggle_theme_mode()
                logging.info(f"Tema modu yüklendi: {'Koyu' if config['is_dark_mode'] else 'Açık'}")

            # Dil ayarı - Temadan sonra yapılması önemli
            if "language" in config and config["language"] in self.languages:
                saved_language = config["language"]
                self.current_language = saved_language
                self.language_var.set(saved_language)
                # Dil değişikliğini hemen uygula
                logging.info(f"Config'den yüklenen dil: {saved_language}")
                # UI dil değişikliğini uygula
                self.update_ui_language()
                # Ana başlıkları güncelle
                self.update_main_titles()
                # Kategorileri güncelle
                # Dil değişikliğinin UI'ye uygulanması için kategorileri yenile
                self.populate_categories()

                # Dil değişikliğinden sonra tema değişikliği uygulandıysa
                # butonların doğru renklenmesi için ikinci kez tema uygula
                if has_theme_setting:
                    # Tema güncellemesi için koruyucu bayrak kontrolü
                    if not hasattr(self, 'theme_update_scheduled') or not self.theme_update_scheduled:
                        # Bayrak ayarla - mükerrer çağrıları önlemek için
                        self.theme_update_scheduled = True
                        # Kısa bir gecikme ekleyerek önce dil değişikliğinin uygulanmasını sağla
                        self.root.after(100, self._delayed_theme_update)

            # Son klasör artık config'den yüklenmiyor - açılışta boş kalacak
            # Kullanıcının klasör seçmesi bekleniyor
            self.folder_path_var.set("")
            if hasattr(self, 'selected_folder_path'):
                delattr(self, 'selected_folder_path')

            # İlk açılışta görünüm modunu 'list' olarak ayarla
            if hasattr(self, 'view_mode_var'):
                self.view_mode_var.set("list")
                self.set_view_mode("list")
                logging.info("İlk açılışta listeleme moduna geçildi")

            # Alt klasörler dahil
            if "include_subfolders" in config:
                self.include_subfolders.set(config["include_subfolders"])

            # Dışa aktarma formatları
            if "export_formats" in config:
                formats = config["export_formats"]
                for fmt in self.export_formats:
                    if fmt in formats:
                        self.export_formats[fmt].set(formats[fmt])

            # Masaüstüne kaydet
            if "save_to_desktop" in config:
                self.save_to_desktop.set(config["save_to_desktop"])

            # Sıralama kriteri
            if "sort_criteria" in config:
                self.selected_sort.set(config["sort_criteria"])

            # Görünüm modu her zaman list modunda başlasın
            # view_mode artık config'den yüklenmiyor
            self.view_mode_var.set("list")
            self.set_view_mode("list")
            logging.info("Program ilk açılışta liste görünümünde başlatıldı")

            logging.info("Ayarlar başarıyla yüklendi")

            # Otomatik güncelleme kontrolü (sessiz mod)
            # Bu işlemi bir thread'de çalıştıralım ki arayüz bloke olmasın
            try:
                update_thread = threading.Thread(target=self.check_for_updates, args=(True,))
                update_thread.daemon = True
                update_thread.start()
                logging.info("Otomatik güncelleme kontrolü başlatıldı")
            except Exception as e:
                logging.error(f"Otomatik güncelleme kontrolü başlatılamadı: {str(e)}")
        except Exception as e:
            logging.error(f"Ayarları yüklerken hata oluştu: {str(e)}")
        finally:
            # Config yükleme işlemini tamamla ve flag'i sıfırla
            self.config_loading_in_progress = False
            logging.info("Config yükleme işlemi tamamlandı")

    def toggle_left_panel(self):
        """Sol paneli aç/kapat"""
        if self.left_panel_visible.get():
            # Panel görünür, gizle
            self.left_column.pack_forget()
            self.toggle_left_panel_btn.config(text="▶")  # Sağ ok işareti (paneli göster)
            self.left_panel_visible.set(False)
        else:
            # Panel gizli, göster
            self.left_column.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10), before=self.left_column.master.winfo_children()[-1])
            self.toggle_left_panel_btn.config(text="◀")  # Sol ok işareti (paneli gizle)
            self.left_panel_visible.set(True)
            self.left_column.pack_propagate(False)  # Prevent shrinking

    def on_close(self):
        # Ask for confirmation before exiting
        response = messagebox.askyesno(
            self.get_text("confirm_exit_title"),
            self.get_text("confirm_exit_message"),
            icon=messagebox.QUESTION
        )

        if response:
            # Ayarları kaydet
            self.save_config()

            # Log application exit
            logging.info("Program sonlandırıldı")
            self.root.destroy()

    def create_file_preview_window(self, file_path, file_index=-1):
        """Create a window to preview the selected file"""
        try:
            # Update the current preview index if provided
            if file_index >= 0 and file_index < len(self.current_preview_files):
                self.current_preview_index = file_index
            elif file_index == -1 and hasattr(self, 'current_preview_files'):
                # Find the file index if not provided
                self.current_preview_index = next((i for i, f in enumerate(self.current_preview_files) 
                                               if f["path"] == file_path), -1)

            # Close existing preview window if one exists
            if hasattr(self, 'preview_window') and self.preview_window and self.preview_window.winfo_exists():
                self.preview_window.destroy()

            # Create a new top-level window
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title(self.get_text("preview_window_title"))
            self.preview_window.geometry("900x700")  # Slightly larger for more info
            self.preview_window.minsize(500, 400)

            # Get file details
            file_name = os.path.basename(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            # Get modification and creation dates
            file_mod_date = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S') if os.path.exists(file_path) else ""

            # Try to get creation date (platform specific)
            try:
                if os.name == 'nt':  # Windows
                    file_creation_date = datetime.datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
                else:  # Unix/Linux/Mac
                    # On Unix, getctime returns status change time, not creation time, since Unix doesn't track creation time
                    # Using stat to get the best approximation
                    stat_info = os.stat(file_path)
                    file_creation_date = datetime.datetime.fromtimestamp(stat_info.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            except:
                file_creation_date = ""

            # Create a container for the info panel at the top
            info_frame = tk.Frame(self.preview_window, bg="#f8f9fa", height=80, relief=tk.GROOVE, bd=1)
            info_frame.pack(fill=tk.X, padx=10, pady=(10, 0))
            info_frame.pack_propagate(False)  # Fixed height

            # Left side info (file name and basic details)
            left_info = tk.Frame(info_frame, bg="#f8f9fa")
            left_info.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=5)

            # File name with icon - use a larger font and make it more prominent
            file_name_frame = tk.Frame(left_info, bg="#f8f9fa")
            file_name_frame.pack(fill=tk.X, anchor=tk.W)

            # File type icon or text based on extension
            if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                icon_text = "IMG"
                icon_bg = "#28a745"  # Green for images
            elif file_extension in ['.pdf']:
                icon_text = "PDF"
                icon_bg = "#dc3545"  # Red for PDFs 
            elif file_extension in ['.psd', '.ai', '.eps']:
                icon_text = "PSD"
                icon_bg = "#6610f2"  # Purple for design files
            elif file_extension in ['.doc', '.docx', '.txt', '.rtf']:
                icon_text = "DOC"
                icon_bg = "#007bff"  # Blue for documents
            elif file_extension in ['.mp3', '.wav', '.flac', '.aac', '.ogg']:
                icon_text = "AUD"
                icon_bg = "#fd7e14"  # Orange for audio
            elif file_extension in ['.mp4', '.mov', '.mkv', '.avi']:
                icon_text = "VID"
                icon_bg = "#6f42c1"  # Purple for video
            else:
                # Get the file extension without dot and make it uppercase
                icon_text = file_extension.upper().replace(".", "") if file_extension else "FILE"
                icon_bg = "#6c757d"  # Gray for other types

            # Create icon label
            icon_label = tk.Label(
                file_name_frame,
                text=icon_text,
                font=("Segoe UI", 10, "bold"),
                bg=icon_bg,
                fg="white",
                width=4,
                padx=5,
                pady=2
            )
            icon_label.pack(side=tk.LEFT, padx=(0, 8))

            # File name label
            name_label = tk.Label(
                file_name_frame,
                text=file_name,
                font=("Segoe UI", 12, "bold"),
                bg="#f8f9fa",
                fg="#212529",
                anchor=tk.W
            )
            name_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

            # File details
            details_frame = tk.Frame(left_info, bg="#f8f9fa")
            details_frame.pack(fill=tk.X, anchor=tk.W, pady=(5, 0))

            # Create two rows for details for better organization
            details_row1 = tk.Frame(details_frame, bg="#f8f9fa")
            details_row1.pack(fill=tk.X, pady=(0, 3))

            details_row2 = tk.Frame(details_frame, bg="#f8f9fa")
            details_row2.pack(fill=tk.X)

            # Row 1: Size and Extension
            size_label = tk.Label(
                details_row1,
                text=f"{self.get_text('file_size')}: {self.format_file_size(file_size)}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W,
                width=25
            )
            size_label.pack(side=tk.LEFT, padx=(0, 10))

            ext_label = tk.Label(
                details_row1,
                text=f"{self.get_text('file_extension')}: {file_extension}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W
            )
            ext_label.pack(side=tk.LEFT, fill=tk.X)

            # Row 2: Creation and Modification dates
            created_label = tk.Label(
                details_row2,
                text=f"{self.get_text('creation_date')}: {file_creation_date}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W,
                width=25
            )
            created_label.pack(side=tk.LEFT, padx=(0, 10))

            modified_label = tk.Label(
                details_row2,
                text=f"{self.get_text('modification_date')}: {file_mod_date}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W
            )
            modified_label.pack(side=tk.LEFT, fill=tk.X)

            # Right side navigation info - only if we have a valid file list and index
            right_info = tk.Frame(info_frame, bg="#f8f9fa")
            right_info.pack(side=tk.RIGHT, fill=tk.Y, padx=10, pady=5)

            # Store navigation label reference
            nav_label = None

            # Show navigation info if we have a valid file list
            if hasattr(self, 'current_preview_files') and len(self.current_preview_files) > 1 and self.current_preview_index >= 0:
                nav_text = f"{self.current_preview_index + 1} / {len(self.current_preview_files)}"
                nav_label = tk.Label(
                    right_info,
                    text=nav_text,
                    font=("Segoe UI", 10),
                    bg="#f8f9fa",
                    fg="#495057"
                )
                nav_label.pack(side=tk.RIGHT)

            # Add a frame for the preview content
            preview_frame = tk.Frame(self.preview_window, bg="#e9ecef")
            preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            # Check file type and generate preview
            # Genişletilmiş görüntü formatları - SVG ve WebP eklendi
            if file_extension.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.webp', '.svg']:
                self.preview_image(preview_frame, file_path)
            elif file_extension.lower() in ['.pdf']:
                self.preview_pdf(preview_frame, file_path)
            elif file_extension.lower() in ['.psd', '.ai', '.eps']:
                self.preview_design_file(preview_frame, file_path)
            else:
                # Show unsupported message
                label = tk.Label(
                    preview_frame, 
                    text=self.get_text("preview_not_supported"),
                    font=("Segoe UI", 12),
                    bg="#e9ecef"
                )
                label.pack(pady=20)

            # Add button frame for action buttons at the bottom
            button_frame = tk.Frame(self.preview_window, bg="#e9ecef")
            button_frame.pack(pady=10, fill=tk.X)

            # Navigation buttons references
            prev_button = None
            next_button = None

            # Navigation buttons - only show if we have multiple files
            if hasattr(self, 'current_preview_files') and len(self.current_preview_files) > 1 and self.current_preview_index >= 0:
                # Previous file button
                prev_button = tk.Button(
                    button_frame, 
                    text="← " + self.get_text("prev_page"),
                    command=self._go_to_prev_file,
                    bg="#17a2b8",
                    fg="white",
                    relief=tk.GROOVE,
                    padx=10,
                    state=tk.NORMAL if self.current_preview_index > 0 else tk.DISABLED
                )
                prev_button.pack(side=tk.LEFT, padx=(10, 5))

                # Next file button
                next_button = tk.Button(
                    button_frame, 
                    text=self.get_text("next_page") + " →",
                    command=self._go_to_next_file,
                    bg="#17a2b8",
                    fg="white",
                    relief=tk.GROOVE,
                    padx=10,
                    state=tk.NORMAL if self.current_preview_index < len(self.current_preview_files) - 1 else tk.DISABLED
                )
                next_button.pack(side=tk.LEFT, padx=(0, 5))

            # Open file directly button
            open_file_button = tk.Button(
                button_frame, 
                text=self.get_text("open_file"),
                command=lambda: self.open_file(file_path),
                bg="#28a745",
                fg="white",
                relief=tk.GROOVE,
                padx=10
            )
            open_file_button.pack(side=tk.LEFT, padx=(10 if not hasattr(self, 'current_preview_files') or len(self.current_preview_files) <= 1 else 0, 5))

            # Open file location button
            open_location_button = tk.Button(
                button_frame, 
                text=self.get_text("open_file_location"),
                command=lambda: self.open_file_location_by_path(file_path),
                bg="#007bff",
                fg="white",
                relief=tk.GROOVE,
                padx=10
            )
            open_location_button.pack(side=tk.LEFT, padx=(0, 5))

            # Close button 
            close_button = tk.Button(
                button_frame, 
                text=self.get_text("exit"),
                command=self.preview_window.destroy,
                bg="#e9ecef",
                relief=tk.GROOVE,
                padx=10
            )
            close_button.pack(side=tk.RIGHT, padx=10)

            # Store widget references for later updates (when navigating between files)
            self._preview_info_widgets = {
                'icon_label': icon_label,
                'name_label': name_label,
                'size_label': size_label,
                'ext_label': ext_label,
                'created_label': created_label,
                'modified_label': modified_label,
                'open_file_button': open_file_button,
                'open_location_button': open_location_button
            }

            # Add navigation widgets if they exist
            if nav_label:
                self._preview_info_widgets['nav_label'] = nav_label
            if prev_button:
                self._preview_info_widgets['prev_button'] = prev_button
            if next_button:
                self._preview_info_widgets['next_button'] = next_button

            # Store reference to content frame for updates
            self._preview_content_frame = preview_frame

            # Center the window on the screen
            self.preview_window.update_idletasks()
            width = self.preview_window.winfo_width()
            height = self.preview_window.winfo_height()
            x = (self.preview_window.winfo_screenwidth() // 2) - (width // 2)
            y = (self.preview_window.winfo_screenheight() // 2) - (height // 2)
            self.preview_window.geometry('{}x{}+{}+{}'.format(width, height, x, y))

        except Exception as e:
            logging.error(f"Error creating preview window: {str(e)}")
            messagebox.showerror(
                self.get_text("error"), 
                f"{self.get_text('preview_error')} {str(e)}"
            )

    def preview_image(self, parent_frame, file_path):
        """
        Gelişmiş görüntü önizleme fonksiyonu.
        Standart resim formatlarının yanı sıra SVG, WebP ve TIF formatlarını destekler.
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            original_image = None

            # SVG dosyaları için özel işleme
            if file_ext == '.svg':
                try:
                    # CairoSVG ile SVG'yi PNG'ye dönüştür
                    import cairosvg
                    import io

                    # Geçici bir bellek akışına PNG olarak dönüştür
                    png_data = io.BytesIO()
                    cairosvg.svg2png(url=file_path, write_to=png_data)
                    png_data.seek(0)

                    # PNG'yi PIL görüntüsüne yükle
                    original_image = Image.open(png_data)

                    # Bellek temizleme işlemi
                    if not hasattr(self, 'temp_files'):
                        self.temp_files = []
                    self.temp_files.append(png_data)

                    logging.info(f"SVG önizleme başarıyla oluşturuldu: {file_path}")
                except Exception as svg_error:
                    logging.warning(f"{self.get_text('svg_conversion_error')}: {str(svg_error)}")
                    try:
                        # Alternatif: PIL ile SVG açmayı dene
                        original_image = Image.open(file_path)
                        logging.info(f"SVG alternatif olarak PIL ile açıldı: {file_path}")
                    except Exception:
                        logging.warning(f"SVG alternatif açılışı da başarısız, ikon gösteriliyor: {file_path}")
                        original_image = self._create_styled_icon(200, 200, "#6c757d", "SVG")

            # WebP dosyaları için özel işleme
            elif file_ext == '.webp':
                try:
                    # PIL artık WebP'yi destekliyor, doğrudan açmayı dene
                    original_image = Image.open(file_path)
                except Exception as webp_error:
                    logging.warning(f"WebP açılışı başarısız: {str(webp_error)}")
                    original_image = self._create_styled_icon(200, 200, "#20c997", "WebP")

            # TIFF/TIF dosyaları için özel işleme
            elif file_ext in ['.tif', '.tiff']:
                try:
                    # PIL TIFF'i destekler, ancak çoklu sayfa olabileceğini unutma
                    original_image = Image.open(file_path)
                    # Birden fazla sayfa varsa ilk sayfayı kullan
                    if hasattr(original_image, 'n_frames') and original_image.n_frames > 1:
                        original_image.seek(0)  # İlk çerçeveye dön
                        # Bir kopya oluştur çünkü çoklu sayfalı görüntüler sorun çıkarabilir
                        original_image = original_image.copy()
                except Exception as tiff_error:
                    logging.warning(f"TIFF açılışı başarısız: {str(tiff_error)}")
                    original_image = self._create_styled_icon(200, 200, "#fd7e14", "TIFF")

            # Normal resim dosyaları için
            if original_image is None:
                try:
                    # Önce standart PIL ile açmayı dene
                    original_image = Image.open(file_path)

                    # Özel format kontrolleri
                    if file_ext == '.psd':
                        # PSD için özellikle ilk katman alma
                        if hasattr(original_image, 'seek'):
                            original_image.seek(0)
                        # Bir kopya oluşturarak düz hale getir
                        original_image = original_image.copy()
                except Exception as img_error:
                    logging.warning(f"PIL ile resim açılamadı: {str(img_error)}")
                    # İkon ile göster
                    original_image = self._create_styled_icon(200, 200, "#17a2b8", file_ext.strip('.').upper())

            # Create a canvas for the image with scrollbars
            canvas_frame = tk.Frame(parent_frame)
            canvas_frame.pack(fill=tk.BOTH, expand=True)

            h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
            v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)

            canvas = tk.Canvas(
                canvas_frame,
                xscrollcommand=h_scrollbar.set,
                yscrollcommand=v_scrollbar.set,
                bg="#ffffff"
            )

            h_scrollbar.config(command=canvas.xview)
            v_scrollbar.config(command=canvas.yview)

            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            # Calculate scaled dimensions (max 700x500 initial view)
            max_width, max_height = 700, 500
            img_width, img_height = original_image.size

            # Calculate scaled size while maintaining aspect ratio
            scale_factor = min(max_width/img_width, max_height/img_height)
            if scale_factor < 1:  # Only scale down, not up
                new_width = int(img_width * scale_factor)
                new_height = int(img_height * scale_factor)
                # Uyumluluk için: LANCZOS yoksa ANTIALIAS kullan
                resize_method = getattr(Image, "LANCZOS", getattr(Image, "ANTIALIAS", Image.BICUBIC))
                display_image = original_image.resize((new_width, new_height), resize_method)
            else:
                display_image = original_image

            # Convert to PhotoImage for Tkinter
            photo = ImageTk.PhotoImage(display_image)

            # Add image to canvas
            canvas.create_image(0, 0, image=photo, anchor=tk.NW)
            canvas.image = photo  # Keep a reference to prevent garbage collection

            # Configure canvas scrollable area
            canvas.config(scrollregion=canvas.bbox(tk.ALL))

            # Add info label with image details
            info_text = f"{img_width}x{img_height} px, {os.path.basename(file_path)}"
            info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
            info_label.pack(pady=5)

        except Exception as e:
            logging.error(f"Error previewing image: {str(e)}")
            error_label = tk.Label(
                parent_frame, 
                text=f"{self.get_text('preview_error')} {str(e)}",
                fg="red", 
                bg="#e9ecef"
            )
            error_label.pack(pady=20)

    def preview_pdf(self, parent_frame, file_path):
        """Display a PDF preview"""
        try:
            # Create a canvas for the PDF with scrollbars
            canvas_frame = tk.Frame(parent_frame)
            canvas_frame.pack(fill=tk.BOTH, expand=True)

            h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
            v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)

            canvas = tk.Canvas(
                canvas_frame,
                xscrollcommand=h_scrollbar.set,
                yscrollcommand=v_scrollbar.set,
                bg="#ffffff"
            )

            h_scrollbar.config(command=canvas.xview)
            v_scrollbar.config(command=canvas.yview)

            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            try:
                # Open the PDF using PyMuPDF
                pdf_document = fitz.open(file_path)

                # Get first page of the PDF
                if pdf_document.page_count > 0:
                    first_page = pdf_document.load_page(0)

                    # Set zoom factor for better quality
                    zoom = 2.0
                    mat = fitz.Matrix(zoom, zoom)

                    # Convert page to an image
                    pix = first_page.get_pixmap(matrix=mat)

                    # Convert to PIL Image
                    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                else:
                    # Boş PDF için uyarı görüntüsü oluştur
                    img = self._create_styled_icon(300, 200, "#f5f5f5", self.get_text("preview_not_available"))
            except Exception as e:
                logging.error(f"PDF preview error: {str(e)}")
                # Hata durumunda uyarı görüntüsü oluştur
                img = self._create_styled_icon(300, 200, "#f5f5f5", self.get_text("preview_not_available"))

            # Scale down if needed
            max_width, max_height = 700, 500
            img_width, img_height = img.size

            # Calculate scaled size while maintaining aspect ratio
            scale_factor = min(max_width/img_width, max_height/img_height)
            if scale_factor < 1:  # Only scale down, not up
                new_width = int(img_width * scale_factor)
                new_height = int(img_height * scale_factor)
                img = img.resize((new_width, new_height), get_pil_resize_method())

            # Convert to PhotoImage for Tkinter
            photo = ImageTk.PhotoImage(img)

            # Add image to canvas
            canvas.create_image(0, 0, image=photo, anchor=tk.NW)
            canvas.image = photo  # Keep a reference to prevent garbage collection

            # Configure canvas scrollable area
            canvas.config(scrollregion=canvas.bbox(tk.ALL))

            # Add info label with PDF details
            info_text = f"PDF: {os.path.basename(file_path)}, {pdf_document.page_count} pages"
            info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
            info_label.pack(pady=5)

            # Close the document when done
            pdf_document.close()

        except Exception as e:
            logging.error(f"Error previewing PDF: {str(e)}")
            error_label = tk.Label(
                parent_frame, 
                text=f"{self.get_text('preview_error')} {str(e)}",
                fg="red", 
                bg="#e9ecef"
            )
            error_label.pack(pady=20)

    def preview_design_file(self, parent_frame, file_path):
        """Display a preview for design files (PSD, AI, EPS)"""
        try:
            # Normalize file path to avoid Windows/Unix path issues
            file_path = os.path.normpath(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Check file size to handle large files appropriately
            file_size = os.path.getsize(file_path)
            is_large_file = file_size > 20 * 1024 * 1024  # 20MB threshold
            
            # Show loading indicator for large files
            loading_frame = None
            if is_large_file:
                loading_frame = tk.Frame(parent_frame, bg="#e9ecef")
                loading_frame.pack(fill=tk.BOTH, expand=True)
                
                loading_label = tk.Label(
                    loading_frame, 
                    text=self.get_text("loading_large_file"),
                    font=("Segoe UI", 12),
                    bg="#e9ecef",
                    fg="#212529"
                )
                loading_label.pack(pady=(50, 10))
                
                size_text = self.format_file_size(file_size)
                file_info = f"{os.path.basename(file_path)} ({size_text})"
                info_label = tk.Label(
                    loading_frame, 
                    text=file_info,
                    font=("Segoe UI", 9),
                    bg="#e9ecef",
                    fg="#6c757d"
                )
                info_label.pack(pady=5)
                
                # Force update to show loading message
                parent_frame.update()
            
            # Handle PSD files - special processing for Photoshop files
            if file_extension == '.psd':
                try:
                    # Dosya boyutu büyükse daha fazla önlem al
                    if is_large_file:
                        # Hata durumunda göstermek için PSD ikonu hazırla
                        fallback_icon = self._create_styled_icon(400, 400, "#E91E63", "PSD")  # Photoshop Pembe
                        
                        # Bellekte daha etkili çalışmak için doğrudan Wand/ImageMagick kullan
                        try:
                            import wand.image
                            with wand.image.Image(filename=file_path, resolution=72) as img:
                                # Çözünürlük sınırla - büyük dosyaları küçük boyuta indirge
                                img.resize(width=600, height=600)
                                
                                # Optimize edilmiş dönüşüm
                                img.format = 'png'
                                img.compression_quality = 75  # Kaliteyi düşür ama hala makul
                                img_blob = img.make_blob()
                                
                                # Geçici dosya üzerinde çalış ve belleği hemen temizle
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                                temp_file.write(img_blob)
                                temp_file.close()
                                
                                # Belleği temizle
                                del img_blob
                                gc.collect()
                                
                                # Geçici dosyayı yükle
                                pil_img = Image.open(temp_file.name)
                                
                                # Önizlemeyi göster ve geçici dosyayı temizle
                                if loading_frame:
                                    loading_frame.destroy()
                                
                                self._display_design_preview(parent_frame, pil_img, file_path)
                                
                                # Geçici dosyayı sil
                                try:
                                    os.unlink(temp_file.name)
                                except:
                                    pass
                                
                                return
                        except ImportError:
                            logging.warning("Wand/ImageMagick not available for large PSD previews")
                        except Exception as wand_err:
                            logging.error(f"Failed to open large PSD with Wand: {str(wand_err)}")
                            
                            # Hata durumunda düşük kalite PIL dene
                            try:
                                psd_img = Image.open(file_path)
                                psd_img.thumbnail((400, 400), get_pil_resize_method())
                                if loading_frame:
                                    loading_frame.destroy()
                                self._display_design_preview(parent_frame, psd_img, file_path)
                                return
                            except:
                                # Son çare - fallback icon kullan
                                if loading_frame:
                                    loading_frame.destroy()
                                self._display_design_preview(parent_frame, fallback_icon, file_path)
                                return
                    
                    # Normal boyutlu dosyalar için standart işleme
                    else:
                        # Try PIL/Pillow first as it's memory efficient
                        try:
                            psd_img = Image.open(file_path)
                            # Orta boy önizleme
                            psd_img.thumbnail((800, 800), get_pil_resize_method())
                            
                            if loading_frame:
                                loading_frame.destroy()
                            self._display_design_preview(parent_frame, psd_img, file_path)
                            return
                        except Exception as pil_err:
                            logging.error(f"Failed to open PSD with PIL: {str(pil_err)}")
                        
                        # Fall back to Wand/ImageMagick for better PSD support if PIL fails
                        try:
                            import wand.image
                            with wand.image.Image(filename=file_path) as img:
                                img.resize(width=800, height=800)
                                img.format = 'png'
                                img_blob = img.make_blob()
                                pil_img = Image.open(io.BytesIO(img_blob))
                                
                                if loading_frame:
                                    loading_frame.destroy()
                                
                                self._display_design_preview(parent_frame, pil_img, file_path)
                                return
                        except ImportError:
                            logging.warning("Wand/ImageMagick not available for PSD previews")
                        except Exception as wand_err:
                            logging.error(f"Failed to open PSD with Wand: {str(wand_err)}")
                    
                except Exception as psd_error:
                    logging.error(f"All PSD preview methods failed: {str(psd_error)}")
                    
                # Son çare - dosya açılamazsa bir ikon göster
                if loading_frame:
                    loading_frame.destroy()
                fallback_icon = self._create_styled_icon(400, 400, "#E91E63", "PSD")  # Photoshop Pembe
                self._display_design_preview(parent_frame, fallback_icon, file_path)
            
            # Handle AI files (Adobe Illustrator)
            elif file_extension == '.ai':
                # AI files are often PDF-compatible, try PyMuPDF
                try:
                    pdf_document = fitz.open(file_path)
                    if pdf_document.page_count > 0:
                        first_page = pdf_document.load_page(0)
                        
                        # Adjust zoom based on file size for memory optimization
                        zoom = 1.0 if is_large_file else 2.0
                        mat = fitz.Matrix(zoom, zoom)
                        
                        pix = first_page.get_pixmap(matrix=mat)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        
                        pdf_document.close()
                        
                        if loading_frame:
                            loading_frame.destroy()
                            
                        self._display_design_preview(parent_frame, img, file_path)
                        return
                    pdf_document.close()
                except Exception as ai_error:
                    logging.error(f"PyMuPDF failed to open AI file: {str(ai_error)}")
                    
                    # Try ImageMagick as fallback for AI files
                    try:
                        import wand.image
                        with wand.image.Image(filename=file_path) as img:
                            # Resize for large files
                            if is_large_file:
                                img.resize(width=800, height=800)
                                
                            img.format = 'png'
                            img_blob = img.make_blob()
                            pil_img = Image.open(io.BytesIO(img_blob))
                            
                            if loading_frame:
                                loading_frame.destroy()
                                
                            self._display_design_preview(parent_frame, pil_img, file_path)
                            return
                    except ImportError:
                        logging.warning("Wand/ImageMagick not available for AI previews")
                    except Exception as wand_ai_err:
                        logging.error(f"Failed to open AI with Wand: {str(wand_ai_err)}")
            
            # Handle EPS files with our specialized function
            elif file_extension == '.eps':
                try:
                    # Use our dedicated EPS preview function with optimized memory handling
                    max_size = 600 if is_large_file else 800
                    preview_result = self._create_eps_preview(file_path, max_size, max_size)
                    
                    if loading_frame:
                        loading_frame.destroy()
                        
                    # Our EPS preview helper might return either an Image or PhotoImage
                    if isinstance(preview_result, Image.Image):
                        self._display_design_preview(parent_frame, preview_result, file_path)
                    elif isinstance(preview_result, ImageTk.PhotoImage):
                        # Create a special display for PhotoImage results
                        canvas_frame = tk.Frame(parent_frame)
                        canvas_frame.pack(fill=tk.BOTH, expand=True)
                        
                        canvas = tk.Canvas(
                            canvas_frame,
                            bg="#ffffff",
                            width=preview_result.width(),
                            height=preview_result.height()
                        )
                        canvas.pack(fill=tk.BOTH, expand=True)
                        canvas.create_image(0, 0, image=preview_result, anchor=tk.NW)
                        canvas.image = preview_result  # Keep reference
                        
                        # Add file info
                        info_text = f"EPS: {os.path.basename(file_path)}, {self.format_file_size(file_size)}"
                        info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
                        info_label.pack(pady=5)
                    
                    return
                except Exception as eps_error:
                    logging.error(f"EPS preview creation failed: {str(eps_error)}")
            
            # For PDF, try to use PDF-specific methods if this was actually a PDF file
            elif file_extension == '.pdf':
                try:
                    self.preview_pdf(parent_frame, file_path)
                    return
                except Exception as pdf_err:
                    logging.error(f"PDF preview failed, falling back: {str(pdf_err)}")
            
            # Try standard PIL as a fallback for any file
            try:
                img = Image.open(file_path)
                
                # For large files, reduce initial size
                if is_large_file:
                    img.thumbnail((800, 800), get_pil_resize_method())
                    
                if loading_frame:
                    loading_frame.destroy()
                    
                self._display_design_preview(parent_frame, img, file_path)
                return
            except Exception as pil_error:
                logging.error(f"PIL fallback could not open file: {str(pil_error)}")
                
                # For AI and EPS, try using PyMuPDF as another fallback
                if file_extension in ['.ai', '.eps']:
                    try:
                        pdf_document = fitz.open(file_path)
                        first_page = pdf_document.load_page(0)
                        
                        # Lower zoom for large files
                        zoom = 1.0 if is_large_file else 2.0
                        mat = fitz.Matrix(zoom, zoom)
                        pix = first_page.get_pixmap(matrix=mat)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        
                        pdf_document.close()
                        
                        if loading_frame:
                            loading_frame.destroy()
                            
                        self._display_design_preview(parent_frame, img, file_path)
                        return
                    except Exception:
                        pass

            # Clean up the loading frame if it exists
            if loading_frame:
                loading_frame.destroy()
                
            # If all else fails, create an enhanced placeholder with more info
            # Choose appropriate colors based on file type for better visual cues
            if file_extension == '.psd':
                bg_color = "#31A8FF"  # Photoshop blue
                text_color = "#FFFFFF"
                file_type_name = "PHOTOSHOP"
            elif file_extension == '.ai':
                bg_color = "#FF9A00"  # Illustrator orange
                text_color = "#330000"
                file_type_name = "ILLUSTRATOR"
            elif file_extension == '.eps':
                bg_color = "#8BC34A"  # Green for EPS
                text_color = "#FFFFFF"
                file_type_name = "EPS VECTOR"
            elif file_extension == '.pdf':
                bg_color = "#F40F02"  # Adobe PDF red
                text_color = "#FFFFFF"
                file_type_name = "PDF DOCUMENT"
            else:
                bg_color = "#f0f0f0"
                text_color = "#2c3e50"
                file_type_name = file_extension.upper().replace(".", "")
            
            # Create enhanced placeholder image with more info
            img = Image.new("RGB", (400, 300), color=bg_color)
            draw = ImageDraw.Draw(img)
            
            # Draw border
            draw.rectangle([10, 10, 390, 290], outline=text_color, width=2)
            
            # Display file type in center with better styling
            if file_extension.startswith('.'):
                file_type = file_extension[1:].upper()
            else:
                file_type = file_extension.upper()
            
            # Draw file type in the center
            draw.text((200, 120), file_type_name, fill=text_color, anchor="mm")
            
            # Draw file size
            size_text = self.format_file_size(file_size)
            draw.text((200, 150), size_text, fill=text_color, anchor="mm")
            
            # Draw file name at the bottom
            file_name = os.path.basename(file_path)
            if len(file_name) > 30:  # Truncate long file names
                file_name = file_name[:27] + "..."
            draw.text((200, 200), file_name, fill=text_color, anchor="mm")
            
            # Add "preview not available" message
            not_available_text = self.get_text("preview_not_available")
            draw.text((200, 240), not_available_text, fill=text_color, anchor="mm")
            
            self._display_design_preview(parent_frame, img, file_path)
            return

        except Exception as e:
            logging.error(f"Error previewing design file: {str(e)}")
            
            # Clean up any loading frame that might exist
            for child in parent_frame.winfo_children():
                if isinstance(child, tk.Frame) and child.winfo_class() == "Frame":
                    child.destroy()
                    
            error_label = tk.Label(
                parent_frame, 
                text=f"{self.get_text('preview_error')} {str(e)}",
                fg="red", 
                bg="#e9ecef"
            )
            error_label.pack(pady=20)

            # Still show file info if possible
            try:
                file_size = os.path.getsize(file_path)
                info_text = f"{os.path.basename(file_path)}, {self.format_file_size(file_size)}"
                info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
                info_label.pack(pady=5)
            except:
                pass

    def _display_design_preview(self, parent_frame, img, file_path):
        """Helper function to display design file previews"""
        # Create a canvas for the image with scrollbars
        canvas_frame = tk.Frame(parent_frame)
        canvas_frame.pack(fill=tk.BOTH, expand=True)

        h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
        v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)

        canvas = tk.Canvas(
            canvas_frame,
            xscrollcommand=h_scrollbar.set,
            yscrollcommand=v_scrollbar.set,
            bg="#ffffff"
        )

        h_scrollbar.config(command=canvas.xview)
        v_scrollbar.config(command=canvas.yview)

        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Scale the image if necessary
        max_width, max_height = 700, 500
        img_width, img_height = img.size

        # Calculate scaled size while maintaining aspect ratio
        scale_factor = min(max_width/img_width, max_height/img_height)
        if scale_factor < 1:  # Only scale down, not up
            new_width = int(img_width * scale_factor)
            new_height = int(img_height * scale_factor)
            img = img.resize((new_width, new_height), get_pil_resize_method())

        # Convert to PhotoImage for Tkinter
        photo = ImageTk.PhotoImage(img)

        # Add image to canvas
        canvas.create_image(0, 0, image=photo, anchor=tk.NW)
        canvas.image = photo  # Keep a reference to prevent garbage collection

        # Configure canvas scrollable area
        canvas.config(scrollregion=canvas.bbox(tk.ALL))

        # Add info label with image details
        file_size = os.path.getsize(file_path)
        info_text = f"{img_width}x{img_height} px, {os.path.basename(file_path)}, {self.format_file_size(file_size)}"
        info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
        info_label.pack(pady=5)
        
    def _display_eps_preview(self, parent_frame, photo_image, file_path):
        """
        Specialized helper function to display EPS file previews when we already have a PhotoImage
        This is needed because the EPS preview function may return a PhotoImage directly
        """
        # Create a canvas for the image with scrollbars
        canvas_frame = tk.Frame(parent_frame)
        canvas_frame.pack(fill=tk.BOTH, expand=True)

        h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
        v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)

        canvas = tk.Canvas(
            canvas_frame,
            xscrollcommand=h_scrollbar.set,
            yscrollcommand=v_scrollbar.set,
            bg="#ffffff"
        )

        h_scrollbar.config(command=canvas.xview)
        v_scrollbar.config(command=canvas.yview)

        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Add the pre-created PhotoImage to canvas
        canvas.create_image(0, 0, image=photo_image, anchor=tk.NW)
        canvas.image = photo_image  # Keep a reference to prevent garbage collection

        # Configure canvas scrollable area
        canvas.config(scrollregion=canvas.bbox(tk.ALL))

        # Add info label with file details
        file_size = os.path.getsize(file_path)
        
        # For PhotoImage we can get dimensions directly
        img_width = photo_image.width()
        img_height = photo_image.height()
        
        info_text = f"EPS: {img_width}x{img_height} px, {os.path.basename(file_path)}, {self.format_file_size(file_size)}"
        info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
        info_label.pack(pady=5)




    def copy_filename_to_clipboard(self):
        """Copy the selected file name to clipboard"""
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection

        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")

        if not values:
            return  # No values found

        # Extract file name
        file_name = values[0]

        # Copy to clipboard
        self.root.clipboard_clear()
        self.root.clipboard_append(file_name)

        # Show a brief status message with translation
        self.update_status(self.get_text("copied_to_clipboard"))

    def copy_filepath_to_clipboard(self):
        """Copy the selected file path to clipboard"""
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection

        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")

        if not values:
            return  # No values found

        # Extract file path
        file_path = values[2]

        # Copy to clipboard
        self.root.clipboard_clear()
        self.root.clipboard_append(file_path)

        # Show a brief status message
        self.update_status("Copied to clipboard")

    def on_drop(self, event):
        """Handle dropped files from external sources"""
        # Get the dropped file path(s)
        try:
            # Process the data - format depends on platform
            # In Windows, it will be in format "{C:/path/to/file}"
            # In Linux/Mac, it will be a normal file path
            data = event.data

            # Clean up the path
            if data.startswith('{') and data.endswith('}'):
                # Windows format
                data = data[1:-1]

            # Remove any quotes
            data = data.replace('"', '')

            paths = data.split()  # Multiple files are space-separated

            # Check if any of the paths are actually files
            valid_files = [path for path in paths if os.path.isfile(path)]

            if valid_files:
                # If files are dropped, we can open them directly (preview first file)
                self.create_file_preview_window(valid_files[0])

                # Show status message
                if len(valid_files) == 1:
                    self.update_status(f"Dosya önizleniyor: {os.path.basename(valid_files[0])}")
                else:
                    self.update_status(f"{len(valid_files)} dosya sürüklendi. İlk dosya önizleniyor.")

            # If a folder is dropped, update the folder selection
            valid_folders = [path for path in paths if os.path.isdir(path)]
            if valid_folders:
                # Use first valid folder
                self.folder_path_var.set(valid_folders[0])
                self.selected_folder_path = valid_folders[0]
                self.update_status(f"Klasör değiştirildi: {valid_folders[0]}")

                # Start folder loading
                self.load_files_thread()

        except Exception as e:
            logging.error(f"Error processing dropped files: {str(e)}")
            self.update_status(f"Sürüklenen dosyaları işlerken hata oluştu: {str(e)}")

    def open_file_location(self):
        """Open the location of the selected file in the file explorer"""
        # Get the selected item
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection

        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")

        if not values:
            return  # No values found

        # Extract file name, extension and directory path
        file_name = values[0]
        file_ext = values[1]
        dir_path = values[2]

        # Construct the full file path to find the correct directory
        # For Windows paths that already include filename, use as-is
        if os.path.basename(dir_path) == file_name:
            file_path = dir_path
        else:
            # Otherwise join directory and filename
            file_path = os.path.join(dir_path, file_name)

        # Open location using common method
        self.open_file_location_by_path(file_path)

    def open_file_location_by_path(self, file_path):
        """Open the location of a file based on its path"""
        # Check if the file exists
        if not os.path.exists(file_path):
            messagebox.showerror(
                self.get_text("error"),
                f"{file_path} not found."
            )
            return

        # Open the directory containing the file
        try:
            # Ensure we're using absolute path to avoid issues
            abs_file_path = os.path.abspath(file_path)

            if os.name == 'nt':  # Windows
                # Highlight the file in Windows Explorer
                subprocess.Popen(f'explorer /select,"{abs_file_path}"')
            elif sys.platform == 'darwin':  # macOS
                subprocess.Popen(['open', '-R', abs_file_path])
            else:  # Linux and other Unix variants
                # Get the directory containing the file - ensure it's absolute
                file_dir = os.path.dirname(abs_file_path)
                # Open the directory in the default file manager
                subprocess.Popen(['xdg-open', file_dir])

            # Log success
            logging.info(f"Opened location for file: {file_path}")
        except Exception as e:
            logging.error(f"Error opening file location: {str(e)}")
            messagebox.showerror(
                self.get_text("error"),
                f"{self.get_text('error_open_location')}: {str(e)}"
            )

    def open_file(self, file_path):
        """Open a file with the default associated application"""
        try:
            # Ensure we're using absolute path to avoid issues
            abs_file_path = os.path.abspath(file_path)

            if os.name == 'nt':  # Windows
                os.startfile(abs_file_path)
            elif sys.platform == 'darwin':  # macOS
                subprocess.Popen(['open', abs_file_path])
            else:  # Linux and other Unix variants
                subprocess.Popen(['xdg-open', abs_file_path])

            # Log success
            logging.info(f"Opened file: {file_path}")
        except Exception as e:
            logging.error(f"Error opening file: {str(e)}")
            messagebox.showerror(
                self.get_text("error"),
                f"{self.get_text('error_open_file')}: {str(e)}"
            )

    def open_website(self, url):
        """Open a website URL in the default browser"""
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception as e:
            logging.error(f"Error opening URL: {str(e)}")
            messagebox.showerror(
                self.get_text("error"),
                f"{self.get_text('error_open_url')}: {str(e)}"
            )

    def show_preview_context_menu(self, event, file_path):
        """Show context menu on right-click in the preview mode"""
        # Store the current file path for context menu actions
        self.current_preview_file_path = file_path

        # Show the context menu
        try:
            self.preview_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            # Make sure to release the grab
            self.preview_context_menu.grab_release()

    def open_preview_file(self):
        """Open the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.open_file(self.current_preview_file_path)

    def open_preview_file_location(self):
        """Open the location of the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.open_file_location_by_path(self.current_preview_file_path)

    def copy_preview_filename_to_clipboard(self):
        """Copy the filename of the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            file_name = os.path.basename(self.current_preview_file_path)
            self.root.clipboard_clear()
            self.root.clipboard_append(file_name)
            self.update_status(f"{self.get_text('copied_to_clipboard')}: {file_name}")

    def copy_preview_filepath_to_clipboard(self):
        """Copy the file path of the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.root.clipboard_clear()
            self.root.clipboard_append(self.current_preview_file_path)
            self.update_status(f"{self.get_text('copied_to_clipboard')}: {self.current_preview_file_path}")
            
    def preview_selected_preview_file(self):
        """Preview the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.create_file_preview_window(self.current_preview_file_path)
            
    def delete_preview_file(self):
        """Delete the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            file_path = self.current_preview_file_path
            file_name = os.path.basename(file_path)
            
            # Ask for confirmation
            if messagebox.askyesno(
                self.get_text("confirm_delete"),
                f"{self.get_text('do_you_want_to_delete')}: {file_name}?"
            ):
                try:
                    # Çöp kutusuna taşıma işlemi (send2trash kütüphanesi mevcut değilse doğrudan sil)
                    try:
                        import send2trash
                        send2trash.send2trash(file_path)
                    except ImportError:
                        os.remove(file_path)
                        
                    self.update_status(f"{self.get_text('file_deleted')}: {file_name}")
                    
                    # Gelişmiş dosya listesi yenileme ve önizleme modu güncellemesi
                    if self.view_mode_var.get() == "preview":
                        try:
                            # Çerçeveyi göster, kullanıcıya geri bildirim ver
                            wait_frame = tk.Frame(self.root, bg="#e9ecef")
                            wait_frame.place(relx=0.5, rely=0.5, anchor=tk.CENTER)
                            wait_label = tk.Label(
                                wait_frame, 
                                text=self.get_text('updating_preview'),
                                font=("Segoe UI", 12),
                                bg="#e9ecef", 
                                fg="#212529"
                            )
                            wait_label.pack(pady=10)
                            
                            # UI'yi hemen güncelle
                            self.root.update_idletasks()
                            
                            # Önce liste görünümüne geçerek bütün önizleme widget'larını temizle
                            self.set_view_mode("list")
                            self.root.update_idletasks()
                            
                            # Şimdi dosya listesini güvenli şekilde yenile
                            self.clear_file_list()
                            self.load_files()
                            
                            # Bekletme çerçevesini kaldır - zaten liste moduna geçtik
                            if wait_frame and wait_frame.winfo_exists():
                                wait_frame.destroy()
                            
                            # Kullanıcıya bilgi ver
                            messagebox.showinfo(
                                self.get_text("information"),
                                self.get_text("file_deleted") + ". " + 
                                self.get_text("view_changed_to_list")
                            )
                            
                        except Exception as e:
                            logging.error(f"Error refreshing preview after delete: {str(e)}")
                            # Hata durumunda silme işleminden sonra liste görünümüne geç
                            if wait_frame and wait_frame.winfo_exists():
                                try:
                                    wait_frame.destroy()
                                except:
                                    pass
                                    
                            messagebox.showinfo(
                                self.get_text("information"),
                                self.get_text("file_deleted")
                            )
                            # Önizleme modunda kal, sadece dosya listesini yenile
                            self.load_files()
                    else:
                        # Normal modda sadece dosya listesini yenile
                        self.load_files()
                except Exception as e:
                    messagebox.showerror(
                        self.get_text("error"),
                        f"{self.get_text('error_deleting_file')}: {str(e)}"
                    )
                    logging.error(f"Error deleting file: {str(e)}")
    
    def copy_preview_file(self):
        """Copy the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            file_path = self.current_preview_file_path
            
            # Ask for destination
            dest_dir = filedialog.askdirectory(
                title=self.get_text("select_destination_folder")
            )
            
            if dest_dir:
                try:
                    import shutil
                    file_name = os.path.basename(file_path)
                    dest_path = os.path.join(dest_dir, file_name)
                    
                    # Check if file already exists
                    if os.path.exists(dest_path):
                        if not messagebox.askyesno(
                            self.get_text("file_exists"),
                            f"{self.get_text('file_already_exists')}: {file_name}. {self.get_text('overwrite')}?"
                        ):
                            return
                    
                    # Copy file
                    shutil.copy2(file_path, dest_path)
                    self.update_status(f"{self.get_text('file_copied')}: {file_name} → {dest_dir}")
                except Exception as e:
                    messagebox.showerror(
                        self.get_text("error"),
                        f"{self.get_text('error_copying_file')}: {str(e)}"
                    )
                    logging.error(f"Error copying file: {str(e)}")
    
    def move_preview_file(self):
        """Move the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            file_path = self.current_preview_file_path
            
            # Ask for destination
            dest_dir = filedialog.askdirectory(
                title=self.get_text("select_destination_folder")
            )
            
            if dest_dir:
                try:
                    import shutil
                    file_name = os.path.basename(file_path)
                    dest_path = os.path.join(dest_dir, file_name)
                    
                    # Check if file already exists
                    if os.path.exists(dest_path):
                        if not messagebox.askyesno(
                            self.get_text("file_exists"),
                            f"{self.get_text('file_already_exists')}: {file_name}. {self.get_text('overwrite')}?"
                        ):
                            return
                    
                    # Move file
                    shutil.move(file_path, dest_path)
                    self.update_status(f"{self.get_text('file_moved')}: {file_name} → {dest_dir}")
                    
                    # Refresh file list
                    self.load_files()
                except Exception as e:
                    messagebox.showerror(
                        self.get_text("error"),
                        f"{self.get_text('error_moving_file')}: {str(e)}"
                    )
                    logging.error(f"Error moving file: {str(e)}")
    
    def select_all_preview_files(self):
        """Select all files in preview mode"""
        # This function selects all visible files in the preview panel
        if hasattr(self, 'current_preview_files') and self.current_preview_files:
            # If we have a treeview selection method for highlighting in preview mode
            # For now, we'll just show a message that all files are selected
            file_count = len(self.current_preview_files)
            self.update_status(f"{self.get_text('selected')}: {file_count} {self.get_text('files')}")
            
            # Bu noktada önizlemelerin tamamını seçili göstermek için 
            # bir görsel işaretleme eklenebilir (ör. çerçeve rengi değiştirme)
            messagebox.showinfo(
                self.get_text("information"),
                f"{file_count} {self.get_text('files')} {self.get_text('selected')}"
            )
    
    def rename_preview_file(self):
        """Rename the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            file_path = self.current_preview_file_path
            dir_path = os.path.dirname(file_path)
            file_name = os.path.basename(file_path)
            
            # Create a dialog to enter new name
            rename_dialog = tk.Toplevel(self.root)
            rename_dialog.title(self.get_text("rename_file"))
            rename_dialog.geometry("400x120")
            rename_dialog.resizable(False, False)
            rename_dialog.transient(self.root)
            rename_dialog.grab_set()
            
            # Apply theme
            rename_dialog.configure(bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
            
            # Create widgets
            tk.Label(
                rename_dialog, 
                text=self.get_text("current_name") + ":", 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
            ).grid(row=0, column=0, padx=10, pady=5, sticky="w")
            
            tk.Label(
                rename_dialog, 
                text=file_name, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
            ).grid(row=0, column=1, padx=10, pady=5, sticky="w")
            
            tk.Label(
                rename_dialog, 
                text=self.get_text("new_name") + ":", 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
            ).grid(row=1, column=0, padx=10, pady=5, sticky="w")
            
            new_name_entry = tk.Entry(
                rename_dialog, 
                width=30,
                bg=LIGHT_MODE_COLORS["entry_bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["entry_bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                insertbackground=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
            )
            new_name_entry.grid(row=1, column=1, padx=10, pady=5, sticky="ew")
            new_name_entry.insert(0, file_name)
            new_name_entry.select_range(0, len(file_name))
            new_name_entry.focus_set()
            
            # Buttons frame
            button_frame = tk.Frame(
                rename_dialog,
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
            )
            button_frame.grid(row=2, column=0, columnspan=2, pady=10)
            
            def on_rename():
                new_name = new_name_entry.get().strip()
                if not new_name:
                    messagebox.showwarning(
                        self.get_text("warning"),
                        self.get_text("filename_cannot_be_empty")
                    )
                    return
                
                if new_name == file_name:
                    rename_dialog.destroy()
                    return
                
                new_path = os.path.join(dir_path, new_name)
                if os.path.exists(new_path):
                    messagebox.showwarning(
                        self.get_text("warning"),
                        self.get_text("file_already_exists")
                    )
                    return
                
                try:
                    # os zaten modülün başında import edildi
                    os.rename(file_path, new_path)
                    self.update_status(f"{self.get_text('file_renamed')}: {file_name} → {new_name}")
                    
                    # Refresh file list
                    self.load_files()
                    rename_dialog.destroy()
                except Exception as e:
                    messagebox.showerror(
                        self.get_text("error"),
                        f"{self.get_text('error_renaming_file')}: {str(e)}"
                    )
                    logging.error(f"Error renaming file: {str(e)}")
            
            def on_cancel():
                rename_dialog.destroy()
            
            tk.Button(
                button_frame, 
                text=self.get_text("rename"),
                command=on_rename,
                bg=LIGHT_MODE_COLORS["btn_bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_bg"],
                fg=LIGHT_MODE_COLORS["btn_fg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_fg"],
                activebackground=LIGHT_MODE_COLORS["btn_active_bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_active_bg"],
                activeforeground=LIGHT_MODE_COLORS["btn_active_fg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_active_fg"]
            ).pack(side=tk.LEFT, padx=5)
            
            tk.Button(
                button_frame, 
                text=self.get_text("cancel"),
                command=on_cancel,
                bg=LIGHT_MODE_COLORS["btn_bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_bg"],
                fg=LIGHT_MODE_COLORS["btn_fg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_fg"],
                activebackground=LIGHT_MODE_COLORS["btn_active_bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_active_bg"],
                activeforeground=LIGHT_MODE_COLORS["btn_active_fg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["btn_active_fg"]
            ).pack(side=tk.LEFT, padx=5)
            
            # Bind Enter key to rename button
            new_name_entry.bind("<Return>", lambda event: on_rename())
            
            # Center the dialog
            rename_dialog.update_idletasks()
            width = rename_dialog.winfo_width()
            height = rename_dialog.winfo_height()
            x = (rename_dialog.winfo_screenwidth() // 2) - (width // 2)
            y = (rename_dialog.winfo_screenheight() // 2) - (height // 2)
            rename_dialog.geometry(f"{width}x{height}+{x}+{y}")
            
            # Wait for the dialog to close
            rename_dialog.wait_window()

    def get_selected_files_paths(self):
        """Seçili dosyaların tam yollarını döndür"""
        selected_items = self.file_tree.selection()
        if not selected_items:
            return []

        file_paths = []
        for item in selected_items:
            values = self.file_tree.item(item, "values")
            if values:
                file_name = values[0]  # İlk sütunda dosya adı var
                file_dir_path = values[2]  # Üçüncü sütunda dosya yolu var

                # Dosya yolu tam yolu içeriyor olabilir
                if os.path.basename(file_dir_path) == file_name:
                    file_path = file_dir_path
                else:
                    file_path = os.path.join(file_dir_path, file_name)

                file_paths.append(file_path)

        return file_paths

    def delete_selected_files(self):
        """Seçili dosyaları sil"""
        try:
            file_paths = self.get_selected_files_paths()
            if not file_paths:
                self.show_error(self.get_text("selection_error"), self.get_text("no_files_to_select"))
                return

            count = len(file_paths)
            # İlk uyarı
            if not messagebox.askyesno(
                self.get_text("delete_files"),
                f"{count} {self.get_text('files')} {self.get_text('confirm_delete_file')}?"
            ):
                return

            # İkinci, daha güçlü uyarı
            if not messagebox.askokcancel(
                self.get_text("warning"),
                f"{self.get_text('warning')}: {self.get_text('action_irreversible')}\n\n{count} {self.get_text('files')} {self.get_text('permanent_delete')}",
                icon=messagebox.WARNING
            ):
                return

            # İlerleme çubuğunu göster
            self.progress_var.set(0)
            self.progress_frame.pack(side=tk.BOTTOM, fill=tk.X)
            self.progress_bar.pack(fill=tk.X, padx=10, pady=5)
            self.update_status(self.get_text("deleting_files"))
            self.root.update()

            # Silme işlemi
            deleted_count = 0
            for i, file_path in enumerate(file_paths):
                try:
                    os.remove(file_path)
                    deleted_count += 1

                    # İlerleme çubuğunu güncelle
                    progress = int((i + 1) / count * 100)
                    self.progress_var.set(progress)
                    self.update_status(f"{self.get_text('deleting_files')} {i+1}/{count}")
                    self.root.update()

                except Exception as e:
                    logging.error(f"{self.get_text('file_delete_error')}: {file_path}, {self.get_text('error')}: {str(e)}")

            # İlerleme çubuğunu tamamla
            self.progress_var.set(100)
            self.root.update()

            # İşlem sonrası bilgilendirme mesajı
            messagebox.showinfo(
                self.get_text("operation_complete"),
                self.get_text("delete_complete").format(deleted_count)
            )

            # İlerleme çubuğunu gizle
            self.progress_bar.pack_forget()
            self.progress_frame.pack_forget()

            # Listeleri yenile
            self.update_status(f"{deleted_count} {self.get_text('files_deleted')}")
            self.load_files()
        except Exception as e:
            # Hata durumunda da ilerleme çubuğunu temizle
            if hasattr(self, 'progress_bar') and hasattr(self, 'progress_frame'):
                self.progress_bar.pack_forget()
                self.progress_frame.pack_forget()
            self.show_error(self.get_text("delete_error"), str(e))

    def copy_selected_files(self):
        """Seçili dosyaları kopyala"""
        try:
            file_paths = self.get_selected_files_paths()
            if not file_paths:
                self.show_error(self.get_text("selection_error"), self.get_text("no_files_to_select"))
                return

            count = len(file_paths)

            # İlk onay
            if not messagebox.askyesno(
                self.get_text("copy_files"),
                f"{count} {self.get_text('files')} {self.get_text('confirm_copy')}?"
            ):
                return

            # Hedef klasörü seç
            target_dir = filedialog.askdirectory(title=self.get_text("select_target_folder"))
            if not target_dir:
                return  # İptal edildi

            # İlerleme çubuğunu göster
            self.progress_var.set(0)
            self.progress_frame.pack(side=tk.BOTTOM, fill=tk.X)
            self.progress_bar.pack(fill=tk.X, padx=10, pady=5)
            self.update_status(self.get_text("copying_files"))
            self.root.update()

            # Kopyalama işlemi
            copied_count = 0

            try:
                for i, file_path in enumerate(file_paths):
                    try:
                        file_name = os.path.basename(file_path)
                        target_path = os.path.join(target_dir, file_name)
                        shutil.copy2(file_path, target_path)
                        copied_count += 1

                        # İlerleme çubuğunu güncelle
                        progress = int((i + 1) / count * 100)
                        self.progress_var.set(progress)
                        self.update_status(f"{self.get_text('copying_files')} {i+1}/{count}")
                        self.root.update()

                    except Exception as e:
                        logging.error(f"{self.get_text('file_copy_error')}: {file_path}, {self.get_text('error')}: {str(e)}")

                # İlerleme çubuğunu tamamla
                self.progress_var.set(100)
                self.root.update()

                # İşlem tamamlandı bildirimi
                messagebox.showinfo(
                    self.get_text("operation_complete"),
                    self.get_text("copy_complete").format(copied_count)
                )
            finally:
                # İlerleme çubuğunu gizle - her durumda çalışacak
                if hasattr(self, 'progress_bar'):
                    self.progress_bar.pack_forget()
                if hasattr(self, 'progress_frame'):
                    self.progress_frame.pack_forget()

            self.update_status(f"{copied_count} dosya kopyalandı")

        except Exception as e:
            self.show_error(self.get_text("copy_error"), str(e))

    def move_selected_files(self):
        """Seçili dosyaları taşı"""
        try:
            file_paths = self.get_selected_files_paths()
            if not file_paths:
                self.show_error(self.get_text("selection_error"), self.get_text("no_files_to_select"))
                return

            count = len(file_paths)

            # İlk onay
            if not messagebox.askyesno(
                self.get_text("move_files"),
                f"{count} {self.get_text('files')} {self.get_text('confirm_move')}?"
            ):
                return

            # Hedef klasörü seç
            target_dir = filedialog.askdirectory(title=self.get_text("select_target_folder"))
            if not target_dir:
                return  # İptal edildi

            # İkinci uyarı - bu geri alınamaz bir işlem
            if not messagebox.askokcancel(
                self.get_text("warning"),
                f"{self.get_text('attention')}: {self.get_text('action_irreversible')}\n\n{count} {self.get_text('files')} {self.get_text('confirm_move_files')}",
                icon=messagebox.WARNING
            ):
                return

            # İlerleme çubuğunu göster
            self.progress_var.set(0)
            self.progress_frame.pack(side=tk.BOTTOM, fill=tk.X)
            self.progress_bar.pack(fill=tk.X, padx=10, pady=5)
            self.update_status(self.get_text("moving_files"))
            self.root.update()

            # Taşıma işlemi
            moved_count = 0
            for i, file_path in enumerate(file_paths):
                try:
                    file_name = os.path.basename(file_path)
                    target_path = os.path.join(target_dir, file_name)
                    shutil.move(file_path, target_path)
                    moved_count += 1

                    # İlerleme çubuğunu güncelle
                    progress = int((i + 1) / count * 100)
                    self.progress_var.set(progress)
                    self.update_status(f"{self.get_text('moving_files')} {i+1}/{count}")
                    self.root.update()

                except Exception as e:
                    logging.error(f"{self.get_text('file_move_error')}: {file_path}, {self.get_text('error')}: {str(e)}")

            # İlerleme çubuğunu tamamla
            self.progress_var.set(100)
            self.root.update()

            # İşlem tamamlandı bildirimi
            messagebox.showinfo(
                self.get_text("operation_complete"),
                self.get_text("move_complete").format(moved_count)
            )

            # İlerleme çubuğunu gizle
            self.progress_bar.pack_forget()
            self.progress_frame.pack_forget()

            # Listeleri yenile
            self.update_status(f"{moved_count} dosya taşındı")
            self.load_files()
        except Exception as e:
            # Hata durumunda da ilerleme çubuğunu temizle
            if hasattr(self, 'progress_bar') and hasattr(self, 'progress_frame'):
                self.progress_bar.pack_forget()
                self.progress_frame.pack_forget()
            self.show_error(self.get_text("move_error"), str(e))

    def cut_selected_files(self):
        """Seçili dosyaları kes (taşıma işleminin başka bir adı)"""
        self.move_selected_files()

    def rename_selected_file(self):
        """Seçili dosyayı yeniden adlandır"""
        try:
            file_paths = self.get_selected_files_paths()
            if not file_paths:
                self.show_error(self.get_text("selection_error"), self.get_text("no_file_selected"))
                return

            if len(file_paths) > 1:
                self.show_error(self.get_text("rename_error"), self.get_text("select_only_one_file"))
                return

            file_path = file_paths[0]
            file_name = os.path.basename(file_path)
            folder_path = os.path.dirname(file_path)

            # Yeniden adlandırma penceresi
            rename_dialog = tk.Toplevel(self.root)
            rename_dialog.title(self.get_text("rename_file"))
            rename_dialog.geometry("500x150")
            rename_dialog.resizable(False, False)
            rename_dialog.transient(self.root)
            rename_dialog.grab_set()

            # Stil renkleri
            bg_color = LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
            fg_color = LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
            rename_dialog.configure(bg=bg_color)

            # Etiket
            tk.Label(
                rename_dialog,
                text=self.get_text("new_name") + ":",
                font=("Arial", 11),
                bg=bg_color,
                fg=fg_color
            ).pack(pady=(15, 5))

            # Giriş alanı
            entry_var = tk.StringVar(value=file_name)
            entry = tk.Entry(
                rename_dialog,
                textvariable=entry_var,
                width=50,
                font=("Arial", 11)
            )
            entry.pack(pady=5, padx=20)
            entry.select_range(0, len(file_name))
            entry.focus_set()

            # Düğme çerçevesi
            button_frame = tk.Frame(rename_dialog, bg=bg_color)
            button_frame.pack(pady=15, fill=tk.X)

            result = [None]

            def on_rename():
                result[0] = entry_var.get().strip()
                rename_dialog.destroy()

            def on_cancel():
                rename_dialog.destroy()

            # İptal ve Tamam düğmeleri
            tk.Button(
                button_frame,
                text=self.get_text("cancel"),
                command=on_cancel,
                width=10,
                font=("Arial", 10)
            ).pack(side=tk.RIGHT, padx=(5, 20))

            tk.Button(
                button_frame,
                text=self.get_text("ok"),
                command=on_rename,
                width=10,
                font=("Arial", 10)
            ).pack(side=tk.RIGHT, padx=5)

            rename_dialog.bind("<Return>", lambda event: on_rename())
            rename_dialog.bind("<Escape>", lambda event: on_cancel())

            self.root.wait_window(rename_dialog)

            new_name = result[0]
            if not new_name or new_name == file_name:
                return  # İptal edildi veya isim aynı

            new_path = os.path.join(folder_path, new_name)
            if os.path.exists(new_path):
                if not messagebox.askyesno(
                    self.get_text("warning"),
                    f"{self.get_text('file_exists')}?",
                    icon=messagebox.WARNING
                ):
                    return

            warning_message = f"{self.get_text('attention')}: {self.get_text('action_irreversible')}\n\n'{file_name}' → '{new_name}'\n\n{self.get_text('confirm_continue')}?"
            if not messagebox.askokcancel(
                self.get_text("warning"),
                warning_message,
                icon=messagebox.WARNING
            ):
                return

            # İlerleme başlat
            self.progress_var.set(0)
            self.progress_frame.pack(side=tk.BOTTOM, fill=tk.X)
            self.progress_bar.pack(fill=tk.X, padx=10, pady=5)
            self.update_status(f"{self.get_text('renaming_file')}...")
            self.root.update()

            # Yeniden adlandır
            os.rename(file_path, new_path)

            # İlerlemeyi tamamla
            self.progress_var.set(100)
            self.root.update()

            # Başarı mesajı
            success_message = f"{self.get_text('file_renamed_successfully')}:\n'{file_name}' → '{new_name}'"
            messagebox.showinfo(self.get_text("operation_complete"), success_message)

            # Arayüzü temizle
            self.progress_bar.pack_forget()
            self.progress_frame.pack_forget()
            self.update_status(f"{self.get_text('file_renamed')}: {file_name} → {new_name}")
            self.load_files()

        except Exception as e:
            self.show_error(self.get_text("rename_error"), str(e))




if __name__ == "__main__":
    root = tk.Tk()
    app = FileManagerApp(root)
    root.mainloop()
