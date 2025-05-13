"""
File list export functionality for ListeKolay application.
Provides functions to export file lists in various formats.
"""

import os
import datetime
import logging
import html
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_file_size_str(size_in_bytes):
    """
    Convert file size in bytes to a human-readable string.
    
    Args:
        size_in_bytes (int): File size in bytes
        
    Returns:
        str: Human-readable file size string
    """
    if size_in_bytes < 1024:
        return f"{size_in_bytes} B"
    elif size_in_bytes < 1024 * 1024:
        return f"{size_in_bytes / 1024:.1f} KB"
    elif size_in_bytes < 1024 * 1024 * 1024:
        return f"{size_in_bytes / (1024 * 1024):.1f} MB"
    else:
        return f"{size_in_bytes / (1024 * 1024 * 1024):.2f} GB"

def get_sort_criteria_text(sort_criteria, get_translation):
    """
    Get the human-readable text for the sort criteria.
    
    Args:
        sort_criteria (str): Sort criteria code (e.g., "name_asc")
        get_translation (function): Function to get translations
        
    Returns:
        str: Human-readable sort criteria text
    """
    criteria_map = {
        "name_asc": get_translation("sort_name_asc", "Dateiname - Aufsteigend"),
        "name_desc": get_translation("sort_name_desc", "Dateiname - Absteigend"),
        "ext_asc": get_translation("sort_ext_asc", "Dateierweiterung - Aufsteigend"),
        "ext_desc": get_translation("sort_ext_desc", "Dateierweiterung - Absteigend"),
        "size_asc": get_translation("sort_size_asc", "Dateigröße - Aufsteigend"),
        "size_desc": get_translation("sort_size_desc", "Dateigröße - Absteigend"),
        "dir_asc": get_translation("sort_dir_asc", "Dateiverzeichnis - Aufsteigend"),
    }
    
    return criteria_map.get(sort_criteria, sort_criteria)

def get_file_info(file_path):
    """
    Get detailed information about a file.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        dict: Dictionary containing file information
    """
    try:
        # Get basic file information
        file_name = os.path.basename(file_path)
        file_dir = os.path.dirname(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Get file size
        try:
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            file_size_str = get_file_size_str(file_size)
        except:
            file_size = 0
            file_size_str = "Unknown"
        
        # Get file timestamps
        try:
            created_time = os.path.getctime(file_path)
            modified_time = os.path.getmtime(file_path)
            created_time_str = datetime.datetime.fromtimestamp(created_time).strftime("%Y-%m-%d %H:%M:%S")
            modified_time_str = datetime.datetime.fromtimestamp(modified_time).strftime("%Y-%m-%d %H:%M:%S")
        except:
            created_time = 0
            modified_time = 0
            created_time_str = "Unknown"
            modified_time_str = "Unknown"
        
        # Return file information as a dictionary
        return {
            "name": file_name,
            "directory": file_dir,
            "extension": file_ext[1:] if file_ext else "",  # Remove leading dot
            "size_bytes": file_size,
            "size_str": file_size_str,
            "created_time": created_time,
            "created_time_str": created_time_str,
            "modified_time": modified_time,
            "modified_time_str": modified_time_str,
            "full_path": file_path
        }
    except Exception as e:
        logging.error(f"Error getting file info for {file_path}: {e}")
        # Return a minimal set of information if an error occurs
        return {
            "name": os.path.basename(file_path),
            "directory": os.path.dirname(file_path),
            "extension": os.path.splitext(file_path)[1].lower()[1:],
            "size_bytes": 0,
            "size_str": "Error",
            "created_time": 0,
            "created_time_str": "Unknown",
            "modified_time": 0,
            "modified_time_str": "Unknown",
            "full_path": file_path
        }

def export_to_excel(file_list, output_file, folder_path, sort_criteria, include_subfolders, get_translation):
    """
    Export the file list to an Excel file.
    
    Args:
        file_list (list): List of file paths
        output_file (str): Path to save the Excel file
        folder_path (str): Path of the source folder
        sort_criteria (str): Sort criteria code
        include_subfolders (bool): Whether subfolders were included
        get_translation (function): Function to get translations
        
    Returns:
        bool: True if export was successful, False otherwise
    """
    try:
        # Create a new workbook
        wb = Workbook()
        ws = wb.active
        ws.title = get_translation("file_list", "Dateiliste")
        
        # Set column widths
        ws.column_dimensions['A'].width = 10  # Row number
        ws.column_dimensions['B'].width = 40  # File name
        ws.column_dimensions['C'].width = 15  # File type
        ws.column_dimensions['D'].width = 60  # File path
        ws.column_dimensions['E'].width = 15  # File size
        ws.column_dimensions['F'].width = 20  # Creation date
        ws.column_dimensions['G'].width = 20  # Modification date
        ws.column_dimensions['H'].width = 15  # File extension
        
        # Define styles
        header_font = Font(name='Arial', bold=True, size=12, color='FFFFFF')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        # Create header row
        headers = [
            get_translation("row_number", "Zeilennummer"),
            get_translation("file_name", "Dateiname"),
            get_translation("file_type", "Dateityp"),
            get_translation("file_path", "Dateipfad"),
            get_translation("file_size", "Dateigröße"),
            get_translation("creation_date", "Erstellungsdatum"),
            get_translation("modification_date", "Änderungsdatum"),
            get_translation("file_extension", "Dateierweiterung")
        ]
        
        for col_idx, header_text in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header_text)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = header_border
            cell.alignment = header_alignment
        
        # Add file data
        for row_idx, file_path in enumerate(file_list, 1):
            # Get file information
            file_info = get_file_info(file_path)
            
            # Add data to worksheet
            row_data = [
                row_idx,  # Row number
                file_info["name"],  # File name
                file_info["extension"].upper(),  # File type
                file_info["full_path"],  # File path
                file_info["size_str"],  # File size
                file_info["created_time_str"],  # Creation date
                file_info["modified_time_str"],  # Modification date
                file_info["extension"]  # File extension
            ]
            
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx + 1, column=col_idx, value=value)
                # Apply alternate row shading
                if row_idx % 2 == 0:
                    cell.fill = PatternFill(start_color='E9EDF4', end_color='E9EDF4', fill_type='solid')
        
        # Add summary information
        summary_row = len(file_list) + 4
        ws.cell(row=summary_row, column=1, value=get_translation("creation_time", "Erstellungszeit:"))
        ws.cell(row=summary_row, column=2, value=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        
        ws.cell(row=summary_row + 1, column=1, value=get_translation("selected_folder", "Ausgewählter Ordner:"))
        ws.cell(row=summary_row + 1, column=2, value=folder_path)
        
        ws.cell(row=summary_row + 2, column=1, value=get_translation("sorted_by", "Sortiert nach:"))
        ws.cell(row=summary_row + 2, column=2, value=get_sort_criteria_text(sort_criteria, get_translation))
        
        ws.cell(row=summary_row + 3, column=1, value=get_translation("subfolders_label", "Alt Klasörler:"))
        ws.cell(row=summary_row + 3, column=2, value=get_translation("include_label", "Dahil Et") if include_subfolders else "")
        
        # Save the workbook
        wb.save(output_file)
        logging.info(f"Excel file created: {output_file}")
        
        return True
        
    except Exception as e:
        logging.error(f"Error exporting to Excel: {e}")
        return False

def export_to_word(file_list, output_file, folder_path, sort_criteria, include_subfolders, get_translation):
    """
    Export the file list to a Word document.
    
    Args:
        file_list (list): List of file paths
        output_file (str): Path to save the Word document
        folder_path (str): Path of the source folder
        sort_criteria (str): Sort criteria code
        include_subfolders (bool): Whether subfolders were included
        get_translation (function): Function to get translations
        
    Returns:
        bool: True if export was successful, False otherwise
    """
    try:
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(get_translation("app_title", "ListeKolay - Dosya Listesi Oluşturucu"), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle with folder path
        subtitle = doc.add_paragraph(folder_path)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(12)
        subtitle_format.italic = True
        
        # Add creation time and sort criteria
        info_text = f"{get_translation('creation_time', 'Erstellungszeit:')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        info_text += f"{get_translation('sorted_by', 'Sortiert nach:')} {get_sort_criteria_text(sort_criteria, get_translation)}\n"
        info_text += f"{get_translation('subfolders_label', 'Alt Klasörler:')} {'✓' if include_subfolders else '✗'}"
        
        info_para = doc.add_paragraph(info_text)
        info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add spacing
        doc.add_paragraph()
        
        # Create table header
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        headers = [
            get_translation("row_number", "Zeilennummer"),
            get_translation("file_name", "Dateiname"),
            get_translation("file_type", "Dateityp"),
            get_translation("file_size", "Dateigröße"),
            get_translation("modification_date", "Änderungsdatum")
        ]
        
        # Set header content and formatting
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            # Apply header formatting
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Add file data
        for row_idx, file_path in enumerate(file_list, 1):
            # Get file information
            file_info = get_file_info(file_path)
            
            # Add a row to the table
            row_cells = table.add_row().cells
            
            # Add data to cells
            row_cells[0].text = str(row_idx)  # Row number
            row_cells[1].text = file_info["name"]  # File name
            row_cells[2].text = file_info["extension"].upper()  # File type
            row_cells[3].text = file_info["size_str"]  # File size
            row_cells[4].text = file_info["modified_time_str"]  # Modification date
        
        # Adjust column widths (based on percentage of page width)
        table.autofit = False
        table.columns[0].width = Inches(0.7)  # Row number
        table.columns[1].width = Inches(3.0)  # File name
        table.columns[2].width = Inches(1.0)  # File type
        table.columns[3].width = Inches(1.0)  # File size
        table.columns[4].width = Inches(1.8)  # Modification date
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"ListeKolay © {datetime.datetime.now().year}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the document
        doc.save(output_file)
        logging.info(f"Word document created: {output_file}")
        
        return True
        
    except Exception as e:
        logging.error(f"Error exporting to Word: {e}")
        return False

def export_to_html(file_list, output_file, folder_path, sort_criteria, include_subfolders, get_translation):
    """
    Export the file list to an HTML file.
    
    Args:
        file_list (list): List of file paths
        output_file (str): Path to save the HTML file
        folder_path (str): Path of the source folder
        sort_criteria (str): Sort criteria code
        include_subfolders (bool): Whether subfolders were included
        get_translation (function): Function to get translations
        
    Returns:
        bool: True if export was successful, False otherwise
    """
    try:
        # Create HTML content
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(get_translation("app_title", "ListeKolay - Dosya Listesi Oluşturucu"))}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 20px;
            color: #333;
        }}
        h1, h2 {{
            text-align: center;
            color: #2c3e50;
        }}
        .info {{
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .info p {{
            margin: 5px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 20px;
        }}
        th, td {{
            padding: 12px 15px;
            border: 1px solid #ddd;
            text-align: left;
        }}
        th {{
            background-color: #4472C4;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        tr:hover {{
            background-color: #e9ecef;
        }}
        .footer {{
            text-align: center;
            margin-top: 30px;
            font-size: 14px;
            color: #6c757d;
        }}
        @media print {{
            .info {{
                background-color: #fff;
                border: 1px solid #ddd;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
        }}
    </style>
</head>
<body>
    <h1>{html.escape(get_translation("app_title", "ListeKolay - Dosya Listesi Oluşturucu"))}</h1>
    <h2>{html.escape(folder_path)}</h2>
    
    <div class="info">
        <p><strong>{html.escape(get_translation("creation_time", "Erstellungszeit:"))}</strong> {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><strong>{html.escape(get_translation("sorted_by", "Sortiert nach:"))}</strong> {html.escape(get_sort_criteria_text(sort_criteria, get_translation))}</p>
        <p><strong>{html.escape(get_translation("subfolders_label", "Alt Klasörler:"))}</strong> {'✓' if include_subfolders else '✗'}</p>
        <p><strong>{html.escape(get_translation("total_files_label", "Toplam Dosya:"))}</strong> {len(file_list)}</p>
    </div>
    
    <table>
        <thead>
            <tr>
                <th>{html.escape(get_translation("row_number", "Zeilennummer"))}</th>
                <th>{html.escape(get_translation("file_name", "Dateiname"))}</th>
                <th>{html.escape(get_translation("file_type", "Dateityp"))}</th>
                <th>{html.escape(get_translation("file_path", "Dateipfad"))}</th>
                <th>{html.escape(get_translation("file_size", "Dateigröße"))}</th>
                <th>{html.escape(get_translation("modification_date", "Änderungsdatum"))}</th>
            </tr>
        </thead>
        <tbody>
"""
        
        # Add file data rows
        for row_idx, file_path in enumerate(file_list, 1):
            # Get file information
            file_info = get_file_info(file_path)
            
            # Add a row for this file
            html_content += f"""            <tr>
                <td>{row_idx}</td>
                <td>{html.escape(file_info["name"])}</td>
                <td>{html.escape(file_info["extension"].upper())}</td>
                <td>{html.escape(file_info["full_path"])}</td>
                <td>{html.escape(file_info["size_str"])}</td>
                <td>{html.escape(file_info["modified_time_str"])}</td>
            </tr>
"""
        
        # Close the table and HTML document
        html_content += f"""        </tbody>
    </table>
    
    <div class="footer">
        ListeKolay © {datetime.datetime.now().year}
    </div>
</body>
</html>
"""
        
        # Write to file
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(html_content)
            
        logging.info(f"HTML file created: {output_file}")
        
        return True
        
    except Exception as e:
        logging.error(f"Error exporting to HTML: {e}")
        return False

def export_to_text(file_list, output_file, folder_path, sort_criteria, include_subfolders, get_translation):
    """
    Export the file list to a text file.
    
    Args:
        file_list (list): List of file paths
        output_file (str): Path to save the text file
        folder_path (str): Path of the source folder
        sort_criteria (str): Sort criteria code
        include_subfolders (bool): Whether subfolders were included
        get_translation (function): Function to get translations
        
    Returns:
        bool: True if export was successful, False otherwise
    """
    try:
        # Create formatted text content
        text_content = f"{get_translation('app_title', 'ListeKolay - Dosya Listesi Oluşturucu')}\n"
        text_content += f"{'-' * 80}\n\n"
        
        text_content += f"{get_translation('selected_folder', 'Ausgewählter Ordner:')} {folder_path}\n"
        text_content += f"{get_translation('creation_time', 'Erstellungszeit:')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        text_content += f"{get_translation('sorted_by', 'Sortiert nach:')} {get_sort_criteria_text(sort_criteria, get_translation)}\n"
        text_content += f"{get_translation('subfolders_label', 'Alt Klasörler:')} {'✓' if include_subfolders else '✗'}\n"
        text_content += f"{get_translation('total_files_label', 'Toplam Dosya:')} {len(file_list)}\n\n"
        
        text_content += f"{'-' * 80}\n"
        
        # Create header row with fixed-width columns
        headers = [
            get_translation("row_number", "Zeilennummer").ljust(6),
            get_translation("file_name", "Dateiname").ljust(40),
            get_translation("file_type", "Dateityp").ljust(12),
            get_translation("file_size", "Dateigröße").ljust(15),
            get_translation("modification_date", "Änderungsdatum")
        ]
        
        text_content += " | ".join(headers) + "\n"
        text_content += f"{'-' * 80}\n"
        
        # Add file data rows
        for row_idx, file_path in enumerate(file_list, 1):
            # Get file information
            file_info = get_file_info(file_path)
            
            # Format row data with fixed-width columns
            row_data = [
                str(row_idx).ljust(6),
                file_info["name"][:38].ljust(40),  # Truncate long filenames
                file_info["extension"].upper().ljust(12),
                file_info["size_str"].ljust(15),
                file_info["modified_time_str"]
            ]
            
            text_content += " | ".join(row_data) + "\n"
        
        text_content += f"{'-' * 80}\n\n"
        
        # Add footer
        text_content += f"ListeKolay © {datetime.datetime.now().year}\n"
        
        # Write to file
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(text_content)
            
        logging.info(f"Text file created: {output_file}")
        
        return True
        
    except Exception as e:
        logging.error(f"Error exporting to text file: {e}")
        return False
