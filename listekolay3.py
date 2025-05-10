import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import datetime
import subprocess
import logging
import json
from openpyxl import Workbook
from docx import Document
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
import threading
import sys
import webbrowser
import tempfile
import collections
import requests
import shutil
import random
import concurrent.futures  # For parallel processing
import time  # For performance measurement
from search_translations import search_translations

# Basit sürükle-bırak desteği için sabit
DND_FILES = "DND_FILES"

# Tema renk sabitleri
# Açık Tema Renkleri
LIGHT_MODE_COLORS = {
    "bg": "#e9ecef",             # Açık gri arkaplan
    "text": "#000000",           # Tüm metin ve etiketler için siyah
    "secondary_text": "#000000", # İkincil metinler de siyah
    "accent": "#007bff",         # Mavi vurgu
    
    # Buton Renkleri (Açık mod)
    "folder_button": "#007bff",  # Klasör seç butonu: Mavi
    "exit_button": "#6c757d",    # Kapat butonu: Gri
    "cancel_button": "#dc3545",  # İptal butonu: Kırmızı  
    "start_button": "#28a745",   # Başlat butonu: Yeşil
    "filter_button": "#17a2b8",  # Filtrele butonu: Turkuaz
    
    # View mode butonları
    "active_view_button": "#17a2b8",   # Aktif görünüm butonu: Turkuaz
    "inactive_view_button": "#6c757d", # Pasif görünüm butonu: Koyu gri
    
    # Ortak renkler
    "button_text": "#ffffff",     # Açık temada buton metinleri beyaz
    "highlight": "#f8f9fa",       # Çok açık gri vurgu
    "border": "#ced4da",          # Açık gri kenarlık
    "error": "#dc3545",           # Kırmızı hata
    "success": "#28a745",         # Yeşil başarı
    "warning": "#ffc107"          # Sarı uyarı
}

# Koyu Tema Renkleri
DARK_MODE_COLORS = {
    "bg": "#212529",             # Koyu arkaplan
    "text": "#ffffff",           # Tüm metin ve etiketler için beyaz
    "secondary_text": "#ffffff", # İkincil metinler de beyaz
    "accent": "#0d6efd",         # Parlak mavi vurgu
    
    # Buton Renkleri (Koyu mod)
    "folder_button": "#007bff",  # Klasör seç butonu: Mavi
    "exit_button": "#6c757d",    # Kapat butonu: Gri
    "cancel_button": "#dc3545",  # İptal butonu: Kırmızı
    "start_button": "#28a745",   # Başlat butonu: Yeşil
    "filter_button": "#17a2b8",  # Filtrele butonu: Turkuaz
    
    # View mode butonları
    "active_view_button": "#17a2b8",   # Aktif görünüm butonu: Turkuaz
    "inactive_view_button": "#6c757d", # Pasif görünüm butonu: Koyu gri
    
    # Ortak renkler
    "button_text": "#ffffff",     # Koyu temada buton metinleri beyaz
    "highlight": "#2b3035",       # Hafif açık koyu gri vurgu
    "border": "#495057",          # Orta koyu gri kenarlık
    "error": "#dc3545",           # Kırmızı hata
    "success": "#28a745",         # Yeşil başarı
    "warning": "#ffc107"          # Sarı uyarı
}

# PIL konfigürasyonu
import warnings
from PIL import Image, ImageTk, ImageDraw
# Devre dışı bırak DecompressionBombWarning (EPS ve büyük resimler için)
warnings.simplefilter('ignore', Image.DecompressionBombWarning)
# PIL maksimum boyut limitini artır
Image.MAX_IMAGE_PIXELS = None

# Ön izlenebilir dosya uzantıları (küçük harflerle)
PREVIEWABLE_EXTENSIONS = ['.jpg', '.jpeg', '.png', '.gif', '.pdf', '.eps', '.ai', '.psd', '.tif', '.tiff', '.bmp', '.ico']

# PIL sürüm uyumluluğu için yardımcı fonksiyon
def get_pil_resize_method():
    """Farklı PIL sürümleri için tutarlı yeniden boyutlandırma yöntemi döndürür"""
    try:
        return Image.Resampling.LANCZOS  # PIL 9.0 ve sonrası
    except (AttributeError, TypeError):
        try:
            return Image.LANCZOS  # PIL 4.0 - 8.x
        except (AttributeError, TypeError):
            return Image.ANTIALIAS  # Eski PIL sürümleri

import fitz  # PyMuPDF
import io
import time
import gc
import pdf2image

# LOG AYARLARI
logging.basicConfig(
    filename='ListeKolay.log',
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logging.info("Program başladı")

# Import language dictionaries
try:
    from new_languages import de_dict, fr_dict, ru_dict, es_dict, it_dict, fa_dict, ur_dict, hi_dict, zh_dict, ja_dict
except ImportError:
    # If module is not in the path, try relative import from current directory
    from attached_assets.new_languages import de_dict, fr_dict, ru_dict, es_dict, it_dict, fa_dict, ur_dict, hi_dict, zh_dict, ja_dict

translations = {
    "tr": {
        "open_file": "Dosyayı Aç",
        "open_file_location": "Dosya Konumunu Aç",
        "copy_filename": "Dosya Adını Kopyala",
        "copy_filepath": "Dosya Yolunu Kopyala",
        "select_folder": "📁 Klasör Seç",
        "no_folder_selected": "Henüz bir klasör seçilmedi",
        "start": "▶️ Başlat",
        "apply_filter": "🔍 Filtreyi Uygula",
        "cancel": "⏹️ İptal",
        "cancelling": "⏹️ İptal Ediliyor...",
        "exit": "✖️ Kapat",
        "select_all": "Tümünü Seç",
        "clear_all": "Temizle",
        "all_files_tip": "Tüm dosyalar uzantılarına bakılmaksızın listelenecek.",
        "filter_tip": "Filtreleme seçenekleri etkin, sadece seçilen uzantılara sahip dosyalar gösterilecek.",
        "calculating_statistics": "İstatistikler hesaplanıyor...",
        "loading_file_list": "Dosya listesi yükleniyor...",
        "tooltip_select": "Dosya listesi oluşturmak için bir klasör seçin",
        "tooltip_start": "Dosya listesini oluştur ve dışa aktar",
        "tooltip_apply": "Seçilen uzantılara göre dosyaları filtrele",
        "tooltip_cancel": "Devam eden işlemi iptal et",
        "tooltip_exit": "Uygulamayı kapat",
        "language": "Dil / Language",
        "view_mode_list": "Listele",
        "view_mode_preview": "Ön İzleme",
        "tooltip_list_view": "Dosyaları liste görünümünde göster",
        "tooltip_preview_view": "Dosyaları önizleme görünümünde göster",
        "files_and_previews": "Dosyalar ve Önizlemeler",
        "search_files": "Dosya ara...",
        "preview_mode_active": "Önizleme modu etkin",
        "loading_preview": "Önizlemeler yükleniyor...",
        "no_preview_available": "Önizlenebilir dosya bulunamadı",
        "preview_file": "Dosyayı Önizle",
        "open_file": "Dosyayı Aç",
        "open_file_location": "Dosya Konumunu Aç",
        "error_open_file": "Dosya açılırken bir hata oluştu",
        "error_open_location": "Dosya konumu açılırken bir hata oluştu",
        "error_open_url": "URL açılırken bir hata oluştu",
        "extension_not_found": "Uzantısı bulunamadı",
        "settings": "Ayarlar",
        "info": "Bilgi",
        "files_filtering": "Dosyalar filtreleniyor, lütfen bekleyin...",
        "files_gathering": "Dosyalar toplanıyor, lütfen bekleyin...",
        "files_loading": "Dosyalar Yükleniyor...",
        "filter_applying": "Filtre uygulanıyor...",
        "filter_cancelled": "Filtreleme işlemi iptal edildi",
        "excel_created": "Excel dosyası başarıyla oluşturuldu...",
        "error_occurred": "Bir hata oluştu: {0}",
        "error": "Hata",
        "ready": "Hazır",
        "operation_cancelled": "İşlem İptal Edildi",
        "file_processed": "İşlenen dosya: {0}/{1} • {2}",
        "prev_page": "Önceki",
        "next_page": "Sonraki",
        "page": "Sayfa",
        "processing": "İşleniyor: %{0:.1f}",
        "folder_loading": "Klasör yükleniyor...",
        "loading_subfolders": "Alt klasörler yükleniyor...",
        "folder_loaded_status": "📁 Klasör içeriği yüklendi. Liste oluşturmak için \"Başlat\" butonuna tıklayın.",
        "create_list_time": "🕒 Liste Oluşturulma Zamanı: {0}",
        "select_folder_first": "Lütfen önce bir klasör seçin!",
        "text_file_error": "Metin dosyası oluşturulamadı: {0}",
        "start_processing": "▶️ Liste oluşturuluyor...",
        "confirm_exit_title": "Çıkış",
        "confirm_exit_message": "Programdan çıkmak istediğinize emin misiniz?",
        "no_files_found": "Seçilen klasörde dosya bulunamadı.",
        "files_loaded_message": "{0} dosya yüklendi.",
        "files_filtered_message": "{0} dosya filtrelendi.",
        "filter_saved_message": "Filtreleme ayarları kaydedildi. Bir klasör seçtiğinizde uygulanacak.",
        "app_title": "ListeKolay - Dosya Listesi Oluşturucu",
        "app_subtitle": "Klasörlerinizdeki dosyaları hızlıca listeyin.",
        "full_window_title": "ListeKolay - Dosya Listesi Oluşturucu",
        "statistics_header": "İstatistikler",
        "total_files_label": "Toplam Dosya:",
        "folder_count_label": "Klasör Sayısı:",
        "total_size_label": "Toplam Boyut:",
        "tips_header": "İpuçları",
        "settings_header": "Ayarlar",
        "operation_status": "İşlem Durumu",
        "subfolders_label": "Alt Klasörler:",
        "include_label": "Dahil Et",
        "list_format_label": "Liste Formatı:",
        "save_location_label": "Kaydetme Yeri:",
        "desktop_label": "Masaüstü",
        "sort_criteria_label": "Sıralama Ölçütü:",
        "text_format_info": "Metin dosyası formatı",
        "excel_format_info": "Excel çalışma kitabı formatı",
        "word_format_info": "Word belgesi formatı",
        "html_format_info": "Web sayfası formatı",
        "filter_label": "Filtrele",
        "kategori_header": "Kategoriler",
        "extensions_header": "Dosya Uzantıları",
        "tip_1": "Alt klasör kutusunu işaretleyerek tüm alt dizinlerdeki dosyaları tarayabilirsiniz.",
        "tip_3": "Dosyaları isim, boyut veya uzantıya göre sıralayabilirsiniz.",
        "tip_4": "Oluşturulan listeler varsayılan olarak program klasörüne kaydedilir.",
        "tip_5": "Masaüstü seçeneğiyle dosyaları doğrudan masaüstüne kaydedebilirsiniz.",
        "tip_6": "İşlemi durdurmak için İptal düğmesini kullanın.",
        "tip_preview_formats": "PDF, JPG, PNG, GIF, PSD, AI, EPS dosyaları ön izleme özelliğine sahiptir.",
        "preview_file": "Dosyayı Önizle",
        "preview_window_title": "Dosya Önizleme",
        "preview_not_supported": "Bu dosya türü için önizleme desteklenmiyor.",
        "preview_error": "Dosya önizleme sırasında bir hata oluştu.",
        "file_list_section": "Dosyalar ve Önizlemeler",
        "list_view": "Listele",
        "preview_view": "Ön İzleme",
        "view_mode": "Görünüm Modu:",
        "preview_mode_active": "Ön izleme modu aktif",
        "no_preview_available": "Önizleme Yok",
        "loading_preview": "Önizleme yükleniyor...",
        "all_files": "Tüm Dosyalar",
        "image_files": "Görsel",
        "audio_files": "Ses",
        "video_files": "Video ve Ses",
        "text_files": "Metin",
        "compressed_files": "Sıkıştırılmış",
        "spreadsheet_files": "Hesap Tablosu",
        "presentation_files": "Sunum",
        "design_files": "Tasarım",
        "sort_name_asc": "Dosya Adı - Artan",
        "sort_name_desc": "Dosya Adı - Azalan",
        "sort_ext_asc": "Dosya Uzantısı - Artan",
        "sort_ext_desc": "Dosya Uzantısı - Azalan",
        "sort_size_asc": "Dosya Boyutu - Artan",
        "sort_size_desc": "Dosya Boyutu - Azalan",
        "sort_dir_asc": "Dosya Dizini - Artan",
        "selected_folder": "Seçilen Klasör:",
        "file_list": "Dosya Listesi",
        "row_number": "Sıra No",
        "file_name": "Dosya Adı",
        "file_type": "Dosya Türü",
        "file_path": "Dosya Yolu",
        "file_size": "Dosya Boyutu",
        "creation_date": "Oluşturulma Tarihi",
        "modification_date": "Değiştirilme Tarihi",
        "file_extension": "Dosya Uzantısı",
        "creation_time": "Oluşturulma Zamanı:",
        "sorted_by": "Sıralama Kriteri:",
        "excel_success": "Excel dosyası başarıyla oluşturuldu",
        "word_success": "Word belgesi başarıyla oluşturuldu",
        "html_success": "HTML dosyası başarıyla oluşturuldu",
        "text_success": "Metin dosyası başarıyla oluşturuldu",
        "open_file_title": "Dosyayı Aç",
        "open_file_message": "Dosyayı açmak ister misiniz?",
        "yes": "Evet",
        "no": "Hayır",
        "tooltip_subfolders": "Alt klasörlerdeki dosyaları da listeler",
        "tooltip_format": "Listenizin hangi formatta kaydedileceğini seçin",
        "tooltip_save_location": "Listenizin nereye kaydedileceğini seçin",
        "tooltip_sort_criteria": "Dosyaların hangi kritere göre sıralanacağını seçin",
        "tooltip_select_all": "Tüm dosya uzantılarını seçer",
        "tooltip_clear_all": "Tüm dosya uzantı seçimlerini temizler",
        "tooltip_file_category": "Dosya kategorilerini görüntülemek için tıklayın",
        "tooltip_file_extension": "Listelenecek dosya uzantılarını seçin",
        "tooltip_filter_apply": "Seçilen filtreleri uygulamak için tıklayın",
        "tooltip_all_files": "Tüm dosya uzantılarını seçer/temizler",
        "tooltip_category_expand": "Bu kategoriyi genişletmek/daraltmak için tıklayın",
        "tooltip_select_category": "Bu kategorideki tüm uzantıları seçer/temizler",
        "select_all_category": "Bu kategorideki tümünü seç",
        "media_files": "Medya Dosyaları (Ses ve Video)",
        "extension_search": "Ara-Bul",
        "copyright_footer": "© {year} Muallimun.Net - ListeKolay",
        "document_files": "Doküman",
        "code_files": "Kodlama",
        "program_files": "Program",
        "filtering_in_progress": "Filtreleme işlemi devam ediyor",
        "filter_complete": "Filtreleme tamamlandı",
        "filter_error": "Filtreleme hatası",
        "filter_error_details": "Filtreleme sırasında hata oluştu",
        "check_updates": "Güncellemeleri Kontrol Et",
        "update_available": "Güncelleme Mevcut",
        "update_available_message": "ListeKolay'ın yeni sürümü mevcut: {0}\nMevcut sürümünüz: {1}\n\nGüncellemeyi indirmek ister misiniz?",
        "no_update_available": "Güncelleme Yok",
        "no_update_available_message": "ListeKolay'ın en son sürümünü kullanıyorsunuz.",
        "update_check_error": "Güncelleme Kontrolü Hatası",
        "update_check_error_message": "Güncellemeler kontrol edilirken bir hata oluştu. Lütfen internet bağlantınızı kontrol edin ve tekrar deneyin.",
        "downloading_update": "Güncelleme İndiriliyor...",
        "downloading_update_message": "{0} sürümü indiriliyor...",
        "download_complete": "İndirme Tamamlandı",
        "download_complete_message": "Güncelleme başarıyla indirildi. Program, güncellemeyi uygulamak için yeniden başlatılacak.",
        "download_error": "İndirme Hatası",
        "download_error_message": "Güncelleme indirilirken hata oluştu: {0}",
        "download_button": "İndir",
        "cancel_button": "İptal",
        "light_mode": "Açık Mod",
        "dark_mode": "Koyu Mod",
        "theme_settings": "Tema Ayarları"
    },
    "en": {
        "select_folder": "📁 Select Folder",
        "extension_not_found": "No extension found",
        "copy_filename": "Copy File Name",
        "copy_filepath": "Copy File Path",
        "no_folder_selected": "No folder selected yet",
        "start": "▶️ Start",
        "apply_filter": "🔍 Apply Filter",
        "cancel": "⏹️ Cancel",
        "cancelling": "⏹️ Cancelling...",
        "exit": "✖️ Exit",
        "select_all": "Select All",
        "clear_all": "Clear All",
        "all_files_tip": "All files will be listed regardless of extension.",
        "filter_tip": "Filter options enabled, only files with selected extensions will be shown.",
        "calculating_statistics": "Calculating statistics...",
        "loading_file_list": "Loading file list...",
        "tooltip_select": "Select a folder to generate the file list",
        "tooltip_start": "Create and export the file list",
        "tooltip_apply": "Filter files by selected extensions",
        "tooltip_cancel": "Cancel ongoing operation",
        "tooltip_exit": "Exit the application",
        "language": "Language / Dil",
        "view_mode_list": "List View",
        "view_mode_preview": "Preview",
        "tooltip_list_view": "Show files in list view",
        "tooltip_preview_view": "Show files with preview thumbnails",
        "files_and_previews": "Files and Previews",
        "search_files": "Search files...",
        "preview_mode_active": "Preview mode active",
        "loading_preview": "Loading previews...",
        "no_preview_available": "No previewable files found",
        "preview_file": "Preview File",
        "open_file": "Open File",
        "open_file_location": "Open File Location",
        "error_open_file": "Error opening file",
        "error_open_location": "Error opening file location",
        "error_open_url": "Error opening URL",
        "settings": "Settings",
        "info": "Info",
        "files_filtering": "Filtering files, please wait...",
        "files_gathering": "Gathering files, please wait...",
        "files_loading": "Loading files...",
        "filter_applying": "Applying filter...",
        "filter_cancelled": "Filtering operation cancelled",
        "excel_created": "Excel file created successfully...",
        "error_occurred": "An error occurred: {0}",
        "error": "Error",
        "ready": "Ready",
        "operation_cancelled": "Operation Cancelled",
        "file_processed": "Processing file: {0}/{1} • {2}",
        "prev_page": "Previous",
        "next_page": "Next",
        "page": "Page",
        "processing": "Processing: %{0:.1f}",
        "folder_loading": "Loading folder...",
        "loading_subfolders": "Loading subfolders...",
        "folder_loaded_status": "📁 Folder loaded. Click \"Start\" to generate the list.",
        "create_list_time": "🕒 List Creation Time: {0}",
        "select_folder_first": "Please select a folder first!",
        "text_file_error": "Text file could not be created: {0}",
        "start_processing": "▶️ Creating list...",
        "confirm_exit_title": "Exit",
        "confirm_exit_message": "Are you sure you want to exit the program?",
        "no_files_found": "No files found in the selected folder.",
        "files_loaded_message": "{0} files loaded.",
        "files_filtered_message": "{0} files filtered.",
        "filter_saved_message": "Filter settings saved. Will be applied when a folder is selected.",
        "app_title": "EasyLister - File List Generator",
        "app_subtitle": "Quickly list the files in your folders.",
        "full_window_title": "EasyLister - File List Generator",
        "statistics_header": "Statistics",
        "total_files_label": "Total Files:",
        "folder_count_label": "Number of Folders:",
        "total_size_label": "Total Size:",
        "tips_header": "Tips",
        "settings_header": "Settings",
        "operation_status": "Operation Status",
        "subfolders_label": "Include Subfolders:",
        "include_label": "Include",
        "list_format_label": "List Format:",
        "save_location_label": "Save Location:",
        "desktop_label": "Desktop",
        "sort_criteria_label": "Sorting Criteria:",
        "text_format_info": "Text file format",
        "excel_format_info": "Excel workbook format",
        "word_format_info": "Word document format",
        "html_format_info": "Web page format",
        "filter_label": "Filter",
        "kategori_header": "Categories",
        "extensions_header": "File Extensions",
        "tip_1": "Check the subfolder box to scan all files in subdirectories.",
        "tip_3": "Files can be sorted by name, size or extension.",
        "tip_4": "Lists are saved to the program folder by default.",
        "tip_5": "Use the desktop option to save files directly to your desktop.",
        "tip_6": "Use the Cancel button to stop any operation immediately.",
        "tip_preview_formats": "PDF, JPG, PNG, GIF, PSD, AI, EPS files support preview functionality.",
        "preview_file": "Preview File",
        "preview_window_title": "File Preview",
        "preview_not_supported": "Preview is not supported for this file type.",
        "preview_error": "An error occurred while previewing the file.",
        "list_view": "List",
        "preview_view": "Preview",
        "view_mode": "View Mode:",
        "preview_mode_active": "Preview mode active",
        "no_preview_available": "No Preview",
        "loading_preview": "Loading preview...",
        "all_files": "All Files",
        "image_files": "Images",
        "audio_files": "Audio",
        "video_files": "Video & Audio",
        "text_files": "Text",
        "program_files": "Program & Archive",
        "compressed_files": "Compressed",
        "document_files": "Documents",
        "spreadsheet_files": "Spreadsheets",
        "presentation_files": "Presentations",
        "sort_name_asc": "File Name - Ascending",
        "sort_name_desc": "File Name - Descending",
        "sort_ext_asc": "File Extension - Ascending",
        "sort_ext_desc": "File Extension - Descending",
        "sort_size_asc": "File Size - Ascending",
        "sort_size_desc": "File Size - Descending",
        "sort_dir_asc": "File Directory - Ascending",
        "selected_folder": "Selected Folder:",
        "file_list": "File List",
        "row_number": "Row No",
        "file_name": "File Name",
        "file_type": "File Type",
        "file_path": "File Path",
        "file_size": "File Size",
        "creation_date": "Creation Date",
        "modification_date": "Modification Date",
        "file_extension": "File Extension",
        "creation_time": "Creation Time:",
        "sorted_by": "Sorted by:",
        "excel_success": "Excel file successfully created",
        "word_success": "Word document successfully created",
        "html_success": "HTML file successfully created",
        "text_success": "Text file successfully created",
        "open_file_title": "Open File",
        "open_file_message": "Would you like to open the file?",
        "yes": "Yes",
        "no": "No",
        "file_list_section": "Files and Previews",
        "tooltip_subfolders": "Also lists files in subfolders",
        "tooltip_format": "Choose the format in which your list will be saved",
        "tooltip_save_location": "Choose where to save your list",
        "tooltip_sort_criteria": "Choose how files will be sorted",
        "tooltip_select_all": "Select all file extensions",
        "tooltip_clear_all": "Clear all file extension selections",
        "tooltip_file_category": "Click to view file categories",
        "tooltip_file_extension": "Select file extensions to be listed",
        "tooltip_filter_apply": "Click to apply selected filters",
        "tooltip_all_files": "Select/deselect all file extensions",
        "tooltip_category_expand": "Click to expand/collapse this category",
        "tooltip_select_category": "Select/deselect all extensions in this category",
        "select_all_category": "Select all in this category",
        "media_files": "Media",
        "code_files": "Code & Web",
        "design_files": "Design",
        "game_files": "Games",
        "extension_search": "Search:",
        "copyright_footer": "© {year} Muallimun.Net - ListeKolay",
        "filtering_in_progress": "Filtering in progress",
        "filter_complete": "Filtering complete",
        "filter_error": "Filtering error",
        "filter_error_details": "Error occurred during filtering",
        "check_updates": "Check for Updates",
        "update_available": "Update Available",
        "update_available_message": "A new version of ListeKolay is available: {0}\nYour current version: {1}\n\nWould you like to download the update?",
        "no_update_available": "No Update Available",
        "no_update_available_message": "You are using the latest version of ListeKolay.",
        "update_check_error": "Update Check Error",
        "update_check_error_message": "An error occurred while checking for updates. Please check your internet connection and try again.",
        "downloading_update": "Downloading Update...",
        "downloading_update_message": "Downloading version {0}...",
        "download_complete": "Download Complete",
        "download_complete_message": "Update has been downloaded successfully. The program will restart to apply the update.",
        "download_error": "Download Error",
        "download_error_message": "Failed to download the update: {0}",
        "download_button": "Download",
        "cancel_button": "Cancel",
        "light_mode": "Light Mode",
        "dark_mode": "Dark Mode",
        "theme_settings": "Theme Settings"
    },
    "ar": {
        "open_file": "فتح الملف",
        "open_file_location": "فتح موقع الملف",
        "copy_filename": "نسخ اسم الملف",
        "copy_filepath": "نسخ مسار الملف",
        "select_folder": "📁 اختر مجلد",
        "no_folder_selected": "لم يتم اختيار مجلد بعد",
        "start": "▶️ ابدأ",
        "apply_filter": "🔍 تطبيق التصفية",
        "cancel": "⏹️ إلغاء",
        "cancelling": "⏹️ جاري الإلغاء...",
        "exit": "✖️ خروج",
        "select_all": "تحديد الكل",
        "clear_all": "مسح الكل",
        "all_files_tip": "سيتم سرد جميع الملفات بغض النظر عن الامتداد.",
        "filter_tip": "خيارات التصفية مفعلة، سيتم عرض الملفات ذات الامتدادات المحددة فقط.",
        "calculating_statistics": "جاري حساب الإحصائيات...",
        "loading_file_list": "جاري تحميل قائمة الملفات...",
        "tooltip_select": "حدد مجلد لإنشاء قائمة الملفات",
        "tooltip_start": "إنشاء وتصدير قائمة الملفات",
        "tooltip_apply": "تصفية الملفات حسب الامتدادات المحددة",
        "tooltip_cancel": "إلغاء العملية الجارية",
        "tooltip_exit": "الخروج من التطبيق",
        "language": "اللغة / Language",
        "search_files": "البحث عن الملفات...",
        "settings": "إعدادات",
        "info": "معلومات",
        "files_filtering": "جاري تصفية الملفات، يرجى الانتظار...",
        "files_gathering": "جاري تجميع الملفات، يرجى الانتظار...",
        "files_loading": "جاري تحميل الملفات...",
        "filter_applying": "جاري تطبيق التصفية...",
        "filter_cancelled": "تم إلغاء عملية التصفية",
        "excel_created": "تم إنشاء ملف إكسل بنجاح...",
        "error_occurred": "حدث خطأ: {0}",
        "error": "خطأ",
        "ready": "جاهز",
        "operation_cancelled": "تم إلغاء العملية",
        "file_processed": "معالجة الملف: {0}/{1} • {2}",
        "prev_page": "السابق",
        "next_page": "التالي",
        "page": "صفحة",
        "processing": "المعالجة: %{0:.1f}",
        "folder_loading": "جاري تحميل المجلد...",
        "loading_subfolders": "جاري تحميل المجلدات الفرعية...",
        "folder_loaded_status": "📁 تم تحميل المجلد. انقر على \"ابدأ\" لإنشاء القائمة.",
        "create_list_time": "🕒 وقت إنشاء القائمة: {0}",
        "select_folder_first": "الرجاء تحديد مجلد أولاً!",
        "text_file_error": "تعذر إنشاء ملف نصي: {0}",
        "start_processing": "▶️ جاري إنشاء القائمة...",
        "confirm_exit_title": "خروج",
        "confirm_exit_message": "هل أنت متأكد أنك تريد الخروج من البرنامج؟",
        "no_files_found": "لم يتم العثور على ملفات في المجلد المحدد.",
        "files_loaded_message": "تم تحميل {0} ملف.",
        "files_filtered_message": "تم تصفية {0} ملف.",
        "filter_saved_message": "تم حفظ إعدادات التصفية. سيتم تطبيقها عند تحديد مجلد.",
        "app_title": "قوائم سهلة - منشئ قوائم الملفات",
        "app_subtitle": "قم بسرد الملفات في مجلداتك بسرعة.",
        "full_window_title": "قوائم سهلة - منشئ قوائم الملفات",
        "statistics_header": "إحصائيات",
        "total_files_label": "إجمالي الملفات:",
        "folder_count_label": "عدد المجلدات:",
        "total_size_label": "الحجم الإجمالي:",
        "tips_header": "نصائح",
        "settings_header": "إعدادات",
        "operation_status": "حالة العملية",
        "subfolders_label": "تضمين المجلدات الفرعية:",
        "include_label": "تضمين",
        "list_format_label": "تنسيق القائمة:",
        "save_location_label": "موقع الحفظ:",
        "desktop_label": "سطح المكتب",
        "sort_criteria_label": "معيار الترتيب:",
        "text_format_info": "تنسيق ملف نصي",
        "excel_format_info": "تنسيق مصنف إكسل",
        "word_format_info": "تنسيق مستند وورد",
        "html_format_info": "تنسيق صفحة ويب",
        "filter_label": "تصفية",
        "kategori_header": "الفئات",
        "extensions_header": "امتدادات الملفات",
        "tip_1": "حدد مربع المجلدات الفرعية لمسح جميع الملفات في الدلائل الفرعية.",
        "tip_3": "يمكن فرز الملفات حسب الاسم أو الحجم أو الامتداد.",
        "tip_4": "يتم حفظ القوائم في مجلد البرنامج بشكل افتراضي.",
        "tip_5": "استخدم خيار سطح المكتب لحفظ الملفات مباشرة على سطح المكتب.",
        "tip_6": "استخدم زر الإلغاء لإيقاف أي عملية على الفور.",
        "tip_preview_formats": "ملفات PDF و JPG و PNG و GIF و PSD و AI و EPS تدعم وظيفة المعاينة.",
        "all_files": "جميع الملفات",
        "image_files": "صور",
        "audio_files": "صوت",
        "video_files": "فيديو و صوت",
        "text_files": "نصوص",
        "code_files": "برمجة و ويب",
        "data_files": "بيانات",
        "document_files": "مستندات",
        "spreadsheet_files": "جداول بيانات",
        "presentation_files": "عروض تقديمية",
        "program_files": "برامج و أرشيف",
        "compressed_files": "ملفات مضغوطة",
        "sort_name_asc": "اسم الملف - تصاعدي",
        "sort_name_desc": "اسم الملف - تنازلي",
        "sort_ext_asc": "امتداد الملف - تصاعدي",
        "sort_ext_desc": "امتداد الملف - تنازلي",
        "sort_size_asc": "حجم الملف - تصاعدي",
        "sort_size_desc": "حجم الملف - تنازلي",
        "sort_dir_asc": "دليل الملف - تصاعدي",
        "selected_folder": "المجلد المحدد:",
        "file_list": "قائمة الملفات",
        "row_number": "رقم الصف",
        "file_name": "اسم الملف",
        "file_type": "نوع الملف",
        "file_path": "مسار الملف",
        "file_size": "حجم الملف",
        "creation_date": "تاريخ الإنشاء",
        "modification_date": "تاريخ التعديل",
        "file_extension": "امتداد الملف",
        "creation_time": "وقت الإنشاء:",
        "sorted_by": "تم الفرز حسب:",
        "excel_success": "تم إنشاء ملف إكسل بنجاح",
        "word_success": "تم إنشاء مستند وورد بنجاح",
        "html_success": "تم إنشاء ملف HTML بنجاح",
        "text_success": "تم إنشاء ملف نصي بنجاح",
        "open_file_title": "فتح الملف",
        "open_file_message": "هل ترغب في فتح الملف؟",
        "yes": "نعم",
        "no": "لا",
        "tooltip_subfolders": "يسرد أيضًا الملفات في المجلدات الفرعية",
        "tooltip_format": "اختر التنسيق الذي سيتم حفظ قائمتك به",
        "tooltip_save_location": "اختر مكان حفظ قائمتك",
        "tooltip_sort_criteria": "اختر كيفية فرز الملفات",
        "tooltip_select_all": "تحديد جميع امتدادات الملفات",
        "tooltip_clear_all": "مسح جميع تحديدات امتداد الملفات",
        "tooltip_file_category": "انقر لعرض فئات الملفات",
        "tooltip_file_extension": "حدد امتدادات الملفات المراد سردها",
        "tooltip_filter_apply": "انقر لتطبيق الفلاتر المحددة",
        "tooltip_all_files": "تحديد/إلغاء تحديد جميع امتدادات الملفات",
        "tooltip_category_expand": "انقر لتوسيع/طي هذه الفئة",
        "tooltip_select_category": "تحديد/إلغاء تحديد جميع الامتدادات في هذه الفئة",
        "select_all_category": "تحديد الكل في هذه الفئة",
        "media_files": "وسائط",
        "code_files": "برمجة و ويب",
        "data_files": "بيانات",
        "design_files": "تصميم",
        "game_files": "ألعاب",
        "extension_search": "بحث:",
        "copyright_footer": "© {year} معلمون.نت - قوائم لسهلة"
    },
    "de": de_dict,
    "fr": fr_dict,
    "ru": ru_dict,
    "es": es_dict,
    "it": it_dict,
    "fa": fa_dict,
    "ur": ur_dict,
    "hi": hi_dict,
    "zh": zh_dict,
    "ja": ja_dict
}


class FileManagerApp:
    def __init__(self, root):
        self.root = root
        self.current_language = "tr"  # Default language is Turkish
        
        # Uygulama sürüm bilgisi
        self.current_version = "5.1.0"
        self.github_version_url = "https://github.com/muallimun/listekolay/raw/main/listekolay_version.txt"
        self.github_download_url = "https://github.com/muallimun/listekolay/releases/latest"
        
        # Tema ayarları (açık/koyu mod)
        self.is_dark_mode = tk.BooleanVar(value=False)  # Varsayılan olarak açık mod
        
        # Add custom translations for pagination
        self.pagination_translations = {
            "tr": {"page": "Sayfa", "prev_page": "Önceki", "next_page": "Sonraki"},
            "en": {"page": "Page", "prev_page": "Previous", "next_page": "Next"},
            "ar": {"page": "صفحة", "prev_page": "السابق", "next_page": "التالي"},
            "de": {"page": "Seite", "prev_page": "Zurück", "next_page": "Weiter"},
            "fr": {"page": "Page", "prev_page": "Précédent", "next_page": "Suivant"},
            "ru": {"page": "Страница", "prev_page": "Предыдущая", "next_page": "Следующая"},
            "es": {"page": "Página", "prev_page": "Anterior", "next_page": "Siguiente"},
            "it": {"page": "Pagina", "prev_page": "Precedente", "next_page": "Successiva"},
            "zh": {"page": "页面", "prev_page": "上一页", "next_page": "下一页"},
            "ja": {"page": "ページ", "prev_page": "前へ", "next_page": "次へ"}
        }

        # Initialize translations from the global translations dict
        self.languages = translations

        self.root.title(self.get_text("full_window_title"))
        self.root.geometry("1024x768")
        # Use normal state instead of zoomed (which doesn't work on some platforms)
        # self.root.state("zoomed")      
        self.root.minsize(800, 600)       # Minimum window size
        self.root.configure(bg="#e9ecef")
        self.root.resizable(True, True)   # Window can be resized
        self.files = []
        self.filtered_files = []          # Store filtered files separately
        self.include_subfolders = tk.BooleanVar(value=False)
        self.selected_folder_path = ""
        self.save_to_desktop = tk.BooleanVar(value=False)
        self.cancel_flag = False
        
        # For file view mode (list or preview)
        self.view_mode_var = tk.StringVar(value="list")  # Default to list view
        
        # OPTIMIZATION: Enhanced preview caching system with LRU (Least Recently Used) behavior
        self.preview_cache = {}  # Cache for previews to improve performance
        self.preview_cache_keys = []  # Keep track of cache access order for LRU implementation
        
        # Önizleme gezinme değişkenleri
        self.current_preview_files = []  # Tüm önizleme dosyalarını saklar
        self.current_preview_index = -1  # Şu anda gösterilen dosyanın indeksi
        
        # OPTIMIZATION: Performance settings for large folders
        # OPTIMIZATION: Increased batch sizes for better performance with large folders
        self.file_loading_batch_size = 1000    # How many files to process at once during loading (increased from 500)
        # OPTIMIZATION: Increase batch sizes for faster processing of large file lists
        self.file_filtering_batch_size = 5000  # How many files to filter at once (increased from 2000)
        self.file_display_batch_size = 2000    # How many files to add to UI at once (increased from 1000)
        # OPTIMIZATION: Increase preview batch size for better parallel processing
        self.preview_batch_size = 200          # How many previews to generate at once (increased from 100)
        self.max_preview_cache_size = 750     # Maximum number of thumbnails to cache (increased from 500)
        self.preview_items_per_page = 150     # Number of preview items per page (increased from 100)
        self.preview_page = 1                 # Current preview page
        
        self.export_formats = {
            "text": tk.BooleanVar(value=True),
            "excel": tk.BooleanVar(value=False),
            "word": tk.BooleanVar(value=False),
            "html": tk.BooleanVar(value=False)
        }
        
        # Category variables for selecting/deselecting all extensions
        self.category_vars = {}
        self.all_files_var = tk.BooleanVar(value=False)
        self.sort_options = [
            "sort_name_asc",
            "sort_name_desc",
            "sort_ext_asc",
            "sort_ext_desc",
            "sort_size_asc",
            "sort_size_desc",
            "sort_dir_asc"
        ]
        self.selected_sort = tk.StringVar(value=self.sort_options[0])
        
# Final: Genişletilmiş ve mantıksal olarak gruplanmış 6 dosya kategorisi
        self.file_categories = {
            # 1. Belge Dosyaları (dokümanlar, tablolar, sunumlar, metinler)
            "document_files": [
                # Belgeler
                ".doc", ".docx", ".rtf", ".odt", ".pdf", ".txt", ".epub", ".mobi", ".tex", 
                ".pages", ".md", ".csv", ".log", ".udf",
                # Tablolar
                ".xls", ".xlsx", ".xlsm", ".ods", ".numbers",
                # Sunumlar
                ".ppt", ".pptx", ".odp", ".key", ".pps", ".ppsx"
            ],
            
            # 2. Görsel Dosyalar (resim formatları, raster-vektörel)
            "image_files": [
                ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".svg", ".ico", 
                ".raw", ".heif", ".cr2",  ".psd", ".ai", ".eps",
            ],
            
            # 3. Video ve Ses Dosyaları (medya formatları)
            "video_files": [
                # Video
                ".mp4", ".avi", ".mov", ".wmv", ".flv", ".mkv", ".webm", ".m4v", ".mpg", ".mpeg", 
                ".3gp", ".ts", ".vob", ".asf", ".ogv", ".m2v",
                # Ses
                ".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".alac", ".aiff", ".opus", ".m4a"
            ],
            
            # 4. Tasarım ve Eğitim İçeriği Dosyaları (grafik, CAD, font, etkileşimli içerik)
            "design_files": [
                # Grafik tasarım ve vektör
                ".psd", ".ai", ".eps", ".xd", ".indd", ".cdr", ".fig", ".afdesign", ".afphoto",
                # Yazı tipleri
                ".ttf", ".otf", ".woff", ".woff2", ".eot", ".fon",
                # 3D / CAD
                ".dwg", ".dxf", ".skp", ".3ds", ".max", ".c4d", ".blend", ".fbx", ".obj", ".stl", ".step", ".stp",
                # Animasyon / Hareketli medya
                ".ae", ".swf",
                # Eğitim / Etkileşimli içerik
                ".h5p", ".scorm", ".xar", ".cptx", ".story", ".ismp", ".quiz", ".interact", ".ao"
            ],
            
            # 5. Kod ve Web Dosyaları (programlama ve betik dosyaları)
            "code_files": [
                # Programlama dilleri
                ".py", ".java", ".c", ".cpp", ".cs", ".php", ".rb", ".go", ".swift", ".ts", ".js",
                # Web dilleri
                ".html", ".css", ".vue", ".jsx", ".ini", ".dat",
                # Veritabanı, yapılandırma
                ".sql", ".json", ".xml", ".yaml", ".yml", ".config", ".mdb", ".mde", ".accdb", ".accdt", ".accde",
                # Scriptler
                ".sh", ".bat", ".cer",
                # Sunucu-tarayıcı
                ".asp", ".aspx", ".jsp",
                # Diğer
                ".h5p"  # Eğer eğitim aracı olarak değilse web içerik olarak burada da olabilir (ama yukarıda da var)
            ],
            
            # 6. Program ve Arşiv Dosyaları (uygulamalar, kurulum ve sıkıştırılmış dosyalar)
            "program_files": [
                # Çalıştırılabilir ve kurulum
                ".exe", ".dll", ".msi", ".app", ".jar", ".dmg", ".apk", ".deb", ".rpm", ".apk",
                # Arşivleme
                ".zip", ".rar", ".7z", ".tar", ".gz", ".iso"
            ]
        }
        
        # Keep track of which extensions are selected
        self.selected_extensions = {}
        for category, extensions in self.file_categories.items():
            for ext in extensions:
                self.selected_extensions[ext] = tk.BooleanVar(value=True)
                
        # Create GUI
        self.create_gui()
        self.update_ui_state()
        
        # Ayarları yükle
        self.load_config()
        
        # Bind close event
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        
        # Show startup message
        self.update_status(self.get_text("ready"))
        
    def show_error(self, error_title, error_message, exception=None):
        """Kullanıcıya hata göster ve loglama yap"""
        # Hata detaylarını logla
        if exception:
            logging.error(f"{error_title}: {str(exception)}")
        else:
            logging.error(error_title)
        
        # Kullanıcıya hata mesajı göster
        messagebox.showerror(
            self.get_text("error"), 
            error_message
        )
        
        # Durum çubuğunda da göster
        self.update_status(error_message)

    def get_text(self, key):
        """Localization helper function"""
        # Special case for search_files - use our centralized search translations
        if key == "search_files" and self.current_language in search_translations:
            return search_translations[self.current_language]
            
        # Special case for pagination - use our custom pagination translations
        if key in ["page", "prev_page", "next_page"] and hasattr(self, 'pagination_translations'):
            if self.current_language in self.pagination_translations and key in self.pagination_translations[self.current_language]:
                return self.pagination_translations[self.current_language][key]
            elif "en" in self.pagination_translations and key in self.pagination_translations["en"]:
                return self.pagination_translations["en"][key]
            
        if key in self.languages[self.current_language]:
            return self.languages[self.current_language][key]
        # Fallback to English if key not found in current language
        elif key in self.languages["en"]:
            return self.languages["en"][key]
        # Return the key itself if not found in any language
        return key

    def create_gui(self):
        # Create base frame with padding
        self.main_frame = tk.Frame(
            self.root, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            padx=10, 
            pady=10
        )
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # Title and subtitle in a more compact layout
        title_frame = tk.Frame(
            self.main_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        title_frame.pack(fill=tk.X, pady=(0, 5))  # Reduced padding
        
        # Header container to align items horizontally
        header_container = tk.Frame(
            title_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        header_container.pack(side=tk.LEFT, fill=tk.X)
        
        # Program title - left aligned
        title_label = tk.Label(
            header_container, 
            text=self.get_text("app_title"), 
            font=("Segoe UI", 16, "bold"),  # Slightly smaller font
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor="w"  # Left aligned text
        )
        title_label.pack(side=tk.LEFT, pady=(0, 2))
        
        # Subtitle - next to title with separator
        self.subtitle_label = tk.Label(
            header_container, 
            text=" - " + self.get_text("app_subtitle"),  # Add separator
            font=("Segoe UI", 10, "italic"),  # Italic for style
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        self.subtitle_label.pack(side=tk.LEFT, padx=(5, 0), pady=(4, 0))  # Align vertically
        
        # Theme settings container - placed on the right side of the title frame
        theme_container = tk.Frame(
            title_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        theme_container.pack(side=tk.RIGHT)
        
        # Theme mode label
        theme_label = tk.Label(
            theme_container, 
            text=self.get_text("theme_settings"),
            font=("Segoe UI", 9),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        theme_label.pack(side=tk.LEFT, padx=(0, 5))
        
        # Light mode (sun icon) radio button
        light_radio = tk.Radiobutton(
            theme_container,
            text="☀️",  # Sun emoji
            variable=self.is_dark_mode,
            value=False,
            command=self.toggle_theme_mode,
            font=("Segoe UI", 9),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            selectcolor="#f8f9fa",
            indicatoron=False,
            width=2,
            bd=1,
            relief=tk.RAISED if not self.is_dark_mode.get() else tk.FLAT
        )
        self.create_tooltip(light_radio, self.get_text("light_mode"))
        light_radio.pack(side=tk.LEFT)
        
        # Dark mode (moon icon) radio button
        dark_radio = tk.Radiobutton(
            theme_container,
            text="🌙",  # Moon emoji
            variable=self.is_dark_mode,
            value=True,
            command=self.toggle_theme_mode,
            font=("Segoe UI", 9),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            selectcolor="#212529",
            indicatoron=False,
            width=2,
            bd=1,
            relief=tk.RAISED if self.is_dark_mode.get() else tk.FLAT
        )
        self.create_tooltip(dark_radio, self.get_text("dark_mode"))
        dark_radio.pack(side=tk.LEFT)

        # Top controls frame (folder selection, action buttons)
        top_frame = tk.Frame(
            self.main_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        top_frame.pack(fill=tk.X, pady=10)
        
        # Folder selection with integrated subfolder option
        folder_frame = tk.Frame(
            top_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        folder_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        folder_label = tk.Label(
            folder_frame, 
            text=self.get_text("selected_folder"), 
            font=("Segoe UI", 9), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        folder_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.folder_path_var = tk.StringVar(value=self.get_text("no_folder_selected"))
        self.folder_path_entry = tk.Entry(
            folder_frame, 
            textvariable=self.folder_path_var, 
            width=35,  # Reduced width
            state="readonly", 
            font=("Segoe UI", 9), 
            readonlybackground=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.folder_path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        
        self.select_folder_btn = tk.Button(
            folder_frame, 
            text=self.get_text("select_folder"), 
            command=self.select_folder, 
            font=("Segoe UI", 9), 
            bg="#007bff",  # Mavi (klasör seçme butonu için uygun)
            fg="white",
            activebackground="#0069d9",
            activeforeground="white", 
            bd=0,
            padx=10
        )
        self.create_tooltip(self.select_folder_btn, self.get_text("tooltip_select"))
        self.select_folder_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Move subfolder option to here beside folder selection
        self.subfolder_cb = tk.Checkbutton(
            folder_frame, 
            text=self.get_text("subfolders_label") + " " + self.get_text("include_label"), 
            variable=self.include_subfolders,
            font=("Segoe UI", 9), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
            selectcolor="#ffffff",  # Beyaz tik kutuları
            command=self.on_subfolder_changed
        )
        self.create_tooltip(self.subfolder_cb, self.get_text("tooltip_subfolders"))
        self.subfolder_cb.pack(side=tk.LEFT)
        
        # Buttons frame
        button_frame = tk.Frame(
            top_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        button_frame.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Buttons frame is now cleaner after moving search to files and previews section
        
        # Language selection
        language_frame = tk.Frame(
            button_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        language_frame.pack(side=tk.LEFT, padx=(0, 10))
        
        language_label = tk.Label(
            language_frame, 
            text=self.get_text("language") + ":", 
            font=("Segoe UI", 9), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        language_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.language_var = tk.StringVar(value=self.current_language)
        self.language_dropdown = ttk.Combobox(
            language_frame, 
            textvariable=self.language_var, 
            values=["tr", "en", "ar", "de", "fr", "ru", "es", "it", "fa", "ur", "hi", "zh", "ja"], 
            state="readonly", 
            width=5
        )
        self.language_dropdown.bind("<<ComboboxSelected>>", self.change_language)
        self.language_dropdown.pack(side=tk.LEFT, padx=(0, 10))
        

        
        # Action buttons
        self.start_btn = tk.Button(
            button_frame, 
            text=self.get_text("start"), 
            command=self.start_processing, 
            font=("Segoe UI", 9), 
            bg="#28a745",  # Yeşil (işlemi başlat butonu)
            fg="white", 
            activebackground="#218838",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.start_btn, self.get_text("tooltip_start"))
        self.start_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.cancel_btn = tk.Button(
            button_frame, 
            text=self.get_text("cancel"), 
            command=self.cancel_operation, 
            font=("Segoe UI", 9), 
            bg="#dc3545",  # Kırmızı (iptal butonu)
            fg="white",
            activebackground="#c82333",
            activeforeground="white",
            bd=0,
            padx=10,
            state=tk.DISABLED
        )
        self.create_tooltip(self.cancel_btn, self.get_text("tooltip_cancel"))
        self.cancel_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.exit_btn = tk.Button(
            button_frame, 
            text=self.get_text("exit"), 
            command=self.on_close, 
            font=("Segoe UI", 9), 
            bg="#6c757d",  # Gri (çıkış butonu için uygun)
            fg="white",
            activebackground="#5a6268",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.exit_btn, self.get_text("tooltip_exit"))
        self.exit_btn.pack(side=tk.LEFT)

        # Main content frame (2 columns)
        content_frame = tk.Frame(self.main_frame, bg="#e9ecef")
        content_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Left column (Settings, Filter, Tips) - Now with increased width
        left_column = tk.Frame(content_frame, bg="#e9ecef", width=450)  # Increased width
        left_column.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10))
        left_column.pack_propagate(False)  # Prevent shrinking
        
        # Settings panel
        self.settings_frame = tk.LabelFrame(
            left_column, 
            text=self.get_text("settings_header"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=10,
            pady=10
        )
        self.settings_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Subfolder setting removed from here - moved to folder selection area
        
        # Export format setting - improved layout with label inline with checkboxes
        format_container = tk.Frame(self.settings_frame, bg="#e9ecef")
        format_container.pack(fill=tk.X, pady=(0, 5))
        
        format_label = tk.Label(
            format_container, 
            text=self.get_text("list_format_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        format_label.pack(side=tk.LEFT, padx=(0, 5))
        
        # Formats frame is now part of the same container
        formats_frame = tk.Frame(format_container, bg="#e9ecef")
        formats_frame.pack(side=tk.LEFT, fill=tk.X)
        
        # Text format
        text_cb = tk.Checkbutton(
            formats_frame, 
            text="TXT", 
            variable=self.export_formats["text"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(text_cb, self.get_text("text_format_info"))
        text_cb.pack(side=tk.LEFT, padx=(0, 5))
        
        # Excel format
        excel_cb = tk.Checkbutton(
            formats_frame, 
            text="Excel", 
            variable=self.export_formats["excel"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(excel_cb, self.get_text("excel_format_info"))
        excel_cb.pack(side=tk.LEFT, padx=(0, 5))
        
        # Word format
        word_cb = tk.Checkbutton(
            formats_frame, 
            text="Word", 
            variable=self.export_formats["word"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(word_cb, self.get_text("word_format_info"))
        word_cb.pack(side=tk.LEFT, padx=(0, 5))
        
        # HTML format
        html_cb = tk.Checkbutton(
            formats_frame, 
            text="HTML", 
            variable=self.export_formats["html"],
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(html_cb, self.get_text("html_format_info"))
        html_cb.pack(side=tk.LEFT)
        
        # Combined save location and sort criteria in one row
        combined_settings_frame = tk.Frame(self.settings_frame, bg="#e9ecef")
        combined_settings_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Save location setting - left side
        save_frame = tk.Frame(combined_settings_frame, bg="#e9ecef")
        save_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        save_label = tk.Label(
            save_frame, 
            text=self.get_text("save_location_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        save_label.pack(side=tk.LEFT)
        
        self.desktop_cb = tk.Checkbutton(
            save_frame, 
            text=self.get_text("desktop_label"), 
            variable=self.save_to_desktop,
            font=("Segoe UI", 9), 
            bg="#e9ecef",
            fg="#495057",
            selectcolor="#ffffff"  # Beyaz tik kutuları
        )
        self.create_tooltip(self.desktop_cb, self.get_text("tooltip_save_location"))
        self.desktop_cb.pack(side=tk.LEFT, padx=(5, 0))
        
        # Small spacer between settings
        spacer = tk.Frame(combined_settings_frame, bg="#e9ecef", width=20)
        spacer.pack(side=tk.LEFT)
        
        # Sort criteria setting - right side
        sort_frame = tk.Frame(combined_settings_frame, bg="#e9ecef")
        sort_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        sort_label = tk.Label(
            sort_frame, 
            text=self.get_text("sort_criteria_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        sort_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.sort_dropdown = ttk.Combobox(
            sort_frame, 
            textvariable=self.selected_sort, 
            state="readonly", 
            font=("Segoe UI", 9), 
            width=15  # Reduced width
        )
        self.create_tooltip(self.sort_dropdown, self.get_text("tooltip_sort_criteria"))
        self.populate_sort_dropdown()
        self.sort_dropdown.pack(side=tk.LEFT)

        # Add Filter button to settings section
        filter_btn_frame = tk.Frame(self.settings_frame, bg="#e9ecef")
        filter_btn_frame.pack(fill=tk.X, pady=(5, 0))
        
        # First row: Filter button and search
        self.show_filter_btn = tk.Button(
            filter_btn_frame, 
            text=self.get_text("filter_label"), 
            command=self.toggle_filter_section, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=10
        )
        self.create_tooltip(self.show_filter_btn, self.get_text("tooltip_file_category"))
        self.show_filter_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Extension search variable still needed for filter panel functionality
        self.extension_search_var = tk.StringVar()
        
        # Filter panel - initially hidden
        self.filter_frame = tk.LabelFrame(
            left_column, 
            text=self.get_text("filter_label"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=10,
            pady=10
        )
        # Don't pack it initially (hidden by default)
        
        # Filter controls
        filter_controls_frame = tk.Frame(self.filter_frame, bg="#e9ecef")
        filter_controls_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.select_all_btn = tk.Button(
            filter_controls_frame, 
            text=self.get_text("select_all"), 
            command=self.select_all_extensions, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.create_tooltip(self.select_all_btn, self.get_text("tooltip_select_all"))
        self.select_all_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.clear_all_btn = tk.Button(
            filter_controls_frame, 
            text=self.get_text("clear_all"), 
            command=self.clear_all_extensions, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.create_tooltip(self.clear_all_btn, self.get_text("tooltip_clear_all"))
        self.clear_all_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.apply_filter_btn = tk.Button(
            filter_controls_frame, 
            text=self.get_text("apply_filter"), 
            command=self.apply_filter, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.create_tooltip(self.apply_filter_btn, self.get_text("tooltip_filter_apply"))
        self.apply_filter_btn.pack(side=tk.LEFT)
        
        # Add search box for extensions
        search_frame = tk.Frame(self.filter_frame, bg="#e9ecef")
        search_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.extension_search_label = tk.Label(
            search_frame,
            text="🔍 " + self.get_text("extension_search"),
            font=("Segoe UI", 9),
            bg="#e9ecef",
            fg="#495057"
        )
        self.extension_search_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.extension_search_var = tk.StringVar()
        # Remove auto-filter on typing
        
        self.extension_search_entry = tk.Entry(
            search_frame,
            textvariable=self.extension_search_var,
            font=("Segoe UI", 9),
            width=15  # Shorter width
        )
        self.extension_search_entry.pack(side=tk.LEFT, padx=(0, 5))
        
        # Add search button
        self.search_button = tk.Button(
            search_frame,
            text="🔍",
            command=self.filter_extensions,
            font=("Segoe UI", 9, "bold"),
            bg="#e9ecef" if not self.is_dark_mode.get() else "#212529",  # Arka plan rengi tema ile aynı
            fg="#000000" if not self.is_dark_mode.get() else "#ffffff",  # Metin rengi siyah (açık tema) veya beyaz (koyu tema)
            activebackground="#d1d1d1" if not self.is_dark_mode.get() else "#34383c",  # Tıklandığında biraz daha koyu
            activeforeground="#000000" if not self.is_dark_mode.get() else "#ffffff",
            bd=1,
            padx=5
        )
        self.search_button.pack(side=tk.LEFT)
        
        # Removed category header to save space
        
        # Filter categories
        self.filter_categories_frame = tk.Frame(self.filter_frame, bg="#e9ecef")
        self.filter_categories_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create a canvas with scrollbar for categories (increased height for full screen usage)
        self.category_canvas = tk.Canvas(self.filter_categories_frame, bg="#e9ecef", highlightthickness=0, height=300)
        category_scrollbar = ttk.Scrollbar(self.filter_categories_frame, orient=tk.VERTICAL, command=self.category_canvas.yview)
        self.category_canvas.configure(yscrollcommand=category_scrollbar.set)
        
        category_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.category_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Frame inside the canvas for categories
        self.category_inner_frame = tk.Frame(self.category_canvas, bg="#e9ecef")
        self.category_canvas.create_window((0, 0), window=self.category_inner_frame, anchor=tk.NW)
        
        # Populate categories
        self.populate_categories()
        
        # Configure canvas
        self.category_inner_frame.bind("<Configure>", lambda e: self.category_canvas.configure(scrollregion=self.category_canvas.bbox("all")))
        self.category_canvas.bind("<Configure>", self.on_category_canvas_configure)

        # Right column (Statistics, File List)
        right_column = tk.Frame(content_frame, bg="#e9ecef")
        right_column.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Statistics Panel
        self.stats_frame = tk.LabelFrame(
            right_column, 
            text=self.get_text("statistics_header"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=10,
            pady=10
        )
        self.stats_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Statistics content - Now in a single row
        stats_content_frame = tk.Frame(self.stats_frame, bg="#e9ecef")
        stats_content_frame.pack(fill=tk.X)
        
        # Total files
        total_files_label = tk.Label(
            stats_content_frame, 
            text=self.get_text("total_files_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W
        )
        total_files_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.total_files_var = tk.StringVar(value="0")
        total_files_value = tk.Label(
            stats_content_frame, 
            textvariable=self.total_files_var, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        total_files_value.pack(side=tk.LEFT, padx=(0, 15))
        
        # Folder count
        folder_count_label = tk.Label(
            stats_content_frame, 
            text=self.get_text("folder_count_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W
        )
        folder_count_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.folder_count_var = tk.StringVar(value="0")
        folder_count_value = tk.Label(
            stats_content_frame, 
            textvariable=self.folder_count_var, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        folder_count_value.pack(side=tk.LEFT, padx=(0, 15))
        
        # Total size
        total_size_label = tk.Label(
            stats_content_frame, 
            text=self.get_text("total_size_label"), 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W
        )
        total_size_label.pack(side=tk.LEFT, padx=(0, 5))
        
        self.total_size_var = tk.StringVar(value="0 MB")
        total_size_value = tk.Label(
            stats_content_frame, 
            textvariable=self.total_size_var, 
            font=("Segoe UI", 9, "bold"), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        total_size_value.pack(side=tk.LEFT)

        # File list - with top bar for title and view mode toggle
        file_list_header_frame = tk.Frame(right_column, bg="#e9ecef", pady=5)
        file_list_header_frame.pack(fill=tk.X)
        
        # Add title and view mode toggles in the same row
        file_list_header_label = tk.Label(
            file_list_header_frame, 
            text=self.get_text("file_list_section"),
            font=("Segoe UI", 10, "bold"), 
            bg="#e9ecef", 
            fg="#343a40"
        )
        file_list_header_label.pack(side=tk.LEFT, padx=(5, 10))
        
        # Add a search frame on the right side before view mode buttons
        search_frame = tk.Frame(file_list_header_frame, bg="#e9ecef")
        search_frame.pack(side=tk.RIGHT, padx=(10, 10))
        
        # Search label
        self.file_search_label = tk.Label(
            search_frame, 
            text=self.get_text("extension_search") + ":", 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg="#495057"
        )
        self.file_search_label.pack(side=tk.LEFT, padx=(0, 5))
        
        # Arama girişi için container - Entry ve temizleme butonu içerir
        search_entry_container = tk.Frame(search_frame, bg="#e9ecef")
        search_entry_container.pack(side=tk.LEFT)
        
        # Add the search entry field
        self.file_search_var = tk.StringVar()
        self.file_search_var.trace("w", self.filter_file_list)
        self.file_search_entry = tk.Entry(
            search_entry_container,
            textvariable=self.file_search_var,
            font=("Segoe UI", 9),
            width=15,  # Biraz daha geniş
            fg="#000000"  # Metin rengini her zaman siyah olarak ayarla
        )
        self.file_search_entry.pack(side=tk.LEFT)
        
        # Temizleme butonu
        self.clear_search_button = tk.Button(
            search_entry_container,
            text="✕",
            font=("Segoe UI", 7),
            bg="#e9ecef",
            fg="#666666",
            relief=tk.FLAT,
            command=self.clear_search_field,
            padx=0,
            pady=0,
            width=2
        )
        self.clear_search_button.pack(side=tk.LEFT, padx=(0, 2))
        self.clear_search_button.config(state=tk.DISABLED)  # Başlangıçta devre dışı
        
        # Create a placeholder for the search entry
        self.file_search_entry.insert(0, self.get_text("search_files"))
        # Change color to gray for placeholder
        self.file_search_entry.config(fg='gray')
        
        # Bind focus events to handle placeholder text
        self.file_search_entry.bind("<FocusIn>", self.on_search_focus_in)
        self.file_search_entry.bind("<FocusOut>", self.on_search_focus_out)
        
        # Arama ipucu ekle
        search_tooltip = "Filtre dosya adları ve uzantılara göre yapılır"
        self.create_tooltip(self.file_search_entry, search_tooltip)
        
        # View mode buttons - compact design
        controls_container = tk.Frame(file_list_header_frame, bg="#e9ecef")
        controls_container.pack(side=tk.RIGHT, padx=(10, 10))
        
        # View mode buttons - compact design (reverted from radio buttons to normal buttons)
        view_modes_container = tk.Frame(controls_container, bg="#e9ecef")
        view_modes_container.pack(side=tk.LEFT)
        
        # Button frame
        button_frame = tk.Frame(view_modes_container, bg="#e9ecef")
        button_frame.pack(side=tk.LEFT)
        
        # List view button
        self.list_view_btn = tk.Button(
            button_frame,
            text=self.get_text("list_view"),
            command=lambda: self.set_view_mode("list"),
            font=("Segoe UI", 9),
            bg="#17a2b8" if self.view_mode_var.get() == "list" else "#e9ecef",  # Turkuaz (aktif) veya gri (pasif)
            fg="white" if self.view_mode_var.get() == "list" else "#495057",
            activebackground="#138496",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.list_view_btn, self.get_text("tooltip_list_view"))
        self.list_view_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Preview view button
        self.preview_view_btn = tk.Button(
            button_frame,
            text=self.get_text("preview_view"),
            command=lambda: self.set_view_mode("preview"),
            font=("Segoe UI", 9),
            bg="#17a2b8" if self.view_mode_var.get() == "preview" else "#e9ecef",  # Turkuaz (aktif) veya gri (pasif)
            fg="white" if self.view_mode_var.get() == "preview" else "#495057",
            activebackground="#138496",
            activeforeground="white",
            bd=0,
            padx=10
        )
        self.create_tooltip(self.preview_view_btn, self.get_text("tooltip_preview_view"))
        self.preview_view_btn.pack(side=tk.LEFT)
        
        # We removed the duplicate header and search is now in the first header
        
        # File list frame
        file_list_frame = tk.LabelFrame(
            right_column, 
            text="",  # No text here anymore
            font=("Segoe UI", 10, "bold"), 
            bg="#e9ecef", 
            fg="#343a40",
            padx=10,
            pady=10
        )
        file_list_frame.pack(fill=tk.BOTH, expand=True)
        
        # Cache for preview images
        
        # Cache for preview images
        self.preview_cache = {}
        
        # Create Treeview for file list - it will be placed in the container
        self.create_file_list_treeview(file_list_frame)

        # Status Bar
        status_frame = tk.Frame(self.main_frame, bg="#e9ecef", height=25)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))
        
        self.status_var = tk.StringVar(value=self.get_text("ready"))
        status_label = tk.Label(
            status_frame, 
            textvariable=self.status_var, 
            font=("Segoe UI", 9), 
            bg="#e9ecef", 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            anchor=tk.W,
            padx=10,
            pady=5
        )
        status_label.pack(fill=tk.X)
        
        # Tips panel - now at the bottom of left column with enhanced styling
        self.tips_frame = tk.LabelFrame(
            left_column, 
            text=self.get_text("tips_header"), 
            font=("Segoe UI", 10, "bold"), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],  # Theme-aware background
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
            padx=12,
            pady=12,
            height=225,  # Reduced height for better UI proportions
            relief=tk.GROOVE,  # Nicer relief style
            bd=2  # Slightly thicker border
        )
        self.tips_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))
        self.tips_frame.pack_propagate(False)  # Prevent the frame from shrinking to fit content
        
        # Inner frame for tips with padding
        tips_inner_frame = tk.Frame(
            self.tips_frame,
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            padx=2
        )
        tips_inner_frame.pack(fill=tk.BOTH, expand=True)
        
        # Tips - reduced to 6 tips with shorter format
        tip_labels = [
            self.get_text("tip_1"),
            self.get_text("tip_3"),
            self.get_text("tip_4"),
            self.get_text("tip_5"),
            self.get_text("tip_6"),
            self.get_text("tip_preview_formats") # Özel ipucu: Desteklenen ön izleme formatları (dil desteğiyle)
        ]
        
        # Calculate maximum width for wrapping - use maximum available width
        frame_width = left_column.winfo_reqwidth() or 300  # Use actual width or default to 300
        wrap_width = frame_width - 20  # Minimum padding for better appearance
        
        for i, tip in enumerate(tip_labels):
            # Create a separator line between tips (except before the first tip)
            if i > 0:
                separator = tk.Frame(
                    tips_inner_frame,
                    height=1,
                    bg="#a0afc0" if not self.is_dark_mode.get() else "#4a5568"  # Theme-aware separator color
                )
                separator.pack(fill=tk.X, pady=3)
            
            tip_label = tk.Label(
                tips_inner_frame, 
                text=f"• {tip}", 
                font=("Segoe UI", 9), 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],  # Match theme background
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],  # Tema rengini kullan
                justify=tk.LEFT,
                anchor=tk.W,
                wraplength=wrap_width  # Use calculated width to make text fill the width
            )
            tip_label.pack(fill=tk.X, anchor=tk.W, pady=(2, 0))
            
        # Add footer with copyright and website link
        footer_frame = tk.Frame(self.main_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(5, 0), before=status_frame)
        
        # Get current year for copyright notice
        current_year = datetime.datetime.now().year
        
        # Left side - Program name and copyright
        program_text = f"© {current_year} Muallimun.Net - ListeKolay"
        
        program_label = tk.Label(
            footer_frame, 
            text=program_text,
            font=("Segoe UI", 8), 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        program_label.pack(side=tk.LEFT, padx=10)
        
        # Version number next to program name
        version_label = tk.Label(
            footer_frame,
            text=f"v{self.current_version}",
            font=("Segoe UI", 8),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
        )
        version_label.pack(side=tk.LEFT, padx=(2, 0))
        
        # Güncelleme ikonu (Label olarak) - Kesinlikle renkli olacak
        update_icon = tk.Label(
            footer_frame,
            text="⟳",
            font=("Segoe UI", 13, "bold"),  # Daha da büyük font
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg="#FF4500" if not self.is_dark_mode.get() else "#FF9800",  # Turuncu-kırmızı renk
            cursor="hand2",  # El işareti
            padx=2
        )
        update_icon.bind("<Button-1>", lambda e: self.check_for_updates(False))  # Tıklama olayını bağla
        update_icon.pack(side=tk.LEFT, padx=(5, 0))
        self.create_tooltip(update_icon, self.get_text("check_updates"))
        
        # Right side - Website link
        website_link = tk.Label(
            footer_frame,
            text="www.muallimun.net",
            font=("Segoe UI", 8, "underline"),
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
            fg=LIGHT_MODE_COLORS["accent"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["accent"],
            cursor="hand2"  # Change cursor to hand when hovering
        )
        website_link.pack(side=tk.RIGHT, padx=10)
        
        # Make the link clickable
        website_link.bind("<Button-1>", lambda e: self.open_website("http://www.muallimun.net"))
        
        # Progress bar
        self.progress_bar = ttk.Progressbar(
            self.main_frame, 
            orient=tk.HORIZONTAL, 
            length=100, 
            mode='determinate'
        )
        self.progress_bar.pack(fill=tk.X, side=tk.BOTTOM, pady=(5, 0))
        self.progress_bar["value"] = 0

    def create_file_list_treeview(self, parent):
        # Create frame for file list
        self.file_list_frame = tk.Frame(
            parent, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.file_list_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create container for the file view (will hold both Tree and Preview panels)
        self.file_view_container = tk.Frame(
            self.file_list_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.file_view_container.pack(fill=tk.BOTH, expand=True)
        
        # Create frame for Treeview with scrollbars
        tree_frame = tk.Frame(
            self.file_view_container, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create scrollbars
        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        h_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
        
        # Create Treeview
        self.file_tree = ttk.Treeview(
            tree_frame,
            columns=("name", "extension", "path", "size", "created", "modified"),
            show="headings",
            selectmode="extended",
            yscrollcommand=v_scrollbar.set,
            xscrollcommand=h_scrollbar.set
        )
        
        # Configure scrollbars
        v_scrollbar.config(command=self.file_tree.yview)
        h_scrollbar.config(command=self.file_tree.xview)
        
        # Pack scrollbars
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Pack Treeview
        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Configure columns with click-to-sort functionality
        self.file_tree.heading("name", text=self.get_text("file_name"), 
                              command=lambda: self.treeview_sort_column("name", False))
        self.file_tree.heading("extension", text=self.get_text("file_extension"),
                              command=lambda: self.treeview_sort_column("extension", False))
        self.file_tree.heading("path", text=self.get_text("file_path"),
                              command=lambda: self.treeview_sort_column("path", False))
        self.file_tree.heading("size", text=self.get_text("file_size"),
                              command=lambda: self.treeview_sort_column("size", False))
        self.file_tree.heading("created", text=self.get_text("creation_date"),
                              command=lambda: self.treeview_sort_column("created", False))
        self.file_tree.heading("modified", text=self.get_text("modification_date"),
                              command=lambda: self.treeview_sort_column("modified", False))
        
        self.file_tree.column("name", width=150)
        self.file_tree.column("extension", width=80)
        self.file_tree.column("path", width=200)
        self.file_tree.column("size", width=80)
        self.file_tree.column("created", width=120)
        self.file_tree.column("modified", width=120)
        
        # Bind right-click event for context menu
        self.file_tree.bind("<Button-3>", self.show_context_menu)
        
        # Bind double-click event to open file
        # Bind double-click event to open file
        self.file_tree.bind("<Double-1>", lambda event: self.open_selected_file())
        
        # Create right-click context menu (enhanced with more options)
        self.context_menu = tk.Menu(
            self.root, 
            tearoff=0, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        
        # Open file option
        self.context_menu.add_command(
            label=self.get_text("open_file"),
            command=self.open_selected_file
        )
        
        # Open file location option
        self.context_menu.add_command(
            label=self.get_text("open_file_location"),
            command=self.open_file_location
        )
        
        # Add separator
        self.context_menu.add_separator()
        
        # Preview file option
        self.context_menu.add_command(
            label=self.get_text("preview_file"),
            command=self.preview_selected_file
        )
        
        # Add separator
        self.context_menu.add_separator()
        
        # Copy filename option
        self.context_menu.add_command(
            label=self.get_text("copy_filename"),
            command=self.copy_filename_to_clipboard
        )
        
        # Copy file path option
        self.context_menu.add_command(
            label=self.get_text("copy_filepath"),
            command=self.copy_filepath_to_clipboard
        )
        
        # Set up context menu for preview mode thumbnails
        self.preview_context_menu = tk.Menu(
            self.root, 
            tearoff=0, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]
        )
        
        # Open file option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("open_file"),
            command=self.open_preview_file
        )
        
        # Open file location option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("open_file_location"),
            command=self.open_preview_file_location
        )
        
        # Add separator
        self.preview_context_menu.add_separator()
        
        # Copy filename option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("copy_filename"),
            command=self.copy_preview_filename_to_clipboard
        )
        
        # Copy file path option for preview
        self.preview_context_menu.add_command(
            label=self.get_text("copy_filepath"),
            command=self.copy_preview_filepath_to_clipboard
        )
        
    def show_context_menu(self, event):
        """Show context menu on right-click in the file treeview"""
        # Select the item under the cursor first
        item = self.file_tree.identify_row(event.y)
        if item:
            # If an item was clicked, select it
            self.file_tree.selection_set(item)
            # Show the context menu
            try:
                self.context_menu.tk_popup(event.x_root, event.y_root)
            finally:
                # Make sure to release the grab
                self.context_menu.grab_release()
                
    def preview_selected_file(self):
        """Open a preview window for the selected file"""
        # Get the selected item
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection
            
        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")
        
        if not values:
            return  # No values found
            
        # Extract file name and path
        file_name = values[0]
        file_ext = values[1]
        file_dir_path = values[2]
        
        # Construct full file path
        # For Windows paths that already include filename, use as-is
        if os.path.basename(file_dir_path) == file_name:
            file_path = file_dir_path
        else:
            # Otherwise join directory and filename
            file_path = os.path.join(file_dir_path, file_name)
        
        # Check if the file exists
        if not os.path.isfile(file_path):
            messagebox.showerror(
                self.get_text("error"),
                f"{file_path} not found."
            )
            return
            
        # Open the preview window
        self.create_file_preview_window(file_path)
    
    def open_selected_file(self):
        """Open the selected file with the default application"""
        # Get the selected item
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection
            
        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")
        
        if not values:
            return  # No values found
            
        # Extract file path and name
        file_name = values[0]
        file_ext = values[1]
        file_dir_path = values[2]
        
        # Construct full file path
        # For Windows paths that already include filename, use as-is
        if os.path.basename(file_dir_path) == file_name:
            file_path = file_dir_path
        else:
            # Otherwise join directory and filename
            file_path = os.path.join(file_dir_path, file_name)
        
        # Check if the file exists
        if not os.path.isfile(file_path):
            messagebox.showerror(
                self.get_text("error"),
                f"{file_path} not found."
            )
            return
            
        # Open the file
        self.open_file(file_path)

    def toggle_filter_section(self):
        """Toggle the visibility of the filter section"""
        if hasattr(self, "filter_frame"):
            # Check if the filter panel is currently visible
            is_visible = False
            for widget in self.filter_frame.master.winfo_children():
                if widget == self.filter_frame and widget.winfo_ismapped():
                    is_visible = True
                    break
                    
            if is_visible:
                # Hide the filter panel
                self.filter_frame.pack_forget()
                self.show_filter_btn.config(bg="#17a2b8")  # Reset button color
                
                # Show the tips frame if it was hidden
                if hasattr(self, 'tips_frame'):
                    self.tips_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=(10, 0))
            else:
                # ALWAYS hide the tips panel when filter is shown - per user request
                if hasattr(self, 'tips_frame'):
                    self.tips_frame.pack_forget()
                
                # Show the filter panel - expand to use the full available height
                self.filter_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0), after=self.settings_frame)
                self.show_filter_btn.config(bg="#138496")  # Change button color to indicate active state
                
                # Clear all extension selections when the filter section is opened
                self.clear_all_extensions()
    
    def filter_extensions(self, *args):
        """Filter the displayed extensions based on search text"""
        search_text = self.extension_search_var.get().lower()
        
        # If search text is empty, just refresh without search
        if not search_text:
            self.populate_categories()
            
            # Hide any previous "not found" message
            if hasattr(self, 'not_found_label') and self.not_found_label:
                self.not_found_label.pack_forget()
            return
            
        # Clear all extension selections when a new search begins
        for ext in self.selected_extensions:
            self.selected_extensions[ext].set(False)
        # Repopulate categories with the search filter
        self.populate_categories(search_filter=search_text)
        
        # Track if we found any matching extensions
        found_match = False
        matching_category = None
        
        # Check each category for matching extensions
        for category in self.file_categories:
            # Get extensions for this category
            category_extensions = self.file_categories[category]
            
            # Check if any extension in this category matches the search
            matching_extensions = [ext for ext in category_extensions if search_text in ext.lower()]
            
            if matching_extensions:
                found_match = True
                matching_category = category
                
                # Set this category as active and expand it
                self.toggle_category_display(category)
                break
        
        # Display "not found" message if no matches
        if not found_match:
            # Create or update "not found" message
            if not hasattr(self, 'not_found_label') or not self.not_found_label:
                self.not_found_label = tk.Label(
                    self.extensions_container,
                    text=f"'{search_text}' not found in any file extension",
                    font=("Segoe UI", 9, "italic"),
                    bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                    fg="#dc3545"  # Keep red text for error message in both themes
                )
            else:
                self.not_found_label.config(text=f"'{search_text}' not found in any file extension")
            
            # Display the message
            self.not_found_label.pack(pady=10)
            
            # Hide any previously shown extension frames
            self.active_category = None
            for cat in self.file_categories.keys():
                if hasattr(self, f"{cat}_extensions_frame"):
                    ext_frame = getattr(self, f"{cat}_extensions_frame")
                    ext_frame.pack_forget()
                    
                    # Reset button appearance
                    if hasattr(self, f"{cat}_btn"):
                        btn = getattr(self, f"{cat}_btn")
                        btn.config(bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], relief=tk.RAISED)
        else:
            # Hide any previous "not found" message
            if hasattr(self, 'not_found_label') and self.not_found_label:
                self.not_found_label.pack_forget()
                
        # Restore the selection state for category checkboxes
        # Create a dictionary to track which categories are visible
        visible_categories = {cat: False for cat in self.file_categories.keys()}
        
        # If we have matching extensions, mark their categories as visible
        for category in self.file_categories.keys():
            # Check if this category has extensions that match the search
            if search_text:  # Using search_text instead of undefined search_filter
                cat_exts = self.file_categories[category]
                matching_exts = [ext for ext in cat_exts if search_text in ext.lower()]
                visible_categories[category] = bool(matching_exts)
            else:
                visible_categories[category] = True
                
        # Update category checkbox state based on visibility
        for category, is_visible in visible_categories.items():
            if category in self.category_vars:
                # Only set to false if not visible, preserve user selection if visible
                if not is_visible:
                    self.category_vars[category].set(False)
    
    def populate_categories(self, search_filter=None):
        # Clear existing widgets
        for widget in self.category_inner_frame.winfo_children():
            widget.destroy()
        
        # Create a frame for the categories - using 3x2 grid layout
        categories_container = tk.Frame(
            self.category_inner_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        categories_container.pack(fill=tk.BOTH, expand=True)
        
        # Process each category from the file_categories dictionary
        all_categories = list(self.file_categories.keys())
        
        # Create category variables if not already created
        if not hasattr(self, 'category_vars'):
            self.category_vars = {}
            
        # Create extensions container (will hold all extension panels)
        self.extensions_container = tk.Frame(
            self.category_inner_frame, 
            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
        )
        self.extensions_container.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        # Track currently active category
        if not hasattr(self, 'active_category'):
            self.active_category = None
        
        # Create checkboxes in a 3x2 grid
        for i, category in enumerate(all_categories):
            # Determine grid position (3 columns, 2 rows)
            row = i // 3
            col = i % 3
            
            # Create a variable for the category checkbox if not exists
            if category not in self.category_vars:
                self.category_vars[category] = tk.BooleanVar(value=False)
            
            # Create a frame for this category in the grid
            category_frame = tk.Frame(
                categories_container, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
                padx=5, 
                pady=5
            )
            category_frame.grid(row=row, column=col, sticky="nsew")
            
            # Create the category button with checkbox appearance
            # Doğrudan kategori adlarını kullan - eşleştirmeye gerek yok
            category_mapped = category
                
            category_label_text = self.get_text(category_mapped)
            category_btn = tk.Button(
                category_frame,
                text=category_label_text,
                font=("Segoe UI", 9, "bold"),
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                bd=1,
                relief=tk.RAISED,
                padx=8,
                pady=3,
                anchor=tk.W,
                command=lambda cat=category: self.toggle_category_display(cat)
            )
            self.create_tooltip(category_btn, self.get_text("tooltip_category_expand"))
            category_btn.pack(side=tk.TOP, anchor=tk.W, fill=tk.X)
            
            # Store reference to category button
            setattr(self, f"{category}_btn", category_btn)
            
            # Create extensions frame for this category (initially hidden)
            extensions_frame = tk.Frame(
                self.extensions_container, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
            )
            setattr(self, f"{category}_extensions_frame", extensions_frame)
            
            # Add a select all checkbox at the top of extensions
            select_all_frame = tk.Frame(
                extensions_frame, 
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
            )
            select_all_frame.pack(fill=tk.X, pady=(0, 5))
            
            select_all_cb = tk.Checkbutton(
                select_all_frame,
                text=self.get_text("select_all_category"),
                variable=self.category_vars[category],
                font=("Segoe UI", 9, "italic"),
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                selectcolor="#ffffff",  # Beyaz tik kutuları
                command=lambda cat=category: self.toggle_all_category_extensions(cat)
            )
            self.create_tooltip(select_all_cb, self.get_text("tooltip_select_category"))
            select_all_cb.pack(side=tk.LEFT, padx=(5, 0))
            
            # Create the extensions within this frame
            self._create_extensions_list(category, extensions_frame, search_filter=search_filter)
    
    def _create_extensions_list(self, category, parent_frame, search_filter=None):
        """Create the list of file extensions for a category"""
        extensions = self.file_categories[category]
        
        # Filter extensions if search is active
        highlighted_extensions = []
        if search_filter:
            # Find extensions that match the search filter
            highlighted_extensions = [ext for ext in extensions if search_filter in ext.lower()]
            
            # If we're not clearing search (empty string), use the filtered list
            if search_filter:
                extensions = highlighted_extensions
            
        # Skip if no extensions match search (but not if clearing search)
        if search_filter and not extensions:
            return
        
        # Create grid layout for extensions
        ext_grid_frame = tk.Frame(parent_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
        ext_grid_frame.pack(fill=tk.X, expand=True)
        
        # Determine optimal number of columns based on extension count
        num_extensions = len(extensions)
        # Daha fazla sütun kullan (3-5) ve her sütunda daha az öğe olsun
        num_columns = min(5, max(2, num_extensions // 4))  # Use 2-5 columns
        extensions_per_column = max(3, (num_extensions + num_columns - 1) // num_columns)  # En az 3 öğe
        
        # Create columns
        column_frames = []
        for i in range(num_columns):
            column_frame = tk.Frame(ext_grid_frame, bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
            column_frame.pack(side=tk.LEFT, fill=tk.Y, expand=True)
            column_frames.append(column_frame)
        
        # Dictionary to store references to highlighted checkboxes
        highlighted_checkboxes = {}
        
        # Add checkboxes for each extension in columns
        for i, ext in enumerate(extensions):
            column_idx = i // extensions_per_column
            if column_idx >= num_columns:
                column_idx = num_columns - 1
                
            ext_frame = tk.Frame(column_frames[column_idx], bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
            ext_frame.pack(fill=tk.X)
            
            # Create checkbox if the extension variable exists, otherwise create it
            if ext not in self.selected_extensions:
                self.selected_extensions[ext] = tk.BooleanVar(value=False)
            
            # Determine if this extension should be highlighted
            should_highlight = search_filter and ext in highlighted_extensions
            
            ext_cb = tk.Checkbutton(
                ext_frame, 
                text=ext, 
                variable=self.selected_extensions[ext],
                font=("Segoe UI", 9, "bold" if should_highlight else "normal"), 
                bg="#e8f4f8" if should_highlight else (LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]),  # Highlight or theme-appropriate background color
                fg="#0056b3" if should_highlight else (LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"]),  # Blue for highlights, otherwise theme-appropriate text color
                selectcolor="#ffffff",  # Beyaz tik kutuları
                anchor=tk.W
            )
            self.create_tooltip(ext_cb, self.get_text("tooltip_file_extension"))
            ext_cb.pack(fill=tk.X)
            
            # Store reference to checkbox if highlighted
            if should_highlight:
                highlighted_checkboxes[ext] = ext_cb
                
                # Make this the last checkbox in the list so it's shown
                self.last_highlighted_checkbox = ext_cb
        
        # Check if any of the extensions in this category are selected
        any_selected = any(self.selected_extensions.get(ext, tk.BooleanVar(value=False)).get() 
                         for ext in self.file_categories.get(category, []))
        if any_selected:
            # Update the category checkbox if any extensions are selected
            self.category_vars[category].set(True)
            
    def toggle_category_display(self, category):
        """Show one category's extensions and hide others"""
        # Check if there's an active category and it's the same as the clicked one
        toggle_off = self.active_category == category
        
        # Hide all extension frames first
        for cat in self.file_categories.keys():
            if hasattr(self, f"{cat}_extensions_frame"):
                ext_frame = getattr(self, f"{cat}_extensions_frame")
                ext_frame.pack_forget()
                
                # Reset button appearance
                if hasattr(self, f"{cat}_btn"):
                    btn = getattr(self, f"{cat}_btn")
                    btn.config(bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], relief=tk.RAISED)
        
        # If we're not toggling off, show the clicked category
        if not toggle_off:
            # Set as active category
            self.active_category = category
            
            # Show extensions frame
            ext_frame = getattr(self, f"{category}_extensions_frame")
            ext_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
            
            # Update button appearance to indicate active state
            btn = getattr(self, f"{category}_btn")
            # Use a light blue for light mode and a darker blue for dark mode as active state
            active_bg = "#d1ecf1" if not self.is_dark_mode.get() else "#0d4b66"
            btn.config(bg=active_bg, relief=tk.SUNKEN)
        else:
            # We're toggling off the current category
            self.active_category = None

    def update_tips(self):
        """Update tips with the current language"""
        if hasattr(self, 'tips_frame'):
            # Clear existing tips
            for widget in self.tips_frame.winfo_children():
                widget.destroy()
                
            # Add new tips with updated language
            tip_labels = [
                self.get_text("tip_1"),
                self.get_text("tip_3"),
                self.get_text("tip_4"),
                self.get_text("tip_5"),
                self.get_text("tip_6"),
                self.get_text("tip_preview_formats")
            ]
            
            for i, tip in enumerate(tip_labels):
                tip_label = tk.Label(
                    self.tips_frame, 
                    text=f"• {tip}", 
                    font=("Segoe UI", 9), 
                    bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"], 
                    fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                    justify=tk.LEFT if self.current_language != "ar" else tk.RIGHT,
                    wraplength=280
                )
                tip_label.pack(anchor=tk.W, pady=(0 if i > 0 else 0, 2))
    
    def toggle_all_category_extensions(self, category):
        """Toggle all extensions in a category"""
        # Get the selection state
        select_all = self.category_vars[category].get()
        
        # Find the extensions for this category
        category_extensions = []
        
        # Yalnızca normal kategorileri kullan, çünkü zaten mantıksal gruplamayı file_categories içinde yaptık
        if category in self.file_categories:
            category_extensions = self.file_categories[category]
        
        # Apply the selection to all extensions in this category
        for ext in category_extensions:
            if ext in self.selected_extensions:
                self.selected_extensions[ext].set(select_all)
        
        # Apply filter
        self.apply_filter()
            
    def toggle_category(self, category, show=None):
        """Toggle category accordion visibility"""
        # Get the button, select frame and extensions frame
        btn = getattr(self, f"{category}_btn")
        select_frame = getattr(self, f"{category}_select_frame")
        frame = getattr(self, f"{category}_frame")
        select_btn = getattr(self, f"{category}_select_btn")
        
        # Check if the frame is packed (visible)
        is_visible = frame.winfo_ismapped()
        
        # If show parameter is provided, use it instead of toggling
        if show is not None:
            should_show = show
        else:
            should_show = not is_visible
        
        # If we're showing this category and not restoring state after filtering
        if should_show and show is None:
            # First, collapse all other categories - selections will be preserved
            for cat in self.category_vars.keys():
                if cat != category:  # Don't affect the clicked category yet
                    try:
                        cat_btn = getattr(self, f"{cat}_btn")
                        cat_select_frame = getattr(self, f"{cat}_select_frame")
                        cat_frame = getattr(self, f"{cat}_frame")
                        
                        # Only hide if currently visible
                        if cat_frame.winfo_ismapped():
                            cat_select_frame.pack_forget()
                            cat_frame.pack_forget()
                            cat_btn.config(text="▶ " + self.get_text(cat))
                    except (AttributeError, Exception) as e:
                        logging.warning(f"Failed to process category {cat}: {str(e)}")
        
        # Now toggle the clicked category
        if not should_show:
            select_frame.pack_forget()
            frame.pack_forget()
            btn.config(text="▶ " + self.get_text(category))
        else:
            select_frame.pack(fill=tk.X, padx=15)
            select_btn.pack(side=tk.LEFT, fill=tk.X)
            frame.pack(fill=tk.X, pady=(0, 5))
            btn.config(text="▼ " + self.get_text(category))
            
        # Update the canvas scrollregion
        self.category_canvas.configure(scrollregion=self.category_canvas.bbox("all"))

    def on_category_canvas_configure(self, event):
        # Update the canvas inner frame width to match the canvas
        self.category_canvas.itemconfig("all", width=event.width)

    def populate_sort_dropdown(self):
        # Clear current items
        self.sort_dropdown['values'] = []
        
        # Add localized options
        localized_options = [self.get_text(option) for option in self.sort_options]
        self.sort_dropdown['values'] = localized_options
        
        # Select the first option by default
        self.sort_dropdown.current(0)

    def on_search_focus_in(self, event):
        """Handle focus in event for search entry"""
        if self.file_search_entry.get() == self.get_text("search_files"):
            self.file_search_entry.delete(0, tk.END)
            # Set text color to always black for better visibility
            self.file_search_entry.config(fg="#000000")
            # Temizleme butonunu devre dışı bırak (placeholder vardı)
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.DISABLED)
        else:
            # İçerik varsa temizleme butonunu etkinleştir
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.NORMAL)
        
        # Highlight search field with a light blue background when focused
        self.file_search_entry.config(bg="#e6f2ff")

    def on_search_focus_out(self, event):
        """Handle focus out event for search entry"""
        if not self.file_search_entry.get():
            self.file_search_entry.insert(0, self.get_text("search_files"))
            self.file_search_entry.config(fg='gray')
            # Temizleme butonunu devre dışı bırak
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.DISABLED)
        else:
            # İçerik varsa temizleme butonunu aktif tut
            if hasattr(self, 'clear_search_button'):
                self.clear_search_button.config(state=tk.NORMAL)
        
        # Reset background color when focus is lost
        self.file_search_entry.config(bg="white")

    def clear_search_field(self):
        """Arama alanını temizler ve tüm dosyaları gösterir"""
        # Placeholder metni olmadan alanı temizle
        self.file_search_entry.delete(0, tk.END)
        # Odağı arama kutusuna getir
        self.file_search_entry.focus_set()
        # Arama alanını odaklandığında olduğu gibi işaretle
        self.file_search_entry.config(bg="#e6f2ff", fg="#000000")
        # Temizleme butonunu devre dışı bırak
        self.clear_search_button.config(state=tk.DISABLED)
        # Dosya listesini güncelle
        self.filter_file_list()
    
    def filter_file_list(self, *args):
        """Filter the file list based on search text"""
        search_text = self.file_search_var.get().lower()
        
        # Skip filtering if the text is the placeholder
        if search_text == self.get_text("search_files").lower():
            return
        
        # Clear the file list
        self.clear_file_list()
        
        # Temizleme butonunun durumunu güncelle (yeni)
        if hasattr(self, 'clear_search_button'):
            if search_text and search_text != self.get_text("search_files").lower():
                self.clear_search_button.config(state=tk.NORMAL)
            else:
                self.clear_search_button.config(state=tk.DISABLED)
            
        # If search text is empty or placeholder, show all files
        if not search_text:
            if hasattr(self, 'all_files'):
                for file_info in self.all_files:
                    self._add_file_to_list(file_info)
            return
        
        # Filter files based on the search text
        if hasattr(self, 'all_files'):
            found_files = []
            for file_info in self.all_files:
                file_name = file_info["name"].lower()
                file_ext = file_info["extension"].lower()
                
                # If the search text is in file name or extension, add it to the list
                if search_text in file_name or search_text in file_ext:
                    self._add_file_to_list(file_info)
                    found_files.append(file_info)
            
            # If in preview mode, refresh the preview panel with filtered files
            if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview":
                self._build_preview_panel(found_files)

    def toggle_theme_mode(self):
        """Koyu/açık mod arası geçiş yap ve temayı uygula"""
        is_dark = self.is_dark_mode.get()
        
        # Tema renk sabitlerini seç
        theme = DARK_MODE_COLORS if is_dark else LIGHT_MODE_COLORS
        
        # Ana arka plan rengi
        self.root.configure(bg=theme["bg"])
        self.main_frame.configure(bg=theme["bg"])
        
        # Gizli sol filtreleme bölümünün temaya uygun renklerini ayarla
        if hasattr(self, 'filter_frame') and self.filter_frame:
            self.filter_frame.configure(bg=theme["bg"])
            
        # Uygulama genelinde tüm widget'ları güncelle
        self._update_widget_colors(self.main_frame, theme)
        
        # Entry widget'ların metin renklerini özellikle güncelle
        if hasattr(self, 'file_search_entry'):
            if self.file_search_entry.get() == self.get_text("search_files"):
                # Placeholder metin gri kalmalı
                self.file_search_entry.config(fg='gray')
            else:
                # Normal metin tema rengine uymalı
                self.file_search_entry.config(fg=theme["text"])
        
        # Config dosyasına kaydet
        self.save_config()
        
    def _update_widget_colors(self, parent, theme):
        """Belirtilen parent widget'ın altındaki tüm widget'ların renklerini güncelle"""
        # Parent widget'ın kendisini güncelle
        if isinstance(parent, (tk.Frame, tk.LabelFrame, tk.Label, tk.Button)):
            parent.configure(bg=theme["bg"])
            
            # LabelFrame başlıkları için özel işlem
            if isinstance(parent, tk.LabelFrame):
                # LabelFrame başlık metninin rengini güncelle
                parent.configure(fg=theme["text"])
            
            # Özel renk ayarları
            elif isinstance(parent, tk.Label):
                # Tüm etiketleri ve tüm metin içeren widget'ları güncelle
                if "fg" in parent.configure():
                    # Açık temada tüm metinler siyah, koyu temada tüm metinler beyaz
                    # Label metnini tema rengine ayarla (gri, koyu gri, açık gri, vs dikkate almadan)
                    # Tooltip rengi veya buton özel renkleri olmadığı sürece tüm metinleri güncelle
                    if parent.cget("background") != "#ffffcc":  # Tooltip rengini kontrol et
                        parent.configure(fg=theme["text"])
            
            # Butonlar için özel ayarlar
            elif isinstance(parent, tk.Button):
                # Metin rengini güncelle - Filtreleme bölümündeki butonlar için özel kontrol
                button_text = parent.cget("text")
                
                # Filtreleme bölümündeki butonlar için özel işlem
                if button_text == self.get_text("select_all") or button_text == self.get_text("clear_all") or button_text == self.get_text("apply_filter") or button_text == "🔍" or button_text == self.get_text("filter_label"):
                    # Bu butonlar için siyah/beyaz metin rengi (temaya bağlı)
                    parent.configure(fg=theme["text"])
                else:
                    # Diğer butonlar için standart buton metin rengi
                    parent.configure(fg=theme["button_text"])
                
                # Buton tipine göre arkaplan rengi atama
                
                # Buton türlerine göre renk atamaları
                if "✖️ Kapat" in button_text or "❌" in button_text:
                    parent.configure(bg=theme["exit_button"])
                elif "📁 Klasör" in button_text:
                    parent.configure(bg=theme["folder_button"])
                elif "⏹️ İptal" in button_text:
                    parent.configure(bg=theme["cancel_button"])
                elif "▶️ Başlat" in button_text:
                    parent.configure(bg=theme["start_button"])
                elif "🔍 Filtre" in button_text or "Filtrele" in button_text:
                    parent.configure(bg=theme["filter_button"])
                elif "Listele" in button_text:
                    # Görünüm modu butonları için özel işlem
                    if self.view_mode_var.get() == "list":
                        parent.configure(bg=theme["active_view_button"])
                    else:
                        parent.configure(bg=theme["inactive_view_button"])
                elif "Ön İzleme" in button_text or "Preview" in button_text:
                    # Görünüm modu butonları için özel işlem
                    if self.view_mode_var.get() == "preview":
                        parent.configure(bg=theme["active_view_button"])
                    else:
                        parent.configure(bg=theme["inactive_view_button"])
                
                # Düz butonlar için (simge butonlar)
                if "relief" in parent.configure() and parent.cget("relief") == tk.FLAT:
                    if "fg" in parent.configure() and parent.cget("fg") == "#007bff":
                        parent.configure(fg=theme["accent"])
        
        # Tüm alt widget'ları yinelemeli olarak güncelle
        for child in parent.winfo_children():
            self._update_widget_colors(child, theme)
    
    def update_ui_state(self):
        # Update UI state based on the current application state
        if not self.selected_folder_path:
            # Disable buttons that need a folder
            self.start_btn.config(state=tk.DISABLED)
            self.subfolder_cb.config(state=tk.DISABLED)
            self.apply_filter_btn.config(state=tk.DISABLED)
        else:
            # Enable buttons when a folder is selected
            self.start_btn.config(state=tk.NORMAL)
            self.subfolder_cb.config(state=tk.NORMAL)
            self.apply_filter_btn.config(state=tk.NORMAL)

    def create_tooltip(self, widget, text):
        def enter(event):
            # Store tooltip as attribute of widget to avoid global reference issues
            widget.tooltip = tk.Toplevel(widget)
            widget.tooltip.overrideredirect(True)
            widget.tooltip.geometry(f"+{event.x_root + 15}+{event.y_root + 10}")
            
            tooltip_label = tk.Label(
                widget.tooltip,
                text=text,
                justify=tk.LEFT,
                background="#ffffcc",
                relief="solid",
                borderwidth=1,
                font=("Segoe UI", 8),
                wraplength=250
            )
            tooltip_label.pack(padx=2, pady=2)
            
            # For Arabic language, switch text direction
            if self.current_language == "ar":
                tooltip_label.config(justify=tk.RIGHT)
            
        def leave(event):
            if hasattr(widget, 'tooltip') and widget.tooltip.winfo_exists():
                widget.tooltip.destroy()
                delattr(widget, 'tooltip')
                
        def motion(event):
            if hasattr(widget, 'tooltip') and widget.tooltip.winfo_exists():
                widget.tooltip.geometry(f"+{event.x_root + 15}+{event.y_root + 10}")
                
        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)
        widget.bind("<Motion>", motion)

    def change_language(self, event=None):
        # Get the selected language
        new_lang = self.language_var.get()
        
        # Update the current language
        if new_lang in self.languages:
            self.current_language = new_lang
            
            # Update the UI with the new language
            self.update_ui_language()
            
            # Force update of all categories and extensions
            self.populate_categories()
            
            # Update main elements and titles
            self.update_main_titles()
            
            # Log the language change
            logging.info(f"Language changed to: {new_lang}")
    
    def update_main_titles(self):
        # Update window title
        self.root.title(self.get_text("full_window_title"))
        
        # Doğrudan alt başlık referansını güncelle
        if hasattr(self, 'subtitle_label'):
            self.subtitle_label.config(text=" - " + self.get_text("app_subtitle"))
        
        # Update title and subtitle
        for widget in self.main_frame.winfo_children():
            if isinstance(widget, tk.Frame):
                for child in widget.winfo_children():
                    # Güvenli bir şekilde widget'ın "text" özelliğine erişmeye çalış
                    try:
                        # Sadece Label widget'ları için
                        if isinstance(child, tk.Label):
                            # Label'ın mevcut metni
                            text = child.cget("text")
                            
                            # Başlık güncelleme
                            if "ListeKolay" in text or "EasyLister" in text:
                                child.config(text=self.get_text("app_title"))
                            # Diğer metinleri güncelleme
                            elif text and not text.startswith(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9")) and not "©" in text:
                                # İlgili çeviri anahtarını bul
                                found_key = None
                                
                                # Tüm dillerde anahtar ara
                                for key in self.languages[self.current_language].keys():
                                    for lang_code in self.languages.keys():
                                        if self.languages[lang_code].get(key, "") == text:
                                            found_key = key
                                            break
                                    
                                    if found_key:
                                        break
                                
                                # Çeviriyi güncelle
                                if found_key:
                                    child.config(text=self.get_text(found_key))
                    except tk.TclError:
                        # Widget'ta "text" özelliği yoksa sessizce devam et
                        continue

    def treeview_sort_column(self, column, reverse):
        """Sort the treeview content when a column header is clicked"""
        # OPTIMIZATION: Enhanced sorting algorithm for large file lists
        # This improved implementation uses better algorithms and data structures
        # to significantly speed up sorting operations
        
        # OPTIMIZATION: Progress bar and cursor change for large lists 
        item_count = len(self.file_tree.get_children(''))
        show_progress = item_count > 2000  # Show progress for moderately large lists
        use_optimized_sort = item_count > 5000  # Use highly optimized algorithm for very large lists
        
        if show_progress:
            self.root.config(cursor="watch")
            self.update_status(self.get_text("sorting_files"))
            self.progress_bar.start(10)
            # Allow UI to update
            self.root.update_idletasks()
        
        try:
            # OPTIMIZATION: Special handling for columns that need special sorting
            # Date columns need datetime parsing
            if column in ["created", "modified"]:
                # Try to handle different date formats
                try:
                    # Use a separate function to parse dates more efficiently
                    def parse_date(date_str):
                        try:
                            # Try common formats
                            for fmt in ["%Y-%m-%d %H:%M:%S", "%d.%m.%Y %H:%M:%S", "%m/%d/%Y %H:%M:%S"]:
                                try:
                                    return datetime.datetime.strptime(date_str, fmt)
                                except ValueError:
                                    continue
                            # If none of the formats match, use a simple string comparison
                            return date_str
                        except Exception:
                            return date_str
                    
                    # Use list comprehension for better performance
                    l = [(parse_date(self.file_tree.set(k, column)), k) for k in self.file_tree.get_children('')]
                except Exception:
                    # Fallback to string sorting if date parsing fails
                    l = [(self.file_tree.set(k, column), k) for k in self.file_tree.get_children('')]
            
            # Size column needs numeric conversion
            elif column == "size":
                # Cache the size conversion function for reuse
                def extract_size_bytes(size_text):
                    if "KB" in size_text:
                        return float(size_text.replace(" KB", "")) * 1024
                    elif "MB" in size_text:
                        return float(size_text.replace(" MB", "")) * 1024 * 1024
                    elif "GB" in size_text:
                        return float(size_text.replace(" GB", "")) * 1024 * 1024 * 1024
                    elif "B" in size_text:
                        return float(size_text.replace(" B", ""))
                    return 0
                
                # OPTIMIZATION: Pre-allocate the list to avoid resizing
                l = [(extract_size_bytes(self.file_tree.set(k, column)), k) for k in self.file_tree.get_children('')]
            
            # Default string sorting
            else:
                l = [(self.file_tree.set(k, column), k) for k in self.file_tree.get_children('')]
            
            # OPTIMIZATION: Use natural sort for filenames and paths
            if column in ["name", "path"]:
                # Natural sort for filenames (1, 2, 10 instead of 1, 10, 2)
                import re
                def natural_sort_key(s):
                    return [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', str(s[0]))]
                
                l.sort(key=natural_sort_key, reverse=reverse)
            else:
                # Regular sort for other columns
                l.sort(reverse=reverse)
            
            # OPTIMIZATION: Batch move items for better performance
            # First detach all items
            if item_count > 1000:
                # For very large lists, detach and reattach all at once
                items = [k for _, k in l]
                self.file_tree.detach(*self.file_tree.get_children(''))
                
                # Reattach in the new order
                for idx, item in enumerate(items):
                    self.file_tree.move(item, '', idx)
            else:
                # For smaller lists, just move items
                for index, (val, k) in enumerate(l):
                    self.file_tree.move(k, '', index)
            
            # Reverse sort next time
            self.file_tree.heading(column, command=lambda: self.treeview_sort_column(column, not reverse))
            
            # Update column headers to show sort indication
            for col in ["name", "extension", "path", "size", "created", "modified"]:
                if col == column:
                    if reverse:
                        self.file_tree.heading(col, text=f"▼ {self.get_text(f'file_{col}' if col != 'extension' else 'file_extension')}")
                    else:
                        self.file_tree.heading(col, text=f"▲ {self.get_text(f'file_{col}' if col != 'extension' else 'file_extension')}")
                else:
                    self.file_tree.heading(col, text=self.get_text(f"file_{col}" if col != 'extension' else "file_extension"))
        
        finally:
            # Always restore cursor and progress bar
            if show_progress:
                self.root.config(cursor="")
                self.progress_bar.stop()
                self.update_status(self.get_text("ready"))
    
    def set_view_mode(self, mode):
        """Set the view mode (list or preview)"""
        # Only change if we're switching modes
        if self.view_mode_var.get() == mode:
            return
            
        self.view_mode_var.set(mode)
        
        # Update button appearance for the view mode buttons
        if hasattr(self, 'list_view_btn') and hasattr(self, 'preview_view_btn'):
            # Update list view button
            self.list_view_btn.config(
                bg="#17a2b8" if mode == "list" else "#e9ecef",
                fg="white" if mode == "list" else "#495057"
            )
            
            # Update preview view button
            self.preview_view_btn.config(
                bg="#17a2b8" if mode == "preview" else "#e9ecef",
                fg="white" if mode == "preview" else "#495057"
            )
        
        # Switch view based on mode
        if mode == "list":
            # Show the normal file list view
            self._switch_to_list_view()
        else:  # preview mode
            # Show the preview view
            self._switch_to_preview_view()
    
    def _switch_to_list_view(self):
        """Switch to normal list view"""
        # Hide preview frame if it exists
        if hasattr(self, 'preview_frame') and self.preview_frame.winfo_exists():
            self.preview_frame.pack_forget()
                
        # Show the file tree
        self.file_tree.master.pack(fill=tk.BOTH, expand=True)
        
        # Update status
        self.update_status(self.get_text("ready"))
    
    def _switch_to_preview_view(self):
        """Switch to preview view mode"""
        # Hide the tree view
        self.file_tree.master.pack_forget()
        
        # Create preview container if needed
        if not hasattr(self, 'preview_frame') or not self.preview_frame.winfo_exists():
            # Main preview area - contains both thumbnail view and pagination
            self.preview_frame = tk.Frame(self.file_view_container, bg="#e9ecef")
            self.preview_frame.pack(fill=tk.BOTH, expand=True)
            
            # Add the main content area first (without pagination)
            # This ensures thumbnails take up most of the space
            self.thumbnails_area = tk.Frame(self.preview_frame, bg="#e9ecef")
            self.thumbnails_area.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
            
            # Add pagination frame at bottom - make it more visible with a subtle border and background
            self.pagination_frame = tk.Frame(
                self.preview_frame, 
                bg="#f0f0f0", 
                height=50,  # Higher height for better visibility
                bd=1,       # Light border
                relief=tk.GROOVE  # Subtle raised effect
            )
            self.pagination_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 5))
            self.pagination_frame.pack_propagate(False)  # Maintain fixed height
            
            # Previous page button - nicer styling
            self.prev_page_btn = tk.Button(
                self.pagination_frame,
                text="◄ " + self.get_text("prev_page"),
                command=self._go_to_prev_page,
                bg="#e0e0e0",
                fg="#495057",
                font=("Segoe UI", 9, "bold"),  # Bold font
                relief=tk.RAISED,
                padx=15,  # More padding
                pady=3,
                bd=1,
                state=tk.DISABLED
            )
            self.prev_page_btn.pack(side=tk.LEFT, padx=15, pady=8)  # More padding
            
            # Page info label - centered with better styling
            self.page_info_label = tk.Label(
                self.pagination_frame,
                text=f"{self.get_text('page')} 1/1",
                bg="#f0f0f0",
                fg="#495057",
                font=("Segoe UI", 10)  # Slightly larger font
            )
            self.page_info_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            # Next page button - matching style with prev button
            self.next_page_btn = tk.Button(
                self.pagination_frame,
                text=self.get_text("next_page") + " ►",
                command=self._go_to_next_page,
                bg="#e0e0e0",
                fg="#495057",
                font=("Segoe UI", 9, "bold"),  # Bold font
                relief=tk.RAISED,
                padx=15,  # More padding
                pady=3,
                bd=1,
                state=tk.DISABLED
            )
            self.next_page_btn.pack(side=tk.RIGHT, padx=15, pady=8)  # More padding
            
            # Create a canvas with scrollbar for thumbnails
            self.preview_canvas = tk.Canvas(self.thumbnails_area, bg="#e9ecef")
            preview_scrollbar = ttk.Scrollbar(self.thumbnails_area, orient=tk.VERTICAL, command=self.preview_canvas.yview)
            self.preview_canvas.config(yscrollcommand=preview_scrollbar.set)
            
            # Pack scrollbar and canvas
            preview_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            self.preview_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            # Create frame inside canvas for thumbnails
            self.thumbnail_container = tk.Frame(self.preview_canvas, bg="#e9ecef")
            self.preview_canvas.create_window((0, 0), window=self.thumbnail_container, anchor=tk.NW)
            
            # Configure canvas to resize with frame
            self.thumbnail_container.bind("<Configure>", lambda e: self.preview_canvas.configure(
                scrollregion=self.preview_canvas.bbox("all")
            ))
            
            # Bind mousewheel event for scrolling
            self.preview_canvas.bind_all("<MouseWheel>", lambda e: self.preview_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        else:
            # Show existing preview frame
            self.preview_frame.pack(fill=tk.BOTH, expand=True)
        
        # Update status
        self.update_status(self.get_text("preview_mode_active"))
        
        # Update with current files
        self._update_preview_panel()
    
    def _update_preview_panel(self):
        """Update the preview panel with current files"""
        # Only update if in preview mode
        if self.view_mode_var.get() != "preview":
            return
            
        # Check if we have files to display
        if not hasattr(self, 'filtered_files') or not self.filtered_files:
            # No files to display
            if hasattr(self, 'thumbnail_container'):
                for widget in self.thumbnail_container.winfo_children():
                    widget.destroy()
                    
                # Show message
                msg_label = tk.Label(
                    self.thumbnail_container,
                    text=self.get_text("no_preview_available"),
                    font=("Segoe UI", 12),
                    bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                    fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
                )
                msg_label.pack(pady=50)
                
                # Reset preview files list
                self.current_preview_files = []
            return
        
        # Filter only previewable files for performance optimization
        previewable_files = []
        
        # Count file types for statistics (only previewable types)
        self.file_type_stats = {}
        
        for file_info in self.filtered_files:
            if file_info.get("is_folder", False):
                continue
            
            # Get file path and extension
            file_path = ""
            if "path" in file_info and "name" in file_info:
                file_path = os.path.join(file_info["path"], file_info["name"])
            else:
                continue
                
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Only include previewable file extensions
            if file_ext in PREVIEWABLE_EXTENSIONS:
                # Add this file to previewable files
                previewable_files.append(file_info)
                
                # Normalize extension for statistics (remove dot and make jpg/jpeg consistent)
                ext = file_ext.replace(".", "")
                if ext == "jpeg":
                    ext = "jpg"
                    
                # Add to statistics
                if ext in self.file_type_stats:
                    self.file_type_stats[ext] += 1
                else:
                    self.file_type_stats[ext] = 1
        
        # Use only previewable files for the rest of the function
        self.previewable_files = previewable_files
            
        # Clear existing thumbnails
        if hasattr(self, 'thumbnail_container'):
            for widget in self.thumbnail_container.winfo_children():
                widget.destroy()
        
        # Start building thumbnails
        self.update_status(self.get_text("loading_preview"))
        
        # This code section is no longer needed since we create the pagination frame in _switch_to_preview_view
        # We'll just update the pagination text
        if hasattr(self, 'pagination_frame'):
            # Update pagination text based on current language
            self.prev_page_btn.config(text=f"◀ {self.get_text('prev_page')}")
            self.next_page_btn.config(text=f"{self.get_text('next_page')} ▶")
        
        # Get the current filtered files and rebuild the preview panel
        if hasattr(self, 'filtered_files') and self.filtered_files:
            # Store all files for preview navigation, but only those that can be previewed
            self.current_preview_files = []
            for file_info in self.filtered_files:
                # Only add files that can be previewed (not folders and only previewable file types)
                if not file_info.get("is_folder", False):
                    # Get file extension
                    file_ext = file_info.get("extension", "").lower().replace(".", "")
                    if not file_ext and "path" in file_info and "name" in file_info:
                        file_path = os.path.join(file_info["path"], file_info["name"])
                        file_ext = os.path.splitext(file_path)[1].lower().replace(".", "")
                    
                    # Only include files that we can actually preview
                    if file_ext in ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "pdf", "psd", "ai", "eps"]:
                        self.current_preview_files.append(file_info)
            
            # Initialize current preview index if needed
            if not hasattr(self, 'current_preview_index'):
                self.current_preview_index = -1
            
            # Store all files
            self.all_preview_files = self.filtered_files
            
            # Calculate total pages
            total_items = len(self.filtered_files)
            self.total_preview_pages = max(1, (total_items + self.preview_items_per_page - 1) // self.preview_items_per_page)
            
            # Adjust current page if needed
            if self.preview_page > self.total_preview_pages:
                self.preview_page = 1
                
            # Update page info
            self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")
            
            # Set button states based on current page
            if self.preview_page <= 1:
                self.prev_page_btn.config(state=tk.DISABLED)
            else:
                self.prev_page_btn.config(state=tk.NORMAL)
                
            if self.preview_page >= self.total_preview_pages:
                self.next_page_btn.config(state=tk.DISABLED)
            else:
                self.next_page_btn.config(state=tk.NORMAL)
                
            # Get current page items from previewable files if available, otherwise from filtered files
            if hasattr(self, 'previewable_files') and self.previewable_files:
                preview_source = self.previewable_files
                total_preview_items = len(self.previewable_files)
                self.total_preview_pages = max(1, (total_preview_items + self.preview_items_per_page - 1) // self.preview_items_per_page)
                
                # Adjust current page if needed
                if self.preview_page > self.total_preview_pages:
                    self.preview_page = 1
                    
                # Update page info
                self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")
                
                start_idx = (self.preview_page - 1) * self.preview_items_per_page
                end_idx = min(start_idx + self.preview_items_per_page, total_preview_items)
                current_page_files = preview_source[start_idx:end_idx]
                
                # Log for debugging
                logging.info(f"Showing preview page {self.preview_page}/{self.total_preview_pages}, items {start_idx+1}-{end_idx} of {total_preview_items} (optimized previewable files)")
            else:
                # Fallback to filtered files
                start_idx = (self.preview_page - 1) * self.preview_items_per_page
                end_idx = min(start_idx + self.preview_items_per_page, total_items)
                current_page_files = self.filtered_files[start_idx:end_idx]
                
                # Log for debugging
                logging.info(f"Showing preview page {self.preview_page}/{self.total_preview_pages}, items {start_idx+1}-{end_idx} of {total_items}")
            
            # Build preview with current page files
            self._build_preview_panel(current_page_files)
        
    def _build_preview_panel(self, files):
        """Build a preview panel showing file thumbnails"""
        # Only proceed if we're in preview mode
        if self.view_mode_var.get() != "preview":
            return
        
        # Initialize pagination attributes if not already set
        if not hasattr(self, 'preview_page'):
            self.preview_page = 1
            self.preview_items_per_page = 100  # Show 100 items per page (increased from 50)
            
        # Initialize or reset cancel flag
        self.cancel_flag = False
        self.enable_cancel_button()
        
        # Filtreleme durumunu izlemek için özel bir değişken ekliyoruz
        self.filtering_complete = False
        
        # "Filtreleme işlemleri devam ediyor" yazısını dönen simge ile göster
        self.update_status(self.get_text("filtering_in_progress") + " ⟳")
            
        # Update status
        self.update_status(self.get_text("loading_preview"))
        
        # OPTIMIZATION: Use batch processing for thumbnails to improve performance with large folders
        self.thumb_batch_size = self.preview_batch_size  # Use the batch size defined in initialization
        
        # Clear existing thumbnails
        if hasattr(self, 'thumbnail_container'):
            for widget in self.thumbnail_container.winfo_children():
                widget.destroy()
        
        # Create image references holder
        if not hasattr(self, 'preview_images'):
            self.preview_images = []
        else:
            self.preview_images.clear()
        
        # Determine preview sizes
        preview_width = 150  # Width for preview images
        preview_height = 150  # Height for preview images
        padding = 10  # Padding between thumbnails
        
        # Increase height for buttons
        preview_frame_height = preview_height + 60  # Add extra height for filename and buttons
        
        # Calculate max columns based on container width
        container_width = self.file_view_container.winfo_width()
        if container_width > 0:
            # Calculate how many thumbnails fit in the available width
            # Each thumbnail takes width + padding on each side
            max_columns = max(1, container_width // (preview_width + padding * 2))
        else:
            # Default if width not yet available
            max_columns = 4  # Default number of columns
        
        # Use the files directly as they should already be filtered for previewable types
        # in _update_preview_panel
        preview_files = files
        
        # Set up progress tracking
        total_files = len(preview_files)
        self.progress_bar["value"] = 0
        self.progress_bar["maximum"] = 100
        
        # If no previewable files, show message
        if not preview_files:
            msg_label = tk.Label(
                self.thumbnail_container,
                text=self.get_text("no_preview_available"),
                font=("Segoe UI", 12),
                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"]
            )
            msg_label.pack(pady=50)
            self.update_status(self.get_text("preview_mode_active"))
            return
            
        # Process thumbnails in a separate thread
        def process_thumbnails():
            try:
                # Dönen simge göster
                self.spinner_chars = ["⟳", "⟲", "↻", "↺"]
                self.spinner_index = 0
                
                # Simge güncelleme fonksiyonu
                def update_spinner():
                    if self.cancel_flag or hasattr(self, 'filtering_complete') and self.filtering_complete:
                        return
                        
                    self.spinner_index = (self.spinner_index + 1) % len(self.spinner_chars)
                    spinner_char = self.spinner_chars[self.spinner_index]
                    self.root.after(0, lambda: self.update_status(f"{self.get_text('filtering_in_progress')} {spinner_char}"))
                    
                    # 200ms sonra tekrar güncelle
                    self.root.after(200, update_spinner)
                
                # Dönen simgeyi başlat
                update_spinner()
                
                # Create a grid for thumbnails
                row = 0
                col = 0
                
                # Initialize cancel flag for this operation specifically
                self.cancel_flag = False
                self.root.after(0, self.enable_cancel_button)
                
                # Define a check for cancellation in the main thread
                def check_cancel():
                    return hasattr(self, 'cancel_flag') and self.cancel_flag
                
                # OPTIMIZATION: Process files in batches for better performance
                batches = [preview_files[i:i+self.thumb_batch_size] for i in range(0, len(preview_files), self.thumb_batch_size)]
                batch_count = 0
                
                # Process each batch with cancellation checks
                for batch_idx, batch in enumerate(batches):
                    # Frequently check if operation is cancelled
                    if check_cancel():
                        self.root.after(0, lambda: self.update_status(self.get_text("operation_cancelled")))
                        self.root.after(0, lambda: self.progress_bar.config(value=0))
                        # Provide visual feedback that cancellation is in progress
                        self.root.after(0, lambda: self.cancel_btn.config(text=self.get_text("cancelling")))
                        self.root.after(0, self.disable_cancel_button)
                        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))
                        return
                    
                    # Process each file in the current batch
                    batch_files = []
                    
                    # Update progress for this batch
                    processed_count = batch_idx * self.thumb_batch_size
                    batch_progress = processed_count / total_files * 100
                    self.root.after(0, lambda p=batch_progress: self.progress_bar.config(value=p))
                    self.root.after(0, lambda c=processed_count, t=total_files: 
                                  self.update_status(f"{self.get_text('generating_previews')} ({c}/{t})"))
                    
                    # Process all files in this batch
                    for file_info in batch:
                        file_name = file_info.get("name", "")
                        file_path = os.path.join(file_info.get("path", ""), file_name)
                        extension = file_info.get("extension", "").lower()
                        
                        # Store thumbnail info for grid placement
                        batch_files.append((file_name, file_path, extension, file_info))
                    
                    # Create thumbnail frame with more details
                    def create_frame(row, col):
                        # Increased height to fit more file details
                        frame = tk.Frame(
                            self.thumbnail_container,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            width=preview_width + padding,
                            height=preview_frame_height + 40,  # Increased height for additional details
                            highlightbackground=LIGHT_MODE_COLORS["border"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["border"],
                            highlightthickness=1
                        )
                        frame.grid(row=row, column=col, padx=padding//2, pady=padding//2, sticky="nsew")
                        frame.grid_propagate(False)  # Keep frame size fixed
                        return frame
                        
                    # OPTIMIZATION: Process thumbnails for all files in the batch
                    thumbnails = []
                    errors = []
                    
                    # Track current position for grid layout
                    idx = processed_count
                    row = idx // max_columns
                    col = idx % max_columns
                    
                    # Process each file in the batch to generate thumbnails
                    for file_name, file_path, extension, file_info in batch_files:
                        try:
                            # OPTIMIZATION: Use lower resolution previews for the general preview page
                            # This significantly improves performance while maintaining usability
                            preview_width_reduced = int(preview_width * 0.7)  # 70% of original size for overview
                            preview_height_reduced = int(preview_height * 0.7)  # 70% of original size for overview
                            preview_img = self._create_file_preview(file_path, preview_width_reduced, preview_height_reduced)
                            
                            # Calculate grid position
                            idx_position = processed_count + len(thumbnails) + len(errors)
                            r = idx_position // max_columns
                            c = idx_position % max_columns
                            
                            if preview_img:
                                # Store thumbnail info for UI update
                                thumbnails.append((r, c, preview_img, file_path, file_name, file_info.get("size", 0)))
                            else:
                                # Store error info for UI update
                                errors.append((r, c, file_path))
                                
                        except Exception as e:
                            logging.error(f"Error creating thumbnail for {file_path}: {str(e)}")
                            
                            # Calculate grid position for error display
                            idx_position = processed_count + len(thumbnails) + len(errors)
                            r = idx_position // max_columns
                            c = idx_position % max_columns
                            
                            # Store error info for UI update
                            errors.append((r, c, file_path))
                    
                    # Add all thumbnails to UI in main thread with improved details
                    def add_thumbnail(r, c, img, path, name, size):
                        frame = create_frame(r, c)
                        
                        # File image/icon at the top
                        if img:
                            # Keep reference
                            self.preview_images.append(img)
                            
                            # Create label for image
                            img_label = tk.Label(frame, image=img, 
                                                bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"])
                            img_label.place(relx=0.5, rely=0.35, anchor=tk.CENTER)  # Moved up to make room for details
                            
                            # Add click event - get file index from current preview files
                            file_index = next((i for i, f in enumerate(self.current_preview_files) if f.get("path") == os.path.dirname(path) and f.get("name") == os.path.basename(path)), -1)
                            img_label.bind("<Button-1>", lambda e, p=path, idx=file_index: self.create_file_preview_window(p, idx))
                            # Add right-click context menu
                            img_label.bind("<Button-3>", lambda e, p=path: self.show_preview_context_menu(e, p))
                        else:
                            # No preview, show extension icon with improved styling
                            ext = os.path.splitext(path)[1].lower().replace(".", "")
                            ext_label = tk.Label(
                                frame,
                                text=ext.upper(),
                                font=("Segoe UI", 16, "bold"),
                                bg="#4285F4",  # Google blue for better visibility
                                fg="white",
                                width=4,
                                height=2,
                                relief=tk.RAISED  # 3D effect
                            )
                            ext_label.place(relx=0.5, rely=0.35, anchor=tk.CENTER)  # Moved up
                            
                            # Add click event - get file index from current preview files
                            file_index = next((i for i, f in enumerate(self.current_preview_files) if f.get("path") == os.path.dirname(path) and f.get("name") == os.path.basename(path)), -1)
                            ext_label.bind("<Button-1>", lambda e, p=path, idx=file_index: self.create_file_preview_window(p, idx))
                            # Add right-click context menu
                            ext_label.bind("<Button-3>", lambda e, p=path: self.show_preview_context_menu(e, p))
                        
                        # File details container to organize information
                        details_frame = tk.Frame(
                            frame, 
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
                        )
                        details_frame.place(relx=0.5, rely=0.75, anchor=tk.CENTER, width=preview_width-10, height=70)
                        
                        # Add filename with better styling
                        display_name = name
                        if len(name) > 20:
                            display_name = name[:17] + "..."
                            
                        name_label = tk.Label(
                            details_frame,
                            text=display_name,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["text"],
                            font=("Segoe UI", 9, "bold"),  # Bold for emphasis
                            wraplength=preview_width-10
                        )
                        name_label.pack(pady=(0, 2))
                        
                        # Add click event to filename - get file index from current preview files
                        file_index = next((i for i, f in enumerate(self.current_preview_files) if f.get("path") == os.path.dirname(path) and f.get("name") == os.path.basename(path)), -1)
                        name_label.bind("<Button-1>", lambda e, p=path, idx=file_index: self.create_file_preview_window(p, idx))
                        
                        # Add file size information
                        size_str = self.format_file_size(size)
                        size_label = tk.Label(
                            details_frame,
                            text=size_str,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8)
                        )
                        size_label.pack(pady=(0, 2))
                        
                        # Add file extension info
                        ext = os.path.splitext(path)[1].lower()
                        ext_info_label = tk.Label(
                            details_frame,
                            text=ext,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8)
                        )
                        ext_info_label.pack(pady=(0, 2))
                    
                    # Function to show error thumbnail with improved styling
                    def show_error(r, c, path):
                        frame = create_frame(r, c)
                        
                        # Improved error icon
                        error_icon_frame = tk.Frame(frame, bg="#dc3545", width=60, height=60, relief=tk.RAISED, bd=2)
                        error_icon_frame.place(relx=0.5, rely=0.35, anchor=tk.CENTER)
                        
                        err_label = tk.Label(
                            error_icon_frame,
                            text="!",
                            font=("Segoe UI", 20, "bold"),
                            bg="#dc3545",
                            fg="white",
                            width=3,
                            height=2
                        )
                        err_label.pack(fill=tk.BOTH, expand=True)
                        
                        # File details container similar to normal thumbnails
                        details_frame = tk.Frame(
                            frame, 
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"]
                        )
                        details_frame.place(relx=0.5, rely=0.75, anchor=tk.CENTER, width=preview_width-10, height=70)
                        
                        # Add filename with error indication
                        name = os.path.basename(path)
                        if len(name) > 20:
                            name = name[:17] + "..."
                            
                        name_label = tk.Label(
                            details_frame,
                            text=name,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg="#dc3545",  # Red text for error (always red regardless of theme)
                            font=("Segoe UI", 9, "bold"),
                            wraplength=preview_width-10
                        )
                        name_label.pack(pady=(0, 2))
                        
                        # Add error message
                        error_label = tk.Label(
                            details_frame,
                            text="Önizleme kullanılamıyor",  # Preview not available
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8),
                            wraplength=preview_width-10
                        )
                        error_label.pack(pady=(0, 2))
                        
                        # Add file extension
                        ext = os.path.splitext(path)[1].lower()
                        ext_label = tk.Label(
                            details_frame,
                            text=ext,
                            bg=LIGHT_MODE_COLORS["bg"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["bg"],
                            fg=LIGHT_MODE_COLORS["secondary_text"] if not self.is_dark_mode.get() else DARK_MODE_COLORS["secondary_text"],
                            font=("Segoe UI", 8)
                        )
                        ext_label.pack(pady=(0, 2))
                    
                    # Update UI with all thumbnails in this batch
                    for r, c, img, path, name, size in thumbnails:
                        self.root.after(0, lambda r=r, c=c, img=img, path=path, name=name, size=size: 
                                       add_thumbnail(r, c, img, path, name, size))
                    
                    # Update UI with all errors in this batch
                    for r, c, path in errors:
                        self.root.after(0, lambda r=r, c=c, path=path: show_error(r, c, path))
                    
                    # İlk batch yüklendiyse (ilk sayfa görüntülenecek durumdaysa)
                    # istatistikleri göster ve filtreleme işleminin tamamlandığını belirt
                    if batch_idx == 0:
                        self.filtering_complete = True
                        stats_message = self._calculate_file_type_statistics()
                        if stats_message:
                            self.root.after(0, lambda msg=stats_message: self.status_var.set(msg))
                            logging.info(f"Displaying initial file stats: {stats_message}")
                    
                    # Update progress after processing this batch
                    batch_count += 1
                    
                    # Update grid position for next batch (if any)
                    processed_count += len(batch_files)
                    row = processed_count // max_columns
                    col = processed_count % max_columns
                    
                    # Brief pause to keep UI responsive
                    self.root.update_idletasks()
                    time.sleep(0.05)  # Small delay to prevent UI freezing
                
                # Configure the scroll region
                self.root.after(0, lambda: self.thumbnail_container.update_idletasks())
                self.root.after(0, lambda: self.preview_canvas.config(scrollregion=self.preview_canvas.bbox("all")))
                
                # Filtreleme ve önizleme yüklenmesi tamamlandı, bayrak ayarla
                self.filtering_complete = True
                
                # İlerleme çubuğunu sıfırla
                self.root.after(0, lambda: self.progress_bar.config(value=0))
                
                # İstatistikleri hesapla ve göster
                stats_message = self._calculate_file_type_statistics()
                if stats_message:
                    self.root.after(0, lambda msg=stats_message: self.status_var.set(msg))
                    logging.info(f"Displaying file stats when thumbnails loaded: {stats_message}")
                else:
                    # Eğer istatistikler hesaplanamazsa, varsayılan durumu göster
                    self.root.after(0, lambda: self.update_status(self.get_text("preview_mode_active")))
                
            except Exception as e:
                logging.error(f"Error building preview panel: {str(e)}")
                self.root.after(0, lambda: self.update_status(f"Error building preview: {str(e)}"))
                self.filtering_complete = True  # Hata olsa bile tamamlandı olarak işaretle
        
        # Start processing in background thread
        threading.Thread(target=process_thumbnails, daemon=True).start()
    
    def _create_eps_preview(self, file_path, max_width, max_height):
        """Specialized function to create a preview for EPS files.
        Uses multiple methods and temporary files to ensure success"""
        
        # Fallback function to create placeholder
        def create_eps_placeholder():
            color = "#8BC34A"  # Green for EPS
            img = Image.new('RGB', (max_width, max_height), color)
            draw = ImageDraw.Draw(img)
            # Add a border and text
            draw.rectangle([(0, 0), (max_width-1, max_height-1)], outline="white", width=2)
            draw.text((max_width//2, max_height//2), "EPS", fill="white", anchor="mm")
            return ImageTk.PhotoImage(img)
        
        # We'll try several methods in sequence, from most reliable to least reliable
        preview_image = None
        
        try:
            # Create a temporary directory for conversion files
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_pdf_path = os.path.join(temp_dir, "temp_eps_preview.pdf")
                
                # METHOD 1: Directly use PIL to open EPS - works with small EPS files
                try:
                    # Set a timeout to prevent hanging on large files
                    img = Image.open(file_path)
                    # Use a smaller target size to prevent decompression bombs
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                    return preview_image
                except Exception as e:
                    logging.info(f"Direct EPS loading failed: {str(e)}")
                
                # METHOD 2: Use pdf2image with specific parameters
                try:
                    # Try to convert EPS directly to image
                    from pdf2image import convert_from_path
                    
                    # Define poppler path to ensure we can find the tools
                    poppler_path = '/nix/store/1f2vbia1rg1rh5cs0ii49v3hln9i36rv-poppler-utils-24.02.0/bin'
                    
                    # Use pdftocairo which often handles EPS better than pdftoppm
                    images = convert_from_path(
                        file_path, 
                        first_page=1, 
                        last_page=1,
                        dpi=72,  # Lower DPI to prevent large image generation
                        size=(max_width, max_height),
                        use_cropbox=True,
                        fmt='ppm',  # Use PPM format which is more reliable
                        poppler_path=poppler_path,
                        use_pdftocairo=True,  # Try pdftocairo instead of pdftoppm
                        timeout=10  # Increase timeout for complex EPS files
                    )
                    
                    if images and len(images) > 0:
                        img = images[0]
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image
                except Exception as e:
                    logging.info(f"pdf2image EPS conversion failed: {str(e)}")
                
                # METHOD 3: Use PyMuPDF (fitz) to open the EPS directly
                try:
                    pdf_doc = fitz.open(file_path)
                    if pdf_doc.page_count > 0:
                        page = pdf_doc[0]
                        # Use a lower zoom factor to prevent large images
                        pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
                        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        pdf_doc.close()
                        return preview_image
                    pdf_doc.close()
                except Exception as e:
                    logging.info(f"PyMuPDF EPS loading failed: {str(e)}")
                
                # METHOD 4: Try using ImageMagick to convert EPS to PNG
                try:
                    # Use ImageMagick convert command
                    convert_path = '/nix/store/1izdxwml9nsifjrh53rdfiglhjmrnx2s-imagemagick-7.1.1-32/bin/convert'
                    
                    # Create temporary output image path
                    temp_image_path = os.path.join(temp_dir, "temp_eps_preview.png")
                    
                    # Convert EPS to PNG using ImageMagick with density parameter for better quality
                    subprocess.run(
                        [convert_path, '-density', '150', '-background', 'white', '-flatten', 
                         file_path, temp_image_path],
                        stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10
                    )
                    
                    # Check if image was created successfully
                    if os.path.exists(temp_image_path) and os.path.getsize(temp_image_path) > 0:
                        img = Image.open(temp_image_path)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image
                        
                    # If that failed, try a simpler conversion method
                    subprocess.run(
                        [convert_path, file_path, temp_image_path],
                        stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10
                    )
                    
                    # Check if image was created with the simpler method
                    if os.path.exists(temp_image_path) and os.path.getsize(temp_image_path) > 0:
                        img = Image.open(temp_image_path)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image
                        
                except Exception as e:
                    logging.info(f"ImageMagick EPS conversion failed: {str(e)}")
        
        except Exception as e:
            logging.error(f"All EPS preview methods failed: {str(e)}")
        
        # If all methods fail or exceptions occur, create a placeholder
        return create_eps_placeholder()
    
    def _create_file_preview(self, file_path, max_width=150, max_height=150):
        """Create a thumbnail preview for a file based on its type"""
        # Normalize file path to avoid Windows/Unix path issues
        file_path = os.path.normpath(file_path)
        
        # Check the file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # OPTIMIZATION: Enhanced LRU cache for thumbnails using multithreading
        cache_key = f"{file_path}_{max_width}_{max_height}"
        
        # Ensure cache structures are initialized (thread-safe)
        if not hasattr(self, 'preview_cache'):
            self.preview_cache = {}
            
        if not hasattr(self, 'preview_cache_keys'):
            self.preview_cache_keys = []
        
        if not hasattr(self, 'preview_cache_lock'):
            self.preview_cache_lock = threading.RLock()
            
        # Check if we have the preview in cache (thread-safe)
        with self.preview_cache_lock:
            if cache_key in self.preview_cache:
                # Update LRU order (move to end of list to mark as recently used)
                if cache_key in self.preview_cache_keys:
                    self.preview_cache_keys.remove(cache_key)
                
                # Add to end (most recently used position)
                self.preview_cache_keys.append(cache_key)
                return self.preview_cache[cache_key]
            
        # Enforce max cache size with LRU eviction policy
        if len(self.preview_cache) >= self.max_preview_cache_size:
            # Remove 20% of least recently used entries
            items_to_remove = max(1, int(self.max_preview_cache_size * 0.2))
            
            # Remove from the beginning of the list (least recently used)
            for _ in range(min(items_to_remove, len(self.preview_cache_keys))):
                if self.preview_cache_keys:
                    oldest_key = self.preview_cache_keys.pop(0)  # Remove and return first item
                    if oldest_key in self.preview_cache:
                        del self.preview_cache[oldest_key]
                        
            # Log cache cleanup for debugging
            logging.info(f"LRU cache cleanup: removed {items_to_remove} cached thumbnails")
            
        preview_image = None
        
        try:
            # Image files - expanded with more formats
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.webp', '.svg', '.ico', '.heic', '.raw', '.cr2', '.nef', '.dng', '.arw']:
                try:
                    # Open and resize the image
                    img = Image.open(file_path)
                    img.thumbnail((max_width, max_height), get_pil_resize_method())
                    preview_image = ImageTk.PhotoImage(img)
                except Exception as e:
                    logging.error(f"Error creating preview for {file_path}: {str(e)}")
                    # Create a placeholder for failed image with file extension
                    img = Image.new("RGB", (max_width, max_height), color="#f0f0f0")
                    draw = ImageDraw.Draw(img)
                    draw.rectangle([10, 10, max_width-10, max_height-10], outline="#dc3545", width=2)
                    # Show only extension without dot and make it uppercase (with fallback to IMG if too long)
                    ext_text = file_ext.upper()[1:] if len(file_ext[1:]) <= 5 else "IMG"
                    draw.text((max_width//2, max_height//2), ext_text, fill="#dc3545", anchor="mm")
                    preview_image = ImageTk.PhotoImage(img)
                
            # PDF files
            elif file_ext == '.pdf':
                # Get the first page of PDF
                try:
                    pdf_doc = fitz.open(file_path)
                    if pdf_doc.page_count > 0:
                        page = pdf_doc[0]
                        pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
                        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                    pdf_doc.close()
                except Exception as e:
                    logging.error(f"Error with PyMuPDF for {file_path}: {str(e)}")
                    # Fall back to pdf2image if fitz fails
                    try:
                        # Define poppler path to ensure we can find the tools
                        poppler_path = '/nix/store/1f2vbia1rg1rh5cs0ii49v3hln9i36rv-poppler-utils-24.02.0/bin'
                        
                        # Use pdftocairo which often produces better quality
                        images = pdf2image.convert_from_path(
                            file_path, 
                            first_page=1, 
                            last_page=1, 
                            size=(max_width, max_height),
                            poppler_path=poppler_path,
                            use_pdftocairo=True
                        )
                        if images:
                            preview_image = ImageTk.PhotoImage(images[0])
                    except Exception as e2:
                        logging.error(f"Error with pdf2image for {file_path}: {str(e2)}")
                        # Create a placeholder
                        preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", "PDF")
            
            # EPS files - use our specialized function
            elif file_ext == '.eps':
                preview_image = self._create_eps_preview(file_path, max_width, max_height)
                    
            # Design files (PSD, AI)
            elif file_ext in ['.psd', '.ai']:
                # Set default placeholder color based on file type
                color = "#1976D2" if file_ext == '.psd' else "#FF5722"  # Blue for PSD, Orange for AI
                file_type = file_ext[1:].upper()
                
                try:
                    # Try to open PSD files with PIL
                    if file_ext == '.psd':
                        img = Image.open(file_path)
                        img.thumbnail((max_width, max_height), get_pil_resize_method())
                        preview_image = ImageTk.PhotoImage(img)
                        return preview_image
                    # Try to open AI files with PyMuPDF (they're often PDF compatible)
                    elif file_ext == '.ai':
                        pdf_doc = fitz.open(file_path)
                        if pdf_doc.page_count > 0:
                            page = pdf_doc[0]
                            pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
                            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                            img.thumbnail((max_width, max_height), get_pil_resize_method())
                            preview_image = ImageTk.PhotoImage(img)
                            pdf_doc.close()
                            return preview_image
                        pdf_doc.close()
                except Exception as e:
                    logging.error(f"Error with design file {file_path}: {str(e)}")
                
                # If we get here, create a placeholder icon
                if not preview_image:
                    preview_image = self._create_styled_icon(max_width, max_height, color, file_type)
                    
            # Office documents - Word (expanded with more formats)
            elif file_ext in ['.doc', '.docx', '.dot', '.dotx', '.dotm', '.rtf', '.odt', '.wpd']:
                try:
                    # Try to generate a thumbnail from Word using docx, if available
                    if file_ext == '.docx' and 'docx' in sys.modules:
                        try:
                            # This is just for potential future expansion - we'll use placeholders for now
                            pass
                        except Exception as word_e:
                            logging.info(f"Failed to generate Word preview: {str(word_e)}")
                    
                    # Create a nice styled Word icon
                    if file_ext in ['.rtf', '.odt', '.wpd']:
                        # Label with the file extension for non-Word formats
                        word_label = file_ext[1:].upper()
                    else:
                        # Just use DOC for all Word formats
                        word_label = "DOC"
                    
                    preview_image = self._create_styled_icon(max_width, max_height, "#2B579A", word_label)  # Word blue
                except Exception as e:
                    logging.error(f"Error with document file {file_path}: {str(e)}")
                    # Fallback icon
                    word_label = file_ext[1:].upper() if len(file_ext) <= 5 else "DOC"
                    preview_image = self._create_styled_icon(max_width, max_height, "#2B579A", word_label)
            
            # Office documents - Excel (expanded with more formats)
            elif file_ext in ['.xls', '.xlsx', '.xlsm', '.xlsb', '.xlt', '.xltx', '.xltm', '.csv', '.ods', '.tsv', '.numbers']:
                try:
                    # For CSV/TSV files, we could potentially show a preview of the data
                    if file_ext in ['.csv', '.tsv']:
                        # Just a placeholder for now
                        pass
                    
                    # Create a nice styled Excel icon
                    if file_ext in ['.csv', '.tsv', '.ods', '.numbers']:
                        # Label with file extension for non-Excel formats
                        excel_label = file_ext[1:].upper()
                    else:
                        # Just use XLS for all Excel formats
                        excel_label = "XLS"
                    
                    preview_image = self._create_styled_icon(max_width, max_height, "#217346", excel_label)  # Excel green
                except Exception as e:
                    logging.error(f"Error with spreadsheet file {file_path}: {str(e)}")
                    # Fallback icon
                    excel_label = file_ext[1:].upper() if len(file_ext) <= 5 else "XLS"
                    preview_image = self._create_styled_icon(max_width, max_height, "#217346", excel_label)
            
            # Office documents - PowerPoint (expanded with more formats)
            elif file_ext in ['.ppt', '.pptx', '.pptm', '.pps', '.ppsx', '.ppsm', '.pot', '.potx', '.potm', '.odp', '.key']:
                try:
                    # Create a nice styled PowerPoint icon
                    if file_ext in ['.odp', '.key']:
                        # Label with file extension for non-PowerPoint formats
                        ppt_label = file_ext[1:].upper()
                    else:
                        # Just use PPT for all PowerPoint formats
                        ppt_label = "PPT"
                    
                    preview_image = self._create_styled_icon(max_width, max_height, "#D24726", ppt_label)  # PowerPoint orange
                except Exception as e:
                    logging.error(f"Error with presentation file {file_path}: {str(e)}")
                    # Fallback icon
                    ppt_label = file_ext[1:].upper() if len(file_ext) <= 5 else "PPT"
                    preview_image = self._create_styled_icon(max_width, max_height, "#D24726", ppt_label)
            
            # Text and code files (expanded with more formats)
            elif file_ext in ['.txt', '.md', '.json', '.xml', '.html', '.htm', '.css', '.js', '.jsx', '.ts', '.tsx', 
                           '.py', '.java', '.c', '.cpp', '.cs', '.h', '.hpp', '.rb', '.php', '.swift', '.go', '.rs',
                           '.pl', '.lua', '.r', '.sh', '.bat', '.ps1', '.yaml', '.yml', '.toml', '.ini', '.cfg',
                           '.conf', '.log', '.sql', '.asm', '.tex', '.srt', '.vtt', '.csv', '.m', '.f90', '.kt']:
                try:
                    # Group files into categories with similar colors
                    if file_ext in ['.html', '.htm', '.xml', '.jsx', '.tsx']:
                        # Markup files - light blue
                        color = "#03A9F4"
                    elif file_ext in ['.js', '.ts', '.py', '.rb', '.php', '.swift', '.java', '.cs', '.go', '.rs', '.kt']:
                        # Popular programming languages - indigo
                        color = "#3F51B5"
                    elif file_ext in ['.c', '.cpp', '.h', '.hpp', '.asm', '.m', '.f90']:
                        # Systems programming - deep blue
                        color = "#1A237E"
                    elif file_ext in ['.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf']:
                        # Config files - teal
                        color = "#009688"
                    elif file_ext in ['.log', '.txt', '.md', '.csv', '.srt', '.vtt']:
                        # Plain text files - gray
                        color = "#607D8B"
                    else:
                        # Other text files - blue-grey
                        color = "#607D8B"
                    
                    # We could potentially show a preview of text content, but for now just an icon
                    txt_label = file_ext[1:].upper() if len(file_ext) <= 5 else "TXT"
                    preview_image = self._create_styled_icon(max_width, max_height, color, txt_label)
                except Exception as e:
                    logging.error(f"Error with text file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#607D8B", "TXT")
            
            # Archive files (expanded with more formats)
            elif file_ext in ['.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz', '.tgz', '.lzma', '.iso', '.cab', 
                           '.dmg', '.img', '.jar', '.war', '.ear', '.bzip2', '.tbz2', '.gzip', '.z', '.lz', '.lz4']:
                try:
                    # Group by archive type
                    if file_ext in ['.iso', '.dmg', '.img']:
                        # Disk images - amber-red
                        color = "#FF8F00"
                        archive_label = file_ext[1:].upper() if len(file_ext) <= 5 else "DISK"
                    elif file_ext in ['.jar', '.war', '.ear']:
                        # Java archives - red
                        color = "#D32F2F"
                        archive_label = file_ext[1:].upper() if len(file_ext) <= 5 else "JAR"
                    else:
                        # Regular archives - amber
                        color = "#FFC107"
                        archive_label = file_ext[1:].upper() if len(file_ext) <= 5 else "ZIP"
                    
                    # Create a styled icon for archives
                    preview_image = self._create_styled_icon(max_width, max_height, color, archive_label)
                except Exception as e:
                    logging.error(f"Error with archive file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#FFC107", "ZIP")
                    
            # Executable and installable files (expanded with more formats)
            elif file_ext in ['.exe', '.msi', '.bat', '.cmd', '.ps1', '.sh', '.bash', '.app', '.run', '.bin', '.deb', '.rpm', 
                           '.dmg', '.pkg', '.appimage', '.apk', '.ipa', '.xap', '.msix', '.dll', '.so', '.dylib', '.com', '.vbs']:
                try:
                    # Group by platform/type for better color coding
                    if file_ext in ['.sh', '.bash', '.run', '.bin']:
                        # Unix executables - deep red
                        color = "#B71C1C"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "UNIX"
                    elif file_ext in ['.app', '.dmg', '.pkg']:
                        # macOS executables - dark red
                        color = "#C62828"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "MAC"
                    elif file_ext in ['.deb', '.rpm', '.appimage']:
                        # Linux packages - brick red
                        color = "#D32F2F"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "LINUX"
                    elif file_ext in ['.apk', '.ipa', '.xap', '.msix']:
                        # Mobile apps - light red
                        color = "#E53935"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "MOBILE"
                    elif file_ext in ['.dll', '.so', '.dylib']:
                        # Library files - orange-red
                        color = "#E64A19"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "LIB"
                    elif file_ext in ['.bat', '.cmd', '.ps1', '.vbs', '.com']:
                        # Scripts and command files - orange
                        color = "#EF6C00"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "SCRIPT"
                    else:
                        # Windows executables - standard red
                        color = "#F44336"
                        exe_label = file_ext[1:].upper() if len(file_ext) <= 5 else "EXE"
                    
                    # Create a styled icon for executables
                    preview_image = self._create_styled_icon(max_width, max_height, color, exe_label)
                except Exception as e:
                    logging.error(f"Error with executable file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#F44336", "EXE")
                    
            # Audio files - expanded with more formats
            elif file_ext in ['.mp3', '.wav', '.ogg', '.flac', '.aac', '.wma', '.m4a', '.aiff', '.ape', '.midi', '.mid', '.amr', '.opus', '.alac', '.aif']:
                try:
                    # Create a styled icon for audio files
                    audio_label = file_ext[1:].upper() if len(file_ext) <= 5 else "AUD"
                    preview_image = self._create_styled_icon(max_width, max_height, "#9C27B0", audio_label)  # Purple
                except Exception as e:
                    logging.error(f"Error with audio file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#9C27B0", "AUD")
                    
            # Video files - expanded with more formats
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg', '.3gp', '.vob', '.ts', '.mts', '.m2ts', '.divx', '.asf', '.ogv']:
                try:
                    # Create a styled icon for video files
                    video_label = file_ext[1:].upper() if len(file_ext) <= 5 else "VID"
                    preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", video_label)  # Deep orange
                except Exception as e:
                    logging.error(f"Error with video file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#FF5722", "VID")
                    
            # Font files
            elif file_ext in ['.ttf', '.otf', '.woff', '.woff2', '.eot']:
                try:
                    # Create a styled icon for font files
                    font_label = file_ext[1:].upper() if len(file_ext) <= 5 else "FONT"
                    preview_image = self._create_styled_icon(max_width, max_height, "#009688", font_label)  # Teal
                except Exception as e:
                    logging.error(f"Error with font file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#009688", "FONT")
                    
            # Database files
            elif file_ext in ['.db', '.sqlite', '.sqlite3', '.mdb', '.accdb', '.sql', '.dbf']:
                try:
                    # Create a styled icon for database files
                    db_label = file_ext[1:].upper() if len(file_ext) <= 5 else "DB"
                    preview_image = self._create_styled_icon(max_width, max_height, "#1565C0", db_label)  # Dark blue
                except Exception as e:
                    logging.error(f"Error with database file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#1565C0", "DB")
                    
            # E-book files
            elif file_ext in ['.epub', '.mobi', '.azw', '.azw3', '.fb2', '.cbz', '.cbr']:
                try:
                    # Create a styled icon for e-book files
                    book_label = file_ext[1:].upper() if len(file_ext) <= 5 else "EBOOK"
                    preview_image = self._create_styled_icon(max_width, max_height, "#4CAF50", book_label)  # Green
                except Exception as e:
                    logging.error(f"Error with e-book file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#4CAF50", "EBOOK")
                    
            # 3D and CAD files
            elif file_ext in ['.obj', '.stl', '.fbx', '.blend', '.3ds', '.dae', '.dwg', '.dxf', '.max', '.c4d']:
                try:
                    # Create a styled icon for 3D files
                    model_label = file_ext[1:].upper() if len(file_ext) <= 5 else "3D"
                    preview_image = self._create_styled_icon(max_width, max_height, "#673AB7", model_label)  # Deep purple
                except Exception as e:
                    logging.error(f"Error with 3D/CAD file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#673AB7", "3D")
                    
            # GIS and Map files
            elif file_ext in ['.shp', '.kml', '.kmz', '.gpx', '.geojson', '.osm', '.mbtiles', '.dem', '.tiff', '.asc']:
                try:
                    # Create a styled icon for GIS files
                    gis_label = file_ext[1:].upper() if len(file_ext) <= 5 else "GIS"
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", gis_label)  # Indigo
                except Exception as e:
                    logging.error(f"Error with GIS file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#3F51B5", "GIS")
                    
            # Certificate and key files
            elif file_ext in ['.pem', '.crt', '.cer', '.key', '.p12', '.pfx']:
                try:
                    # Create a styled icon for certificate files
                    cert_label = file_ext[1:].upper() if len(file_ext) <= 5 else "CERT"
                    preview_image = self._create_styled_icon(max_width, max_height, "#795548", cert_label)  # Brown
                except Exception as e:
                    logging.error(f"Error with certificate file {file_path}: {str(e)}")
                    preview_image = self._create_styled_icon(max_width, max_height, "#795548", "CERT")
            
            # If no specific handler, create a generic icon
            elif not preview_image:
                # Get just the extension without the dot
                ext = file_ext[1:].upper() if len(file_ext) > 1 else "?"
                # Truncate if too long
                if len(ext) > 5:
                    ext = ext[:4] + "…"
                preview_image = self._create_styled_icon(max_width, max_height, "#9E9E9E", ext)  # Grey
                
        except Exception as e:
            print(f"Error creating preview for {file_path}: {str(e)}")
            # Return a placeholder for errors
            img = Image.new('RGB', (max_width, max_height), "#F44336")
            draw = ImageDraw.Draw(img)
            draw.rectangle([(0, 0), (max_width-1, max_height-1)], outline="white", width=1)
            draw.text((max_width//2, max_height//2 - 10), "!", fill="white", anchor="mm")
            preview_image = ImageTk.PhotoImage(img)
            
        # If we got a preview, cache it with LRU tracking
        if preview_image:
            try:
                # Store in cache
                self.preview_cache[cache_key] = preview_image
                
                # Update LRU tracking list
                if cache_key in self.preview_cache_keys:
                    self.preview_cache_keys.remove(cache_key)
                    
                # Add to end (most recently used position)
                self.preview_cache_keys.append(cache_key)
                
                # Debug log
                if len(self.preview_cache) % 100 == 0:
                    logging.info(f"Preview cache size: {len(self.preview_cache)} items")
            except Exception as e:
                logging.error(f"Error updating preview cache: {str(e)}")
            
            return preview_image
        else:
            return None
            
    def _create_styled_icon(self, width, height, color, text):
        """
        Create a styled icon with the given color and text
        This is used for file types that don't have actual visual previews
        """
        try:
            # Create a new image with the specified color
            img = Image.new('RGB', (width, height), color)
            draw = ImageDraw.Draw(img)
            
            # Log for debugging
            logging.info(f"Creating styled icon with color {color} and text {text}")
            
            # Add an outer border
            draw.rectangle([(0, 0), (width-1, height-1)], outline="white", width=2)
            
            # Draw a file icon shape (a rectangle with a folded corner)
            padding = 10
            draw.rectangle(
                [(padding, padding), (width-padding, height-padding)], 
                fill="white", outline=color, width=1
            )
            
            # Fold the top-right corner
            corner_size = 15
            draw.polygon(
                [
                    (width-padding-corner_size, padding),
                    (width-padding, padding),
                    (width-padding, padding+corner_size)
                ],
                fill=color
            )
            
            # Draw lines to represent text in the "document"
            line_padding = 8
            line_height = 5
            num_lines = 3
            line_width = width - (padding*2 + line_padding*2)
            line_start_y = padding + 20
            
            for i in range(num_lines):
                y_pos = line_start_y + (i * (line_height + 5))
                draw.rectangle(
                    [(padding + line_padding, y_pos), 
                     (padding + line_padding + line_width, y_pos + line_height)], 
                    fill=color
                )
            
            # Add text centered at the bottom
            # In older versions of PIL, anchor="mm" might not be supported
            # So let's calculate the text position manually if needed
            try:
                draw.text(
                    (width//2, height - padding - 8), 
                    text, 
                    fill=color, 
                    anchor="mm"
                )
            except TypeError:
                # Older PIL versions don't support anchor
                # We'll need to center manually
                text_bbox = draw.textbbox((0, 0), text)
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                text_x = (width - text_width) // 2
                text_y = height - padding - 8 - text_height // 2
                draw.text((text_x, text_y), text, fill=color)
            
            photo_img = ImageTk.PhotoImage(img)
            return photo_img
            
        except Exception as e:
            logging.error(f"Error creating styled icon: {str(e)}")
            # Fallback to a simple colored rectangle with text
            try:
                img = Image.new('RGB', (width, height), color)
                draw = ImageDraw.Draw(img)
                draw.rectangle([(0, 0), (width-1, height-1)], outline="white", width=1)
                try:
                    draw.text((width//2, height//2), text, fill="white", anchor="mm")
                except TypeError:
                    # Center manually for older PIL versions
                    text_bbox = draw.textbbox((0, 0), text)
                    text_width = text_bbox[2] - text_bbox[0]
                    text_height = text_bbox[3] - text_bbox[1]
                    text_x = (width - text_width) // 2
                    text_y = (height - text_height) // 2
                    draw.text((text_x, text_y), text, fill="white")
                return ImageTk.PhotoImage(img)
            except Exception as e2:
                logging.error(f"Fallback icon creation also failed: {str(e2)}")
                # Last resort fallback
                img = Image.new('RGB', (width, height), "#F44336")  # Red
                return ImageTk.PhotoImage(img)
            
    def update_ui_language(self):
        # Update window title
        self.root.title(self.get_text("full_window_title"))
        
        # Update all widgets with text
        # Title section
        for widget in self.main_frame.winfo_children():
            if isinstance(widget, tk.LabelFrame):
                try:
                    current_text = widget.cget("text")
                    # Try to find a matching key in the language dictionary
                    found_key = None
                    for key in self.languages[self.current_language].keys():
                        # Check if this text matches any language's value for this key
                        for lang_code in self.languages.keys():
                            if self.languages[lang_code].get(key, "") == current_text:
                                found_key = key
                                break
                        
                        if found_key:
                            break
                    
                    # If found, update with current language version
                    if found_key:
                        widget.config(text=self.get_text(found_key))
                except Exception as e:
                    logging.warning(f"Failed to update label text: {str(e)}")
        
        # Update buttons
        self.select_folder_btn.config(text=self.get_text("select_folder"))
        self.start_btn.config(text=self.get_text("start"))
        self.cancel_btn.config(text=self.get_text("cancel"))
        self.exit_btn.config(text=self.get_text("exit"))
        self.select_all_btn.config(text=self.get_text("select_all"))
        self.clear_all_btn.config(text=self.get_text("clear_all"))
        self.apply_filter_btn.config(text=self.get_text("apply_filter"))
        self.show_filter_btn.config(text=self.get_text("filter_label"))
        
        # Update view mode buttons if they exist
        if hasattr(self, 'list_view_btn'):
            self.list_view_btn.config(text=self.get_text("list_view"))
        if hasattr(self, 'preview_view_btn'):
            self.preview_view_btn.config(text=self.get_text("preview_view"))
            
        # Update search labels if they exist
        if hasattr(self, 'extension_search_label'):
            self.extension_search_label.config(text="🔍 " + self.get_text("extension_search"))
        if hasattr(self, 'file_search_label'):
            self.file_search_label.config(text=self.get_text("extension_search") + ":")
            
        # Update search placeholder text if entry exists
        if hasattr(self, 'file_search_entry') and not self.file_search_var.get():
            self.file_search_entry.delete(0, 'end')
            self.file_search_entry.insert(0, self.get_text("search_files"))
            self.file_search_entry.config(fg='gray')
            
        # Update desktop checkbox
        if hasattr(self, 'desktop_cb'):
            self.desktop_cb.config(text=self.get_text("desktop_label"))
        
        # Update checkbox texts
        self.subfolder_cb.config(text=self.get_text("include_label"))
        
        # Update setting labels
        for frame_name in ["settings_frame", "filter_frame", "tips_frame", "stats_frame", "file_list_frame"]:
            if hasattr(self, frame_name):
                frame = getattr(self, frame_name)
                
                # Check if frame is a LabelFrame before trying to access text property
                try:
                    if isinstance(frame, tk.LabelFrame):
                        # Update frame title
                        current_text = frame.cget("text")
                        found_key = None
                        
                        # Look for this text in all languages
                        for key in ["settings_header", "filter_label", "tips_header", "statistics_header", "file_list"]:
                            # Check if this text matches any language's value for this key
                            for lang_code in self.languages.keys():
                                if self.languages[lang_code].get(key, "") == current_text:
                                    found_key = key
                                    break
                            
                            if found_key:
                                break
                        
                        # If found, update with current language version
                        if found_key:
                            frame.config(text=self.get_text(found_key))
                except Exception as e:
                    logging.warning(f"Failed to update frame title: {str(e)}")
        
        # Update labels in settings section
        for widget in self.settings_frame.winfo_children():
            if isinstance(widget, tk.Frame):
                for child in widget.winfo_children():
                    if isinstance(child, tk.Label):
                        current_text = child.cget("text")
                        found_key = None
                        
                        # Check against all languages
                        for key in ["subfolders_label", "list_format_label", "save_location_label", "sort_criteria_label"]:
                            # Check if this text matches any language's value for this key
                            for lang_code in self.languages.keys():
                                if self.languages[lang_code].get(key, "") == current_text:
                                    found_key = key
                                    break
                            
                            if found_key:
                                break
                        
                        # If found, update with current language version
                        if found_key:
                            child.config(text=self.get_text(found_key))
                    elif isinstance(child, tk.Checkbutton) and child == self.subfolder_cb:
                        child.config(text=self.get_text("include_label"))
        
        # Update all static labels
        self.update_all_static_labels(self.main_frame)
        
        # Update tooltip texts
        self.create_tooltip(self.select_folder_btn, self.get_text("tooltip_select"))
        self.create_tooltip(self.start_btn, self.get_text("tooltip_start"))
        self.create_tooltip(self.cancel_btn, self.get_text("tooltip_cancel"))
        self.create_tooltip(self.exit_btn, self.get_text("tooltip_exit"))
        self.create_tooltip(self.subfolder_cb, self.get_text("tooltip_subfolders"))
        self.create_tooltip(self.select_all_btn, self.get_text("tooltip_select_all"))
        self.create_tooltip(self.clear_all_btn, self.get_text("tooltip_clear_all"))
        self.create_tooltip(self.apply_filter_btn, self.get_text("tooltip_filter_apply"))
        
        # Update tip texts
        self.update_tips()
        
        # Update footer text with current language and year
        for widget in self.main_frame.winfo_children():
            if isinstance(widget, tk.Frame) and widget.winfo_children():
                for child in widget.winfo_children():
                    if isinstance(child, tk.Label) and "©" in child.cget("text"):
                        # Update copyright footer text
                        current_year = datetime.datetime.now().year
                        copyright_text = self.get_text("copyright_footer").format(year=current_year)
                        child.config(text=copyright_text)
        
        # Update categories
        self.populate_categories()
        
        # Update sorting options
        self.populate_sort_dropdown()
        
        # Update file tree headers with click-to-sort functionality
        self.file_tree.heading("name", text=self.get_text("file_name"), 
                              command=lambda: self.treeview_sort_column("name", False))
        self.file_tree.heading("extension", text=self.get_text("file_extension"),
                              command=lambda: self.treeview_sort_column("extension", False))
        self.file_tree.heading("path", text=self.get_text("file_path"),
                              command=lambda: self.treeview_sort_column("path", False))
        self.file_tree.heading("size", text=self.get_text("file_size"),
                              command=lambda: self.treeview_sort_column("size", False))
        self.file_tree.heading("created", text=self.get_text("creation_date"),
                              command=lambda: self.treeview_sort_column("created", False))
        self.file_tree.heading("modified", text=self.get_text("modification_date"),
                              command=lambda: self.treeview_sort_column("modified", False))
        
        # Update status text if it's the default ready message
        if self.status_var.get() == self.get_text("ready") or self.status_var.get() == "Ready":
            self.update_status(self.get_text("ready"))
            
        # Update the folder path display
        if not self.selected_folder_path:
            self.folder_path_var.set(self.get_text("no_folder_selected"))
            
        # Force a redraw
        self.root.update_idletasks()
    
    def update_all_static_labels(self, parent_widget):
        """Recursively update all Label widgets with translated text"""
        for widget in parent_widget.winfo_children():
            try:
                if isinstance(widget, tk.Label) and hasattr(widget, 'cget'):
                    try:
                        current_text = widget.cget("text")
                        if current_text:  # Ensure text is not empty
                            # Skip dynamic content like file counts, etc.
                            if current_text.isdigit() or current_text.startswith(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9")):
                                continue
                            
                            # Try to find a matching key in the language dictionary
                            found_key = None
                            for key in self.languages[self.current_language].keys():
                                # Check if this text matches any language's value for this key
                                for lang_code in self.languages.keys():
                                    if self.languages[lang_code].get(key, "") == current_text:
                                        found_key = key
                                        break
                                
                                if found_key:
                                    break
                            
                            # If found, update with current language version
                            if found_key:
                                widget.config(text=self.get_text(found_key))
                    except Exception as e:
                        logging.warning(f"Error updating label text: {str(e)}")
                
                # Recursively process child widgets
                if widget.winfo_children():
                    self.update_all_static_labels(widget)
            except Exception as e:
                logging.warning(f"Error processing widget: {str(e)}")
                
    def update_tips(self):
        """Update tips with the current language"""
        # Find tips frame and update all tip labels
        if hasattr(self, "tips_frame"):
            # Update tip labels
            tip_texts = [
                self.get_text("tip_1"),
                self.get_text("tip_3"),
                self.get_text("tip_4"),
                self.get_text("tip_5"),
                self.get_text("tip_6"),
                self.get_text("tip_preview_formats")
            ]
            
            # Find all tip labels recursively
            tip_labels = []
            
            def find_tip_labels(parent):
                for widget in parent.winfo_children():
                    try:
                        if isinstance(widget, tk.Label) and hasattr(widget, "cget") and (
                            widget.cget("text").startswith("• ") or 
                            any(widget.cget("text").endswith(t) for t in self.languages["en"].values())
                        ):
                            tip_labels.append(widget)
                    except Exception as e:
                        logging.warning(f"Error checking tip label: {str(e)}")
                        
                    if widget.winfo_children():
                        find_tip_labels(widget)
            
            find_tip_labels(self.tips_frame)
            
            # Update the tip texts
            for i, label in enumerate(tip_labels[:len(tip_texts)]):
                label.config(text=f"• {tip_texts[i]}")

    def select_folder(self):
        folder_path = filedialog.askdirectory(title=self.get_text("select_folder"))
        if folder_path:
            # Always switch back to list mode automatically when selecting a new folder
            # This ensures a consistent experience and better performance
            self.view_mode_var.set("list")
            self._switch_to_list_view()
                
            # Store the selected folder and update display
            self.selected_folder_path = folder_path
            self.folder_path_var.set(folder_path)
            
            # Update UI state now that a folder is selected
            self.update_ui_state()
            
            # Load files from the selected folder
            self.load_files_thread()

    def load_files_thread(self):
        # Clear current files
        self.files = []
        self.filtered_files = []
        
        # Update UI
        self.clear_file_list()
        self.update_status(self.get_text("folder_loading"))
        
        # Start loading files in a separate thread
        self.cancel_flag = False
        loading_thread = threading.Thread(target=self.load_files)
        loading_thread.daemon = True
        loading_thread.start()

    def load_files(self):
        try:
            # Clear current files
            self.files = []
            folder_count = 0
            total_size = 0
            
            # Update the UI to show loading state
            self.root.after(0, lambda: self.update_status(self.get_text("files_loading")))
            self.root.after(0, lambda: self.progress_bar.start(5))
            self.root.after(0, lambda: self.enable_cancel_button())
            
            # OPTIMIZATION: Create a batch processing approach
            file_batch = []
            total_processed = 0
            total_estimated_files = 0
            last_ui_update_time = time.time()
            
            # OPTIMIZATION: Pre-count files for better progress indication (for large directories)
            # Only count if top folder has more than certain number of direct files
            try:
                # Quick check of top directory first
                top_files_count = len([f for f in os.listdir(self.selected_folder_path) 
                                    if os.path.isfile(os.path.join(self.selected_folder_path, f))])
                
                # If we have many files or include subfolders, do a rough estimation
                if top_files_count > 1000 or self.include_subfolders.get():
                    self.update_status(f"{self.get_text('files_gathering')}...")
                    total_estimated_files = self._estimate_file_count()
                    logging.info(f"Estimated file count: {total_estimated_files}")
            except Exception as e:
                logging.warning(f"Error estimating file count: {str(e)}")
                total_estimated_files = 0
                
            # Create a deque for efficient batch processing
            file_batch = collections.deque(maxlen=self.file_loading_batch_size)
            
            # Walk through the folder structure
            for root, dirs, files in os.walk(self.selected_folder_path):
                # Check if the operation was cancelled
                if self.cancel_flag:
                    return self.handle_cancellation()
                
                # If we're only processing the top level, don't go into subdirectories
                if not self.include_subfolders.get() and root != self.selected_folder_path:
                    continue
                
                folder_count += 1
                
                # OPTIMIZATION: Parallel processing to calculate file sizes for faster sorting
                # This provides much faster feedback to the user for large directories
                try:
                    files_with_sizes = []
                    
                    # Distribute work across multiple cores using a thread pool
                    def get_file_size(file):
                        if self.cancel_flag:
                            return None
                        try:
                            file_path = os.path.join(root, file)
                            size = os.path.getsize(file_path)
                            return (file, size)
                        except:
                            return (file, 0)
                    
                    # Use ThreadPoolExecutor for parallelism with a reasonable number of workers
                    with concurrent.futures.ThreadPoolExecutor(max_workers=16) as executor:
                        # Submit all files for processing
                        future_to_file = {executor.submit(get_file_size, file): file for file in files}
                        
                        # Collect results as they complete
                        for future in concurrent.futures.as_completed(future_to_file):
                            if self.cancel_flag:
                                executor.shutdown(wait=False)
                                return self.handle_cancellation()
                            
                            result = future.result()
                            if result is not None:
                                files_with_sizes.append(result)
                except Exception as file_e:
                    logging.warning(f"Error getting file size: {str(file_e)}")
                    # Skip this file from the sorting process
                
                # Sort by size, smallest first
                try:
                    if files_with_sizes:  # Verify we have something to sort
                        files_with_sizes.sort(key=lambda x: x[1])
                        files = [f[0] for f in files_with_sizes]
                except Exception as e:
                    logging.warning(f"Error sorting files by size: {str(e)}")
                    # Continue with unsorted files
                
                # Process files in the current directory
                for file in files:
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        return self.handle_cancellation()
                    
                    file_path = os.path.join(root, file)
                    try:
                        # Get file info
                        file_stats = os.stat(file_path)
                        file_size = file_stats.st_size
                        total_size += file_size
                        
                        # Extract file extension
                        _, file_extension = os.path.splitext(file)
                        
                        # Format dates - OPTIMIZATION: lazy loading of date formatting
                        # Defer expensive datetime operations
                        file_info = {
                            "name": file,
                            "extension": file_extension.lower(),
                            "path": root,
                            "full_path": file_path,
                            "size": file_size,
                            "ctime": file_stats.st_ctime,  # Store raw timestamp
                            "mtime": file_stats.st_mtime,  # Store raw timestamp
                            "created": None,  # Will be formatted when needed
                            "modified": None  # Will be formatted when needed
                        }
                        
                        # Format the dates now (we need them for the UI)
                        file_info["created"] = datetime.datetime.fromtimestamp(file_info["ctime"]).strftime('%Y-%m-%d %H:%M:%S')
                        file_info["modified"] = datetime.datetime.fromtimestamp(file_info["mtime"]).strftime('%Y-%m-%d %H:%M:%S')
                        
                        # Add file to the batch
                        file_batch.append(file_info)
                        total_processed += 1
                        
                        # When batch is full, update UI and clear batch
                        if len(file_batch) >= self.file_loading_batch_size:
                            # Update the file list in batches for better performance
                            batch_list = list(file_batch)
                            self.files.extend(batch_list)
                            
                            # Update progress calculations
                            progress_percentage = 0
                            if total_estimated_files > 0:
                                progress_percentage = min(95, (total_processed / total_estimated_files) * 100)
                                self.root.after(0, lambda p=progress_percentage: self.progress_bar.config(value=p))
                            
                            # Only update UI every 0.5 seconds to avoid UI freeze with rapid updates
                            current_time = time.time()
                            if current_time - last_ui_update_time >= 0.5:
                                last_ui_update_time = current_time
                                # Update progress periodically to keep UI responsive
                                progress_text = f"{self.get_text('files_loading')} ({total_processed})"
                                if total_estimated_files > 0:
                                    progress_text += f" - {progress_percentage:.1f}%"
                                self.root.after(0, lambda msg=progress_text: self.update_status(msg))
                                self.root.update_idletasks()  # Force UI update
                            
                            # Clear the batch for the next round
                            file_batch.clear()
                            
                    except Exception as e:
                        logging.error(f"Error processing file {file_path}: {str(e)}")
            
            # Add any remaining files in the final batch
            if file_batch:
                batch_list = list(file_batch)
                self.files.extend(batch_list)
            
            # Save a copy of all files for search functionality
            # We need a copy here since filtering modifies elements
            self.all_files = self.files.copy()
            
            # Sort files before applying filter to ensure consistent results
            self.sort_files()
            
            # Set flag to indicate this is the first load
            self.is_first_load = True
            
            # Apply the initial filter immediately
            self.apply_filter_internal()
            
            # Update statistics
            self.root.after(0, lambda: self.update_statistics(len(self.filtered_files), folder_count, total_size))
            
            # Update UI
            self.root.after(0, lambda: self.update_status(
                self.get_text("folder_loaded_status")
            ))
            self.root.after(0, lambda: self.progress_bar.stop())
            def update_progress_value():
                self.progress_bar["value"] = 100
            self.root.after(0, update_progress_value)
            self.root.after(0, lambda: self.disable_cancel_button())
            
            # Preview button reference no longer needed with radio buttons
            
            # Run garbage collection to free memory
            self._cleanup_memory()
            
            # Log the operation
            logging.info(f"Loaded {len(self.files)} files from {self.selected_folder_path}")
            
        except Exception as e:
            # Handle errors
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            self.root.after(0, lambda: self.progress_bar.stop())
            self.root.after(0, lambda: self.disable_cancel_button())
            logging.error(f"Error loading files: {str(e)}")
            messagebox.showerror(self.get_text("error"), error_message)
            
    def _estimate_file_count(self):
        """
        OPTIMIZATION: Quickly estimate the number of files to be processed.
        This provides better progress indication for very large folders.
        Uses optimized sampling and memory-efficient processing to avoid full directory traversal.
        """
        try:
            # Use the selected folder path rather than looking it up again
            selected_path = getattr(self, 'selected_folder_path', self.folder_path_var.get())
            
            # Quick check if path is valid
            if not os.path.isdir(selected_path):
                return 100  # Default estimate for invalid path
                
            # For non-recursive mode, count top-level files efficiently
            if not self.include_subfolders.get():
                # Use scandir() instead of listdir() for better performance with large directories
                file_count = 0
                with os.scandir(selected_path) as entries:
                    for entry in entries:
                        if entry.is_file():
                            file_count += 1
                            # Cap the estimation time for very large directories
                            if file_count > 10000:
                                return int(file_count * 1.2)  # Return with a small buffer
                return file_count
            
            # For recursive mode, use advanced sampling for better efficiency
            total_files = 0
            total_dirs = 0
            sampled_dirs = 0
            max_time = 0.5  # Max seconds to spend on estimation
            start_time = time.time()
            
            # Use an efficient sample-based approach for large directories
            # Start with the top level
            top_level_files = 0
            top_level_dirs = []
            
            with os.scandir(selected_path) as entries:
                for entry in entries:
                    if self.cancel_flag:
                        return 100  # Cancelled, return reasonable default
                        
                    if entry.is_file():
                        top_level_files += 1
                    elif entry.is_dir():
                        top_level_dirs.append(entry.path)
            
            # Add top level files to our total
            total_files += top_level_files
            total_dirs += 1
            
            # If we have lots of subdirectories, sample a subset for efficiency
            if len(top_level_dirs) > 20:
                # Prioritize directories with common names that often have many files
                common_large_folders = ["documents", "downloads", "pictures", "videos", "music", 
                                      "photos", "images", "docs", "media"]
                
                # Sort directories to prioritize sampling known large folder types
                priority_dirs = []
                other_dirs = []
                
                for dir_path in top_level_dirs:
                    dirname = os.path.basename(dir_path).lower()
                    if any(common in dirname for common in common_large_folders):
                        priority_dirs.append(dir_path)
                    else:
                        other_dirs.append(dir_path)
                
                # Randomize the order to avoid bias in the same types of directories
                random.shuffle(priority_dirs)
                random.shuffle(other_dirs)
                
                # Take all priority directories up to 5, then fill with other directories
                sample_dirs = priority_dirs[:5]
                remaining_slots = 10 - len(sample_dirs)  # Sample up to 10 directories total
                
                if remaining_slots > 0:
                    sample_dirs.extend(other_dirs[:remaining_slots])
            else:
                # If few directories, process all of them
                sample_dirs = top_level_dirs
            
            # Sample the selected directories to depth 1 only
            sample_dir_files = 0
            sample_dir_subdirs = 0
            
            # Process each directory in our sample
            for dir_path in sample_dirs:
                # Check timing constraint
                if time.time() - start_time > max_time:
                    # Time limit reached, extrapolate from what we've seen
                    break
                    
                # Check cancellation
                if self.cancel_flag:
                    break
                    
                # Count files and subdirectories in this directory
                dir_file_count = 0
                dir_subdir_count = 0
                
                try:
                    with os.scandir(dir_path) as entries:
                        for entry in entries:
                            if entry.is_file():
                                dir_file_count += 1
                            elif entry.is_dir():
                                dir_subdir_count += 1
                except (PermissionError, FileNotFoundError):
                    # Skip inaccessible directories
                    continue
                
                # Update counts
                sample_dir_files += dir_file_count
                sample_dir_subdirs += dir_subdir_count
                total_dirs += 1
                total_files += dir_file_count
            
            # Use our sample to extrapolate the full estimate
            # Base case - what we've counted so far
            estimate = total_files
            
            # If we have subfolders, extrapolate based on our sample
            if top_level_dirs:
                if sample_dirs:  # If we sampled directories
                    # Calculate average files per sampled directory
                    avg_files_per_dir = sample_dir_files / max(1, len(sample_dirs))
                    
                    # Estimate for unsampled top-level directories
                    unsampled_count = len(top_level_dirs) - len(sample_dirs)
                    if unsampled_count > 0:
                        estimate += unsampled_count * avg_files_per_dir
                    
                    # Add estimate for subdirectories (second level)
                    if sample_dir_subdirs > 0:
                        # Assume each second-level directory has similar file count as first level
                        # but with a dampening factor since deeper directories often have fewer files
                        avg_files_per_subdir = avg_files_per_dir * 0.7
                        estimate += sample_dir_subdirs * avg_files_per_subdir
                
                # Add a buffer for very deep directory structures
                if self.include_subfolders.get() and (sample_dir_subdirs > 20 or len(top_level_dirs) > 20):
                    # More aggressive multiplier for folders with many subfolders
                    depth_multiplier = 1.0 + min(1.0, (sample_dir_subdirs + len(top_level_dirs)) / 100)
                    estimate *= depth_multiplier
            
            # Ensure we return a reasonable minimum
            return max(100, int(estimate))
            
        except Exception as e:
            logging.error(f"Error estimating file count: {str(e)}")
            return 100  # Default estimate on error

    def enable_cancel_button(self):
        self.cancel_btn.config(state=tk.NORMAL)

    def disable_cancel_button(self):
        self.cancel_btn.config(state=tk.DISABLED)

    def handle_cancellation(self):
        """Common method to handle cancellation UI updates across the application"""
        self.root.after(0, lambda: self.update_status(self.get_text("operation_cancelled")))
        self.root.after(0, lambda: self.cancel_btn.config(text=self.get_text("cancelling")))
        self.root.after(0, lambda: self.progress_bar.stop())
        self.root.after(500, self.disable_cancel_button)
        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))
        logging.info("Handling cancellation in progress")
        
        # OPTIMIZATION: Clean up memory after cancelled operation
        self.root.after(1500, self._cleanup_memory)
        
        return False  # Common return value for operation cancellation
        
    def _cleanup_memory(self):
        """
        OPTIMIZATION: Enhanced memory management for large operations
        This helps prevent memory leaks and improves performance
        for operations that process many files by implementing a more
        aggressive and proactive memory management strategy.
        """
        try:
            # Thread-safety for cache operations
            if hasattr(self, 'preview_cache_lock'):
                cache_lock = self.preview_cache_lock
            else:
                cache_lock = threading.RLock()
                self.preview_cache_lock = cache_lock
            
            with cache_lock:
                # OPTIMIZATION: More aggressive cache management 
                # Clear the preview cache more aggressively based on usage patterns
                if hasattr(self, 'preview_cache'):
                    cache_size = len(self.preview_cache)
                    
                    # If cache is very large, be more aggressive in pruning
                    if cache_size > self.max_preview_cache_size * 0.8:  # > 80% of max
                        # Determine how many items to keep based on current memory pressure
                        # Keep fewer items when cache is larger (adaptive strategy)
                        keep_count = min(100, max(25, int(self.max_preview_cache_size * 0.15)))
                        
                        if hasattr(self, 'preview_cache_keys') and self.preview_cache_keys:
                            # Keep the most recently used items, scaled by current cache size
                            keys_to_keep = self.preview_cache_keys[-keep_count:] if len(self.preview_cache_keys) > keep_count else self.preview_cache_keys
                            
                            # Create new cache with only those items (more efficient than deleting)
                            new_cache = {}
                            for key in keys_to_keep:
                                if key in self.preview_cache:
                                    new_cache[key] = self.preview_cache[key]
                            
                            # Log memory cleanup for better diagnostics
                            items_removed = cache_size - len(new_cache)
                            if items_removed > 0:
                                logging.info(f"Memory optimization: Released {items_removed} cached previews from memory")
                            
                            # Replace old cache with smaller one
                            self.preview_cache = new_cache
                            self.preview_cache_keys = keys_to_keep
                        else:
                            # If no LRU tracking, clear most of the cache 
                            self.preview_cache = {}
            
            # Clear any temporary references that might be holding large objects
            if hasattr(self, 'preview_images'):
                self.preview_images = []
                
            # Clear references to large file lists when appropriate
            if hasattr(self, 'cancel_flag') and self.cancel_flag:
                if hasattr(self, 'files'):
                    self.files = []
                if hasattr(self, 'filtered_files'):
                    self.filtered_files = []
                
            # Clear thumbnail references
            if hasattr(self, 'preview_thumbnails'):
                self.preview_thumbnails = []
                
            # Explicitly force multiple passes of garbage collection to free memory
            for _ in range(3):  # Multiple passes are more effective for complex reference cycles
                gc.collect()
                
            # Encourage Python to return memory to OS
            if hasattr(sys, 'pypy_version_info'):
                # PyPy specific
                if hasattr(gc, 'collect'):
                    gc.collect()
            else:
                # CPython
                try:
                    import ctypes
                    ctypes.pythonapi.PyGC_Collect()
                    ctypes.pythonapi.malloc_trim(0)
                except (ImportError, AttributeError):
                    pass
                
            logging.info("Enhanced memory cleanup completed")
        except Exception as e:
            logging.error(f"Error during memory cleanup: {str(e)}")
        
    def cancel_operation(self):
        """Cancel the current operation and update UI accordingly"""
        # Set the cancellation flag
        self.cancel_flag = True
        
        # Update the status and progress immediately
        self.update_status(self.get_text("operation_cancelled"))
        self.progress_bar.stop()
        self.progress_bar["value"] = 0
        
        # Temporarily disable further cancellations to prevent multiple clicks
        self.cancel_btn.config(state=tk.DISABLED)
        
        # Visual feedback during cancellation
        self.cancel_btn.config(text=self.get_text("cancelling"))
        self.root.update_idletasks()  # Force UI update
        
        # Schedule re-enabling of the UI after a short delay
        self.root.after(500, self.disable_cancel_button)
        
        # Reset UI state once operation is fully cancelled
        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))
        
        # Log the action
        logging.info("Operation cancelled by user")

    def _calculate_file_type_statistics(self):
        """Dosya uzantı istatistiklerini hesaplar ve statusbar için formatlı metni döndürür"""
        if not hasattr(self, 'filtered_files') or not self.filtered_files:
            return None
            
        # Her tür istatistik için yeni bir sözlük oluştur
        temp_stats = {}
        
        # Önizlemesi desteklenen dosya türleri
        supported_preview_extensions = [
            "jpg", "jpeg", "png", "gif", "bmp", "tiff", "tif", "svg", "ico", 
            "pdf", "eps", "psd", "ai"
        ]
        
        # Tüm dosyaları tekrar sayalım
        for file_info in self.filtered_files:
            if file_info.get("is_folder", False):
                continue
                
            # Dosya uzantısını al
            file_ext = file_info.get("extension", "").lower().replace(".", "")
            if not file_ext and "path" in file_info and "name" in file_info:
                file_path = os.path.join(file_info["path"], file_info["name"])
                file_ext = os.path.splitext(file_path)[1].lower().replace(".", "")
            
            # jpeg ve jpg uzantılarını birleştir
            if file_ext == "jpeg":
                file_ext = "jpg"
            
            # Sadece önizleme desteği olan uzantıları say
            if file_ext in supported_preview_extensions:
                # İstatistikleri güncelle
                if file_ext in temp_stats:
                    temp_stats[file_ext] += 1
                else:
                    temp_stats[file_ext] = 1
        
        # Geçici istatistikleri ana sözlüğe atayalım ve formatlı metni hazırlayalım
        self.file_type_stats = temp_stats
        
        if not self.file_type_stats:
            return None
            
        # Format statistics: sort by count (descending)
        stats_sorted = sorted(self.file_type_stats.items(), key=lambda x: x[1], reverse=True)
        
        # Ön izleme modunda top 10, liste modunda ya hiç gösterme ya da tüm dosya türleri
        if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview":
            # Ön izleme modunda en çok bulunan 10 dosya türünü göster
            top_stats = stats_sorted[:10]
            
            stats_text = []
            for ext, count in top_stats:
                # Uppercase the extension for better visibility
                stats_text.append(f"{ext.upper()}: {count}")
            
            # Create a nice statistics message including file count
            file_count_text = f"{len(self.filtered_files)} {self.get_text('files')}"
            stats_message = f"{file_count_text} | " + " | ".join(stats_text)
        else:
            # Liste modunda sadece dosya sayısı bilgisini göster
            file_count_text = f"{len(self.filtered_files)} {self.get_text('files')}"
            stats_message = file_count_text
        
        logging.info(f"Calculated stats: {stats_message}")
        return stats_message
        
    def update_status(self, message):
        """Update the status message in the status bar"""
        # If filtering is complete and we're trying to show filtering message, show statistics instead
        if hasattr(self, 'filtering_complete') and self.filtering_complete:
            # Skip filtering messages if filtering is already complete
            if message.startswith(self.get_text("filtering_in_progress")) or message == self.get_text("filtering"):
                # Filtreleme tamamlandıysa, istatistikleri göster
                stats_message = self._calculate_file_type_statistics()
                if stats_message:
                    self.status_var.set(stats_message)
                    logging.info(f"Showing statistics instead of filtering message: {stats_message}")
                    return
                
        # Check if we have file type statistics to display when in preview mode
        if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview" and message == self.get_text("preview_mode_active"):
            # Loglama ekleyelim
            logging.info("Updating status with preview mode stats")
            
            # İstatistik hesaplayıp gösterelim
            stats_message = self._calculate_file_type_statistics()
            if stats_message:
                self.status_var.set(stats_message)
                return
            
        # Default behavior (no statistics or not in preview mode)
        self.status_var.set(message)
        self.root.update_idletasks()

    def update_statistics(self, file_count, folder_count, total_size):
        # Update file count with optional subfolder indicator
        self.total_files_var.set(str(file_count))
        
        # Update folder count with clearer wording if subfolders are included
        if self.include_subfolders.get() and folder_count > 1:
            subfolder_count = folder_count - 1  # Main folder + subfolders
            folder_label = f"{folder_count} ({subfolder_count} {self.get_text('subfolders_label')})"
            self.folder_count_var.set(folder_label)
        else:
            self.folder_count_var.set(str(folder_count))
        
        # Format total size with appropriate units for better readability
        if total_size < 1024:
            size_str = f"{total_size} B"
        elif total_size < 1024 * 1024:
            size_str = f"{total_size/1024:.2f} KB"
        elif total_size < 1024 * 1024 * 1024:
            size_str = f"{total_size/(1024*1024):.2f} MB"
        else:
            size_str = f"{total_size/(1024*1024*1024):.2f} GB"
            
        self.total_size_var.set(size_str)
        
        # Log statistics for debugging
        logging.info(f"Statistics updated: {file_count} files, {folder_count} folders, {size_str} total size")

    def clear_file_list(self):
        # Clear all items in the file tree
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)

    def update_file_list(self, files):
        # Clear current list
        self.clear_file_list()
        
        # Store the original files list for filtering
        self.all_files = files.copy()
        
        # No files to display
        if not files:
            return
        
        # OPTIMIZATION: Insert files in batches for better performance with large file lists
        total_files = len(files)
        # Use the batch size defined in initialization
        batch_size = self.file_display_batch_size
        
        # Process in batches
        for i in range(0, total_files, batch_size):
            # Get the current batch
            batch = files[i:min(i+batch_size, total_files)]
            
            # Create a temporary list for batch updates
            batch_values = []
            for file_info in batch:
                values = (
                    file_info["name"],
                    file_info["extension"],
                    file_info["path"],
                    self.format_file_size(file_info["size"]),
                    file_info["created"],
                    file_info["modified"]
                )
                batch_values.append(values)
            
            # Insert the entire batch at once
            for values in batch_values:
                self.file_tree.insert("", tk.END, values=values)
                
            # Update the UI after each batch to maintain responsiveness
            self.root.update_idletasks()
            
        # If in preview mode, refresh the preview panel
        if hasattr(self, 'view_mode_var') and self.view_mode_var.get() == "preview":
            self._build_preview_panel(files)
            
    def _add_file_to_list(self, file_info):
        """Helper method to add a file to the list"""
        values = (
            file_info["name"],
            file_info["extension"],
            file_info["path"],
            self.format_file_size(file_info["size"]),
            file_info["created"],
            file_info["modified"]
        )
        self.file_tree.insert("", tk.END, values=values)

    def format_file_size(self, size):
        # Format file size with appropriate units
        if size < 1024:
            return f"{size} B"
        elif size < 1024 * 1024:
            return f"{size/1024:.2f} KB"
        elif size < 1024 * 1024 * 1024:
            return f"{size/(1024*1024):.2f} MB"
        else:
            return f"{size/(1024*1024*1024):.2f} GB"

    def select_all_extensions(self):
        # Set all extension checkboxes to True
        for ext in self.selected_extensions:
            self.selected_extensions[ext].set(True)

    def clear_all_extensions(self):
        """Clear all file extension selections and reset the search filter"""
        # Clear all extension checkboxes
        for ext in self.selected_extensions:
            self.selected_extensions[ext].set(False)
            
        # Clear category selection checkboxes
        for category in self.category_vars:
            if category in self.category_vars:
                self.category_vars[category].set(False)
                
        # Reset search filter if it exists
        if hasattr(self, 'extension_search_var'):
            self.extension_search_var.set("")
            # Reset categories display based on cleared search
            self.filter_extensions()

    def apply_filter(self):
        # Start filtering in a separate thread
        self.cancel_flag = False
        self.update_status(self.get_text("filter_applying"))
        self.progress_bar.start(5)
        self.enable_cancel_button()
        
        filtering_thread = threading.Thread(target=self.apply_filter_internal)
        filtering_thread.daemon = True
        filtering_thread.start()

    def apply_filter_internal(self):
        try:
            # Set filtering status flag at the beginning
            self.filtering_complete = False
            
            # Filtreleme işlemi başladığında durum çubuğunda bir dönen simge göster
            self.update_status(self.get_text("filtering_in_progress") + " ⟳")
            self.root.config(cursor="watch")  # İmleç değiştir
            
            # Güncelleme fonksiyonu - dönen simge efekti için
            self.spinner_chars = ["⟳", "⟲", "↻", "↺"]
            self.spinner_index = 0
            
            def update_spinner():
                if self.cancel_flag:
                    return
                
                # Dönen simge karakterini değiştir
                self.spinner_index = (self.spinner_index + 1) % len(self.spinner_chars)
                self.update_status(self.get_text("filtering_in_progress") + " " + 
                                 self.spinner_chars[self.spinner_index])
                
                # Her 200ms'de bir güncelle
                self.root.after(200, update_spinner)
            
            # Dönen simgeyi başlat
            update_spinner()
            
            # Check if any extensions are selected
            any_selected = any(self.selected_extensions[ext].get() for ext in self.selected_extensions)
            
            # OPTIMIZATION: If this is the first load (right after load_files), show all files
            if hasattr(self, 'is_first_load') and self.is_first_load:
                self.is_first_load = False  # Reset flag after first use
                any_selected = False  # Force showing all files at first load
                
            # Create a set of selected extensions for faster lookup
            # OPTIMIZATION: Use set for O(1) lookup time instead of linear search
            selected_extensions_set = {ext.lower() for ext in self.selected_extensions 
                                      if ext in self.selected_extensions and self.selected_extensions[ext].get()}
            
            # If no extensions are selected, don't filter anything
            if not any_selected:
                self.filtered_files = self.files.copy()
            else:
                # OPTIMIZATION: Process in batches to improve responsiveness with large file lists
                self.filtered_files = []
                total_files = len(self.files)
                processed_count = 0
                last_ui_update_time = time.time()
                
                # OPTIMIZATION: Use deque for more efficient append operations
                filtered_batch = collections.deque(maxlen=self.file_filtering_batch_size * 2)
                
                # Process each file with efficient batching
                for i in range(0, total_files, self.file_filtering_batch_size):
                    # Get the current batch
                    end_idx = min(i + self.file_filtering_batch_size, total_files)
                    batch = self.files[i:end_idx]
                    
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        # Use special filter cancelled text for this particular case
                        self.root.after(0, lambda: self.update_status(self.get_text("filter_cancelled")))
                        self.root.after(0, lambda: self.cancel_btn.config(text=self.get_text("cancelling")))
                        self.root.after(0, lambda: self.progress_bar.stop())
                        self.root.after(500, self.disable_cancel_button)
                        self.root.after(1000, lambda: self.cancel_btn.config(text=self.get_text("cancel")))
                        return
                    
                    # Process each file in the current batch
                    # OPTIMIZATION: Avoid function calls in inner loop
                    for file_info in batch:
                        ext = file_info["extension"].lower()
                        
                        # Check if the extension is in the selected extensions set (or if we're showing all files)
                        if not any_selected:
                            # If no extension is selected, show all files including ones without extension
                            if not ext:
                                # Handle files with no extension - create a copy to avoid modifying original
                                file_copy = file_info.copy()
                                file_copy["extension"] = self.get_text("extension_not_found")
                                filtered_batch.append(file_copy)
                            else:
                                filtered_batch.append(file_info)
                        else:
                            # If extensions are selected for filtering, only show files with those extensions
                            if ext in selected_extensions_set:
                                filtered_batch.append(file_info)
                            # Files with no extension are excluded when filtering by extension
                    
                    # Only update UI every 0.3 seconds to improve performance
                    processed_count += len(batch)
                    current_time = time.time()
                    if current_time - last_ui_update_time >= 0.3:
                        last_ui_update_time = current_time
                        progress = (processed_count / total_files) * 100
                        self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                        self.root.after(0, lambda c=processed_count, t=total_files: 
                                      self.update_status(f"{self.get_text('filter_applying')} ({c}/{t})"))
                        self.root.update_idletasks()
                
                # Convert deque to list for the filtered files
                self.filtered_files = list(filtered_batch)
            
            # Sort files according to selected criteria
            self.sort_files()
            
            # OPTIMIZATION: Incrementally update UI for very large file lists
            # This avoids UI freezes when updating the treeview with thousands of items
            if len(self.filtered_files) > 5000:
                # For extremely large datasets, update in chunks
                self.update_status(f"{self.get_text('files_filtered_message').format(len(self.filtered_files))} - {self.get_text('loading_file_list')}")
                
                # First clear the list
                self.root.after(0, self.clear_file_list)
                self.root.update_idletasks()
                
                # Then update in chunks of the display batch size
                chunk_size = self.file_display_batch_size
                for i in range(0, len(self.filtered_files), chunk_size):
                    end_idx = min(i + chunk_size, len(self.filtered_files))
                    chunk = self.filtered_files[i:end_idx]
                    # Pass a copy of the chunk to avoid reference issues
                    self.root.after(0, lambda files=chunk.copy(): self._update_file_list_chunk(files))
                    # Allow small pauses for UI to remain responsive
                    self.root.update_idletasks()
            else:
                # For smaller datasets, update all at once
                self.root.after(0, lambda: self.update_file_list(self.filtered_files))
            
            # Update statistics
            # OPTIMIZATION: Calculate statistics in a more efficient way
            total_size = 0
            folder_paths = set()
            
            # Reset file type statistics for the filtered files
            self.file_type_stats = {}
            
            for file in self.filtered_files:
                total_size += file["size"]
                folder_paths.add(file["path"])
                
                # Update file type statistics
                if not file.get("is_folder", False):
                    # Get file extension, use a placeholder for files without extension
                    file_ext = file.get("extension", "").lower()
                    if not file_ext:
                        file_ext = self.get_text("extension_not_found")
                    
                    # Update the count for this extension
                    if file_ext in self.file_type_stats:
                        self.file_type_stats[file_ext] += 1
                    else:
                        self.file_type_stats[file_ext] = 1
                
            folder_count = len(folder_paths)
            self.root.after(0, lambda: self.update_statistics(len(self.filtered_files), folder_count, total_size))
            
            # Update status
            self.root.after(0, lambda: self.update_status(
                self.get_text("files_filtered_message").format(len(self.filtered_files))
            ))
            self.root.after(0, lambda: self.progress_bar.stop())
            def update_progress_value():
                self.progress_bar["value"] = 100
            self.root.after(0, update_progress_value)
            self.root.after(0, lambda: self.disable_cancel_button())
            
            # Run garbage collection to free memory after filtering
            self._cleanup_memory()
            
            # İşlem bitince imleci normale döndür ve statüyü güncelle
            self.root.config(cursor="")
            # Set filtering complete flag to true
            self.filtering_complete = True
            # Update status with completion message
            self.update_status(self.get_text("filter_complete"))
            
        except Exception as e:
            # Hata göster ve loglama yap
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            self.root.after(0, lambda: self.progress_bar.stop())
            self.root.after(0, lambda: self.disable_cancel_button())
            logging.error(f"Error applying filter: {str(e)}")
            
            # Hata mesajını göster
            self.show_error(
                self.get_text("filter_error"), 
                f"{self.get_text('filter_error_details')}: {str(e)}", 
                e
            )
            
            # İmleci normale döndür
            self.root.config(cursor="")
            # Ensure we mark filtering as complete, even in case of error
            self.filtering_complete = True
            
    def _update_file_list_chunk(self, files_chunk):
        """Helper method to update the file list with a chunk of files
        Used for incremental updates with very large file lists"""
        # OPTIMIZATION: Batch insert records using a list comprehension
        # This is faster than inserting one by one for large chunks
        
        # First, prepare all values
        values_list = []
        for file_info in files_chunk:
            values = (
                file_info["name"],
                file_info["extension"],
                file_info["path"],
                self.format_file_size(file_info["size"]),
                file_info["created"],
                file_info["modified"]
            )
            values_list.append(values)
            
        # If we have a small number of files, insert directly
        if len(values_list) <= 100:
            for values in values_list:
                self.file_tree.insert("", tk.END, values=values)
        else:
            # For larger chunks, disable UI updates temporarily for better performance
            self.file_tree.config(takefocus=0)  # Temporarily disable focus
            
            try:
                # Insert all records
                for values in values_list:
                    self.file_tree.insert("", tk.END, values=values)
            finally:
                # Re-enable focus
                self.file_tree.config(takefocus=1)
                # Update the view
                self.file_tree.update()

    def sort_files(self):
        # Get the sort criteria
        sort_option = self.sort_options[self.sort_dropdown.current()]
        
        # Sort the filtered files
        if sort_option == "sort_name_asc":
            self.filtered_files.sort(key=lambda x: x["name"].lower())
        elif sort_option == "sort_name_desc":
            self.filtered_files.sort(key=lambda x: x["name"].lower(), reverse=True)
        elif sort_option == "sort_ext_asc":
            self.filtered_files.sort(key=lambda x: x["extension"].lower())
        elif sort_option == "sort_ext_desc":
            self.filtered_files.sort(key=lambda x: x["extension"].lower(), reverse=True)
        elif sort_option == "sort_size_asc":
            self.filtered_files.sort(key=lambda x: x["size"])
        elif sort_option == "sort_size_desc":
            self.filtered_files.sort(key=lambda x: x["size"], reverse=True)
        elif sort_option == "sort_dir_asc":
            self.filtered_files.sort(key=lambda x: x["path"].lower())

    def start_processing(self):
        # Check if a folder is selected
        if not self.selected_folder_path:
            messagebox.showwarning(
                self.get_text("error"),
                self.get_text("select_folder_first")
            )
            return
        
        # Check if any export format is selected
        if not any(self.export_formats[fmt].get() for fmt in self.export_formats):
            messagebox.showwarning(
                self.get_text("error"),
                self.get_text("select_format_first")
            )
            return
        
        # Start processing in a separate thread
        self.cancel_flag = False
        self.update_status(self.get_text("start_processing"))
        self.progress_bar["value"] = 0
        self.enable_cancel_button()
        
        processing_thread = threading.Thread(target=self.process_files)
        processing_thread.daemon = True
        processing_thread.start()

    def process_files(self):
        try:
            # Check if we have files to process
            if not self.filtered_files:
                self.root.after(0, lambda: messagebox.showinfo(
                    self.get_text("info"),
                    self.get_text("no_files_found")
                ))
                self.root.after(0, lambda: self.update_status(self.get_text("ready")))
                self.root.after(0, lambda: self.disable_cancel_button())
                return
            
            # Create timestamp for file names
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            folder_name = os.path.basename(self.selected_folder_path)
            base_filename = f"ListeKolay_{folder_name}_{timestamp}"
            
            # Determine save location
            if self.save_to_desktop.get():
                desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                # Create ListeKolay folder on desktop if it doesn't exist
                listekolay_folder = os.path.join(desktop_path, "ListeKolay")
                if not os.path.exists(listekolay_folder):
                    try:
                        os.makedirs(listekolay_folder)
                        logging.info(f"Created ListeKolay folder at {listekolay_folder}")
                    except Exception as e:
                        logging.error(f"Failed to create ListeKolay folder: {str(e)}")
                save_path = listekolay_folder
            else:
                save_path = os.path.dirname(os.path.abspath(__file__))
            
            created_files = []
            
            # Process text format
            if self.export_formats["text"].get():
                text_file = os.path.join(save_path, base_filename + ".txt")
                if self.export_text_file(text_file):
                    created_files.append(("text", text_file))
            
            # Process Excel format
            if self.export_formats["excel"].get():
                excel_file = os.path.join(save_path, base_filename + ".xlsx")
                if self.export_excel_file(excel_file):
                    created_files.append(("excel", excel_file))
            
            # Process Word format
            if self.export_formats["word"].get():
                word_file = os.path.join(save_path, base_filename + ".docx")
                if self.export_word_file(word_file):
                    created_files.append(("word", word_file))
            
            # Process HTML format
            if self.export_formats["html"].get():
                html_file = os.path.join(save_path, base_filename + ".html")
                if self.export_html_file(html_file):
                    created_files.append(("html", html_file))
            
            # Update UI
            def update_progress_value():
                self.progress_bar["value"] = 100
            self.root.after(0, update_progress_value)
            self.root.after(0, lambda: self.update_status(self.get_text("create_list_time").format(
                datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )))
            self.root.after(0, lambda: self.disable_cancel_button())
            
            # Ask user if they want to open the first created file
            if created_files:
                file_type, file_path = created_files[0]
                self.root.after(0, lambda: self.ask_to_open_file(file_type, file_path))
            
        except Exception as e:
            # Handle errors
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            self.root.after(0, lambda: self.disable_cancel_button())
            logging.error(f"Error processing files: {str(e)}")
            self.root.after(0, lambda: messagebox.showerror(self.get_text("error"), error_message))

    def export_text_file(self, file_path):
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write(f"{self.get_text('file_list')} - {self.selected_folder_path}\n")
                f.write(f"{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                
                # Add sort information
                sort_option = self.selected_sort.get()
                sort_name = self.get_text(sort_option)
                f.write(f"{self.get_text('sorted_by')}: {sort_name}\n")
                
                f.write("=" * 80 + "\n\n")
                
                # Make sure files are sorted according to the selected criteria
                self.sort_files()
                
                # Write file information
                for i, file_info in enumerate(self.filtered_files, 1):
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        return self.handle_cancellation()
                    
                    # Update progress
                    progress = (i / len(self.filtered_files)) * 100
                    self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                    self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                                 self.update_status(self.get_text("file_processed").format(i, total, name)))
                    
                    # Write file details
                    f.write(f"{i}. {file_info['name']}\n")
                    f.write(f"   {self.get_text('file_path')}: {file_info['path']}\n")
                    f.write(f"   {self.get_text('file_extension')}: {file_info['extension']}\n")
                    f.write(f"   {self.get_text('file_size')}: {self.format_file_size(file_info['size'])}\n")
                    f.write(f"   {self.get_text('creation_date')}: {file_info['created']}\n")
                    f.write(f"   {self.get_text('modification_date')}: {file_info['modified']}\n")
                    f.write("\n")
            
            # Log success
            logging.info(f"Created text file: {file_path}")
            return True
            
        except Exception as e:
            # Log error
            error_message = self.get_text("text_file_error").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating text file: {str(e)}")
            return False

    def export_excel_file(self, file_path):
        try:
            # Create a new workbook and select the active sheet
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = self.get_text("file_list")
            
            # Create a more attractive header
            # Title row
            title_cell = sheet.cell(row=1, column=1)
            title_cell.value = self.get_text("file_list")
            title_cell.font = Font(name='Arial', size=14, bold=True, color="0000FF")
            sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
            title_cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Folder path row
            folder_cell = sheet.cell(row=2, column=1)
            folder_cell.value = f"{self.get_text('selected_folder')}: {self.selected_folder_path}"
            folder_cell.font = Font(name='Arial', size=10, bold=True)
            sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=7)
            
            # Creation time row
            time_cell = sheet.cell(row=3, column=1)
            time_cell.value = f"{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            time_cell.font = Font(name='Arial', size=10)
            sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=7)
            
            # Sort information row
            sort_option = self.selected_sort.get()
            sort_name = self.get_text(sort_option)
            sort_cell = sheet.cell(row=4, column=1)
            sort_cell.value = f"{self.get_text('sorted_by')} {sort_name}"
            sort_cell.font = Font(name='Arial', size=10, italic=True)
            sheet.merge_cells(start_row=4, start_column=1, end_row=4, end_column=7)
            
            # Add a bit of space before the table
            # Row 5 is empty for spacing
            
            # Make sure files are sorted according to the selected criteria
            self.sort_files()
            
            # Add header row
            headers = [
                self.get_text("row_number"),
                self.get_text("file_name"),
                self.get_text("file_extension"),
                self.get_text("file_path"),
                self.get_text("file_size"),
                self.get_text("creation_date"),
                self.get_text("modification_date")
            ]
            
            # Style for the header row
            header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            header_border = Border(
                left=Side(style='thin'), 
                right=Side(style='thin'), 
                top=Side(style='thin'), 
                bottom=Side(style='thin')
            )
            
            # Apply header styles
            for col, header in enumerate(headers, 1):
                cell = sheet.cell(row=6, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = header_border
            
            # Style for data rows
            data_border = Border(
                left=Side(style='thin'), 
                right=Side(style='thin'), 
                top=Side(style='thin'), 
                bottom=Side(style='thin')
            )
            even_row_fill = PatternFill(start_color="E9EDF4", end_color="E9EDF4", fill_type="solid")
            odd_row_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
            
            # Add file information
            for i, file_info in enumerate(self.filtered_files, 1):
                # Check if the operation was cancelled
                if self.cancel_flag:
                    return self.handle_cancellation()
                
                # Update progress
                progress = (i / len(self.filtered_files)) * 100
                self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                             self.update_status(self.get_text("file_processed").format(i, total, name)))
                
                # Add file details
                row = i + 6  # Data starts at row 7 (6+1)
                
                # Apply alternating row colors
                row_fill = even_row_fill if i % 2 == 0 else odd_row_fill
                
                # Set cell values and styles
                for col in range(1, 8):
                    cell = sheet.cell(row=row, column=col)
                    cell.border = data_border
                    cell.fill = row_fill
                    
                    # Set alignment based on column
                    if col == 1:  # Row number - center align
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                
                # Set cell values
                sheet.cell(row=row, column=1).value = i
                sheet.cell(row=row, column=2).value = file_info["name"]
                sheet.cell(row=row, column=3).value = file_info["extension"]
                sheet.cell(row=row, column=4).value = file_info["path"]
                sheet.cell(row=row, column=5).value = self.format_file_size(file_info["size"])
                sheet.cell(row=row, column=6).value = file_info["created"]
                sheet.cell(row=row, column=7).value = file_info["modified"]
            
            # Auto-adjust column widths
            for col in range(1, len(headers) + 1):
                max_length = 0
                for row in range(1, len(self.filtered_files) + 7):  # +7 for the header rows
                    cell_value = sheet.cell(row=row, column=col).value
                    if cell_value:
                        max_length = max(max_length, len(str(cell_value)))
                # Set minimum and maximum widths
                adjusted_width = min(50, max(12, max_length + 2))
                sheet.column_dimensions[chr(64 + col)].width = adjusted_width
            
            # Apply auto-filter to the header row to make columns filterable
            sheet.auto_filter.ref = f"A6:G{len(self.filtered_files) + 6}"
            
            # Freeze the header row so it stays visible when scrolling
            sheet.freeze_panes = "A7"
            
            # Save the workbook
            workbook.save(file_path)
            
            # Log success
            logging.info(f"Created Excel file: {file_path}")
            return True
            
        except Exception as e:
            # Log error
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating Excel file: {str(e)}")
            return False

    def export_word_file(self, file_path):
        try:
            # Create a new document
            document = Document()
            
            # Add title
            document.add_heading(self.get_text("file_list"), level=1)
            
            # Add folder information
            document.add_paragraph(f"{self.get_text('selected_folder')}: {self.selected_folder_path}")
            document.add_paragraph(f"{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add sort information
            sort_option = self.selected_sort.get()
            sort_name = self.get_text(sort_option)
            document.add_paragraph(f"{self.get_text('sorted_by')}: {sort_name}")
            
            # Make sure files are sorted according to the selected criteria
            self.sort_files()
            
            # Add a horizontal line
            document.add_paragraph("_" * 50)
            
            # Add file information
            for i, file_info in enumerate(self.filtered_files, 1):
                # Check if the operation was cancelled
                if self.cancel_flag:
                    return self.handle_cancellation()
                
                # Update progress
                progress = (i / len(self.filtered_files)) * 100
                self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                             self.update_status(self.get_text("file_processed").format(i, total, name)))
                
                # Add file details
                document.add_heading(f"{i}. {file_info['name']}", level=2)
                table = document.add_table(rows=5, cols=2)
                table.style = "Table Grid"
                
                # Fill the table with file details
                cells = table.rows[0].cells
                cells[0].text = self.get_text("file_path")
                cells[1].text = file_info["path"]
                
                cells = table.rows[1].cells
                cells[0].text = self.get_text("file_extension")
                cells[1].text = file_info["extension"]
                
                cells = table.rows[2].cells
                cells[0].text = self.get_text("file_size")
                cells[1].text = self.format_file_size(file_info["size"])
                
                cells = table.rows[3].cells
                cells[0].text = self.get_text("creation_date")
                cells[1].text = file_info["created"]
                
                cells = table.rows[4].cells
                cells[0].text = self.get_text("modification_date")
                cells[1].text = file_info["modified"]
                
                # Add a space after each file
                document.add_paragraph()
            
            # Save the document
            document.save(file_path)
            
            # Log success
            logging.info(f"Created Word file: {file_path}")
            return True
            
        except Exception as e:
            # Log error
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating Word file: {str(e)}")
            return False

    def export_html_file(self, file_path):
        try:
            # Make sure files are sorted according to the selected criteria
            self.sort_files()
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # Write HTML header
                f.write(f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{self.get_text('file_list')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1, h2 {{ color: #333; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ padding: 8px; text-align: left; border: 1px solid #ddd; }}
        th {{ background-color: #f2f2f2; }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        .info {{ color: #666; margin-bottom: 20px; }}
    </style>
</head>
<body>
    <h1>{self.get_text('file_list')}</h1>
    <div class="info">
        <p>{self.get_text('selected_folder')}: {self.selected_folder_path}</p>
        <p>{self.get_text('creation_time')} {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p>{self.get_text('sorted_by')}: {self.get_text(self.selected_sort.get())}</p>
    </div>
    <table>
        <tr>
            <th>{self.get_text('row_number')}</th>
            <th>{self.get_text('file_name')}</th>
            <th>{self.get_text('file_extension')}</th>
            <th>{self.get_text('file_path')}</th>
            <th>{self.get_text('file_size')}</th>
            <th>{self.get_text('creation_date')}</th>
            <th>{self.get_text('modification_date')}</th>
        </tr>
""")
                
                # Add file information
                for i, file_info in enumerate(self.filtered_files, 1):
                    # Check if the operation was cancelled
                    if self.cancel_flag:
                        return self.handle_cancellation()
                    
                    # Update progress
                    progress = (i / len(self.filtered_files)) * 100
                    self.root.after(0, lambda p=progress: self.progress_bar.config(value=p))
                    self.root.after(0, lambda i=i, total=len(self.filtered_files), name=file_info["name"]: 
                                 self.update_status(self.get_text("file_processed").format(i, total, name)))
                    
                    # Write file details
                    f.write(f"""
        <tr>
            <td>{i}</td>
            <td>{file_info['name']}</td>
            <td>{file_info['extension']}</td>
            <td>{file_info['path']}</td>
            <td>{self.format_file_size(file_info['size'])}</td>
            <td>{file_info['created']}</td>
            <td>{file_info['modified']}</td>
        </tr>""")
                
                # Write HTML footer
                f.write("""
    </table>
</body>
</html>""")
            
            # Log success
            logging.info(f"Created HTML file: {file_path}")
            return True
            
        except Exception as e:
            # Log error
            error_message = self.get_text("error_occurred").format(str(e))
            self.root.after(0, lambda: self.update_status(error_message))
            logging.error(f"Error creating HTML file: {str(e)}")
            return False

    def ask_to_open_file(self, file_type, file_path):
        # Format success message based on file type
        success_message = self.get_text(f"{file_type}_success")
        self.update_status(success_message)
        
        # Ask user if they want to open the file
        response = messagebox.askyesno(
            self.get_text("open_file_title"),
            self.get_text("open_file_message"),
            icon=messagebox.QUESTION
        )
        
        if response:
            try:
                # Open the file with the default application
                if sys.platform == "win32":
                    os.startfile(file_path)
                elif sys.platform == "darwin":  # macOS
                    subprocess.call(["open", file_path])
                else:  # Linux and other Unix-like
                    subprocess.call(["xdg-open", file_path])
                    
                logging.info(f"Opened file: {file_path}")
            except Exception as e:
                logging.error(f"Error opening file {file_path}: {str(e)}")
                messagebox.showerror(
                    self.get_text("error"),
                    self.get_text("error_occurred").format(str(e))
                )

    def _go_to_prev_page(self):
        """Navigate to the previous page of thumbnails"""
        if hasattr(self, 'preview_page') and self.preview_page > 1:
            # Update status
            self.update_status(self.get_text("loading_preview"))
            
            # Show loading indicator
            self.progress_bar.start(10)
            
            # Change page
            self.preview_page -= 1
            
            # Log for debugging
            logging.info(f"Moving to previous page: {self.preview_page}")
            
            # Update preview panel
            self._update_preview_panel()
            
            # Update page info
            if hasattr(self, 'page_info_label') and hasattr(self, 'total_preview_pages'):
                self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")
            
            # Disable prev button if on first page
            if self.preview_page == 1:
                self.prev_page_btn.config(state=tk.DISABLED)
            
            # Always enable next button if not on last page
            if hasattr(self, 'total_preview_pages') and self.preview_page < self.total_preview_pages:
                self.next_page_btn.config(state=tk.NORMAL)
            
    def _go_to_next_page(self):
        """Navigate to the next page of thumbnails"""
        if hasattr(self, 'preview_page') and hasattr(self, 'total_preview_pages') and self.preview_page < self.total_preview_pages:
            # Update status
            self.update_status(self.get_text("loading_preview"))
            
            # Show loading indicator
            self.progress_bar.start(10)
            
            # Change page
            self.preview_page += 1
            
            # Log for debugging
            logging.info(f"Moving to next page: {self.preview_page}/{self.total_preview_pages}")
            
            # Update preview panel
            self._update_preview_panel()
            
            # Update page info
            if hasattr(self, 'page_info_label'):
                self.page_info_label.config(text=f"{self.get_text('page')} {self.preview_page}/{self.total_preview_pages}")
            
            # Disable next button if on last page
            if self.preview_page == self.total_preview_pages:
                self.next_page_btn.config(state=tk.DISABLED)
            
            # Always enable prev button if not on first page
            if self.preview_page > 1:
                self.prev_page_btn.config(state=tk.NORMAL)
    
    def _update_preview_content(self, file_path):
        """Update the content of the preview window without destroying it"""
        if not hasattr(self, 'preview_window') or not self.preview_window:
            return False
            
        try:
            # Get file details
            file_name = os.path.basename(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            
            # Get modification and creation dates
            file_mod_date = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S') if os.path.exists(file_path) else ""
            
            # Try to get creation date (platform specific)
            try:
                if os.name == 'nt':  # Windows
                    file_creation_date = datetime.datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
                else:  # Unix/Linux/Mac
                    # On Unix, getctime returns status change time, not creation time, since Unix doesn't track creation time
                    # Using stat to get the best approximation
                    stat_info = os.stat(file_path)
                    file_creation_date = datetime.datetime.fromtimestamp(stat_info.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            except:
                file_creation_date = ""
            
            # Update window title (optional)
            self.preview_window.title(self.get_text("preview_window_title"))
            
            # Find and update all widgets in the info frame
            if hasattr(self, '_preview_info_widgets'):
                # Update file name
                if 'name_label' in self._preview_info_widgets:
                    self._preview_info_widgets['name_label'].config(text=file_name)
                
                # Update file type icon
                if 'icon_label' in self._preview_info_widgets:
                    # File type icon or text based on extension
                    if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                        icon_text = "IMG"
                        icon_bg = "#28a745"  # Green for images
                    elif file_extension in ['.pdf']:
                        icon_text = "PDF"
                        icon_bg = "#dc3545"  # Red for PDFs 
                    elif file_extension in ['.psd', '.ai', '.eps']:
                        icon_text = "PSD"
                        icon_bg = "#6610f2"  # Purple for design files
                    elif file_extension in ['.doc', '.docx', '.txt', '.rtf']:
                        icon_text = "DOC"
                        icon_bg = "#007bff"  # Blue for documents
                    elif file_extension in ['.mp3', '.wav', '.flac', '.aac', '.ogg']:
                        icon_text = "AUD"
                        icon_bg = "#fd7e14"  # Orange for audio
                    elif file_extension in ['.mp4', '.mov', '.mkv', '.avi']:
                        icon_text = "VID"
                        icon_bg = "#6f42c1"  # Purple for video
                    else:
                        # Get the file extension without dot and make it uppercase
                        icon_text = file_extension.upper().replace(".", "") if file_extension else "FILE"
                        icon_bg = "#6c757d"  # Gray for other types
                        
                    self._preview_info_widgets['icon_label'].config(text=icon_text, bg=icon_bg)
                
                # Update size label
                if 'size_label' in self._preview_info_widgets:
                    self._preview_info_widgets['size_label'].config(text=f"{self.get_text('file_size')}: {self.format_file_size(file_size)}")
                
                # Update extension label
                if 'ext_label' in self._preview_info_widgets:
                    self._preview_info_widgets['ext_label'].config(text=f"{self.get_text('file_extension')}: {file_extension}")
                
                # Update created date label
                if 'created_label' in self._preview_info_widgets:
                    self._preview_info_widgets['created_label'].config(text=f"{self.get_text('creation_date')}: {file_creation_date}")
                
                # Update modified date label
                if 'modified_label' in self._preview_info_widgets:
                    self._preview_info_widgets['modified_label'].config(text=f"{self.get_text('modification_date')}: {file_mod_date}")
                
                # Update navigation label
                if 'nav_label' in self._preview_info_widgets and hasattr(self, 'current_preview_files') and len(self.current_preview_files) > 1 and self.current_preview_index >= 0:
                    nav_text = f"{self.current_preview_index + 1} / {len(self.current_preview_files)}"
                    self._preview_info_widgets['nav_label'].config(text=nav_text)
                
                # Update navigation buttons state
                if 'prev_button' in self._preview_info_widgets:
                    self._preview_info_widgets['prev_button'].config(state=tk.NORMAL if self.current_preview_index > 0 else tk.DISABLED)
                
                if 'next_button' in self._preview_info_widgets:
                    self._preview_info_widgets['next_button'].config(state=tk.NORMAL if self.current_preview_index < len(self.current_preview_files) - 1 else tk.DISABLED)
            
            # Clear preview frame and add new content
            if hasattr(self, '_preview_content_frame'):
                # Remove all widgets from preview frame
                for widget in self._preview_content_frame.winfo_children():
                    widget.destroy()
                
                # Check file type and generate preview
                if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                    self.preview_image(self._preview_content_frame, file_path)
                elif file_extension in ['.pdf']:
                    self.preview_pdf(self._preview_content_frame, file_path)
                elif file_extension in ['.psd', '.ai', '.eps']:
                    self.preview_design_file(self._preview_content_frame, file_path)
                else:
                    # Show unsupported message
                    label = tk.Label(
                        self._preview_content_frame, 
                        text=self.get_text("preview_not_supported"),
                        font=("Segoe UI", 12),
                        bg="#e9ecef"
                    )
                    label.pack(pady=20)
            
            return True
        except Exception as e:
            logging.error(f"Error updating preview content: {str(e)}")
            return False

    def _go_to_prev_file(self):
        """Navigate to the previous file in the preview window"""
        if not hasattr(self, 'current_preview_files') or not self.current_preview_files:
            return
            
        if self.current_preview_index > 0:
            # Get previous file
            self.current_preview_index -= 1
            prev_file = self.current_preview_files[self.current_preview_index]
            
            # Open preview for that file
            if "path" in prev_file and "name" in prev_file:
                # Get full file path
                file_path = os.path.join(prev_file["path"], prev_file["name"])
                
                # Always try to update existing preview window first
                result = self._update_preview_content(file_path)
                if not result:
                    # Only if updating fails, create a new window
                    logging.debug("Falling back to creating a new preview window")
                    self.create_file_preview_window(file_path, self.current_preview_index)
            
    def _go_to_next_file(self):
        """Navigate to the next file in the preview window"""
        if not hasattr(self, 'current_preview_files') or not self.current_preview_files:
            return
            
        if self.current_preview_index < len(self.current_preview_files) - 1:
            # Get next file
            self.current_preview_index += 1
            next_file = self.current_preview_files[self.current_preview_index]
            
            # Open preview for that file
            if "path" in next_file and "name" in next_file:
                # Get full file path
                file_path = os.path.join(next_file["path"], next_file["name"])
                
                # Always try to update existing preview window first
                result = self._update_preview_content(file_path)
                if not result:
                    # Only if updating fails, create a new window
                    logging.debug("Falling back to creating a new preview window")
                    self.create_file_preview_window(file_path, self.current_preview_index)
            
    def on_subfolder_changed(self):
        """Called when the Include Subfolders checkbox state changes"""
        # If a folder is selected, reload files when subfolder setting changes
        if self.selected_folder_path:
            # Show loading indicator
            self.progress_bar.start(10)
            
            # Update status bar with appropriate message
            if self.include_subfolders.get():
                status_msg = self.get_text("loading_subfolders") if "loading_subfolders" in self.languages[self.current_language] else "Loading subfolders..."
            else:
                status_msg = self.get_text("folder_loading")
            self.update_status(status_msg)
            
            # Start a new thread to reload the files with the new subfolder setting
            self.load_files_thread()
            
            # Log the change for debugging
            logging.info(f"Subfolder option changed to: {self.include_subfolders.get()}")
            
            # Change the tooltip to reflect current state
            subfolders_tooltip = self.get_text("tooltip_subfolders")
            if self.include_subfolders.get():
                subfolders_tooltip += " ✓"  # Add a checkmark to indicate it's enabled
            self.create_tooltip(self.subfolder_cb, subfolders_tooltip)

    def open_website(self, url):
        """Open a website URL in the default browser"""
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception as e:
            logging.error(f"Failed to open website: {e}")
            messagebox.showerror("Error", f"Failed to open website: {e}")
            
    def check_for_updates(self, silent=False):
        """
        GitHub'dan son sürüm bilgisini alıp mevcut sürümle karşılaştırır
        silent=True ise sadece güncelleme varsa bildirim yapar
        """
        try:
            # Sürüm kontrolü için HTTP isteği gönder
            response = requests.get(self.github_version_url, timeout=5)
            
            if response.status_code == 200:
                # Uzaktaki sürüm bilgisini al (boşlukları temizle)
                github_version = response.text.strip()
                
                # Sürüm karşılaştırması yap
                if github_version != self.current_version:
                    # Güncelleme mevcut
                    if messagebox.askyesno(
                        self.get_text("update_available"),
                        self.get_text("update_available_message").format(github_version, self.current_version)
                    ):
                        # Replit ortamında doğrudan indirme devre dışı, sayfayı aç
                        self.open_website(self.github_download_url)
                        # Normal sistemlerde indirme aşağıdaki şekilde olacak
                        # self.download_update(github_version)
                elif not silent:  
                    # Zaten son sürüm kullanılıyor ve sessiz mod değilse bildirim yap
                    messagebox.showinfo(
                        self.get_text("no_update_available"),
                        self.get_text("no_update_available_message")
                    )
                    
                return github_version != self.current_version
            else:
                if not silent:
                    # Hata durumunda bildirim yap (sessiz mod değilse)
                    self.show_error(
                        self.get_text("update_check_error"),
                        self.get_text("update_check_error_message")
                    )
                return False
                
        except Exception as e:
            if not silent:
                # Bağlantı hatası bildirim yap (sessiz mod değilse)
                self.show_error(
                    self.get_text("update_check_error"),
                    f"{self.get_text('update_check_error_message')} ({str(e)})"
                )
            logging.error(f"Güncelleme kontrolü hatası: {str(e)}")
            return False
            
    def download_update(self, new_version):
        """
        Yeni sürümü GitHub'dan indir ve otomatik olarak güncelle
        """
        try:
            # İndirme URL'ini oluştur (GitHub releases sayfasından indirme)
            download_url = f"{self.github_download_url}/download/v{new_version}/ListeKolay_v{new_version}.zip"
            
            # İndirme ilerleme penceresini oluştur
            download_window = tk.Toplevel(self.root)
            download_window.title(self.get_text("downloading_update"))
            download_window.geometry("400x150")
            download_window.resizable(False, False)
            download_window.configure(bg="#e9ecef")
            download_window.transient(self.root)  # Ana pencereye bağlı
            download_window.grab_set()  # Diğer işlemleri engelle
            
            # Pencere merkezi konumlandırma
            download_window.update_idletasks()
            width = download_window.winfo_width()
            height = download_window.winfo_height()
            x = (download_window.winfo_screenwidth() // 2) - (width // 2)
            y = (download_window.winfo_screenheight() // 2) - (height // 2)
            download_window.geometry(f"+{x}+{y}")
            
            # İndirme durumu etiketi
            status_label = tk.Label(
                download_window, 
                text=self.get_text("downloading_update_message").format(new_version),
                font=("Segoe UI", 10),
                bg="#e9ecef",
                fg="#212529"
            )
            status_label.pack(pady=(20, 10))
            
            # İndirme ilerleme çubuğu
            progress_bar = ttk.Progressbar(
                download_window, 
                orient=tk.HORIZONTAL, 
                length=350, 
                mode='indeterminate'
            )
            progress_bar.pack(pady=10, padx=25)
            progress_bar.start(10)
            
            # İptal butonu
            cancel_button = tk.Button(
                download_window,
                text=self.get_text("cancel_button"),
                command=download_window.destroy,
                font=("Segoe UI", 9),
                bg="#dc3545",
                fg="white",
                activebackground="#c82333",
                activeforeground="white",
                bd=0,
                padx=10
            )
            cancel_button.pack(pady=10)
            
            # İndirme fonksiyonu (thread içinde çalışacak)
            def download_thread():
                try:
                    # Geçici dosya oluştur
                    temp_dir = tempfile.gettempdir()
                    output_file = os.path.join(temp_dir, f"ListeKolay_v{new_version}.zip")
                    
                    # İndirme işlemini başlat
                    response = requests.get(download_url, stream=True)
                    
                    if response.status_code == 200:
                        # Dosyayı kaydet
                        with open(output_file, "wb") as f:
                            for chunk in response.iter_content(chunk_size=1024):
                                if chunk:  # Boş paketleri filtrele
                                    f.write(chunk)
                        
                        # İndirme penceresini kapat
                        download_window.after(0, download_window.destroy)
                        
                        # İndirme tamamlandı mesajı
                        if messagebox.showinfo(
                            self.get_text("download_complete"),
                            self.get_text("download_complete_message")
                        ):
                            # Yeni sürümü başlat
                            self.launch_updated_version(output_file)
                    else:
                        # İndirme hatası
                        download_window.after(0, download_window.destroy)
                        messagebox.showerror(
                            self.get_text("download_error"),
                            self.get_text("download_error_message").format(f"HTTP {response.status_code}")
                        )
                except Exception as e:
                    # İndirme sırasında hata
                    logging.error(f"Download error: {e}")
                    download_window.after(0, download_window.destroy)
                    messagebox.showerror(
                        self.get_text("download_error"),
                        self.get_text("download_error_message").format(str(e))
                    )
                    
            # İndirme thread'ini başlat
            download_thread_obj = threading.Thread(target=download_thread)
            download_thread_obj.daemon = True
            download_thread_obj.start()
                
        except Exception as e:
            # Genel hata durumu
            logging.error(f"Update download error: {e}")
            messagebox.showerror(
                self.get_text("download_error"),
                self.get_text("download_error_message").format(str(e))
            )
            
    def launch_updated_version(self, zip_file):
        """
        İndirilen zip dosyasını çıkart ve yeni sürümü başlat
        """
        try:
            # Programın mevcut konumunu al
            current_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            
            # Zip dosyasını çıkartma ve yeni sürümü başlatma işlemleri için yardımcı betik oluştur
            updater_script = os.path.join(tempfile.gettempdir(), "listekolay_updater.py")
            
            # Yolları uygun şekilde formatlayalım
            safe_zip_path = zip_file.replace('\\', '\\\\')
            safe_current_dir = current_dir.replace('\\', '\\\\')
            safe_python_path = os.path.join(current_dir, 'listekolay.py').replace('\\', '\\\\')
            
            # Dosyayı normal string oluşturarak yazalım
            updater_content = """
import os
import sys
import time
import zipfile
import shutil
import subprocess

def update_app():
    # Orijinal uygulamanın kapanması için bekle
    time.sleep(2)
    
    try:
        # Zip dosyasını çıkart
        with zipfile.ZipFile(r"{0}", "r") as zip_ref:
            zip_ref.extractall(r"{1}")
        
        # Yeni sürümü başlat
        subprocess.Popen(["python", r"{2}"])
        
        return True
    except Exception as e:
        print(f"Update error: {{e}}")
        return False

if __name__ == "__main__":
    update_app()
"""
            # Format ile değerleri ekle
            formatted_content = updater_content.format(safe_zip_path, safe_current_dir, safe_python_path)
            
            # Dosyaya yaz
            with open(updater_script, "w", encoding="utf-8") as f:
                f.write(formatted_content)
            
            # Yardımcı betiği başlat
            subprocess.Popen([sys.executable, updater_script])
            
            # Mevcut uygulamayı kapat
            self.on_close()
            
        except Exception as e:
            # Güncelleme hatası
            logging.error(f"Update launch error: {e}")
            messagebox.showerror(
                self.get_text("download_error"),
                self.get_text("download_error_message").format(str(e))
            )
    def save_config(self):
        """Kullanıcı ayarlarını config.json dosyasına kaydet"""
        try:
            config = {
                "language": self.current_language,
                # "last_folder" değeri artık kaydedilmiyor
                "include_subfolders": self.include_subfolders.get(),
                "export_formats": {
                    "text": self.export_formats["text"].get(),
                    "excel": self.export_formats["excel"].get(),
                    "word": self.export_formats["word"].get(),
                    "html": self.export_formats["html"].get()
                },
                "save_to_desktop": self.save_to_desktop.get(),
                "sort_criteria": self.selected_sort.get(),
                # "view_mode" değeri artık kaydedilmiyor
                "is_dark_mode": self.is_dark_mode.get()
            }
            
            # Config dosyasını oluştur
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
            
            # Config dosyasının bir yedeğini oluştur (kaydetmeden önce)
            try:
                backup_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json.bak")
                if os.path.exists(config_path):
                    shutil.copy2(config_path, backup_path)
                    logging.info("Config dosyası yedeklendi")
            except Exception as backup_error:
                logging.error(f"Config dosyası yedeklenirken hata oluştu: {str(backup_error)}")
            
            # Asıl kayıt işlemi
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=4, ensure_ascii=False)
                
            logging.info("Ayarlar başarıyla kaydedildi")
        except Exception as e:
            logging.error(f"Ayarları kaydederken hata oluştu: {str(e)}")
            
            # 1. Yedekten geri yüklemeyi dene
            try:
                backup_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json.bak")
                config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
                if os.path.exists(backup_path):
                    if os.path.exists(config_path):
                        os.remove(config_path)
                    shutil.copy2(backup_path, config_path)
                    logging.info("Config dosyası yedekten geri yüklendi")
                    return
            except Exception as restore_error:
                logging.error(f"Config yedekten geri yüklenirken hata oluştu: {str(restore_error)}")
            
            # 2. Yöntem: Dosya yazma hatası olursa, tekrar deneme yaparak veri kaybını önleyelim
            try:
                # Önce temp dosyaya yaz, sonra adını değiştir (daha güvenli yaklaşım)
                temp_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config_temp.json")
                with open(temp_path, 'w', encoding='utf-8') as f:
                    json.dump(config, f, indent=4, ensure_ascii=False)
                
                # Başarıyla yazıldıysa, asıl dosyanın yerine koy
                if os.path.exists(temp_path):
                    if os.path.exists(config_path):
                        os.remove(config_path)
                    os.rename(temp_path, config_path)
                    logging.info("İkinci denemede ayarlar başarıyla kaydedildi")
            except Exception as e2:
                logging.error(f"Ayarları tekrar kaydederken ikinci hata oluştu: {str(e2)}")

    def load_config(self):
        """config.json dosyasından kullanıcı ayarlarını yükle"""
        try:
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
            
            # Dosya yoksa oluştur ve varsayılan ayarlarla devam et
            if not os.path.exists(config_path):
                logging.info("Yapılandırma dosyası bulunamadı, varsayılan ayarlarla oluşturuluyor")
                self.save_config()  # Varsayılan ayarlarla config.json dosyası oluştur
                return
                
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            
            # Dil ayarı
            if "language" in config and config["language"] in self.languages:
                saved_language = config["language"]
                self.current_language = saved_language
                self.language_var.set(saved_language)
                # Dil değişikliğini hemen uygula
                logging.info(f"Config'den yüklenen dil: {saved_language}")
                # UI dil değişikliğini uygula
                self.update_ui_language()
                # Ana başlıkları güncelle
                self.update_main_titles()
                # Kategorileri güncelle
                self.populate_categories()
                
            # Son klasör artık config'den yüklenmiyor - açılışta boş kalacak
            # Kullanıcının klasör seçmesi bekleniyor
            self.folder_path_var.set("")
            if hasattr(self, 'selected_folder_path'):
                delattr(self, 'selected_folder_path')
            
            # İlk açılışta görünüm modunu 'list' olarak ayarla
            if hasattr(self, 'view_mode_var'):
                self.view_mode_var.set("list")
                self.set_view_mode("list")
                logging.info("İlk açılışta listeleme moduna geçildi")
                
            # Alt klasörler dahil
            if "include_subfolders" in config:
                self.include_subfolders.set(config["include_subfolders"])
                
            # Dışa aktarma formatları
            if "export_formats" in config:
                formats = config["export_formats"]
                for fmt in self.export_formats:
                    if fmt in formats:
                        self.export_formats[fmt].set(formats[fmt])
                        
            # Masaüstüne kaydet
            if "save_to_desktop" in config:
                self.save_to_desktop.set(config["save_to_desktop"])
                
            # Sıralama kriteri
            if "sort_criteria" in config:
                self.selected_sort.set(config["sort_criteria"])
                
            # Tema modu (açık/koyu)
            if "is_dark_mode" in config:
                self.is_dark_mode.set(config["is_dark_mode"])
                # Tema modunu hemen uygula
                self.toggle_theme_mode()
                logging.info(f"Tema modu yüklendi: {'Koyu' if config['is_dark_mode'] else 'Açık'}")
                
            # Görünüm modu her zaman list modunda başlasın
            # view_mode artık config'den yüklenmiyor
            self.view_mode_var.set("list")
            self.set_view_mode("list")
            logging.info("Program ilk açılışta liste görünümünde başlatıldı")
                
            logging.info("Ayarlar başarıyla yüklendi")
            
            # Otomatik güncelleme kontrolü (sessiz mod)
            # Bu işlemi bir thread'de çalıştıralım ki arayüz bloke olmasın
            try:
                update_thread = threading.Thread(target=self.check_for_updates, args=(True,))
                update_thread.daemon = True
                update_thread.start()
                logging.info("Otomatik güncelleme kontrolü başlatıldı")
            except Exception as e:
                logging.error(f"Otomatik güncelleme kontrolü başlatılamadı: {str(e)}")
        except Exception as e:
            logging.error(f"Ayarları yüklerken hata oluştu: {str(e)}")
            
    def on_close(self):
        # Ask for confirmation before exiting
        response = messagebox.askyesno(
            self.get_text("confirm_exit_title"),
            self.get_text("confirm_exit_message"),
            icon=messagebox.QUESTION
        )
        
        if response:
            # Ayarları kaydet
            self.save_config()
            
            # Log application exit
            logging.info("Program sonlandırıldı")
            self.root.destroy()
            
    def create_file_preview_window(self, file_path, file_index=-1):
        """Create a window to preview the selected file"""
        try:
            # Update the current preview index if provided
            if file_index >= 0 and file_index < len(self.current_preview_files):
                self.current_preview_index = file_index
            elif file_index == -1 and hasattr(self, 'current_preview_files'):
                # Find the file index if not provided
                self.current_preview_index = next((i for i, f in enumerate(self.current_preview_files) 
                                               if f["path"] == file_path), -1)
            
            # Close existing preview window if one exists
            if hasattr(self, 'preview_window') and self.preview_window and self.preview_window.winfo_exists():
                self.preview_window.destroy()
                
            # Create a new top-level window
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title(self.get_text("preview_window_title"))
            self.preview_window.geometry("900x700")  # Slightly larger for more info
            self.preview_window.minsize(500, 400)
            
            # Get file details
            file_name = os.path.basename(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            
            # Get modification and creation dates
            file_mod_date = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S') if os.path.exists(file_path) else ""
            
            # Try to get creation date (platform specific)
            try:
                if os.name == 'nt':  # Windows
                    file_creation_date = datetime.datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
                else:  # Unix/Linux/Mac
                    # On Unix, getctime returns status change time, not creation time, since Unix doesn't track creation time
                    # Using stat to get the best approximation
                    stat_info = os.stat(file_path)
                    file_creation_date = datetime.datetime.fromtimestamp(stat_info.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            except:
                file_creation_date = ""
            
            # Create a container for the info panel at the top
            info_frame = tk.Frame(self.preview_window, bg="#f8f9fa", height=80, relief=tk.GROOVE, bd=1)
            info_frame.pack(fill=tk.X, padx=10, pady=(10, 0))
            info_frame.pack_propagate(False)  # Fixed height
            
            # Left side info (file name and basic details)
            left_info = tk.Frame(info_frame, bg="#f8f9fa")
            left_info.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=5)
            
            # File name with icon - use a larger font and make it more prominent
            file_name_frame = tk.Frame(left_info, bg="#f8f9fa")
            file_name_frame.pack(fill=tk.X, anchor=tk.W)
            
            # File type icon or text based on extension
            if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                icon_text = "IMG"
                icon_bg = "#28a745"  # Green for images
            elif file_extension in ['.pdf']:
                icon_text = "PDF"
                icon_bg = "#dc3545"  # Red for PDFs 
            elif file_extension in ['.psd', '.ai', '.eps']:
                icon_text = "PSD"
                icon_bg = "#6610f2"  # Purple for design files
            elif file_extension in ['.doc', '.docx', '.txt', '.rtf']:
                icon_text = "DOC"
                icon_bg = "#007bff"  # Blue for documents
            elif file_extension in ['.mp3', '.wav', '.flac', '.aac', '.ogg']:
                icon_text = "AUD"
                icon_bg = "#fd7e14"  # Orange for audio
            elif file_extension in ['.mp4', '.mov', '.mkv', '.avi']:
                icon_text = "VID"
                icon_bg = "#6f42c1"  # Purple for video
            else:
                # Get the file extension without dot and make it uppercase
                icon_text = file_extension.upper().replace(".", "") if file_extension else "FILE"
                icon_bg = "#6c757d"  # Gray for other types
            
            # Create icon label
            icon_label = tk.Label(
                file_name_frame,
                text=icon_text,
                font=("Segoe UI", 10, "bold"),
                bg=icon_bg,
                fg="white",
                width=4,
                padx=5,
                pady=2
            )
            icon_label.pack(side=tk.LEFT, padx=(0, 8))
            
            # File name label
            name_label = tk.Label(
                file_name_frame,
                text=file_name,
                font=("Segoe UI", 12, "bold"),
                bg="#f8f9fa",
                fg="#212529",
                anchor=tk.W
            )
            name_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            # File details
            details_frame = tk.Frame(left_info, bg="#f8f9fa")
            details_frame.pack(fill=tk.X, anchor=tk.W, pady=(5, 0))
            
            # Create two rows for details for better organization
            details_row1 = tk.Frame(details_frame, bg="#f8f9fa")
            details_row1.pack(fill=tk.X, pady=(0, 3))
            
            details_row2 = tk.Frame(details_frame, bg="#f8f9fa")
            details_row2.pack(fill=tk.X)
            
            # Row 1: Size and Extension
            size_label = tk.Label(
                details_row1,
                text=f"{self.get_text('file_size')}: {self.format_file_size(file_size)}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W,
                width=25
            )
            size_label.pack(side=tk.LEFT, padx=(0, 10))
            
            ext_label = tk.Label(
                details_row1,
                text=f"{self.get_text('file_extension')}: {file_extension}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W
            )
            ext_label.pack(side=tk.LEFT, fill=tk.X)
            
            # Row 2: Creation and Modification dates
            created_label = tk.Label(
                details_row2,
                text=f"{self.get_text('creation_date')}: {file_creation_date}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W,
                width=25
            )
            created_label.pack(side=tk.LEFT, padx=(0, 10))
            
            modified_label = tk.Label(
                details_row2,
                text=f"{self.get_text('modification_date')}: {file_mod_date}",
                font=("Segoe UI", 9),
                bg="#f8f9fa",
                fg="#495057",
                anchor=tk.W
            )
            modified_label.pack(side=tk.LEFT, fill=tk.X)
            
            # Right side navigation info - only if we have a valid file list and index
            right_info = tk.Frame(info_frame, bg="#f8f9fa")
            right_info.pack(side=tk.RIGHT, fill=tk.Y, padx=10, pady=5)
            
            # Store navigation label reference
            nav_label = None
            
            # Show navigation info if we have a valid file list
            if hasattr(self, 'current_preview_files') and len(self.current_preview_files) > 1 and self.current_preview_index >= 0:
                nav_text = f"{self.current_preview_index + 1} / {len(self.current_preview_files)}"
                nav_label = tk.Label(
                    right_info,
                    text=nav_text,
                    font=("Segoe UI", 10),
                    bg="#f8f9fa",
                    fg="#495057"
                )
                nav_label.pack(side=tk.RIGHT)
            
            # Add a frame for the preview content
            preview_frame = tk.Frame(self.preview_window, bg="#e9ecef")
            preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # Check file type and generate preview
            if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                self.preview_image(preview_frame, file_path)
            elif file_extension in ['.pdf']:
                self.preview_pdf(preview_frame, file_path)
            elif file_extension in ['.psd', '.ai', '.eps']:
                self.preview_design_file(preview_frame, file_path)
            else:
                # Show unsupported message
                label = tk.Label(
                    preview_frame, 
                    text=self.get_text("preview_not_supported"),
                    font=("Segoe UI", 12),
                    bg="#e9ecef"
                )
                label.pack(pady=20)
            
            # Add button frame for action buttons at the bottom
            button_frame = tk.Frame(self.preview_window, bg="#e9ecef")
            button_frame.pack(pady=10, fill=tk.X)
            
            # Navigation buttons references
            prev_button = None
            next_button = None
            
            # Navigation buttons - only show if we have multiple files
            if hasattr(self, 'current_preview_files') and len(self.current_preview_files) > 1 and self.current_preview_index >= 0:
                # Previous file button
                prev_button = tk.Button(
                    button_frame, 
                    text="← " + self.get_text("prev_page"),
                    command=self._go_to_prev_file,
                    bg="#17a2b8",
                    fg="white",
                    relief=tk.GROOVE,
                    padx=10,
                    state=tk.NORMAL if self.current_preview_index > 0 else tk.DISABLED
                )
                prev_button.pack(side=tk.LEFT, padx=(10, 5))
                
                # Next file button
                next_button = tk.Button(
                    button_frame, 
                    text=self.get_text("next_page") + " →",
                    command=self._go_to_next_file,
                    bg="#17a2b8",
                    fg="white",
                    relief=tk.GROOVE,
                    padx=10,
                    state=tk.NORMAL if self.current_preview_index < len(self.current_preview_files) - 1 else tk.DISABLED
                )
                next_button.pack(side=tk.LEFT, padx=(0, 5))
            
            # Open file directly button
            open_file_button = tk.Button(
                button_frame, 
                text=self.get_text("open_file"),
                command=lambda: self.open_file(file_path),
                bg="#28a745",
                fg="white",
                relief=tk.GROOVE,
                padx=10
            )
            open_file_button.pack(side=tk.LEFT, padx=(10 if not hasattr(self, 'current_preview_files') or len(self.current_preview_files) <= 1 else 0, 5))
            
            # Open file location button
            open_location_button = tk.Button(
                button_frame, 
                text=self.get_text("open_file_location"),
                command=lambda: self.open_file_location_by_path(file_path),
                bg="#007bff",
                fg="white",
                relief=tk.GROOVE,
                padx=10
            )
            open_location_button.pack(side=tk.LEFT, padx=(0, 5))
            
            # Close button 
            close_button = tk.Button(
                button_frame, 
                text=self.get_text("exit"),
                command=self.preview_window.destroy,
                bg="#e9ecef",
                relief=tk.GROOVE,
                padx=10
            )
            close_button.pack(side=tk.RIGHT, padx=10)
            
            # Store widget references for later updates (when navigating between files)
            self._preview_info_widgets = {
                'icon_label': icon_label,
                'name_label': name_label,
                'size_label': size_label,
                'ext_label': ext_label,
                'created_label': created_label,
                'modified_label': modified_label,
                'open_file_button': open_file_button,
                'open_location_button': open_location_button
            }
            
            # Add navigation widgets if they exist
            if nav_label:
                self._preview_info_widgets['nav_label'] = nav_label
            if prev_button:
                self._preview_info_widgets['prev_button'] = prev_button
            if next_button:
                self._preview_info_widgets['next_button'] = next_button
                
            # Store reference to content frame for updates
            self._preview_content_frame = preview_frame
            
            # Center the window on the screen
            self.preview_window.update_idletasks()
            width = self.preview_window.winfo_width()
            height = self.preview_window.winfo_height()
            x = (self.preview_window.winfo_screenwidth() // 2) - (width // 2)
            y = (self.preview_window.winfo_screenheight() // 2) - (height // 2)
            self.preview_window.geometry('{}x{}+{}+{}'.format(width, height, x, y))
            
        except Exception as e:
            logging.error(f"Error creating preview window: {str(e)}")
            messagebox.showerror(
                self.get_text("error"), 
                f"{self.get_text('preview_error')} {str(e)}"
            )
    
    def preview_image(self, parent_frame, file_path):
        """Display an image preview"""
        try:
            # Load the image
            original_image = Image.open(file_path)
            
            # Create a canvas for the image with scrollbars
            canvas_frame = tk.Frame(parent_frame)
            canvas_frame.pack(fill=tk.BOTH, expand=True)
            
            h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
            v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
            
            canvas = tk.Canvas(
                canvas_frame,
                xscrollcommand=h_scrollbar.set,
                yscrollcommand=v_scrollbar.set,
                bg="#ffffff"
            )
            
            h_scrollbar.config(command=canvas.xview)
            v_scrollbar.config(command=canvas.yview)
            
            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            # Calculate scaled dimensions (max 700x500 initial view)
            max_width, max_height = 700, 500
            img_width, img_height = original_image.size
            
            # Calculate scaled size while maintaining aspect ratio
            scale_factor = min(max_width/img_width, max_height/img_height)
            if scale_factor < 1:  # Only scale down, not up
                new_width = int(img_width * scale_factor)
                new_height = int(img_height * scale_factor)
                display_image = original_image.resize((new_width, new_height), Image.LANCZOS)
            else:
                display_image = original_image
                
            # Convert to PhotoImage for Tkinter
            photo = ImageTk.PhotoImage(display_image)
            
            # Add image to canvas
            canvas.create_image(0, 0, image=photo, anchor=tk.NW)
            canvas.image = photo  # Keep a reference to prevent garbage collection
            
            # Configure canvas scrollable area
            canvas.config(scrollregion=canvas.bbox(tk.ALL))
            
            # Add info label with image details
            info_text = f"{img_width}x{img_height} px, {os.path.basename(file_path)}"
            info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
            info_label.pack(pady=5)
            
        except Exception as e:
            logging.error(f"Error previewing image: {str(e)}")
            error_label = tk.Label(
                parent_frame, 
                text=f"{self.get_text('preview_error')} {str(e)}",
                fg="red", 
                bg="#e9ecef"
            )
            error_label.pack(pady=20)
    
    def preview_pdf(self, parent_frame, file_path):
        """Display a PDF preview"""
        try:
            # Create a canvas for the PDF with scrollbars
            canvas_frame = tk.Frame(parent_frame)
            canvas_frame.pack(fill=tk.BOTH, expand=True)
            
            h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
            v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
            
            canvas = tk.Canvas(
                canvas_frame,
                xscrollcommand=h_scrollbar.set,
                yscrollcommand=v_scrollbar.set,
                bg="#ffffff"
            )
            
            h_scrollbar.config(command=canvas.xview)
            v_scrollbar.config(command=canvas.yview)
            
            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            # Open the PDF using PyMuPDF
            pdf_document = fitz.open(file_path)
            
            # Get first page of the PDF
            first_page = pdf_document.load_page(0)
            
            # Set zoom factor for better quality
            zoom = 2.0
            mat = fitz.Matrix(zoom, zoom)
            
            # Convert page to an image
            pix = first_page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Scale down if needed
            max_width, max_height = 700, 500
            img_width, img_height = img.size
            
            # Calculate scaled size while maintaining aspect ratio
            scale_factor = min(max_width/img_width, max_height/img_height)
            if scale_factor < 1:  # Only scale down, not up
                new_width = int(img_width * scale_factor)
                new_height = int(img_height * scale_factor)
                img = img.resize((new_width, new_height), get_pil_resize_method())
            
            # Convert to PhotoImage for Tkinter
            photo = ImageTk.PhotoImage(img)
            
            # Add image to canvas
            canvas.create_image(0, 0, image=photo, anchor=tk.NW)
            canvas.image = photo  # Keep a reference to prevent garbage collection
            
            # Configure canvas scrollable area
            canvas.config(scrollregion=canvas.bbox(tk.ALL))
            
            # Add info label with PDF details
            info_text = f"PDF: {os.path.basename(file_path)}, {pdf_document.page_count} pages"
            info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
            info_label.pack(pady=5)
            
            # Close the document when done
            pdf_document.close()
            
        except Exception as e:
            logging.error(f"Error previewing PDF: {str(e)}")
            error_label = tk.Label(
                parent_frame, 
                text=f"{self.get_text('preview_error')} {str(e)}",
                fg="red", 
                bg="#e9ecef"
            )
            error_label.pack(pady=20)
    
    def preview_design_file(self, parent_frame, file_path):
        """Display a preview for design files (PSD, AI, EPS)"""
        try:
            # Normalize file path to avoid Windows/Unix path issues
            file_path = os.path.normpath(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Try to open the file with PIL first
            try:
                img = Image.open(file_path)
                self._display_design_preview(parent_frame, img, file_path)
                return
            except Exception as pil_error:
                logging.error(f"PIL could not open file: {str(pil_error)}")
                
                # For AI and EPS, try using PyMuPDF (many AI files are PDF-compatible)
                if file_extension in ['.ai', '.eps', '.pdf']:
                    try:
                        pdf_document = fitz.open(file_path)
                        first_page = pdf_document.load_page(0)
                        
                        zoom = 2.0
                        mat = fitz.Matrix(zoom, zoom)
                        pix = first_page.get_pixmap(matrix=mat)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        
                        pdf_document.close()
                        self._display_design_preview(parent_frame, img, file_path)
                        return
                    except Exception as pdf_error:
                        logging.error(f"PyMuPDF could not open file: {str(pdf_error)}")
                        
                        # For EPS files that failed with PyMuPDF, try pdf2image
                        if file_extension == '.eps':
                            try:
                                # Convert EPS to images using pdf2image
                                from pdf2image import convert_from_path
                                images = convert_from_path(file_path, first_page=1, last_page=1)
                                if images and len(images) > 0:
                                    self._display_design_preview(parent_frame, images[0], file_path)
                                    return
                            except Exception as eps_error:
                                logging.error(f"pdf2image could not convert EPS: {str(eps_error)}")
                
                # If all else fails, create a placeholder
                bg_color = "#f0f0f0"
                text_color = "#2c3e50"
                border_color = "#2c3e50"
                img = Image.new("RGB", (400, 300), color=bg_color)
                draw = ImageDraw.Draw(img)
                draw.rectangle([20, 20, 380, 280], outline=border_color, width=2)
                
                # Display file type in center
                if file_extension.startswith('.'):
                    file_type = file_extension[1:].upper()
                else:
                    file_type = file_extension.upper()
                
                # Draw file type in the center of the placeholder
                draw.text((200, 150), file_type, fill=text_color)
                
                # Draw file name at the bottom
                file_name = os.path.basename(file_path)
                if len(file_name) > 30:  # Truncate long file names
                    file_name = file_name[:27] + "..."
                draw.text((200, 220), file_name, fill=text_color)
                
                self._display_design_preview(parent_frame, img, file_path)
                return
            
        except Exception as e:
            logging.error(f"Error previewing design file: {str(e)}")
            error_label = tk.Label(
                parent_frame, 
                text=f"{self.get_text('preview_error')} {str(e)}",
                fg="red", 
                bg="#e9ecef"
            )
            error_label.pack(pady=20)
            
            # Still show file info if possible
            try:
                file_size = os.path.getsize(file_path)
                info_text = f"{os.path.basename(file_path)}, {self.format_file_size(file_size)}"
                info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
                info_label.pack(pady=5)
            except:
                pass
    
    def _display_design_preview(self, parent_frame, img, file_path):
        """Helper function to display design file previews"""
        # Create a canvas for the image with scrollbars
        canvas_frame = tk.Frame(parent_frame)
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
        v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
        
        canvas = tk.Canvas(
            canvas_frame,
            xscrollcommand=h_scrollbar.set,
            yscrollcommand=v_scrollbar.set,
            bg="#ffffff"
        )
        
        h_scrollbar.config(command=canvas.xview)
        v_scrollbar.config(command=canvas.yview)
        
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Scale the image if necessary
        max_width, max_height = 700, 500
        img_width, img_height = img.size
        
        # Calculate scaled size while maintaining aspect ratio
        scale_factor = min(max_width/img_width, max_height/img_height)
        if scale_factor < 1:  # Only scale down, not up
            new_width = int(img_width * scale_factor)
            new_height = int(img_height * scale_factor)
            img = img.resize((new_width, new_height), get_pil_resize_method())
        
        # Convert to PhotoImage for Tkinter
        photo = ImageTk.PhotoImage(img)
        
        # Add image to canvas
        canvas.create_image(0, 0, image=photo, anchor=tk.NW)
        canvas.image = photo  # Keep a reference to prevent garbage collection
        
        # Configure canvas scrollable area
        canvas.config(scrollregion=canvas.bbox(tk.ALL))
        
        # Add info label with image details
        file_size = os.path.getsize(file_path)
        info_text = f"{img_width}x{img_height} px, {os.path.basename(file_path)}, {self.format_file_size(file_size)}"
        info_label = tk.Label(parent_frame, text=info_text, bg="#e9ecef", fg="#495057")
        info_label.pack(pady=5)



    
    def copy_filename_to_clipboard(self):
        """Copy the selected file name to clipboard"""
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection
            
        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")
        
        if not values:
            return  # No values found
            
        # Extract file name
        file_name = values[0]
        
        # Copy to clipboard
        self.root.clipboard_clear()
        self.root.clipboard_append(file_name)
        
        # Show a brief status message
        self.update_status("Copied to clipboard")
        
    def copy_filepath_to_clipboard(self):
        """Copy the selected file path to clipboard"""
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection
            
        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")
        
        if not values:
            return  # No values found
            
        # Extract file path
        file_path = values[2]
        
        # Copy to clipboard
        self.root.clipboard_clear()
        self.root.clipboard_append(file_path)
        
        # Show a brief status message
        self.update_status("Copied to clipboard")
        
    def on_drop(self, event):
        """Handle dropped files from external sources"""
        # Get the dropped file path(s)
        try:
            # Process the data - format depends on platform
            # In Windows, it will be in format "{C:/path/to/file}"
            # In Linux/Mac, it will be a normal file path
            data = event.data
            
            # Clean up the path
            if data.startswith('{') and data.endswith('}'):
                # Windows format
                data = data[1:-1]
            
            # Remove any quotes
            data = data.replace('"', '')
            
            paths = data.split()  # Multiple files are space-separated
            
            # Check if any of the paths are actually files
            valid_files = [path for path in paths if os.path.isfile(path)]
            
            if valid_files:
                # If files are dropped, we can open them directly (preview first file)
                self.create_file_preview_window(valid_files[0])
                
                # Show status message
                if len(valid_files) == 1:
                    self.update_status(f"Dosya önizleniyor: {os.path.basename(valid_files[0])}")
                else:
                    self.update_status(f"{len(valid_files)} dosya sürüklendi. İlk dosya önizleniyor.")
            
            # If a folder is dropped, update the folder selection
            valid_folders = [path for path in paths if os.path.isdir(path)]
            if valid_folders:
                # Use first valid folder
                self.folder_var.set(valid_folders[0])
                self.update_status(f"Klasör değiştirildi: {valid_folders[0]}")
                
                # Start folder loading
                self.load_files_thread()
        
        except Exception as e:
            logging.error(f"Error processing dropped files: {str(e)}")
            self.update_status(f"Sürüklenen dosyaları işlerken hata oluştu: {str(e)}")
        
    def open_file_location(self):
        """Open the location of the selected file in the file explorer"""
        # Get the selected item
        selected_items = self.file_tree.selection()
        if not selected_items:
            return  # No selection
            
        # Get the first selected item
        item = selected_items[0]
        # Get the values for this item
        values = self.file_tree.item(item, "values")
        
        if not values:
            return  # No values found
            
        # Extract file name, extension and directory path
        file_name = values[0]
        file_ext = values[1]
        dir_path = values[2]
        
        # Construct the full file path to find the correct directory
        # For Windows paths that already include filename, use as-is
        if os.path.basename(dir_path) == file_name:
            file_path = dir_path
        else:
            # Otherwise join directory and filename
            file_path = os.path.join(dir_path, file_name)
        
        # Open location using common method
        self.open_file_location_by_path(file_path)
        
    def open_file_location_by_path(self, file_path):
        """Open the location of a file based on its path"""
        # Check if the file exists
        if not os.path.exists(file_path):
            messagebox.showerror(
                self.get_text("error"),
                f"{file_path} not found."
            )
            return
            
        # Open the directory containing the file
        try:
            # Ensure we're using absolute path to avoid issues
            abs_file_path = os.path.abspath(file_path)
            
            if os.name == 'nt':  # Windows
                # Highlight the file in Windows Explorer
                subprocess.Popen(f'explorer /select,"{abs_file_path}"')
            elif sys.platform == 'darwin':  # macOS
                subprocess.Popen(['open', '-R', abs_file_path])
            else:  # Linux and other Unix variants
                # Get the directory containing the file - ensure it's absolute
                file_dir = os.path.dirname(abs_file_path)
                # Open the directory in the default file manager
                subprocess.Popen(['xdg-open', file_dir])
                
            # Log success
            logging.info(f"Opened location for file: {file_path}")
        except Exception as e:
            logging.error(f"Error opening file location: {str(e)}")
            messagebox.showerror(
                self.get_text("error"),
                f"{self.get_text('error_open_location')}: {str(e)}"
            )
            
    def open_file(self, file_path):
        """Open a file with the default associated application"""
        try:
            # Ensure we're using absolute path to avoid issues
            abs_file_path = os.path.abspath(file_path)
            
            if os.name == 'nt':  # Windows
                os.startfile(abs_file_path)
            elif sys.platform == 'darwin':  # macOS
                subprocess.Popen(['open', abs_file_path])
            else:  # Linux and other Unix variants
                subprocess.Popen(['xdg-open', abs_file_path])
                
            # Log success
            logging.info(f"Opened file: {file_path}")
        except Exception as e:
            logging.error(f"Error opening file: {str(e)}")
            messagebox.showerror(
                self.get_text("error"),
                f"{self.get_text('error_open_file')}: {str(e)}"
            )
            
    def open_website(self, url):
        """Open a website URL in the default browser"""
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception as e:
            logging.error(f"Error opening URL: {str(e)}")
            messagebox.showerror(
                self.get_text("error"),
                f"{self.get_text('error_open_url')}: {str(e)}"
            )
            
    def show_preview_context_menu(self, event, file_path):
        """Show context menu on right-click in the preview mode"""
        # Store the current file path for context menu actions
        self.current_preview_file_path = file_path
        
        # Show the context menu
        try:
            self.preview_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            # Make sure to release the grab
            self.preview_context_menu.grab_release()
    
    def open_preview_file(self):
        """Open the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.open_file(self.current_preview_file_path)
            
    def open_preview_file_location(self):
        """Open the location of the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.open_file_location_by_path(self.current_preview_file_path)
            
    def copy_preview_filename_to_clipboard(self):
        """Copy the filename of the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            file_name = os.path.basename(self.current_preview_file_path)
            self.root.clipboard_clear()
            self.root.clipboard_append(file_name)
            self.update_status(f"{self.get_text('copied_to_clipboard')}: {file_name}")
            
    def copy_preview_filepath_to_clipboard(self):
        """Copy the file path of the file that was right-clicked in preview mode"""
        if hasattr(self, 'current_preview_file_path') and os.path.isfile(self.current_preview_file_path):
            self.root.clipboard_clear()
            self.root.clipboard_append(self.current_preview_file_path)
            self.update_status(f"{self.get_text('copied_to_clipboard')}: {self.current_preview_file_path}")


if __name__ == "__main__":
    root = tk.Tk()
    app = FileManagerApp(root)
    root.mainloop()
